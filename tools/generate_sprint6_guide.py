"""Generates 'Sprint 6 - Implementation Guide.docx' covering the fees,
invoices, receipts, and bursar-workflow domain delivered in sprint 6
(FeeSchedule, FeeScheduleItem, Invoice, InvoiceLine, Payment,
PaymentAllocation, the four finance lookup tables, the matching services,
and the Razor pages that drive the schedule -> invoice -> receipt
pipeline).

Long-form edition. Code blocks embed actual source files so the guide
stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint6_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 6 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260512224258_FeesInvoicesPayments.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 6 — Fees, Invoices, Receipts & Bursar Workflows")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Fee schedules · Invoice issuance · Payment recording · Receipts · Bursar dashboard")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/6-fees-bursar")
    meta.add_run("\nBuilt on: Sprints 1–5b (identity, academic domain, students & parents, attendance, results & report cards, pupil photos)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 6 in context", 1)
    add_para(doc,
        "Sprint 6 lays the financial spine of the school. The first five "
        "sprints established who the people are, what the calendar looks "
        "like, which pupil is in which class, who showed up on which day, "
        "and how each pupil performed. Sprint 6 turns its attention to the "
        "exchange of money that pays for all of that: the fee schedule the "
        "head teacher publishes at the start of every term, the per-pupil "
        "invoice that drops into the bursar's queue when the term begins, "
        "the cash, transfer, POS, cheque, mobile-money or online payment "
        "that comes back in, and the receipts that go home in the school "
        "bag.")
    add_para(doc,
        "Functionally, the sprint introduces a three-stage pipeline. Stage "
        "one is the schedule: a SuperAdmin or HeadTeacher (or the bursar "
        "herself) creates one FeeSchedule per (Term, ClassLevel) and lists "
        "its line items, each one keyed to a FeeCategory (Tuition, "
        "Development Levy, Examination, Books, Uniforms, Transport, Meals, "
        "Boarding, PTA Levy, Other) with a published amount. Stage two is "
        "issuance: a single click on a published schedule fans out into "
        "one Invoice per actively-enrolled pupil in the chosen class, each "
        "invoice carrying a copy of the schedule's items so the bill stays "
        "stable even if the schedule is later refined. Stage three is "
        "collection: the bursar records a Payment against the pupil, "
        "allocates it across one or more outstanding invoices, and the "
        "system stamps the receipt number, updates each invoice's status, "
        "and produces a printable receipt view.")
    add_para(doc,
        "Once this sprint ships, the school has a complete-enough money "
        "loop to operate without spreadsheets. The bursar dashboard "
        "summarises total invoiced, total collected, total outstanding, "
        "and a per-method and per-category breakdown for the active term. "
        "The Family → Students profile and the Report Card detail page "
        "will, in a later sprint, surface the same data on the pupil "
        "side; the model is already in place.")
    add_para(doc,
        "This document is a long-form implementation guide in the tone of "
        "the sprint 5 guide. An engineer who has read sprints 1–5b and has "
        "the codebase checked out can recreate every change here without "
        "referring to the diff. The structure mirrors the build order: "
        "design decisions first, Domain entities next, Application "
        "contracts after that, Infrastructure (DbContext, services, "
        "seeder, migration) in the middle, then the Razor UI and "
        "navigation. Smoke-test, troubleshooting, and forward-"
        "compatibility chapters round it off.")

    add_heading(doc, "1.1 Where this sits relative to sprint 5b", 2)
    add_para(doc, "Every load-bearing piece of the earlier sprints is reused:")
    add_bullets(doc, [
        "BaseEntity — every new entity in this sprint inherits Guid Id, "
        "IAuditable, and ISoftDelete from it.",
        "ApplicationDbContext.SaveChanges — the override stamps audit "
        "columns and rewrites Delete to IsDeleted = true. The finance "
        "rows therefore inherit the same audit and soft-delete machinery "
        "as everything else.",
        "Global query filters — every new entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries automatically.",
        "OperationResult / OperationResult<T> — every new service uses "
        "this for predictable success/failure responses.",
        "ILookupService — already had twenty methods. Sprint 6 adds four "
        "more (fee categories, payment methods, invoice statuses, payment "
        "statuses) without rewriting the existing ones.",
        "Student, Term, SchoolClass, ClassLevel — pick up collection "
        "navigations only. No scalar columns change on any of those "
        "tables, so existing pages still render unchanged.",
        "StudentAvatar (sprint 5b) — already lives in Components.Shared "
        "and is reused on every invoice and payment row so a pupil's "
        "face follows them through the finance views.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_numbered(doc, [
        "Create a FeeSchedule for any (Term, ClassLevel), give it a "
        "title and optional notes, then add line items keyed to a "
        "FeeCategory.",
        "Publish a schedule. Once published, the schedule and its items "
        "are locked from edits; unpublishing is allowed only while no "
        "invoices have been issued from it.",
        "From the Issue Invoices page, pick a published schedule and a "
        "matching class. The system creates one Invoice per actively-"
        "enrolled pupil with a copy of every schedule line item, assigns "
        "a sequential InvoiceNumber (NPS/INV/<year>/<seq>), and stamps "
        "the issued-on date.",
        "Apply a per-line discount on an invoice. Subtotal, AmountDue "
        "and the invoice status are recomputed automatically.",
        "Record a Payment for a pupil, pick a method (Cash, Bank "
        "Transfer, POS, Cheque, Mobile Money, Online Payment), enter the "
        "amount, optionally an auto-allocate across outstanding invoices "
        "oldest-first, or hand-allocate row-by-row.",
        "Open a payment to see the receipt: receipt number, method, "
        "reference, allocations, and the running balance each invoice "
        "moved to.",
        "Refund a payment. Allocations are released, the receipt stays "
        "on file with a 'Refunded' badge, and the touched invoices flip "
        "back to their pre-payment state.",
        "Open the Bursar dashboard for a term to see total invoiced, "
        "total collected, total outstanding, and the breakdown of "
        "collections by payment method and invoiced amounts by fee "
        "category.",
        "Soft-delete any of the above with guards: cannot delete a "
        "schedule that has issued invoices, cannot delete an invoice "
        "with payments applied, cannot delete a payment that still has "
        "allocations.",
    ])

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_bullets(doc, [
        "Printable PDF invoices and receipts. The DTOs carry every "
        "field a layout would need; the Razor pages already render a "
        "human-readable receipt view. A QuestPDF-driven endpoint is a "
        "follow-on lift.",
        "Payment gateway integration. Cards, USSD, and online "
        "checkout would each plug in as a new PaymentMethod row plus "
        "a webhook endpoint. The internal model — Payment + "
        "PaymentAllocation — is ready to receive callbacks unchanged.",
        "Multi-currency. Amounts are stored as decimal(12,2). The "
        "school is single-currency (NGN) today; a Currency lookup "
        "would be added by future maintainers without schema gymnastics.",
        "Scholarships and waivers as a first-class concept. Today the "
        "discount column on InvoiceLine handles per-pupil reductions; a "
        "dedicated Scholarship table could later be introduced.",
        "Parent / student visibility. Parents and pupils still see "
        "placeholder navigation. The portal sprint will expose existing "
        "invoices and the unallocated payments using the Invoice.Notes "
        "+ Payment.ReceiptNumber the bursar already produces.",
        "Aged-debt reports and dunning notices. The data the report "
        "needs (DueDate + Balance per pupil) is in place; the report "
        "itself is a feature-add.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers:")
    add_bullets(doc, [
        "10 new domain entities under src/NaijaPrimeSchool.Domain/Finance/.",
        "4 collection navigations on existing entities (Student, Term, "
        "SchoolClass, ClassLevel).",
        "4 DTO files under src/NaijaPrimeSchool.Application/Finance/Dtos/.",
        "3 new service contracts under src/NaijaPrimeSchool.Application/Finance/.",
        "3 service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "4 new methods on ILookupService (and the matching LookupService).",
        "1 EF Core migration introducing 10 new tables and the indexes "
        "that go with them.",
        "1 DatabaseInitializer extension seeding FeeCategories, "
        "PaymentMethods, InvoiceStatuses and PaymentStatuses.",
        "8 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Finance/.",
        "1 navigation menu rewrite: the previously-disabled 'Finance' "
        "placeholder is replaced with a six-item panel gated to "
        "SuperAdmin + HeadTeacher + SchoolBursar.",
    ])
    add_para(doc,
        "Everything compiles with zero warnings on .NET 10. The code "
        "follows the patterns already accepted in sprints 1–5b.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)

    add_heading(doc, "2.1 Schedule → Invoice → Payment, three tables", 2)
    add_para(doc,
        "The pipeline could in theory be collapsed: 'the bursar enters a "
        "payment, the system computes the bill from a schedule, sends the "
        "receipt.' Real schools do not work that way. The bill goes home "
        "in week one of term so the family can plan; the payment trickles "
        "in over weeks; receipts go back as evidence. Persisting all "
        "three stages — schedule, invoice, payment — costs four extra "
        "tables (schedule, item, line, allocation) and earns back, in "
        "exchange:")
    add_bullets(doc, [
        "Stability. A schedule edited after invoices have been issued "
        "does not retroactively change the invoiced amounts.",
        "Auditability. Each invoice carries a copy of the schedule item "
        "it came from, plus a FeeScheduleItemId pointing back at the "
        "source — so a future audit can answer 'where did this line "
        "come from?'.",
        "Bookkeeping. A payment can sit unallocated as credit on the "
        "pupil's ledger, or fan across several invoices via "
        "PaymentAllocation rows; the schema does not force the bursar "
        "into a one-to-one shape that real life does not respect.",
    ])

    add_heading(doc, "2.2 No enums — four lookup tables", 2)
    add_para(doc,
        "The rule from earlier sprints holds. FeeCategory, PaymentMethod, "
        "InvoiceStatus, and PaymentStatus are all proper entities "
        "derived from BaseEntity, seeded on first run, and editable from "
        "the database without a redeploy. Each one carries a Code column "
        "(TUI, CASH, ISSUED, CONFIRMED, …) that the service layer keys "
        "off when it transitions a row between states — so renaming "
        "'Issued' to 'Active' in the row would not break logic.")

    add_heading(doc, "2.3 Decimal precision, not floats", 2)
    add_para(doc,
        "Every money column is decimal(12,2) — twelve significant digits "
        "before the decimal point is plenty for a single school's "
        "invoices (₦9,999,999,999.99) and two after preserves kobo "
        "resolution. EF Core's HasPrecision attribute is used on every "
        "Amount, Discount, AmountDue, AmountPaid, and AmountApplied "
        "column. Float would have introduced rounding errors that "
        "compound through allocations.")

    add_heading(doc, "2.4 Receipt and invoice numbering", 2)
    add_para(doc,
        "Both InvoiceNumber and ReceiptNumber follow a "
        "year-prefixed pattern: NPS/INV/<year>/<4-digit-sequence> and "
        "NPS/RCP/<year>/<4-digit-sequence>. The services compute the "
        "next sequence by reading the largest existing number for the "
        "year and adding one. The cost is one extra query per issuance "
        "run; the benefit is that auditors can read a number and know, "
        "without lookups, which year it belongs to. The format is "
        "constants in the service classes and easy to change.")

    add_heading(doc, "2.5 Allocations as their own table", 2)
    add_para(doc,
        "PaymentAllocation is a separate entity with a unique "
        "(PaymentId, InvoiceId) index. A payment can apply to many "
        "invoices; an invoice can be paid down by many payments. The "
        "alternative — putting an InvoiceId column on Payment — does "
        "not allow either, and falls apart the first time a parent pays "
        "₦100,000 toward two children's fees in one transfer.")

    add_heading(doc, "2.6 Soft delete, again, with guards", 2)
    add_para(doc, "Every new entity implements ISoftDelete via BaseEntity. "
        "What is new is the operation guards:")
    add_bullets(doc, [
        "FeeSchedule cannot be edited or items added/removed while it is "
        "published. Unpublish first.",
        "FeeSchedule cannot be unpublished if any invoice has ever been "
        "issued from it — the auditable link from invoice line back to "
        "the schedule item would otherwise become misleading.",
        "FeeSchedule cannot be deleted if it has been used to issue "
        "invoices.",
        "Invoice cannot be edited (lines, discounts) once cancelled.",
        "Invoice cannot be cancelled while payments are applied; refund "
        "those first.",
        "Invoice cannot be deleted while payments are applied.",
        "Payment cannot be deleted while it has allocations; refund first "
        "so the allocations are released, then delete.",
    ])

    add_heading(doc, "2.7 Unique indexes", 2)
    add_bullets(doc, [
        "Unique (TermId, ClassLevelId) on FeeSchedule — at most one "
        "schedule per (term, level). A school cannot accidentally "
        "publish two competing schedules.",
        "Unique InvoiceNumber, unique ReceiptNumber — auditable.",
        "Unique (PaymentId, InvoiceId) on PaymentAllocation — a payment "
        "cannot 'double-pay' the same invoice in a single transaction.",
        "Plus unique Name/Code on every lookup table.",
    ])

    add_heading(doc, "2.8 Status recomputation on every money change", 2)
    add_para(doc,
        "InvoiceService.RecomputeInvoiceTotalsAsync is the single seam "
        "where Subtotal, DiscountTotal, AmountDue, AmountPaid, and "
        "InvoiceStatusId are brought into agreement with the underlying "
        "lines and allocations. It is called from "
        "SetLineDiscountAsync after a discount changes and from "
        "PaymentService.RecordAsync / RefundAsync after allocations "
        "are created or released. A CANCELLED invoice is never "
        "automatically demoted from cancellation — the override is "
        "explicit in the helper.")

    add_heading(doc, "2.9 PaymentService depends on the concrete InvoiceService", 2)
    add_para(doc,
        "RecomputeInvoiceTotalsAsync is an internal helper, not part of "
        "the IInvoiceService contract. PaymentService takes the "
        "concrete InvoiceService in its constructor so it can call the "
        "helper directly. DI registers InvoiceService twice — once as "
        "the concrete class for PaymentService, once behind "
        "IInvoiceService for the UI. The pattern keeps the public "
        "service surface small without forcing a public 'Recompute' "
        "method that any caller could misuse.")

    add_heading(doc, "2.10 Inline forms, the same UI rhythm", 2)
    add_para(doc,
        "Every Razor page in this sprint follows the same shape as the "
        "earlier sprints: a filter card on top, a paged grid in the "
        "middle, an inline RadzenCard form below the grid for new-or-"
        "edit, and a confirm-then-act dialog for destructive operations. "
        "Two pages — IssueInvoices and RecordPayment — are dedicated "
        "full-page forms because their workflows have enough fields and "
        "side-effects that an inline form would have been cramped.")

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)
    add_para(doc,
        "Every sprint-6 entity lives in a single new folder, "
        "src/NaijaPrimeSchool.Domain/Finance/. Entities are anaemic — no "
        "domain methods, no validation logic. Validation lives in DTOs "
        "(DataAnnotations) and services.")

    add_heading(doc, "3.1 Folder layout", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "├── Finance/                       <- (new in sprint 6)\n"
        "│   ├── FeeCategory.cs             <- lookup\n"
        "│   ├── PaymentMethod.cs           <- lookup\n"
        "│   ├── InvoiceStatus.cs           <- lookup\n"
        "│   ├── PaymentStatus.cs           <- lookup\n"
        "│   ├── FeeSchedule.cs             <- per (term, class level)\n"
        "│   ├── FeeScheduleItem.cs         <- line item on a schedule\n"
        "│   ├── Invoice.cs                 <- bill issued to a pupil\n"
        "│   ├── InvoiceLine.cs             <- line item on an invoice\n"
        "│   ├── Payment.cs                 <- receipt header\n"
        "│   └── PaymentAllocation.cs       <- payment x invoice split\n"
        "├── Results/                       <- from sprint 5\n"
        "├── Attendance/                    <- from sprint 4\n"
        "├── Family/                        <- from sprint 3\n"
        "├── Academics/                     <- from sprint 2\n"
        "├── Common/                        <- from sprint 1\n"
        "└── Identity/                      <- from sprint 1\n")

    add_heading(doc, "3.2 The four lookup entities", 2)
    add_heading(doc, "3.2.1 FeeCategory.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/FeeCategory.cs")
    add_heading(doc, "3.2.2 PaymentMethod.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/PaymentMethod.cs")
    add_heading(doc, "3.2.3 InvoiceStatus.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/InvoiceStatus.cs")
    add_heading(doc, "3.2.4 PaymentStatus.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/PaymentStatus.cs")

    add_heading(doc, "3.3 The six core entities", 2)
    add_heading(doc, "3.3.1 FeeSchedule.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/FeeSchedule.cs")
    add_heading(doc, "3.3.2 FeeScheduleItem.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/FeeScheduleItem.cs")
    add_heading(doc, "3.3.3 Invoice.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/Invoice.cs")
    add_para(doc,
        "Balance is a computed property — not persisted. It is just "
        "AmountDue minus AmountPaid for the convenience of callers; "
        "the persisted Subtotal/DiscountTotal/AmountDue/AmountPaid "
        "fields are kept in sync by RecomputeInvoiceTotalsAsync.")
    add_heading(doc, "3.3.4 InvoiceLine.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/InvoiceLine.cs")
    add_heading(doc, "3.3.5 Payment.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/Payment.cs")
    add_heading(doc, "3.3.6 PaymentAllocation.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Finance/PaymentAllocation.cs")

    add_heading(doc, "3.4 Relationships at a glance", 2)
    add_code(doc,
        "                       FeeCategory (lookup)\n"
        "                                |\n"
        "                                v 1..N\n"
        "    Term --------+------- FeeSchedule -------+\n"
        "                 |              | 1          |\n"
        "    ClassLevel --+              |            |\n"
        "                                v N          |\n"
        "                          FeeScheduleItem ---+\n"
        "                                |\n"
        "                                | source\n"
        "                                v\n"
        "    Student + Term + SchoolClass --> Invoice ---+\n"
        "                                          |     |\n"
        "                                          |     v N\n"
        "                              InvoiceLine <-+   v 1..N\n"
        "                                          |\n"
        "                                          v N\n"
        "                                   PaymentAllocation N <-- Payment <-- Student\n"
        "                                                                |\n"
        "                                                                v\n"
        "                                                       PaymentMethod (lookup)\n"
        "                                                       PaymentStatus (lookup)\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)

    add_heading(doc, "4.1 FeeScheduleDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/Dtos/FeeScheduleDtos.cs")
    add_heading(doc, "4.2 InvoiceDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/Dtos/InvoiceDtos.cs")
    add_para(doc,
        "InvoiceDto carries the sprint-5b photo plumbing "
        "(StudentPhotoUrl, StudentFirstName, StudentLastName) so the "
        "<StudentAvatar /> component renders pupil photographs on the "
        "invoice list and detail without a second round-trip.")
    add_heading(doc, "4.3 PaymentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/Dtos/PaymentDtos.cs")
    add_heading(doc, "4.4 StudentLedgerDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/Dtos/StudentLedgerDtos.cs")

    add_heading(doc, "4.5 Service contracts", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/IFeeScheduleService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/IInvoiceService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Finance/IPaymentService.cs")

    add_heading(doc, "4.6 ILookupService extension", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Application/Users/ILookupService.cs",
        "Task<IReadOnlyList<LookupDto>> GetFeeCategoriesAsync",
        end_marker="Task<IReadOnlyList<LookupDto>> GetPaymentStatusesAsync",
        caption="Excerpt — ILookupService.cs (sprint 6 additions)")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)

    add_heading(doc, "5.1 Ten new DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<FeeCategory>",
        end_marker="public DbSet<PaymentAllocation>",
        lines_after_start=11,
        caption="Excerpt — the ten finance DbSets")

    add_heading(doc, "5.2 ConfigureFinance", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureFinance",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureFinance")
    add_para(doc, "Highlights:")
    add_bullets(doc, [
        "HasPrecision(12, 2) on every money column.",
        "Cascade delete on FeeScheduleItem.FeeScheduleId, "
        "InvoiceLine.InvoiceId, PaymentAllocation.PaymentId — the "
        "child rows always belong to their parent.",
        "Restrict on every cross-aggregate FK (StudentId, TermId, "
        "SchoolClassId, ClassLevelId, FeeCategoryId, InvoiceId on the "
        "allocation side, …) so a schema-level constraint reinforces "
        "the soft-delete guards in the service layer.",
        "Unique indexes on Invoice.InvoiceNumber, Payment.ReceiptNumber, "
        "FeeSchedule (TermId, ClassLevelId), and PaymentAllocation "
        "(PaymentId, InvoiceId).",
        "b.Ignore(i => i.Balance) and b.Ignore(l => l.LineTotal) so EF "
        "Core does not try to materialise the computed properties as "
        "columns.",
    ])

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)

    add_heading(doc, "6.1 FeeScheduleService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/FeeScheduleService.cs")

    add_heading(doc, "6.2 InvoiceService.cs", 2)
    add_para(doc,
        "InvoiceService owns the schedule-to-invoice fan-out, per-line "
        "discount application, cancellation, and the per-student "
        "ledger. The internal helper RecomputeInvoiceTotalsAsync is "
        "shared with PaymentService via the concrete-class injection "
        "described in chapter 2.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/InvoiceService.cs")

    add_heading(doc, "6.3 PaymentService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/PaymentService.cs")

    add_heading(doc, "6.4 DI registration", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IFeeScheduleService",
        end_marker="return services;",
        caption="Excerpt — DependencyInjection.cs (sprint 6 additions)")

    add_heading(doc, "6.5 LookupService — four new methods", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",
        "public async Task<IReadOnlyList<LookupDto>> GetFeeCategoriesAsync",
        lines_after_start=30,
        caption="Excerpt — LookupService.cs (sprint 6 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_code(doc,
        "dotnet ef migrations add FeesInvoicesPayments \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the finance lookups", 1)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedFinanceLookupsAsync",
        end_marker="private static async Task SeedResultsLookupsAsync",
        caption="Excerpt — SeedFinanceLookupsAsync")
    add_para(doc, "What gets seeded:")
    add_bullets(doc, [
        "FeeCategories — Tuition (TUI), Development Levy (DEV), "
        "Examination (EXAM), Books (BOOK), Uniforms (UNIF), Transport "
        "(TRAN), Meals (MEAL), Boarding (BRD), PTA Levy (PTA), Other "
        "(OTH). Each has an IsMandatoryByDefault flag.",
        "PaymentMethods — Cash (CASH), Bank Transfer (BANK), POS "
        "(POS), Cheque (CHQ), Mobile Money (MMNY), Online Payment "
        "(ONL). Each has a RequiresReference flag that the UI can use "
        "to demand a reference number.",
        "InvoiceStatuses — Draft (DRAFT), Issued (ISSUED), "
        "Partially Paid (PARTIAL), Paid (PAID), Overdue (OVERDUE), "
        "Cancelled (CANCELLED). The service layer keys off Code, not "
        "Id, when transitioning.",
        "PaymentStatuses — Pending (PENDING), Confirmed (CONFIRMED), "
        "Bounced (BOUNCED), Refunded (REFUNDED).",
    ])

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Finance/\n"
        "├── FeeSchedules.razor           <- /fees\n"
        "├── FeeScheduleDetail.razor      <- /fees/{id}\n"
        "├── Invoices.razor               <- /invoices\n"
        "├── IssueInvoices.razor          <- /invoices/issue\n"
        "├── InvoiceDetail.razor          <- /invoices/{id}\n"
        "├── Payments.razor               <- /payments\n"
        "├── RecordPayment.razor          <- /payments/new\n"
        "├── PaymentDetail.razor          <- /payments/{id}\n"
        "└── FinanceDashboard.razor       <- /finance\n")
    add_para(doc,
        "Every page is gated to SuperAdmin + HeadTeacher + SchoolBursar. "
        "Sprint 6 is the first sprint where the SchoolBursar role moves "
        "from placeholder navigation to a fully realised workspace.")

    add_heading(doc, "9.1 FeeSchedules.razor — the schedule list", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/FeeSchedules.razor")

    add_heading(doc, "9.2 FeeScheduleDetail.razor — schedule + items", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/FeeScheduleDetail.razor")

    add_heading(doc, "9.3 Invoices.razor — the invoice list", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/Invoices.razor")

    add_heading(doc, "9.4 IssueInvoices.razor — fan-out from a schedule", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/IssueInvoices.razor")

    add_heading(doc, "9.5 InvoiceDetail.razor — the invoice view", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/InvoiceDetail.razor")

    add_heading(doc, "9.6 Payments.razor — the receipts list", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/Payments.razor")

    add_heading(doc, "9.7 RecordPayment.razor — the bursar's collection workflow", 2)
    add_para(doc,
        "The most complex page in the sprint. Picking a pupil reloads "
        "their ledger of outstanding invoices; the bursar then either "
        "auto-allocates oldest-first or hand-types per-row amounts. The "
        "total allocated, the amount unallocated, and a live running tally "
        "are visible above the table so the bursar can sanity-check the "
        "math before saving.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/RecordPayment.razor")

    add_heading(doc, "9.8 PaymentDetail.razor — the receipt view", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/PaymentDetail.razor")

    add_heading(doc, "9.9 FinanceDashboard.razor — bursar at a glance", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Finance/FinanceDashboard.razor")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and authorization", 1)
    add_para(doc,
        "The previously-disabled 'Finance' nav placeholder is replaced "
        "with a six-item panel: Bursar dashboard, Fee schedules, "
        "Invoices, Issue invoices, Payments, Record payment. The panel "
        "is wrapped in an AuthorizeView that grants access to "
        "SuperAdmin, HeadTeacher, and SchoolBursar.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Finance\"",
        end_marker="</AuthorizeView>",
        caption="Excerpt — NavMenu.razor")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Finance",
        lines_after_start=2,
        caption="Excerpt — _Imports.razor")

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of a finance flow", 1)

    add_heading(doc, "11.1 Publishing a fee schedule", 2)
    add_bullets(doc, [
        "HeadTeacher opens /fees, clicks New schedule.",
        "Picks the current Term and the target ClassLevel.",
        "Adds line items (Tuition ₦80,000; Development Levy ₦15,000; "
        "Examination ₦5,000; Books ₦12,000).",
        "Clicks Publish. The service refuses if no items exist; "
        "otherwise IsPublished flips to true and the schedule is ready "
        "to be used.",
    ])

    add_heading(doc, "11.2 Issuing invoices", 2)
    add_bullets(doc, [
        "Open /invoices/issue.",
        "Pick the term, the published schedule, the matching class, an "
        "issued-on date, and (optionally) a due date.",
        "Click Issue invoices. InvoiceService loads every actively-"
        "enrolled pupil in the class, skips pupils that already have an "
        "invoice in that (term, class), and creates a fresh Invoice + "
        "InvoiceLine[] for everyone else.",
        "Invoice numbers are NPS/INV/<year>/<seq>; sequence starts from "
        "the highest existing number for the year + 1.",
        "The Issued status is stamped from the seeded InvoiceStatus "
        "row with Code = 'ISSUED'.",
    ])

    add_heading(doc, "11.3 Recording a payment", 2)
    add_bullets(doc, [
        "A parent comes to the bursar's window with ₦60,000 cash.",
        "Bursar opens /payments/new, picks the pupil from the dropdown.",
        "The ledger of outstanding invoices loads in the table.",
        "Bursar enters method Cash, amount 60,000.",
        "Clicks Auto-allocate oldest first — the service walks the "
        "outstanding invoices oldest-first and allocates the cash row "
        "by row.",
        "Clicks Save payment. PaymentService validates allocations "
        "against each invoice's balance, generates the receipt number "
        "(NPS/RCP/<year>/<seq>), and calls "
        "InvoiceService.RecomputeInvoiceTotalsAsync for every touched "
        "invoice — flipping ISSUED to PARTIAL or PAID as appropriate.",
        "User lands on /payments/{id} — the receipt view.",
    ])

    add_heading(doc, "11.4 Refunding a bounced cheque", 2)
    add_bullets(doc, [
        "Bursar opens the original receipt, clicks Refund.",
        "PaymentService deletes the allocations, flips the payment "
        "status to REFUNDED, and recomputes the touched invoices.",
        "Invoices return to PARTIAL or ISSUED status depending on what "
        "other payments may still be applied.",
        "The receipt row stays on file with the Refunded badge for the "
        "audit trail.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "First run applies the FeesInvoicesPayments migration and "
        "seeds the four new lookup tables. Sign in as the SuperAdmin.")

    add_heading(doc, "12.2 Verify navigation and lookups", 2)
    add_bullets(doc, [
        "The Finance panel is now an active six-item dropdown rather "
        "than a disabled placeholder.",
        "SELECT Name, Code FROM FeeCategories ORDER BY DisplayOrder; "
        "shows ten rows.",
        "SELECT Name, Code FROM PaymentMethods ORDER BY DisplayOrder; "
        "shows six rows.",
        "SELECT Name, Code FROM InvoiceStatuses ORDER BY DisplayOrder; "
        "shows six rows.",
        "SELECT Name, Code FROM PaymentStatuses ORDER BY DisplayOrder; "
        "shows four rows.",
    ])

    add_heading(doc, "12.3 End-to-end happy path", 2)
    add_numbered(doc, [
        "Finance → Fee schedules → New. Term: current. Class level: "
        "Primary 1. Title: 'Primary 1 — First Term fees'. Save.",
        "Add three items (Tuition 80,000; Development Levy 15,000; "
        "Examination 5,000).",
        "Click Publish.",
        "Finance → Issue invoices. Pick the term, the schedule, a "
        "Primary 1 class with enrolled pupils. Issue.",
        "Finance → Invoices. Confirm one invoice per actively-enrolled "
        "pupil in the chosen class. Each should show "
        "AmountDue = 100,000 and balance = 100,000.",
        "Click Record payment on any invoice row.",
        "Method: Cash. Amount: 50,000. Auto-allocate. Save.",
        "Open the resulting receipt. Status should be Confirmed; the "
        "invoice it touched should now show 50,000 / 100,000 with a "
        "Partial badge.",
        "Open /finance. Totals should reflect the issuance and the "
        "payment.",
    ])

    add_heading(doc, "12.4 Error paths", 2)
    add_numbered(doc, [
        "Try to publish a schedule with no items. Refused.",
        "Try to unpublish a schedule after invoices have been issued. "
        "Refused.",
        "Try to record a payment whose allocations exceed the amount. "
        "Refused.",
        "Try to record an allocation that exceeds the invoice's "
        "outstanding balance. Refused.",
        "Refund a payment, then try to record allocations against "
        "the now-released invoice — the balance should reflect the "
        "released amount.",
    ])

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)

    add_heading(doc, "13.1 'A fee schedule already exists for this (term, class level)'", 2)
    add_para(doc,
        "The composite unique index on FeeSchedule.(TermId, ClassLevelId) "
        "is a deliberate guard: a school should not accidentally publish "
        "two competing schedules for the same audience. If you need to "
        "change a schedule, open the existing one and edit it.")

    add_heading(doc, "13.2 'Publish the fee schedule before issuing invoices from it'", 2)
    add_para(doc,
        "Draft schedules are deliberately not allowed to fan out into "
        "invoices — the audit trail would otherwise be confusing. "
        "Click Publish on the schedule first.")

    add_heading(doc, "13.3 'No actively-enrolled pupils in the selected class'", 2)
    add_para(doc,
        "An enrolment counts as active when WithdrawnOn is null. If "
        "every pupil's enrolment has been withdrawn, the issuance "
        "produces no invoices. Re-enrol pupils, or use the family "
        "pages to confirm the class's roster.")

    add_heading(doc, "13.4 Allocations exceed the payment amount", 2)
    add_para(doc,
        "Both the UI and the service layer guard against this. The "
        "live running total on RecordPayment shows the math; if "
        "TotalAllocated > Amount, the save button is still clickable but "
        "the service rejects the request with a friendly message.")

    add_heading(doc, "13.5 Cannot delete a payment", 2)
    add_para(doc,
        "If the payment still has allocations, soft-delete is refused. "
        "Refund it first — that releases the allocations and flips the "
        "status to REFUNDED — then delete. The receipt history is "
        "preserved either way.")

    add_heading(doc, "13.6 'Invoice statuses are not seeded' / 'Cancelled status is not seeded'", 2)
    add_para(doc,
        "The service keys off the seeded Code values. If you cleared "
        "the InvoiceStatuses table by hand, restart the app so "
        "DatabaseInitializer re-seeds it.")

    add_heading(doc, "13.7 The IdentityRole migration warning", 2)
    add_para(doc,
        "Same pre-existing warning that has accompanied every sprint "
        "since sprint 1. Harmless at runtime; will be addressed in a "
        "future role-administration sprint.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_bullets(doc, [
        "Parent and student portals (later sprint) will read Invoices "
        "and Payments via the existing IInvoiceService and "
        "IPaymentService contracts. The IsPublished flag does not "
        "exist on invoices — schools issue immediately — but the "
        "InvoiceStatus.Code is the gate the portal will key off.",
        "Aged debt reports can be added without schema changes: a new "
        "DTO + a service method that groups Invoices by (PupilId, "
        "OutstandingBalance, DaysOverdue).",
        "Online payments gateway: add a new PaymentMethod row, plus a "
        "minimal-API endpoint that POSTs back from the gateway, calls "
        "PaymentService.RecordAsync with the same shape any in-person "
        "payment uses, and emails the receipt URL to the parent.",
        "Discounts beyond the per-line discount column can grow into "
        "a dedicated Scholarship table without breaking existing "
        "invoices.",
        "Multi-currency: a Currency table plus a CurrencyId on Invoice "
        "and Payment is the cleanest path; the decimal precision is "
        "ready.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 6", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Finance/FeeCategory.cs",     "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Finance/PaymentMethod.cs",   "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Finance/InvoiceStatus.cs",   "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Finance/PaymentStatus.cs",   "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Finance/FeeSchedule.cs",     "Per-(term, level) schedule."),
        ("src/NaijaPrimeSchool.Domain/Finance/FeeScheduleItem.cs", "Line item on a schedule."),
        ("src/NaijaPrimeSchool.Domain/Finance/Invoice.cs",         "Pupil bill."),
        ("src/NaijaPrimeSchool.Domain/Finance/InvoiceLine.cs",     "Line item on an invoice."),
        ("src/NaijaPrimeSchool.Domain/Finance/Payment.cs",         "Receipt header."),
        ("src/NaijaPrimeSchool.Domain/Finance/PaymentAllocation.cs","Payment x invoice split."),
        ("Domain layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Domain/Family/Student.cs",          "Added Invoices + Payments."),
        ("src/NaijaPrimeSchool.Domain/Academics/Term.cs",          "Added FeeSchedules + Invoices."),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs",   "Added Invoices."),
        ("src/NaijaPrimeSchool.Domain/Academics/ClassLevel.cs",    "Added FeeSchedules."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Finance/Dtos/FeeScheduleDtos.cs",   "Schedule DTOs."),
        ("src/NaijaPrimeSchool.Application/Finance/Dtos/InvoiceDtos.cs",       "Invoice DTOs."),
        ("src/NaijaPrimeSchool.Application/Finance/Dtos/PaymentDtos.cs",       "Payment / allocation DTOs."),
        ("src/NaijaPrimeSchool.Application/Finance/Dtos/StudentLedgerDtos.cs", "Student ledger + dashboard DTOs."),
        ("src/NaijaPrimeSchool.Application/Finance/IFeeScheduleService.cs",    "Schedule service contract."),
        ("src/NaijaPrimeSchool.Application/Finance/IInvoiceService.cs",        "Invoice service contract."),
        ("src/NaijaPrimeSchool.Application/Finance/IPaymentService.cs",        "Payment service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",           "Added 4 new lookup methods."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/FeeScheduleService.cs", "Schedule + items CRUD."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/InvoiceService.cs",     "Issue, discount, cancel, ledger, status recompute."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/PaymentService.cs",     "Record, refund, dashboard summary."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 10 tables."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",         "Registered 3 new services (InvoiceService twice)."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs", "Added 10 DbSets, ConfigureFinance."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs", "Seeded the 4 new lookup tables."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",      "Added 4 new lookup methods."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/FeeSchedules.razor",      "Schedule list + new form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/FeeScheduleDetail.razor", "Edit schedule + items."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/Invoices.razor",          "Invoice list."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/IssueInvoices.razor",     "Fan out from a schedule."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/InvoiceDetail.razor",     "Invoice view with discount edit."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/Payments.razor",          "Payments list."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/RecordPayment.razor",     "Pupil ledger + record + allocate."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/PaymentDetail.razor",     "Receipt view + refund."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Finance/FinanceDashboard.razor",  "Bursar summary."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                        "Added Finance + Finance.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                  "Replaced disabled Finance placeholder with full panel."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint6_guide.py",                                           "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 6 implementation guide. With the money "
        "loop closed, the next sprint can turn to inventory and the "
        "storekeeper's workflows, confident that the financial side "
        "of the school is in good order.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
