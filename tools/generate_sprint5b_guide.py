"""Generates 'Sprint 5b - Implementation Guide.docx' covering the
student-photo work delivered in sprint 5b (IStudentPhotoService,
StudentPhotoService, the reusable StudentAvatar.razor component,
the per-DTO PhotoUrl plumbing, and the Photo tab on the Edit
Student page).

Run from the repo root:  python tools/generate_sprint5b_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 5b - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 5b — Student Photographs")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Upload pipeline · Reusable avatar · Pupil photos everywhere")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/5b-student-photo")
    meta.add_run("\nBuilt on: Sprints 1–5 (identity, academic domain, students & parents, attendance, results & report cards)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 5b in context", 1)
    add_para(doc,
        "Sprint 5b is a focused mid-sprint top-up. Sprint 3 created the "
        "Student entity with a PhotoUrl column but deliberately left the "
        "upload pipeline for later. Sprints 4 and 5 then went on to layer "
        "attendance and results on top of it without anyone ever actually "
        "putting a face to a name. This sprint closes that gap: it adds "
        "the upload pipeline, the reusable avatar component, the small "
        "wiring on every pupil-facing page that already existed, and the "
        "storage hygiene that keeps uploaded files out of the repository.")
    add_para(doc,
        "Functionally, after this sprint a SuperAdmin or HeadTeacher signing "
        "in to the application can pick any pupil from the Students "
        "directory, open the new Photo tab on their edit page, choose a "
        "JPG / PNG / WebP image up to 5 MB, and watch it appear in every "
        "list and detail view that previously showed the pupil's name "
        "alone — Students, Enrolments, daily and subject attendance "
        "registers, the assessment score sheet, the report-cards list, and "
        "the big avatar on the report-card detail header.")
    add_para(doc,
        "Pupils without a photo continue to render cleanly: the avatar "
        "falls back to the pupil's first-and-last initials on a coloured "
        "tile, so a row never looks broken just because nobody has gotten "
        "around to taking a picture.")
    add_para(doc,
        "This document is a long-form implementation guide in the tone of "
        "the sprint 4 and sprint 5 guides. An engineer who has read those "
        "and has the codebase checked out can recreate every change here "
        "without referring to the diff.")

    add_heading(doc, "1.1 Where this sits relative to sprint 5", 2)
    add_para(doc,
        "Nothing about the database schema changes. The Student.PhotoUrl "
        "column has been waiting since sprint 3. Nothing about the service "
        "lifetimes changes either — IStudentPhotoService is just one more "
        "scoped service registered in DI alongside IStudentService.")
    add_bullets(doc, [
        "BaseEntity / SaveChanges audit + soft-delete machinery — reused; "
        "the photo service does its work inside the regular SaveChangesAsync "
        "and Student.ModifiedOn/By are stamped automatically.",
        "OperationResult / OperationResult<T> — same predictable success/"
        "failure shape for the new service.",
        "Radzen Blazor + the green/gold app.css — the new avatar component "
        "leans on the same colour tokens (the initials tile fills with "
        "--nps-green-200 and writes initials in --nps-green-900).",
        "DTO denormalisation rule — every per-row DTO that already carried "
        "StudentName and StudentAdmissionNumber gains StudentPhotoUrl plus "
        "StudentFirstName / StudentLastName, so the grid renders without an "
        "extra round-trip per row.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_numbered(doc, [
        "Open a pupil's profile, switch to the new Photo tab, and click "
        "'Choose new photo'. A standard file picker pops up; pick an "
        "image and watch it round-trip to the server and refresh in the "
        "preview within a couple of seconds.",
        "Re-upload over the top to replace the photo. The previous file "
        "on disk is removed first, so a switch from JPG to PNG does not "
        "leave an orphan behind.",
        "Click 'Remove photo' to wipe the photo entirely. The pupil "
        "falls back to the initials tile.",
        "See the photo everywhere the pupil is listed — Students, "
        "Enrolments, Daily attendance, Subject attendance, Score sheet, "
        "Report cards list, and as a 96-pixel avatar at the top of the "
        "Report card detail view.",
        "A teacher who replaces a photo sees the new image immediately "
        "rather than the browser's cached version — a `?v=<ticks>` "
        "cache-buster is appended to every avatar URL.",
    ])

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_bullets(doc, [
        "Image resizing or thumbnail generation. The stored file is the "
        "raw upload; CSS handles cropping into a circle. A future sprint "
        "may introduce a 'thumbs/' folder of pre-sized variants if the "
        "raw photos grow large enough to slow down the lists.",
        "Bulk import. The picture of every pupil in a school's existing "
        "database is a one-shot ETL task and lives outside this app.",
        "External storage (S3 / Azure Blob). The wwwroot path is the "
        "simplest thing that works for a small Nigerian primary school "
        "running locally. Swapping the StudentPhotoService for a blob "
        "variant later is a single class change.",
        "Photos for parents and staff. Sprint 5b is pupil-only on "
        "purpose — the Parent entity already has UserId but the parent "
        "portal sprint will likely surface staff/parent avatars from "
        "the existing ApplicationUser.ProfilePhotoUrl column.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers, this sprint adds:")
    add_bullets(doc, [
        "1 new application-layer contract (IStudentPhotoService).",
        "1 new infrastructure-layer service (StudentPhotoService).",
        "1 new reusable Razor component (StudentAvatar.razor).",
        "1 new tab (Photo) on the existing EditStudent.razor page.",
        "5 DTOs gain three new fields each (StudentPhotoUrl, "
        "StudentFirstName, StudentLastName).",
        "5 service projections updated to populate the new fields.",
        "7 Razor pages gain the StudentAvatar component in their pupil "
        "column (or header).",
        "1 set of CSS additions to app.css for .nps-avatar / .nps-cell-"
        "with-avatar / .nps-photo-panel.",
        "1 .gitkeep + a 2-line .gitignore addition to keep the uploads "
        "folder present-but-empty.",
    ])
    add_para(doc,
        "No schema migration, no new lookup tables, no domain entity "
        "changes. Everything compiles with zero warnings on .NET 10.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)

    add_heading(doc, "2.1 Stable file name keyed by Student.Id", 2)
    add_para(doc,
        "Each pupil's photo is stored at "
        "wwwroot/uploads/students/{studentId}{ext}. The file name is "
        "fully determined by the pupil's Id and the chosen extension. "
        "The benefits compound:")
    add_bullets(doc, [
        "Replace-on-upload is automatic — saving 6f9c…3a.png over 6f9c…3a.jpg "
        "does the right thing as long as the old extension is cleaned up first.",
        "No second 'CurrentPhotoFileName' column is needed; Student.PhotoUrl "
        "is the single source of truth.",
        "Direct static-file serving — Blazor's MapStaticAssets serves "
        "/uploads/students/{id}.jpg without any controller or endpoint, so "
        "the request path is one hop short of file IO.",
        "Predictable cleanup — when a pupil is hard-deleted (which we "
        "currently refuse from the UI), a single Directory.EnumerateFiles "
        "by pattern wipes any extension.",
    ])
    add_para(doc,
        "The trade-off is that two pupils with the same id never exist, "
        "which is fine because Student.Id is a Guid.")

    add_heading(doc, "2.2 Cache-busting on the URL, not the filename", 2)
    add_para(doc,
        "Because the filename is stable per pupil, browsers will happily "
        "cache /uploads/students/6f9c.jpg even after a teacher uploads a "
        "replacement. The StudentAvatar component appends "
        "?v={DateTime.UtcNow.Ticks} to every URL on render, so each "
        "fresh render emits a unique URL and the browser bypasses cache. "
        "The query string never reaches the file system; it's purely a "
        "browser-side eviction trigger.")

    add_heading(doc, "2.3 5 MB cap, three formats", 2)
    add_para(doc,
        "JPG, PNG, and WebP are accepted; GIF / TIFF / BMP / SVG are "
        "rejected. The cap is 5 MB. The combination is "
        "deliberately conservative:")
    add_bullets(doc, [
        "A modern phone takes pupil photos at 2–4 MB when set to a "
        "reasonable resolution.",
        "Larger uploads slow down the page and crowd disk space without "
        "improving the displayed thumbnail.",
        "WebP support means a school using a phone with a 'save as WebP' "
        "setting (most current Android phones do this) doesn't need an "
        "intermediate conversion step.",
        "The cap also defeats trivial 'upload a 4K video as my photo' "
        "mischief.",
    ])
    add_para(doc,
        "The cap and the format list are constants in StudentPhotoService "
        "— a school that wants to raise the cap edits one number, "
        "rebuilds, and ships.")

    add_heading(doc, "2.4 InputFile, not RadzenUpload", 2)
    add_para(doc,
        "Blazor's built-in <InputFile> is paired with a hidden input "
        "element and a Radzen-styled 'Choose new photo' button that "
        "click-proxies the input via a one-line JS interop. The "
        "alternative — RadzenUpload — wants a server endpoint URL to "
        "POST to, which would mean adding an MVC controller or a "
        "minimal-API route just to forward bytes to the service we "
        "already have. The hidden-input approach keeps everything in the "
        "Blazor circuit, where authentication, anti-forgery, and the "
        "service injection just work.")

    add_heading(doc, "2.5 Denormalising photo URL into per-row DTOs", 2)
    add_para(doc,
        "Five DTOs gain a StudentPhotoUrl plus FirstName / LastName: "
        "StudentDto already had them, but EnrolmentDto, the daily and "
        "subject attendance entry DTOs, AssessmentScoreDto, and "
        "ReportCardDto did not. Adding the columns to the projection is "
        "cheap (an extra column or two on a join EF Core already "
        "performs) and saves the alternative of N+1 queries to fetch "
        "the photo URL per row when the page renders.")
    add_para(doc,
        "First and last name are added alongside the photo URL because "
        "the avatar's initials fallback needs them. Splitting "
        "'StudentName' back into first/last on the client would have "
        "been error-prone for pupils with hyphenated last names or "
        "multiple given names.")

    add_heading(doc, "2.6 .gitignore the uploads folder", 2)
    add_para(doc,
        "Uploaded photos are real school data — usually photographs of "
        "minors — and have no business in a public Git repository. The "
        "gitignore rule keeps the entire wwwroot/uploads/students/ tree "
        "out of commits, with a single .gitkeep file allowed through so "
        "the folder exists on a fresh clone. The StudentPhotoService "
        "creates the folder on demand if it ever happens to be missing "
        "(Directory.CreateDirectory is idempotent), but the .gitkeep is "
        "still useful to clarify intent.")

    add_page_break(doc)


def chapter3_application(doc):
    add_heading(doc, "3. Application layer — the photo service contract", 1)
    add_para(doc,
        "One small, focused interface. Both methods take just enough "
        "data to do the job and return the same OperationResult shape "
        "every other service in the codebase returns.")
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/IStudentPhotoService.cs")
    add_para(doc,
        "UploadAsync takes the stream, the MIME content type, and the "
        "byte count. Three reasons the count is a parameter rather than "
        "Stream.Length:")
    add_bullets(doc, [
        "Blazor's IBrowserFile gives the length directly without seeking, "
        "and seeking on a forward-only stream costs.",
        "The caller has cheaper access to the count than the service "
        "(IBrowserFile.Size is a property; reading the stream to compute "
        "length forces a buffer).",
        "Passing it explicitly makes the contract honest — the service "
        "is allowed to short-circuit on size without touching the "
        "payload.",
    ])
    add_para(doc,
        "The Upload return value is OperationResult<string> where the "
        "string is the public-facing URL — the same URL persisted to "
        "Student.PhotoUrl — so callers that want to display the new "
        "photo immediately can do so without another GetByIdAsync.")

    add_page_break(doc)


def chapter4_infrastructure(doc):
    add_heading(doc, "4. Infrastructure — StudentPhotoService", 1)
    add_para(doc,
        "The implementation is short enough to read top-to-bottom. The "
        "two halves are similar: upload validates, writes, and stamps "
        "Student.PhotoUrl; remove finds the file, deletes it, and clears "
        "Student.PhotoUrl.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/StudentPhotoService.cs")
    add_para(doc, "Four details worth dwelling on:")
    add_bullets(doc, [
        "MaxPhotoBytes and AllowedTypes are public static — a future "
        "feature (a settings page, perhaps) can read them at startup "
        "and surface them to the UI without round-tripping through the "
        "service.",
        "AllowedTypes is a Dictionary keyed on Content-Type with the "
        "extension as the value. The content type comes directly from "
        "the browser's IBrowserFile.ContentType; mapping to a fixed "
        "extension avoids trusting the user-provided file name "
        "(which can be anything).",
        "Cleanup loop on upload (Directory.EnumerateFiles by pattern) "
        "removes any other-extension file for the same Student.Id "
        "before writing the new one. Best-effort — file-system "
        "exceptions are swallowed because they should not block the "
        "upload, just leave an orphan.",
        "PhotoUrl is set to a relative path starting with '/'. Blazor's "
        "MapStaticAssets serves wwwroot files at that URL, so no "
        "controller is needed and the browser never makes a Razor "
        "circuit call to fetch the image.",
    ])

    add_heading(doc, "4.1 DI registration", 2)
    add_para(doc, "One line in DependencyInjection.cs:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IStudentPhotoService",
        end_marker="services.AddScoped<IDailyAttendanceService",
        caption="Excerpt — DependencyInjection.cs (sprint 5b addition)")

    add_page_break(doc)


def chapter5_avatar(doc):
    add_heading(doc, "5. The reusable StudentAvatar component", 1)
    add_para(doc,
        "Every pupil-facing page in the application now leans on a "
        "single Razor component. Three parameters drive it: PhotoUrl, "
        "FirstName, LastName. A fourth parameter (Size) picks the "
        "circle dimension — small, medium, or large.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Shared/StudentAvatar.razor")
    add_para(doc, "Two patterns worth noting:")
    add_bullets(doc, [
        "Photo-or-initials fallback — when PhotoUrl is empty, the "
        "component renders a coloured tile with the pupil's first and "
        "last initials. Pupils with single-name records still render "
        "(a '?' is the last-resort fallback).",
        "CacheBustedUrl() — if the URL doesn't already contain a query "
        "string the component appends ?v={ticks}. The Ticks value "
        "changes on every render pass so a re-upload mid-session is "
        "picked up immediately.",
    ])

    add_heading(doc, "5.1 _Imports.razor", 2)
    add_para(doc,
        "Components.Shared is imported at the Components folder level so "
        "every page can drop in <StudentAvatar /> without a per-file "
        "@using directive.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Web.Components.Shared",
        lines_after_start=1,
        caption="Excerpt — _Imports.razor")

    add_heading(doc, "5.2 The CSS that backs it", 2)
    add_para(doc,
        "The .nps-avatar block in app.css carries the colour tokens, the "
        "circle radius, the size variants, and a couple of helper "
        "classes used by the surrounding cells.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/wwwroot/app.css",
        "Student avatars (sprint 5b)",
        end_marker=".nps-hidden-input",
        lines_after_start=80,
        caption="Excerpt — app.css (sprint 5b additions)")

    add_page_break(doc)


def chapter6_upload_ui(doc):
    add_heading(doc, "6. The upload UI on the Edit Student page", 1)
    add_para(doc,
        "EditStudent.razor picks up a fourth tab — Photo — between "
        "Profile and Parents. The tab body is a single RadzenCard with "
        "a large avatar preview on the left and the action stack on "
        "the right.")
    add_para(doc, "The Razor markup, in three pieces:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor",
        "<RadzenTabsItem Text=\"Photo\">",
        end_marker="<RadzenTabsItem Text=\"@($\"Parents",
        caption="Excerpt — EditStudent.razor (Photo tab markup)")

    add_para(doc, "And the C# handlers behind it:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor",
        "private async Task TriggerPhotoPicker",
        end_marker="private async Task RemovePhotoAsync",
        caption="Excerpt — EditStudent.razor (photo handlers)")

    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor",
        "private async Task RemovePhotoAsync",
        end_marker="}\n}\n",
        caption="Excerpt — EditStudent.razor (RemovePhotoAsync)")

    add_para(doc, "Three pieces of plumbing are worth calling out:")
    add_bullets(doc, [
        "The <InputFile> sits hidden (display: none, via .nps-hidden-input) "
        "and is clicked programmatically via a one-line IJSRuntime call. "
        "This lets the Radzen-styled button trigger a native file picker "
        "without forcing the user through Radzen's upload widget.",
        "HandlePhotoSelected reads the stream into a MemoryStream before "
        "calling UploadAsync. The intermediate buffer is what lets the "
        "service compute Length and set the stream Position cleanly on "
        "behalf of the writer.",
        "RemovePhotoAsync routes through DialogService.Confirm — there "
        "is no Undo for a removed photo, so the dialog is the safety "
        "net.",
    ])

    add_page_break(doc)


def chapter7_display(doc):
    add_heading(doc, "7. Showing photos everywhere a pupil appears", 1)
    add_para(doc,
        "Seven Razor pages get the StudentAvatar treatment. The diff in "
        "each is the same shape: replace the plain '<strong>{Name}</strong>' "
        "cell with a flex container that holds the avatar and the name "
        "stack. The CSS class .nps-cell-with-avatar provides the spacing.")

    add_heading(doc, "7.1 Students list", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/Students.razor",
        "<RadzenDataGridColumn TItem=\"StudentDto\" Title=\"Pupil\"",
        end_marker="</RadzenDataGridColumn>",
        caption="Excerpt — Students.razor")

    add_heading(doc, "7.2 Enrolments", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/Enrolments.razor",
        "<RadzenDataGridColumn TItem=\"EnrolmentDto\" Title=\"Pupil\"",
        end_marker="</RadzenDataGridColumn>",
        caption="Excerpt — Enrolments.razor")

    add_heading(doc, "7.3 Daily attendance", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/DailyAttendance.razor",
        "<RadzenDataGridColumn TItem=\"DailyAttendanceEntryDto\" Title=\"Pupil\"",
        end_marker="</RadzenDataGridColumn>",
        caption="Excerpt — DailyAttendance.razor")

    add_heading(doc, "7.4 Subject attendance", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/SubjectAttendance.razor",
        "<RadzenDataGridColumn TItem=\"SubjectAttendanceEntryDto\" Title=\"Pupil\"",
        end_marker="</RadzenDataGridColumn>",
        caption="Excerpt — SubjectAttendance.razor")

    add_heading(doc, "7.5 Score sheet", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Results/AssessmentScores.razor",
        "<div class=\"nps-cell-with-avatar\">",
        lines_after_start=10,
        caption="Excerpt — AssessmentScores.razor")

    add_heading(doc, "7.6 Report cards list", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCards.razor",
        "<RadzenDataGridColumn TItem=\"ReportCardDto\" Title=\"Pupil\"",
        end_marker="</RadzenDataGridColumn>",
        caption="Excerpt — ReportCards.razor")

    add_heading(doc, "7.7 Report card detail (large avatar)", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCardDetail.razor",
        "<div class=\"nps-card-headline\">",
        lines_after_start=12,
        caption="Excerpt — ReportCardDetail.razor")

    add_page_break(doc)


def chapter8_dto_plumbing(doc):
    add_heading(doc, "8. DTO and projection plumbing", 1)
    add_para(doc,
        "Each per-row DTO that the seven pages above consume now carries "
        "three extra fields: StudentPhotoUrl, StudentFirstName, "
        "StudentLastName. The corresponding service projections were "
        "updated in lock-step. The diff is the same shape every time — "
        "one trio of property assignments inside an existing Select.")

    add_heading(doc, "8.1 The DTOs", 2)
    add_bullets(doc, [
        "EnrolmentDto — Family/Dtos/EnrolmentDtos.cs",
        "DailyAttendanceEntryDto — Attendance/Dtos/DailyAttendanceDtos.cs",
        "SubjectAttendanceEntryDto — Attendance/Dtos/SubjectAttendanceDtos.cs",
        "AssessmentScoreDto — Results/Dtos/AssessmentScoreDtos.cs",
        "ReportCardDto — Results/Dtos/ReportCardDtos.cs",
        "(StudentDto already had PhotoUrl + FirstName + LastName from sprint 3.)",
    ])

    add_heading(doc, "8.2 Projection snippets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/EnrolmentService.cs",
        "StudentPhotoUrl",
        lines_after_start=3,
        caption="Excerpt — EnrolmentService.Project")

    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/DailyAttendanceService.cs",
        "StudentPhotoUrl",
        lines_after_start=3,
        caption="Excerpt — DailyAttendanceService.ProjectEntries")

    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/SubjectAttendanceService.cs",
        "StudentPhotoUrl",
        lines_after_start=3,
        caption="Excerpt — SubjectAttendanceService.ProjectEntries")

    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/AssessmentService.cs",
        "StudentPhotoUrl = s.PhotoUrl",
        lines_after_start=3,
        caption="Excerpt — AssessmentService.GetScoreSheetAsync")

    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/ReportCardService.cs",
        "StudentPhotoUrl = c.Student",
        lines_after_start=3,
        caption="Excerpt — ReportCardService.ProjectCard")

    add_page_break(doc)


def chapter9_storage(doc):
    add_heading(doc, "9. Storage layout and git hygiene", 1)
    add_para(doc,
        "Uploaded files live under "
        "src/NaijaPrimeSchool.Web/wwwroot/uploads/students/. The folder "
        "exists on a fresh clone via a sentinel .gitkeep file; everything "
        "else in the folder is gitignored so school photographs never end "
        "up in the repository.")
    add_excerpt(doc, ".gitignore", "Uploaded pupil photos",
                lines_after_start=4,
                caption="Excerpt — .gitignore (sprint 5b addition)")
    add_para(doc,
        "If the folder is ever deleted by accident, the next upload re-"
        "creates it: StudentPhotoService.UploadAsync calls "
        "Directory.CreateDirectory(folder) on every save. The .gitkeep "
        "is documentation-as-data — it keeps the intent visible to "
        "anyone browsing the source tree.")

    add_page_break(doc)


def chapter10_lifecycle(doc):
    add_heading(doc, "10. Lifecycle of a photo upload", 1)
    add_para(doc,
        "Walking one upload end-to-end is the clearest way to see how "
        "all the pieces co-operate.")

    add_heading(doc, "10.1 Admin opens the Photo tab", 2)
    add_bullets(doc, [
        "Browser fetches /students/{id} (server-side Blazor circuit).",
        "EditStudent.razor calls LoadStudentAsync, which calls "
        "StudentService.GetByIdAsync(Id). The DTO carries PhotoUrl, "
        "FirstName, LastName.",
        "photoBust is initialised to PhotoUrl + '?v={ticks}' so the "
        "current photo, if any, shows immediately without a cache hit.",
        "User clicks the Photo tab. The hidden <InputFile> and the "
        "Radzen 'Choose new photo' button render.",
    ])

    add_heading(doc, "10.2 Admin picks a file", 2)
    add_bullets(doc, [
        "User clicks 'Choose new photo'. The TriggerPhotoPicker handler "
        "calls JS.InvokeVoidAsync to click the hidden input.",
        "The native OS file picker opens. User picks a JPG.",
        "Blazor's <InputFile> raises OnChange, populating an "
        "InputFileChangeEventArgs with an IBrowserFile.",
        "HandlePhotoSelected validates the size client-side (a fast "
        "guard) and opens a read stream from the browser file.",
        "The stream is copied into an in-memory MemoryStream so the "
        "service can re-read Length and Position.",
    ])

    add_heading(doc, "10.3 The service does its work", 2)
    add_bullets(doc, [
        "StudentPhotoService.UploadAsync revalidates length and "
        "content type against MaxPhotoBytes and AllowedTypes.",
        "Loads the Student entity by Id.",
        "Ensures wwwroot/uploads/students/ exists.",
        "Deletes any existing file matching {studentId}.* — handles "
        "the JPG-to-PNG case cleanly.",
        "Streams the new bytes to {studentId}{ext} on disk.",
        "Sets Student.PhotoUrl to '/uploads/students/{studentId}{ext}' "
        "and calls SaveChangesAsync.",
        "ApplicationDbContext's SaveChanges override stamps "
        "Student.ModifiedOn / Student.ModifiedBy from the current user.",
        "Returns OperationResult<string>.Success with the public URL.",
    ])

    add_heading(doc, "10.4 The UI refreshes", 2)
    add_bullets(doc, [
        "HandlePhotoSelected sees Succeeded == true, raises a Radzen "
        "success notification, calls LoadStudentAsync again to refresh "
        "the StudentDto with the new PhotoUrl, and updates photoBust.",
        "StudentAvatar re-renders with the new URL plus the cache-"
        "buster query string. The browser fetches the new file.",
        "Every other page in the app — Students, Enrolments, Report "
        "cards, etc. — will pick up the new photo the next time it "
        "renders. The cache-buster ?v=... is added at render time, "
        "so even a tab that was open from before the upload picks up "
        "the new image on its next reactive re-render.",
    ])

    add_page_break(doc)


def chapter11_smoketest(doc):
    add_heading(doc, "11. Smoke-test walkthrough", 1)

    add_heading(doc, "11.1 Build and run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "No new migration is needed — the Student.PhotoUrl column has "
        "existed since sprint 3. Sign in as the SuperAdmin.")

    add_heading(doc, "11.2 Upload your first photo", 2)
    add_numbered(doc, [
        "Family → Students. Pick any enrolled pupil.",
        "Click the Photo tab.",
        "Click Choose new photo. Pick a JPG / PNG / WebP up to 5 MB.",
        "Watch the preview refresh within a second or two.",
        "Open the same pupil's Students-list row in another tab. The "
        "thumbnail is there.",
    ])

    add_heading(doc, "11.3 Verify the photo flows through every page", 2)
    add_bullets(doc, [
        "Family → Enrolments. The pupil's avatar appears next to "
        "their name.",
        "Attendance → Daily attendance. Open a register for the "
        "pupil's class.",
        "Attendance → Subject attendance. Pick a lesson; the pupil's "
        "row shows the avatar.",
        "Results & Reports → Assessments → open any score sheet. The "
        "score-sheet table now has avatars on every row.",
        "Results & Reports → Report cards. The list shows thumbnail "
        "avatars; clicking a card shows a 96-pixel avatar in the page "
        "header.",
    ])

    add_heading(doc, "11.4 Verify the error paths", 2)
    add_numbered(doc, [
        "Try uploading a 10 MB image. The Photo handler refuses with a "
        "'Too large' notification.",
        "Try uploading a .gif or .svg. The service refuses with "
        "'Unsupported image format'.",
        "Click Remove photo. The confirmation dialog appears; confirm. "
        "The avatar reverts to the initials tile.",
    ])

    add_heading(doc, "11.5 Verify storage hygiene", 2)
    add_para(doc,
        "After uploading a photo, run `git status` from the repo root. "
        "The uploaded file should NOT appear in untracked files — the "
        ".gitignore rule keeps it out. The .gitkeep file should be the "
        "only entry under wwwroot/uploads/students/ that git knows "
        "about.")

    add_page_break(doc)


def chapter12_troubleshooting(doc):
    add_heading(doc, "12. Troubleshooting and gotchas", 1)

    add_heading(doc, "12.1 'Unsupported image format'", 2)
    add_para(doc,
        "Only JPG, PNG, and WebP pass the MIME-type check. If the file "
        "is a HEIC straight from an iPhone, the browser will report "
        "image/heic and the upload will fail. Either change the iPhone's "
        "'Most Compatible' camera setting, or convert the photo to JPG "
        "first.")

    add_heading(doc, "12.2 Photo doesn't refresh after a re-upload", 2)
    add_para(doc,
        "The cache-buster (?v={ticks}) is added on every avatar render, "
        "so a browser that still shows the old photo is either showing "
        "a server-rendered page from before the upload completed, or is "
        "running an aggressive service worker. Reload the page; if that "
        "doesn't help, the file on disk in "
        "wwwroot/uploads/students/{id}.* is the source of truth — open "
        "it directly in another browser tab.")

    add_heading(doc, "12.3 'The file might still be in use by another process'", 2)
    add_para(doc,
        "This occurs on Windows when the same file is open in Word, "
        "Paint, or another image viewer while the service tries to "
        "overwrite it. Close the viewer and re-upload. The service does "
        "not loop or retry — file-system contention is rare enough that "
        "a clearer error is more useful than silent retries.")

    add_heading(doc, "12.4 Folder vanished after a wwwroot rebuild", 2)
    add_para(doc,
        "Some IDE 'clean' actions wipe wwwroot/uploads/students/. "
        "Don't worry: StudentPhotoService.UploadAsync calls "
        "Directory.CreateDirectory on every save, so the very next "
        "upload re-creates the folder. The .gitkeep file is checked in "
        "to make the directory's intent visible.")

    add_heading(doc, "12.5 Photo upload during a long-running save", 2)
    add_para(doc,
        "The Photo tab disables the upload buttons while uploading is "
        "true. A rapid double-click is therefore harmless. If a teacher "
        "navigates away mid-upload, the browser cancels the fetch and "
        "the service receives a cancelled stream — Student.PhotoUrl is "
        "left untouched.")

    add_page_break(doc)


def chapter13_forward(doc):
    add_heading(doc, "13. Forward-compatibility, today", 1)
    add_bullets(doc, [
        "Parent and staff avatars. Parent.PhotoUrl does not exist yet, "
        "but the StudentAvatar component is named generically enough "
        "that the next portal sprint can reuse it (rename to "
        "PersonAvatar with a one-line refactor).",
        "External storage. StudentPhotoService is the single seam to "
        "swap. A future Azure-Blob variant lives in the Infrastructure "
        "layer, gets registered behind the same IStudentPhotoService "
        "interface, and every consumer keeps working unchanged.",
        "Image moderation / face cropping. Could wrap UploadAsync with "
        "an image-processing step before persistence. The service is "
        "small enough to add a private CropAndCompressAsync without "
        "the call sites noticing.",
        "Bulk import. A school migrating from an existing system can "
        "drop pre-named files (named by Student.Id) straight into the "
        "uploads folder, then run a one-line SQL UPDATE setting "
        "Student.PhotoUrl. The current schema already supports this.",
    ])

    add_page_break(doc)


def chapter14_appendix(doc):
    add_heading(doc, "14. Appendix — files added or changed in sprint 5b", 1)
    entries = [
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Family/IStudentPhotoService.cs", "Photo upload / remove contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Family/Dtos/EnrolmentDtos.cs",            "Added StudentPhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Application/Attendance/Dtos/DailyAttendanceDtos.cs",  "Added StudentPhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Application/Attendance/Dtos/SubjectAttendanceDtos.cs","Added StudentPhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/AssessmentScoreDtos.cs",     "Added StudentPhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/ReportCardDtos.cs",          "Added StudentPhotoUrl + first/last name."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/StudentPhotoService.cs",      "Validate, write to disk, stamp Student.PhotoUrl."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",               "Registered IStudentPhotoService."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/EnrolmentService.cs",         "Projection includes PhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/DailyAttendanceService.cs",   "Projection includes PhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SubjectAttendanceService.cs", "Projection includes PhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/AssessmentService.cs",        "Score-sheet projection includes PhotoUrl + first/last name."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/ReportCardService.cs",        "ReportCard projection includes PhotoUrl + first/last name."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Shared/StudentAvatar.razor",           "Reusable circular avatar (photo or initials fallback)."),
        ("src/NaijaPrimeSchool.Web/wwwroot/uploads/students/.gitkeep",               "Sentinel so the uploads folder exists on a fresh clone."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                       "Added Components.Shared @using."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor",      "Added Photo tab with upload + remove."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/Students.razor",         "Avatar in the Pupil column."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/Enrolments.razor",       "Avatar in the Pupil column."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Attendance/DailyAttendance.razor",   "Avatar in the Pupil column."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Attendance/SubjectAttendance.razor", "Avatar in the Pupil column."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/AssessmentScores.razor",     "Avatar in the score sheet row."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCards.razor",          "Avatar in the Pupil column."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCardDetail.razor",     "Large avatar in the page header."),
        ("src/NaijaPrimeSchool.Web/wwwroot/app.css",                                 "Avatar / cell / photo-panel styles."),
        ("Repo root (modified)", "—"),
        (".gitignore",                                                               "Ignore everything under uploads/students except .gitkeep."),
        ("README.md",                                                                "Sprint 5b section + roadmap + project-map updates."),
        ("NaijaPrimeSchool.slnx",                                                    "Added 'Sprint 5b - Implementation Guide.docx' to Solution Items."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint5b_guide.py",                                         "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 5b implementation guide. With pupils now "
        "showing their faces across every list and detail page, the "
        "next sprint can turn to fees, invoices, and the bursar's "
        "workflows with confidence that the rest of the app is "
        "already humane to look at.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_application(doc)
    chapter4_infrastructure(doc)
    chapter5_avatar(doc)
    chapter6_upload_ui(doc)
    chapter7_display(doc)
    chapter8_dto_plumbing(doc)
    chapter9_storage(doc)
    chapter10_lifecycle(doc)
    chapter11_smoketest(doc)
    chapter12_troubleshooting(doc)
    chapter13_forward(doc)
    chapter14_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
