"""Generates 'Sprint 1 - Implementation Guide.docx' describing how the
NaijaPrimeSchool Blazor codebase was built step-by-step in VS Code.

This script is a one-shot document generator; it is not part of the app.
Run from the repo root:  python tools/generate_sprint1_guide.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUTPUT = "Sprint 1 - Implementation Guide.docx"

# ---------- Styling helpers ----------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_code(doc, text, language_hint=""):
    """Insert a code block as a single-cell shaded table so it stands out."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)

    # Remove default paragraph margins in the cell
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    # Trailing spacing paragraph
    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)
    trailing.paragraph_format.space_before = Pt(0)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)  # deep green
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_toc(doc):
    """Insert a Word table-of-contents field. User presses F9 to update, or
    Word offers 'Update fields' when the document opens."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._element
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def configure_document(doc):
    """Tweak page size, margins, default font."""
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Trigger update of fields on open so the TOC builds itself
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


# ---------- Document content ---------------------------------------------------

def build_title_page(doc):
    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    run = t.add_run("Naija Prime School")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Sprint 1 — Identity & User Management")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)  # gold

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run("Implementation walk-through in Visual Studio Code")
    r2.font.size = Pt(14)
    r2.italic = True

    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto interactivity), "
                 "Entity Framework Core 10, SQL Server, ASP.NET Core Identity, "
                 "Radzen Blazor 10")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")

    add_page_break(doc)


def build_toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def build_overview(doc):
    add_heading(doc, "1. Overview", 1)
    add_para(doc,
        "This document is the narrative I wrote while building Sprint 1 of the Naija "
        "Prime School management system end-to-end in Visual Studio Code. It walks "
        "through every meaningful decision — how I scaffolded the solution, how the "
        "four Clean Architecture layers fit together, how ASP.NET Core Identity is "
        "wired into Blazor, how the Radzen UI is assembled, and how the whole thing "
        "finally lands on GitHub. A developer following the chapters in order should "
        "be able to recreate the codebase from an empty folder.")

    add_heading(doc, "1.1 What Sprint 1 delivers", 2)
    add_bullets(doc, [
        "Cookie-based authentication backed by ASP.NET Core Identity, with Guid keys.",
        "Seven seeded application roles: SuperAdmin, HeadTeacher, Teacher, "
        "SchoolBursar, SchoolStoreKeeper, Parent, Student.",
        "Role-based authorization policies and a role-scoped sidebar.",
        "Lockout on repeated failed sign-ins (5 attempts per 15 minutes).",
        "Revalidating authentication state — deactivated users are signed out within "
        "the revalidation window without needing to restart the app.",
        "A SuperAdmin-only user management workspace: list, search, filter, "
        "create, edit, activate / deactivate, reassign roles, and reset password.",
        "Auditing (CreatedOn/By, ModifiedOn/By, DeletedOn/By) applied automatically "
        "inside SaveChanges.",
        "Soft delete with EF Core global query filters so deleted rows are "
        "invisible to every ordinary query.",
        "Lookups (titles, genders, roles) as first-class tables — no enums.",
        "A custom Nigerian-themed green + gold Radzen skin.",
    ])

    add_heading(doc, "1.2 Architecture at a glance", 2)
    add_para(doc,
        "The solution follows a classic Clean Architecture layering. Each layer "
        "references only the layers strictly inward of it, so the Domain layer is "
        "pure POCOs with no knowledge of ASP.NET, EF Core, or the UI.")
    add_code(doc,
        "src/\n"
        "├── NaijaPrimeSchool.Domain/          # Entities, base types, interfaces\n"
        "├── NaijaPrimeSchool.Application/     # DTOs, service contracts, abstractions\n"
        "├── NaijaPrimeSchool.Infrastructure/  # EF Core, Identity stores, service impls\n"
        "├── NaijaPrimeSchool.Web/             # Blazor server host, auth, pages\n"
        "└── NaijaPrimeSchool.Web.Client/      # Blazor WebAssembly client (Auto mode)\n")
    add_para(doc,
        "Dependency direction flows inward: Web → Infrastructure → Application → Domain. "
        "The Web.Client project is a peer of Web; the Web project references it so "
        "that Blazor Auto interactivity can share components between server and WASM.")


def build_prereqs(doc):
    add_heading(doc, "2. Prerequisites and VS Code setup", 1)
    add_para(doc,
        "Before I opened VS Code for the first line of code I made sure the tools "
        "below were on my machine. Everything listed is free.")

    add_heading(doc, "2.1 Tooling", 2)
    add_bullets(doc, [
        ".NET 10 SDK (10.0.201 or newer) — the project targets net10.0.",
        "SQL Server — any of SQL Server LocalDB, Developer / Express / Standard, "
        "or Azure SQL. I used LocalDB on my laptop and a local SQL Server instance "
        "on my development box.",
        "Git for Windows (brings the credential manager used when pushing to GitHub).",
        "GitHub CLI (gh) — optional, but handy for creating the remote and "
        "running pull-request commands from the terminal.",
        "Visual Studio Code.",
    ])

    add_heading(doc, "2.2 VS Code extensions", 2)
    add_para(doc,
        "In the Extensions side-bar I installed these five, in this order:")
    add_numbered(doc, [
        "C# Dev Kit (Microsoft) — brings the C# language server, a solution "
        "explorer, and test explorer. This is the one essential extension.",
        "C# (ms-dotnettools.csharp) — installed automatically with C# Dev Kit.",
        ".NET Install Tool — manages .NET runtimes the Dev Kit needs.",
        "IntelliCode for C# Dev Kit — AI-assisted completions.",
        "EditorConfig for VS Code — honours the .editorconfig (if one exists).",
    ])
    add_para(doc,
        "I did not install any Blazor-specific extension: VS Code recognises "
        ".razor files through C# Dev Kit, and Radzen ships with its own "
        "IntelliSense via the NuGet package.")

    add_heading(doc, "2.3 Verifying the toolchain", 2)
    add_para(doc, "In a VS Code integrated terminal I confirmed everything:")
    add_code(doc,
        "dotnet --version         # 10.0.201 or higher\n"
        "dotnet --list-sdks       # net10.0 listed\n"
        "git --version\n"
        "gh --version             # optional\n"
        "sqlcmd -S .\\SQLEXPRESS -Q \"SELECT @@VERSION\"   # or (localdb)\\MSSQLLocalDB\n")


def build_scaffolding(doc):
    add_heading(doc, "3. Scaffolding the solution", 1)
    add_para(doc,
        "VS Code does not have the project wizards Visual Studio does, so every "
        "project was created from the integrated terminal using the .NET CLI. "
        "The advantage is that the commands are reproducible and scriptable.")

    add_heading(doc, "3.1 Creating the root folder and solution file", 2)
    add_code(doc,
        "mkdir NaijaPrimeSchool\n"
        "cd NaijaPrimeSchool\n"
        "code .                                 # open this folder in VS Code\n"
        "dotnet new sln --format slnx -n NaijaPrimeSchool\n")
    add_para(doc,
        "I used the modern XML-flavoured .slnx format rather than the legacy .sln; "
        "both load in Visual Studio and in the C# Dev Kit solution explorer.")

    add_heading(doc, "3.2 Creating the four projects", 2)
    add_para(doc,
        "Each project lives under src/. I used the classlib template for pure "
        "libraries and the Blazor Web App template for the front door.")
    add_code(doc,
        "mkdir src\n"
        "cd src\n"
        "\n"
        "dotnet new classlib -n NaijaPrimeSchool.Domain         -f net10.0\n"
        "dotnet new classlib -n NaijaPrimeSchool.Application    -f net10.0\n"
        "dotnet new classlib -n NaijaPrimeSchool.Infrastructure -f net10.0\n"
        "\n"
        "# Blazor Web App with Auto interactivity — produces both Web and Web.Client\n"
        "dotnet new blazor -n NaijaPrimeSchool.Web \\\n"
        "    --interactivity Auto \\\n"
        "    --auth Individual \\\n"
        "    --use-local-db false \\\n"
        "    -f net10.0\n"
        "\n"
        "cd ..\n"
        "dotnet sln NaijaPrimeSchool.slnx add \\\n"
        "    src/NaijaPrimeSchool.Domain \\\n"
        "    src/NaijaPrimeSchool.Application \\\n"
        "    src/NaijaPrimeSchool.Infrastructure \\\n"
        "    src/NaijaPrimeSchool.Web \\\n"
        "    src/NaijaPrimeSchool.Web.Client\n")
    add_para(doc,
        "The blazor template generated both NaijaPrimeSchool.Web (the server host) "
        "and NaijaPrimeSchool.Web.Client (the WebAssembly counterpart) side-by-side. "
        "I chose --auth Individual so the template scaffolded a ready-to-customise "
        "Identity UI, and I set --use-local-db false because I prefer to wire the "
        "connection string explicitly.")

    add_heading(doc, "3.3 Project references", 2)
    add_para(doc, "The dependency graph is strictly inward:")
    add_code(doc,
        "cd src\n"
        "dotnet add NaijaPrimeSchool.Application    reference NaijaPrimeSchool.Domain\n"
        "dotnet add NaijaPrimeSchool.Infrastructure reference NaijaPrimeSchool.Application\n"
        "dotnet add NaijaPrimeSchool.Web            reference NaijaPrimeSchool.Infrastructure\n"
        "dotnet add NaijaPrimeSchool.Web            reference NaijaPrimeSchool.Web.Client\n"
        "cd ..\n")

    add_heading(doc, "3.4 NuGet packages", 2)
    add_para(doc,
        "These are the package references I added to each project. Versions follow "
        "the net10.0 release line.")
    add_para(doc, "Domain", bold=True)
    add_code(doc,
        "dotnet add src/NaijaPrimeSchool.Domain package Microsoft.Extensions.Identity.Stores --version 10.0.0\n")
    add_para(doc, "Infrastructure", bold=True)
    add_code(doc,
        "dotnet add src/NaijaPrimeSchool.Infrastructure package Microsoft.EntityFrameworkCore            --version 10.0.0\n"
        "dotnet add src/NaijaPrimeSchool.Infrastructure package Microsoft.EntityFrameworkCore.SqlServer  --version 10.0.0\n"
        "dotnet add src/NaijaPrimeSchool.Infrastructure package Microsoft.EntityFrameworkCore.Tools      --version 10.0.0\n"
        "dotnet add src/NaijaPrimeSchool.Infrastructure package Microsoft.AspNetCore.Identity.EntityFrameworkCore --version 10.0.0\n")
    add_para(doc,
        "I also added <FrameworkReference Include=\"Microsoft.AspNetCore.App\" /> to "
        "the Infrastructure csproj so that ASP.NET Identity abstractions are "
        "available without pulling the full web SDK.")
    add_para(doc, "Web", bold=True)
    add_code(doc,
        "dotnet add src/NaijaPrimeSchool.Web package Radzen.Blazor                                   --version 10.3.1\n"
        "dotnet add src/NaijaPrimeSchool.Web package Microsoft.AspNetCore.Components.WebAssembly.Server\n"
        "dotnet add src/NaijaPrimeSchool.Web package Microsoft.AspNetCore.Diagnostics.EntityFrameworkCore\n"
        "dotnet add src/NaijaPrimeSchool.Web package Microsoft.AspNetCore.Identity.EntityFrameworkCore\n")
    add_para(doc, "Web.Client", bold=True)
    add_code(doc,
        "dotnet add src/NaijaPrimeSchool.Web.Client package Radzen.Blazor                                           --version 10.3.1\n"
        "dotnet add src/NaijaPrimeSchool.Web.Client package Microsoft.AspNetCore.Components.WebAssembly.Authentication\n")

    add_heading(doc, "3.5 Restore and first build", 2)
    add_code(doc, "dotnet restore\ndotnet build")
    add_para(doc,
        "Everything compiled, leaving me with an empty but structurally correct "
        "Clean Architecture solution ready to grow.")


def build_domain(doc):
    add_heading(doc, "4. The Domain layer", 1)
    add_para(doc,
        "Domain contains the company's language: entities, value types, and the "
        "interfaces that describe cross-cutting concerns like auditing and soft "
        "delete. It has a single package reference — Microsoft.Extensions.Identity.Stores "
        "— which gives me access to IdentityUser<TKey> and IdentityRole<TKey> base "
        "classes without pulling in EF Core.")

    add_heading(doc, "4.1 Auditing & soft-delete contracts", 2)
    add_para(doc,
        "I defined two tiny interfaces under Common/ so any entity can opt in "
        "independently:")
    add_code(doc,
        "// Common/IAuditable.cs\n"
        "public interface IAuditable\n"
        "{\n"
        "    DateTimeOffset CreatedOn { get; set; }\n"
        "    string? CreatedBy { get; set; }\n"
        "    DateTimeOffset? ModifiedOn { get; set; }\n"
        "    string? ModifiedBy { get; set; }\n"
        "}")
    add_code(doc,
        "// Common/ISoftDelete.cs\n"
        "public interface ISoftDelete\n"
        "{\n"
        "    bool IsDeleted { get; set; }\n"
        "    DateTimeOffset? DeletedOn { get; set; }\n"
        "    string? DeletedBy { get; set; }\n"
        "}")
    add_para(doc,
        "Then BaseEntity gives ordinary entities a Guid key and implements both "
        "interfaces so I do not need to repeat the boilerplate on every table.")
    add_code(doc,
        "public abstract class BaseEntity : IAuditable, ISoftDelete\n"
        "{\n"
        "    public Guid Id { get; set; } = Guid.NewGuid();\n"
        "    public DateTimeOffset CreatedOn { get; set; } = DateTimeOffset.UtcNow;\n"
        "    public string? CreatedBy { get; set; }\n"
        "    public DateTimeOffset? ModifiedOn { get; set; }\n"
        "    public string? ModifiedBy { get; set; }\n"
        "    public bool IsDeleted { get; set; }\n"
        "    public DateTimeOffset? DeletedOn { get; set; }\n"
        "    public string? DeletedBy { get; set; }\n"
        "}")

    add_heading(doc, "4.2 Identity entities", 2)
    add_para(doc,
        "ApplicationUser extends IdentityUser<Guid> and additionally implements "
        "IAuditable and ISoftDelete so the same audit columns flow through the Users "
        "table. I added the profile fields the UI asks for (FirstName, LastName, "
        "MiddleName, TitleId, GenderId, DateOfBirth, Address, ProfilePhotoUrl), "
        "an IsActive flag, and a small pair of deactivation fields. I also "
        "exposed a computed FullName property that the UI uses everywhere.")
    add_code(doc,
        "public class ApplicationUser : IdentityUser<Guid>, IAuditable, ISoftDelete\n"
        "{\n"
        "    public string FirstName { get; set; } = string.Empty;\n"
        "    public string LastName  { get; set; } = string.Empty;\n"
        "    public string? MiddleName { get; set; }\n"
        "\n"
        "    public Guid?   TitleId { get; set; }\n"
        "    public Title?  Title   { get; set; }\n"
        "    public Guid?   GenderId { get; set; }\n"
        "    public Gender? Gender   { get; set; }\n"
        "\n"
        "    public DateTime? DateOfBirth { get; set; }\n"
        "    public string?   Address     { get; set; }\n"
        "    public string?   ProfilePhotoUrl { get; set; }\n"
        "\n"
        "    public bool IsActive { get; set; } = true;\n"
        "    public DateTimeOffset? DeactivatedOn { get; set; }\n"
        "    public string?         DeactivationReason { get; set; }\n"
        "\n"
        "    /* IAuditable + ISoftDelete members, omitted for brevity */\n"
        "\n"
        "    public string FullName => string.IsNullOrWhiteSpace(MiddleName)\n"
        "        ? $\"{FirstName} {LastName}\".Trim()\n"
        "        : $\"{FirstName} {MiddleName} {LastName}\".Trim();\n"
        "}")
    add_para(doc,
        "ApplicationRole and ApplicationUserRole are the same pattern: extend the "
        "Identity base types, add description / audit fields, and (for the join "
        "table) track AssignedOn and AssignedBy so there is a breadcrumb for who "
        "added a user to a role and when.")

    add_heading(doc, "4.3 Lookup tables — no enums", 2)
    add_para(doc,
        "Anything that would normally be a C# enum I modelled as a database table. "
        "This makes the data translatable, extensible by data rather than by code, "
        "and visible in reports. For Sprint 1 those tables are Title and Gender:")
    add_code(doc,
        "public class Gender : BaseEntity\n"
        "{\n"
        "    public string Name        { get; set; } = string.Empty;\n"
        "    public string Code        { get; set; } = string.Empty;\n"
        "    public int    DisplayOrder { get; set; }\n"
        "    public ICollection<ApplicationUser> Users { get; set; } = [];\n"
        "}")
    add_para(doc, "Title follows the same shape minus the Code column.")

    add_heading(doc, "4.4 Role names — constants, not an enum", 2)
    add_para(doc,
        "Role membership is validated against strings everywhere in ASP.NET "
        "(the Authorize attribute and policies expect string role names), but I "
        "did not want magic strings scattered through the code. So Roles is a "
        "static class whose members are const strings:")
    add_code(doc,
        "public static class Roles\n"
        "{\n"
        "    public const string SuperAdmin         = nameof(SuperAdmin);\n"
        "    public const string HeadTeacher        = nameof(HeadTeacher);\n"
        "    public const string Teacher            = nameof(Teacher);\n"
        "    public const string SchoolBursar       = nameof(SchoolBursar);\n"
        "    public const string SchoolStoreKeeper  = nameof(SchoolStoreKeeper);\n"
        "    public const string Parent             = nameof(Parent);\n"
        "    public const string Student            = nameof(Student);\n"
        "\n"
        "    public static readonly IReadOnlyList<string> All =\n"
        "        [SuperAdmin, HeadTeacher, Teacher, SchoolBursar,\n"
        "         SchoolStoreKeeper, Parent, Student];\n"
        "}")
    add_para(doc,
        "The seeder iterates Roles.All to create exactly these rows on every "
        "startup, and the Razor pages spell role names as Roles.SuperAdmin so the "
        "compiler catches typos.")


def build_application(doc):
    add_heading(doc, "5. The Application layer", 1)
    add_para(doc,
        "Application holds contracts — service interfaces and DTOs — with no "
        "EF Core, no ASP.NET, and no Razor. Its sole dependency is Domain.")

    add_heading(doc, "5.1 Cross-cutting abstractions", 2)
    add_para(doc,
        "Two helpers are used by almost every feature:")
    add_code(doc,
        "public interface ICurrentUser\n"
        "{\n"
        "    Guid?   UserId { get; }\n"
        "    string? UserName { get; }\n"
        "    bool    IsAuthenticated { get; }\n"
        "}")
    add_para(doc,
        "ICurrentUser is implemented in the Web project (where HttpContext is "
        "available) but consumed by ApplicationDbContext in Infrastructure so that "
        "audit columns can be stamped with the signed-in user's name without "
        "leaking HttpContext into EF Core.")
    add_code(doc,
        "public class OperationResult\n"
        "{\n"
        "    public bool Succeeded { get; init; }\n"
        "    public IReadOnlyList<string> Errors { get; init; } = [];\n"
        "\n"
        "    public static OperationResult Success() => new() { Succeeded = true };\n"
        "    public static OperationResult Failure(params string[] errors) =>\n"
        "        new() { Succeeded = false, Errors = errors };\n"
        "}\n"
        "\n"
        "public class OperationResult<T> : OperationResult\n"
        "{\n"
        "    public T? Data { get; init; }\n"
        "    public static OperationResult<T> Success(T data) =>\n"
        "        new() { Succeeded = true, Data = data };\n"
        "}")
    add_para(doc,
        "OperationResult lets services return a predictable shape — success + data "
        "or failure + an array of messages — without throwing exceptions for "
        "expected validation failures. The Blazor pages turn those messages into "
        "RadzenNotification pop-ups.")

    add_heading(doc, "5.2 Data transfer objects", 2)
    add_para(doc,
        "Each feature has small request / response DTOs in Application/Users/Dtos/. "
        "DataAnnotations attributes live on the requests so that Blazor's "
        "DataAnnotationsValidator can use them without re-declaring the rules on "
        "the page.")
    add_bullets(doc, [
        "UserDto — flat projection returned to the UI, including TitleName, "
        "GenderName, FullName, IsActive, and the role names.",
        "CreateUserRequest — Username, Email, Password, ConfirmPassword, profile "
        "fields, IsActive, and Roles (a List<string>). Roles carries [Required] "
        "and [MinLength(1)] so the form cannot be submitted without a role.",
        "UpdateUserRequest — profile + contact fields, but no password and no "
        "roles (role changes go through the dedicated assign-roles page).",
        "UserListFilter — paging inputs: SearchTerm, Role, IsActive, Skip, Take, "
        "OrderBy. Plus a generic PagedResult<T>.",
        "LookupDto / RoleDto — simple Id/Name (+ Code, Description) shapes for "
        "dropdowns and tables.",
    ])

    add_heading(doc, "5.3 Service contracts", 2)
    add_para(doc,
        "Two interfaces own the entire sprint's behaviour:")
    add_code(doc,
        "public interface IUserService\n"
        "{\n"
        "    Task<PagedResult<UserDto>> ListAsync(UserListFilter filter, CancellationToken ct = default);\n"
        "    Task<UserDto?> GetByIdAsync(Guid id, CancellationToken ct = default);\n"
        "    Task<OperationResult<Guid>> CreateAsync(CreateUserRequest request, CancellationToken ct = default);\n"
        "    Task<OperationResult>       UpdateAsync(UpdateUserRequest request, CancellationToken ct = default);\n"
        "    Task<OperationResult>       SetActiveAsync(Guid userId, bool isActive, string? reason, CancellationToken ct = default);\n"
        "    Task<OperationResult>       SoftDeleteAsync(Guid userId, CancellationToken ct = default);\n"
        "    Task<OperationResult>       AssignRolesAsync(Guid userId, IEnumerable<string> roles, CancellationToken ct = default);\n"
        "    Task<OperationResult>       SetPasswordAsync(Guid userId, string newPassword, CancellationToken ct = default);\n"
        "    Task<IReadOnlyList<string>> GetRolesAsync(Guid userId, CancellationToken ct = default);\n"
        "}")
    add_code(doc,
        "public interface ILookupService\n"
        "{\n"
        "    Task<IReadOnlyList<LookupDto>> GetGendersAsync(CancellationToken ct = default);\n"
        "    Task<IReadOnlyList<LookupDto>> GetTitlesAsync(CancellationToken ct = default);\n"
        "    Task<IReadOnlyList<RoleDto>>   GetRolesAsync(CancellationToken ct = default);\n"
        "}")
    add_para(doc,
        "Notice that every method takes a CancellationToken with a default value. "
        "Blazor Server components pass the component's cancellation token "
        "naturally, so long-running queries stop when the user navigates away.")


def build_infrastructure(doc):
    add_heading(doc, "6. The Infrastructure layer", 1)
    add_para(doc,
        "Infrastructure does the heavy lifting: it owns the EF Core DbContext, "
        "all persistence, and the real implementations of the service interfaces. "
        "It depends on Application and therefore transitively on Domain.")

    add_heading(doc, "6.1 ApplicationDbContext", 2)
    add_para(doc,
        "I started with IdentityDbContext because Identity needs its own DbContext "
        "with the UserStore and RoleStore relationships pre-configured. I used the "
        "eight-generic-argument overload so I could supply my own ApplicationUserRole "
        "join entity (which carries AssignedOn / AssignedBy):")
    add_code(doc,
        "public class ApplicationDbContext(\n"
        "    DbContextOptions<ApplicationDbContext> options,\n"
        "    ICurrentUser currentUser)\n"
        "    : IdentityDbContext<\n"
        "        ApplicationUser,\n"
        "        ApplicationRole,\n"
        "        Guid,\n"
        "        IdentityUserClaim<Guid>,\n"
        "        ApplicationUserRole,\n"
        "        IdentityUserLogin<Guid>,\n"
        "        IdentityRoleClaim<Guid>,\n"
        "        IdentityUserToken<Guid>>(options)\n"
        "{\n"
        "    public DbSet<Gender> Genders => Set<Gender>();\n"
        "    public DbSet<Title>  Titles  => Set<Title>();")
    add_para(doc,
        "The constructor takes ICurrentUser alongside the options so that the audit "
        "logic can reach the signed-in user without static state. I pass it into a "
        "primary constructor (a C# 12 feature) to keep the plumbing tight.")

    add_heading(doc, "6.2 Model configuration", 2)
    add_para(doc,
        "In OnModelCreating I do three things: rename the Identity tables to "
        "friendlier names, apply column constraints, and attach soft-delete query "
        "filters. A small extract:")
    add_code(doc,
        "builder.Entity<ApplicationUser>(b =>\n"
        "{\n"
        "    b.ToTable(\"Users\");\n"
        "    b.Property(u => u.FirstName).HasMaxLength(80).IsRequired();\n"
        "    b.Property(u => u.LastName).HasMaxLength(80).IsRequired();\n"
        "    b.Ignore(u => u.FullName);\n"
        "\n"
        "    b.HasOne(u => u.Title).WithMany(t => t.Users)\n"
        "        .HasForeignKey(u => u.TitleId)\n"
        "        .OnDelete(DeleteBehavior.SetNull);\n"
        "\n"
        "    b.HasOne(u => u.Gender).WithMany(g => g.Users)\n"
        "        .HasForeignKey(u => u.GenderId)\n"
        "        .OnDelete(DeleteBehavior.SetNull);\n"
        "\n"
        "    b.HasIndex(u => u.IsDeleted);\n"
        "    b.HasQueryFilter(u => !u.IsDeleted);\n"
        "});")
    add_para(doc,
        "The HasQueryFilter call is the critical line — every LINQ query against "
        "Users automatically filters out soft-deleted rows. Admin reports that need "
        "to see deleted rows can call IgnoreQueryFilters() explicitly.")
    add_para(doc,
        "For the lookup tables I wrote a tiny ConfigureLookup<T> helper that "
        "applies the same pattern (table name, audit column lengths, soft-delete "
        "filter) so I do not repeat myself across Gender and Title.")

    add_heading(doc, "6.3 Auditing and soft delete inside SaveChanges", 2)
    add_para(doc,
        "Rather than stamp CreatedBy, ModifiedBy, DeletedBy manually in every "
        "service, I override SaveChanges / SaveChangesAsync to do it in one place:")
    add_code(doc,
        "private void ApplyAuditAndSoftDelete()\n"
        "{\n"
        "    var now      = DateTimeOffset.UtcNow;\n"
        "    var userName = currentUser.UserName ?? \"system\";\n"
        "\n"
        "    foreach (var entry in ChangeTracker.Entries())\n"
        "    {\n"
        "        if (entry.Entity is IAuditable a)\n"
        "        {\n"
        "            if (entry.State == EntityState.Added)\n"
        "            {\n"
        "                a.CreatedOn = a.CreatedOn == default ? now : a.CreatedOn;\n"
        "                a.CreatedBy ??= userName;\n"
        "            }\n"
        "            else if (entry.State == EntityState.Modified)\n"
        "            {\n"
        "                a.ModifiedOn = now;\n"
        "                a.ModifiedBy = userName;\n"
        "            }\n"
        "        }\n"
        "\n"
        "        if (entry.Entity is ISoftDelete s && entry.State == EntityState.Deleted)\n"
        "        {\n"
        "            entry.State = EntityState.Modified;  // turn Delete into Update\n"
        "            s.IsDeleted = true;\n"
        "            s.DeletedOn = now;\n"
        "            s.DeletedBy = userName;\n"
        "        }\n"
        "    }\n"
        "}")
    add_para(doc,
        "Rewriting Deleted → Modified is the trick that turns every db.Remove() "
        "call into a soft delete without services needing to know about it.")

    add_heading(doc, "6.4 Dependency injection — AddInfrastructure", 2)
    add_para(doc,
        "I expose a single extension method that the Web project calls from "
        "Program.cs. It registers the DbContext, configures Identity with the "
        "password and lockout rules, and wires up the service implementations.")
    add_code(doc,
        "public static IServiceCollection AddInfrastructure(\n"
        "    this IServiceCollection services, IConfiguration configuration)\n"
        "{\n"
        "    var connectionString = configuration.GetConnectionString(\"DefaultConnection\")\n"
        "        ?? throw new InvalidOperationException(\"Connection string 'DefaultConnection' not found.\");\n"
        "\n"
        "    services.AddDbContext<ApplicationDbContext>(options =>\n"
        "        options.UseSqlServer(connectionString, sql =>\n"
        "            sql.MigrationsAssembly(typeof(ApplicationDbContext).Assembly.FullName)));\n"
        "\n"
        "    services\n"
        "        .AddIdentityCore<ApplicationUser>(o =>\n"
        "        {\n"
        "            o.SignIn.RequireConfirmedAccount = false;\n"
        "            o.User.RequireUniqueEmail        = true;\n"
        "            o.Password.RequireDigit          = true;\n"
        "            o.Password.RequireLowercase      = true;\n"
        "            o.Password.RequireUppercase      = true;\n"
        "            o.Password.RequireNonAlphanumeric = true;\n"
        "            o.Password.RequiredLength         = 8;\n"
        "            o.Lockout.AllowedForNewUsers      = true;\n"
        "            o.Lockout.MaxFailedAccessAttempts = 5;\n"
        "            o.Lockout.DefaultLockoutTimeSpan  = TimeSpan.FromMinutes(15);\n"
        "        })\n"
        "        .AddRoles<ApplicationRole>()\n"
        "        .AddEntityFrameworkStores<ApplicationDbContext>()\n"
        "        .AddSignInManager()\n"
        "        .AddDefaultTokenProviders();\n"
        "\n"
        "    services.AddScoped<IUserService,   UserService>();\n"
        "    services.AddScoped<ILookupService, LookupService>();\n"
        "\n"
        "    return services;\n"
        "}")

    add_heading(doc, "6.5 UserService", 2)
    add_para(doc,
        "UserService uses the Identity UserManager / RoleManager for the "
        "operations they know about (create, password, role membership) and drops "
        "down to ApplicationDbContext directly for list queries that need flexible "
        "filtering. Notable decisions:")
    add_bullets(doc, [
        "ListAsync joins UserRoles -> Roles manually rather than calling "
        "userManager.GetRolesAsync per user, because the latter would be N+1 "
        "queries for a paged list.",
        "CreateAsync validates uniqueness of UserName and Email up front, checks "
        "every requested role actually exists, then creates the user and assigns "
        "roles in two calls.",
        "SetActiveAsync flips IsActive, sets LockoutEnd to DateTimeOffset.MaxValue "
        "when deactivating (so the cookie is rejected immediately) and calls "
        "UpdateSecurityStampAsync so the revalidating provider will log the user "
        "out within its 30-minute window.",
        "SoftDeleteAsync does the same plus sets IsDeleted = true; the global "
        "query filter on Users means the row vanishes from ordinary queries.",
        "AssignRolesAsync diffs the current role set against the requested set "
        "and only adds / removes the difference, so the audit log stays clean.",
        "SetPasswordAsync uses a generated reset token rather than passing "
        "Identity's internal password-hash APIs — more conservative.",
    ])

    add_heading(doc, "6.6 LookupService", 2)
    add_para(doc,
        "LookupService is a thin projection layer — Genders ordered by "
        "DisplayOrder, Titles the same, and Roles ordered by Name. Each method "
        "returns LookupDto / RoleDto so the UI never sees the raw entities.")

    add_heading(doc, "6.7 DatabaseInitializer — migrate and seed", 2)
    add_para(doc,
        "A single static class drives first-run bootstrapping. Called from "
        "Program.cs after the DI container is built, it:")
    add_numbered(doc, [
        "Applies any outstanding EF Core migrations (db.Database.MigrateAsync).",
        "Seeds Genders (Male, Female) and nine Titles (Mr., Mrs., Miss, Ms., "
        "Dr., Prof., Chief, Alhaji, Alhaja) if those tables are empty.",
        "Iterates Roles.All and creates any missing ApplicationRole rows with "
        "descriptions and IsSystemRole = true.",
        "Creates the default SuperAdmin user (superadmin@naijaprimeschool.ng / "
        "Admin@12345) if none exists, and adds them to the SuperAdmin role.",
    ])
    add_para(doc,
        "This is what makes the app runnable on a fresh database without any "
        "manual SQL. The README warns administrators to reset that initial "
        "password on first sign-in.")

    add_heading(doc, "6.8 Creating the first migration", 2)
    add_para(doc,
        "I installed dotnet-ef globally (once per machine) and then added the "
        "initial migration. The migration lives in the Infrastructure project "
        "but is driven by the Web project because that is where the connection "
        "string lives:")
    add_code(doc,
        "dotnet tool install --global dotnet-ef --version 10.0.0\n"
        "\n"
        "dotnet ef migrations add InitialCreate \\\n"
        "    --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "    --startup-project src/NaijaPrimeSchool.Web \\\n"
        "    --output-dir Persistence/Migrations\n")
    add_para(doc,
        "Because DatabaseInitializer.InitializeAsync runs on every startup and "
        "calls MigrateAsync, I did not need a separate 'dotnet ef database update' "
        "step — but it is still available if I want to apply migrations "
        "out-of-band.")


def build_web_host(doc):
    add_heading(doc, "7. The Web host — wiring Blazor + Identity", 1)
    add_para(doc,
        "The Web project is the composition root. It boots Blazor, registers "
        "Radzen, turns on authentication, builds the authorization policies, and "
        "finally runs the database initializer.")

    add_heading(doc, "7.1 Program.cs walk-through", 2)
    add_para(doc,
        "Reading Program.cs top-to-bottom tells you almost everything:")
    add_code(doc,
        "var builder = WebApplication.CreateBuilder(args);\n"
        "\n"
        "builder.Services.AddRazorComponents()\n"
        "    .AddInteractiveServerComponents()\n"
        "    .AddInteractiveWebAssemblyComponents()\n"
        "    .AddAuthenticationStateSerialization();\n"
        "\n"
        "builder.Services.AddCascadingAuthenticationState();\n"
        "builder.Services.AddScoped<IdentityRedirectManager>();\n"
        "builder.Services.AddScoped<AuthenticationStateProvider,\n"
        "    IdentityRevalidatingAuthenticationStateProvider>();\n"
        "\n"
        "builder.Services.AddHttpContextAccessor();\n"
        "builder.Services.AddScoped<ICurrentUser, CurrentUserAccessor>();\n"
        "\n"
        "builder.Services.AddAuthentication(o =>\n"
        "{\n"
        "    o.DefaultScheme       = IdentityConstants.ApplicationScheme;\n"
        "    o.DefaultSignInScheme = IdentityConstants.ExternalScheme;\n"
        "}).AddIdentityCookies();\n"
        "\n"
        "builder.Services.ConfigureApplicationCookie(o =>\n"
        "{\n"
        "    o.LoginPath         = \"/Account/Login\";\n"
        "    o.AccessDeniedPath  = \"/Account/AccessDenied\";\n"
        "    o.LogoutPath        = \"/Account/Logout\";\n"
        "    o.SlidingExpiration = true;\n"
        "    o.ExpireTimeSpan    = TimeSpan.FromHours(8);\n"
        "});\n"
        "\n"
        "builder.Services.AddInfrastructure(builder.Configuration);\n"
        "\n"
        "builder.Services.AddAuthorization(o =>\n"
        "{\n"
        "    o.AddPolicy(\"ManageUsers\", p => p.RequireRole(Roles.SuperAdmin));\n"
        "});\n"
        "\n"
        "builder.Services.AddRadzenComponents();")
    add_para(doc, "After app = builder.Build() the request pipeline is:")
    add_code(doc,
        "if (app.Environment.IsDevelopment())\n"
        "    app.UseWebAssemblyDebugging();\n"
        "else\n"
        "{\n"
        "    app.UseExceptionHandler(\"/Error\", createScopeForErrors: true);\n"
        "    app.UseHsts();\n"
        "}\n"
        "\n"
        "app.UseStatusCodePagesWithReExecute(\"/not-found\", createScopeForStatusCodePages: true);\n"
        "app.UseHttpsRedirection();\n"
        "app.UseAntiforgery();\n"
        "\n"
        "app.MapStaticAssets();\n"
        "app.MapRazorComponents<App>()\n"
        "   .AddInteractiveServerRenderMode()\n"
        "   .AddInteractiveWebAssemblyRenderMode()\n"
        "   .AddAdditionalAssemblies(typeof(NaijaPrimeSchool.Web.Client._Imports).Assembly);\n"
        "\n"
        "app.MapAdditionalIdentityEndpoints();\n"
        "\n"
        "using (var scope = app.Services.CreateScope())\n"
        "{\n"
        "    await DatabaseInitializer.InitializeAsync(scope.ServiceProvider);\n"
        "}\n"
        "\n"
        "app.Run();")
    add_para(doc,
        "The single authorization policy, ManageUsers, is what guards every user "
        "management page with [Authorize(Roles = Roles.SuperAdmin)] — I also kept "
        "the direct role form on the pages as a defensive belt-and-braces.")

    add_heading(doc, "7.2 CurrentUserAccessor — bridging HttpContext to DI", 2)
    add_para(doc,
        "ICurrentUser is implemented in the Web project because only here is "
        "HttpContext available:")
    add_code(doc,
        "public sealed class CurrentUserAccessor(IHttpContextAccessor ctx) : ICurrentUser\n"
        "{\n"
        "    public Guid? UserId\n"
        "    {\n"
        "        get\n"
        "        {\n"
        "            var raw = ctx.HttpContext?.User.FindFirstValue(ClaimTypes.NameIdentifier);\n"
        "            return Guid.TryParse(raw, out var id) ? id : null;\n"
        "        }\n"
        "    }\n"
        "    public string? UserName => ctx.HttpContext?.User.Identity?.Name;\n"
        "    public bool    IsAuthenticated => ctx.HttpContext?.User.Identity?.IsAuthenticated == true;\n"
        "}")

    add_heading(doc, "7.3 Logout endpoint", 2)
    add_para(doc,
        "I mapped a POST /Account/Logout endpoint so the header logout button is "
        "a traditional form submission (anti-forgery token included). It calls "
        "SignInManager.SignOutAsync and redirects back to login:")
    add_code(doc,
        "accountGroup.MapPost(\"/Logout\", async (\n"
        "    [FromServices] SignInManager<ApplicationUser> signInManager,\n"
        "    [FromForm] string? returnUrl) =>\n"
        "{\n"
        "    await signInManager.SignOutAsync();\n"
        "    return TypedResults.LocalRedirect($\"~/{returnUrl ?? \"Account/Login\"}\");\n"
        "});")

    add_heading(doc, "7.4 Revalidating authentication state", 2)
    add_para(doc,
        "A plain cookie would let a deactivated user keep clicking around until "
        "the cookie expired. I do not want that. IdentityRevalidatingAuthenticationStateProvider "
        "extends the framework's RevalidatingServerAuthenticationStateProvider and "
        "re-checks the user against the database every 30 minutes — if the user "
        "is deleted, deactivated, or their security stamp has changed, the "
        "provider returns false and Blazor forces a sign-out:")
    add_code(doc,
        "protected override TimeSpan RevalidationInterval => TimeSpan.FromMinutes(30);\n"
        "\n"
        "protected override async Task<bool> ValidateAuthenticationStateAsync(\n"
        "    AuthenticationState authenticationState, CancellationToken ct)\n"
        "{\n"
        "    await using var scope = scopeFactory.CreateAsyncScope();\n"
        "    var userManager = scope.ServiceProvider.GetRequiredService<UserManager<ApplicationUser>>();\n"
        "    var user = await userManager.GetUserAsync(authenticationState.User);\n"
        "    if (user is null || user.IsDeleted || !user.IsActive) return false;\n"
        "    /* check security stamp matches current stamp ... */\n"
        "    return principalStamp == userStamp;\n"
        "}")


def build_ui(doc):
    add_heading(doc, "8. The Razor UI", 1)
    add_para(doc,
        "Every page is an InteractiveServer Blazor component. I kept things "
        "server-rendered because Sprint 1 is administrative and small, and "
        "InteractiveServer avoids the 'load WASM runtime' latency. The Web.Client "
        "project is present and wired so future features can opt into Auto "
        "interactivity per page.")

    add_heading(doc, "8.1 App and Routes", 2)
    add_para(doc,
        "App.razor is the html shell. It pre-loads Radzen's material-base "
        "stylesheet, my app.css, and the Blazor framework script.")
    add_code(doc,
        "<link rel=\"stylesheet\" href=\"_content/Radzen.Blazor/css/material-base.css\" />\n"
        "<link rel=\"stylesheet\" href=\"@Assets[\"app.css\"]\" />\n"
        "<link rel=\"stylesheet\" href=\"@Assets[\"NaijaPrimeSchool.Web.styles.css\"]\" />\n"
        "...\n"
        "<Routes />\n"
        "<script src=\"_content/Radzen.Blazor/Radzen.Blazor.js\"></script>\n"
        "<script src=\"@Assets[\"_framework/blazor.web.js\"]\"></script>")
    add_para(doc,
        "Routes.razor uses AuthorizeRouteView so an unauthenticated user hitting "
        "a protected page gets sent to the login flow (via RedirectToLogin — a "
        "tiny client component that calls NavigationManager.NavigateTo with "
        "forceLoad: true).")

    add_heading(doc, "8.2 _Imports.razor", 2)
    add_para(doc,
        "Every page shares one set of usings, so most components do not need any "
        "@using directives of their own:")
    add_code(doc,
        "@using Microsoft.AspNetCore.Authorization\n"
        "@using Microsoft.AspNetCore.Components.Authorization\n"
        "@using Microsoft.AspNetCore.Components.Forms\n"
        "@using Radzen\n"
        "@using Radzen.Blazor\n"
        "@using NaijaPrimeSchool.Domain.Identity\n"
        "@using NaijaPrimeSchool.Application.Common\n"
        "@using NaijaPrimeSchool.Application.Users\n"
        "@using NaijaPrimeSchool.Application.Users.Dtos\n"
        "@using NaijaPrimeSchool.Web.Components.Account\n"
        "@using NaijaPrimeSchool.Web.Components.Layout")

    add_heading(doc, "8.3 MainLayout and NavMenu", 2)
    add_para(doc,
        "MainLayout wraps authenticated pages with a RadzenLayout: header on top "
        "(brand, sign-out form), sidebar on the left (NavMenu), and a RadzenBody "
        "for page content. The whole tree is inside <AuthorizeView> so the "
        "sidebar is hidden on login / access-denied screens.")
    add_para(doc,
        "NavMenu renders different items depending on which roles the signed-in "
        "user has. Only SuperAdmins see the User Management group; other role "
        "groups are stubbed as Disabled='true' items, reserving the space in the "
        "sidebar for future sprints.")

    add_heading(doc, "8.4 Sign-in page — Login.razor", 2)
    add_para(doc,
        "The login page is a classic server-rendered EditForm with FormName="
        "\"login\" so Blazor binds the POST back to the same component. It "
        "accepts either a username or an email address, checks that the user is "
        "not soft-deleted or deactivated, and calls SignInManager.PasswordSignInAsync "
        "with lockoutOnFailure: true so the Identity lockout policy applies.")
    add_code(doc,
        "var user = await UserManager.FindByNameAsync(Input.Email)\n"
        "       ?? await UserManager.FindByEmailAsync(Input.Email);\n"
        "\n"
        "if (user is null || user.IsDeleted) { errorMessage = \"Invalid sign in attempt.\"; return; }\n"
        "if (!user.IsActive) { errorMessage = \"Your account has been deactivated...\"; return; }\n"
        "\n"
        "var result = await SignInManager.PasswordSignInAsync(\n"
        "    user.UserName!, Input.Password, Input.RememberMe, lockoutOnFailure: true);")
    add_para(doc,
        "On success the page redirects to the ReturnUrl query-string parameter "
        "or to the dashboard. Lockouts surface a friendly message.")

    add_heading(doc, "8.5 The Users list — Users.razor", 2)
    add_para(doc,
        "The list is a RadzenDataGrid with server-side paging. Above the grid I "
        "placed a filter bar — a text box bound to SearchTerm, a role dropdown, "
        "and an Active/Inactive dropdown. Every change calls LoadAsync, which "
        "rebuilds a UserListFilter and hands it to IUserService.ListAsync.")
    add_bullets(doc, [
        "Rows show FullName with the email underneath it (via the Template "
        "slot), followed by username, a horizontal stack of role badges, and a "
        "green / grey status badge.",
        "Row actions live in the right-most column: Edit, Assign roles, and "
        "either Deactivate or Activate depending on current status.",
        "Deactivate asks for confirmation via DialogService.Confirm before "
        "calling IUserService.SetActiveAsync — no accidental clicks.",
        "After every action I reload the page via LoadAsync and push a Radzen "
        "notification so the admin gets feedback.",
    ])

    add_heading(doc, "8.6 Add user — CreateUser.razor", 2)
    add_para(doc,
        "The add-user form is a RadzenTemplateForm bound to a CreateUserRequest. "
        "It uses <DataAnnotationsValidator /> so all the [Required], [EmailAddress], "
        "[StringLength] and [Compare] attributes I put on the DTO are enforced "
        "with no extra ceremony. The role selector is a RadzenCheckBoxList bound "
        "directly to model.Roles (via Value / ValueChanged / ValueExpression) — "
        "binding to model.Roles, not a separate field, is important so that "
        "DataAnnotationsValidator sees the selection and the [MinLength(1)] "
        "validation on Roles actually prevents submit.")
    add_para(doc,
        "On submit the page calls IUserService.CreateAsync, shows a Radzen "
        "notification on success or spreads the OperationResult.Errors into a "
        "red notification on failure, then navigates back to /users.")

    add_heading(doc, "8.7 Edit user — EditUser.razor", 2)
    add_para(doc,
        "The edit page routes on /users/{Id:guid}, loads the user via "
        "IUserService.GetByIdAsync, and presents the same profile fields as the "
        "add form (minus username, password, and roles, which are handled "
        "elsewhere). The header carries Activate / Deactivate buttons and a "
        "shortcut to the role-assignment page.")

    add_heading(doc, "8.8 Assign roles — UserRoles.razor", 2)
    add_para(doc,
        "This page is intentionally small: a RadzenCheckBoxList of every role "
        "bound to the user's current role set, plus Save / Cancel. On save it "
        "calls IUserService.AssignRolesAsync, which diffs and applies only the "
        "change set so audit history stays clean.")

    add_heading(doc, "8.9 Roles catalogue — RolesPage.razor", 2)
    add_para(doc,
        "A read-only grid at /roles lists every role with its description and a "
        "System / Custom badge — useful for admins who want to see the full "
        "taxonomy at a glance. All rows are System because Sprint 1 seeds them "
        "from code.")

    add_heading(doc, "8.10 Reset password — ResetPasswordDialog.razor", 2)
    add_para(doc,
        "Admin-initiated password reset is a Radzen dialog hosted inside "
        "EditUser, not a separate route. It asks for a new password (twice), "
        "validates the confirmation, and calls IUserService.SetPasswordAsync, "
        "which internally generates an Identity reset token and applies it — "
        "far safer than touching the password hash directly.")

    add_heading(doc, "8.11 Home dashboard", 2)
    add_para(doc,
        "The home page shows the signed-in user's name and four stat tiles "
        "(total users, roles configured, active users, deactivated users). It "
        "also renders SuperAdmin-only quick-action buttons for Manage users and "
        "Add user. All stats are computed once on OnInitializedAsync with the "
        "existing services — no new endpoint was needed.")


def build_theme(doc):
    add_heading(doc, "9. Theming — the Nigerian green + gold skin", 1)
    add_para(doc,
        "Radzen ships with neutral base stylesheets. I layered a thin custom "
        "skin on top (wwwroot/app.css) using CSS custom properties so every page "
        "can reference --nps-green / --nps-gold / --nps-ink-500 and stay "
        "consistent. Only three concerns matter:")
    add_bullets(doc, [
        "Colour tokens — deep Nigerian green for primary, gold for accents, "
        "a muted ink palette for text and borders.",
        "Layout primitives — .nps-page-container, .nps-page-header, .nps-card, "
        ".nps-form-grid (two-column responsive grid), .nps-form-actions. Used "
        "on every page for a consistent rhythm.",
        "The auth page — .nps-auth-page and .nps-auth-card give the login "
        "screen a centred, shadowed card with a brand mark.",
    ])
    add_para(doc,
        "I deliberately avoided tailwind / SASS for Sprint 1 because the Radzen "
        "+ CSS-variables approach is enough for a small admin area, and it means "
        "there is no additional build step for the front-end.")


def build_run_locally(doc):
    add_heading(doc, "10. Running the application locally", 1)
    add_heading(doc, "10.1 appsettings.json", 2)
    add_code(doc,
        "{\n"
        "  \"ConnectionStrings\": {\n"
        "    \"DefaultConnection\": \"Server=.;Database=NaijaPrimeSchool;Trusted_Connection=True;MultipleActiveResultSets=true;TrustServerCertificate=True\"\n"
        "  },\n"
        "  \"Logging\": {\n"
        "    \"LogLevel\": { \"Default\": \"Information\", \"Microsoft.AspNetCore\": \"Warning\" }\n"
        "  },\n"
        "  \"AllowedHosts\": \"*\"\n"
        "}")
    add_para(doc,
        "For SQL Server LocalDB change Server=. to Server=(localdb)\\MSSQLLocalDB. "
        "The connection string never contains real credentials in the committed "
        "file — production secrets live in user-secrets or environment variables.")

    add_heading(doc, "10.2 Restore, build, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "On first run the DatabaseInitializer creates the schema, seeds the "
        "lookups, roles, and the default SuperAdmin. The console prints the URL "
        "to open (typically https://localhost:7xxx). I signed in with:")
    add_code(doc,
        "Email:    superadmin@naijaprimeschool.ng\n"
        "Password: Admin@12345   ← change immediately in Users → Edit → Reset password\n")

    add_heading(doc, "10.3 A suggested VS Code launch configuration", 2)
    add_para(doc,
        "C# Dev Kit auto-generates .vscode/launch.json the first time I press "
        "F5, but I keep a hand-written version so everyone on the team gets the "
        "same experience:")
    add_code(doc,
        "{\n"
        "  \"version\": \"0.2.0\",\n"
        "  \"configurations\": [\n"
        "    {\n"
        "      \"name\": \".NET Core Launch (web)\",\n"
        "      \"type\": \"coreclr\",\n"
        "      \"request\": \"launch\",\n"
        "      \"preLaunchTask\": \"build\",\n"
        "      \"program\": \"${workspaceFolder}/src/NaijaPrimeSchool.Web/bin/Debug/net10.0/NaijaPrimeSchool.Web.dll\",\n"
        "      \"cwd\":      \"${workspaceFolder}/src/NaijaPrimeSchool.Web\",\n"
        "      \"stopAtEntry\": false,\n"
        "      \"serverReadyAction\": { \"action\": \"openExternally\", \"pattern\": \"\\\\bNow listening on:\\\\s+(https?://\\\\S+)\" },\n"
        "      \"env\": { \"ASPNETCORE_ENVIRONMENT\": \"Development\" }\n"
        "    }\n"
        "  ]\n"
        "}")


def build_git(doc):
    add_heading(doc, "11. Source control and pushing to GitHub", 1)
    add_para(doc,
        "The final step is putting the sprint under version control. I did "
        "everything from the VS Code terminal.")

    add_heading(doc, "11.1 Writing a sensible .gitignore", 2)
    add_para(doc,
        "I hand-rolled a .gitignore covering the usual .NET build outputs "
        "(bin/, obj/), Visual Studio and Rider user state (.vs/, .vscode/, "
        ".idea/, *.suo, *.user), NuGet caches, publish profiles, appsettings.*.Local.json "
        "and appsettings.Production.json, plus secrets like .env, *.pfx, and "
        "secrets.json. Nothing sensitive leaves the repo.")

    add_heading(doc, "11.2 First commit on the sprint branch", 2)
    add_code(doc,
        "git init -b main\n"
        "git remote add origin https://github.com/benjaminsqlserver/NaijaPrimeSchool.git\n"
        "git checkout -b sprint/1-identity-and-user-management\n"
        "\n"
        "git add .gitignore README.md NaijaPrimeSchool.slnx src\n"
        "git status --short              # sanity check — no bin/, obj/, or .vs/\n"
        "\n"
        "git commit -m \"Sprint 1: identity and user management foundation\"\n")
    add_para(doc,
        "I avoided git add -A so nothing unexpected would slip in. The status "
        "output confirmed only source, the solution file, the README, and the "
        ".gitignore were staged.")

    add_heading(doc, "11.3 Pushing to the empty GitHub repo", 2)
    add_para(doc,
        "I had already created an empty public repository at "
        "github.com/benjaminsqlserver/NaijaPrimeSchool in the GitHub UI (without "
        "a README or .gitignore — the local commit is the initial state). Then:")
    add_code(doc,
        "git push -u origin sprint/1-identity-and-user-management\n"
        "\n"
        "# Populate main with the same commit so GitHub's default branch is not empty\n"
        "git branch main sprint/1-identity-and-user-management\n"
        "git push -u origin main\n")
    add_para(doc,
        "Git Credential Manager prompted me once for GitHub credentials via a "
        "browser pop-up; subsequent pushes re-use the stored token. Sprint 1 is "
        "now live on both main and sprint/1-identity-and-user-management.")


def build_map(doc):
    add_heading(doc, "12. Appendix — project map", 1)
    add_para(doc, "For reference, the key files produced during Sprint 1 and what each is for:")
    entries = [
        ("src/NaijaPrimeSchool.Domain/Common/BaseEntity.cs",
         "Base id + audit + soft-delete fields."),
        ("src/NaijaPrimeSchool.Domain/Identity/ApplicationUser.cs",
         "IdentityUser<Guid> with profile, audit, and soft-delete."),
        ("src/NaijaPrimeSchool.Domain/Identity/Roles.cs",
         "String constants for role names — not an enum."),
        ("src/NaijaPrimeSchool.Application/Common/ICurrentUser.cs",
         "Abstraction over the signed-in user, implemented in Web."),
        ("src/NaijaPrimeSchool.Application/Common/OperationResult.cs",
         "Success/failure shape returned by services."),
        ("src/NaijaPrimeSchool.Application/Users/IUserService.cs",
         "User management contract (CRUD + roles + password)."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
         "IdentityDbContext + query filters + SaveChanges auditing."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
         "Migrate + seed lookups, roles, SuperAdmin."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/UserService.cs",
         "Implementation using UserManager / RoleManager / DbContext."),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
         "AddInfrastructure — DI composition for the app."),
        ("src/NaijaPrimeSchool.Web/Program.cs",
         "Host wiring: Identity, cookies, Radzen, policies, endpoints, DB init."),
        ("src/NaijaPrimeSchool.Web/Services/CurrentUserAccessor.cs",
         "HttpContext-backed ICurrentUser implementation."),
        ("src/NaijaPrimeSchool.Web/Components/Account/Pages/Login.razor",
         "Sign-in page (EditForm + SignInManager)."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Users/",
         "List / Create / Edit / Roles / Reset password."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/MainLayout.razor",
         "Radzen shell with sidebar + header."),
        ("src/NaijaPrimeSchool.Web/wwwroot/app.css",
         "Green + gold school skin."),
    ]
    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Purpose"
    for run in hdr[0].paragraphs[0].runs + hdr[1].paragraphs[0].runs:
        run.bold = True
    for i, (path, purpose) in enumerate(entries, start=1):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose

    doc.add_paragraph()
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of Sprint 1 implementation guide. Sprint 2 picks up the academic "
        "domain (sessions, terms, classes, subjects).")
    r.italic = True


# ---------- Main ---------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    build_title_page(doc)
    build_toc_page(doc)
    build_overview(doc)
    build_prereqs(doc)
    build_scaffolding(doc)
    build_domain(doc)
    build_application(doc)
    build_infrastructure(doc)
    build_web_host(doc)
    build_ui(doc)
    build_theme(doc)
    build_run_locally(doc)
    build_git(doc)
    build_map(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
