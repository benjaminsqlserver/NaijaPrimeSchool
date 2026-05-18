"""Generates 'Sprint 10 - Implementation Guide.docx' covering the
parent-deletion hotfix: ParentService.SoftDeleteAsync now counts active
StudentParent links via a fresh database query (immune to EF Core
relationship fix-up surfacing soft-deleted rows from the change tracker)
and retires the linked ApplicationUser so the deleted parent can no
longer sign in to the portal.

Run from the repo root:  python tools/generate_sprint10_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 10 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 10 — Parent-Deletion Hotfix")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Fresh DB-driven link count · Retire the auto-provisioned ApplicationUser")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/10-parent-delete-fix")
    meta.add_run("\nBuilt on: Sprints 1–9 (identity, academic domain, students & parents, "
                 "attendance, results & report cards, pupil photos, fees & bursar workflows, "
                 "store & inventory, parent & student portals + announcements, "
                 "auto-provisioned portal accounts)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


# ---------- Chapter 1 -------------------------------------------------------

def chapter1_overview(doc):
    add_heading(doc, "1. Why this hotfix exists", 1)
    add_para(doc,
        "Sprint 9 closed the create-side loop on portal accounts. "
        "Sprint 10 closes a smaller but very visible delete-side gap. "
        "The school office reported that after unlinking every pupil "
        "from a parent (via Edit student → Parents → Unlink), pressing "
        "Delete on the parent in Parents still produced \"Cannot "
        "delete a parent who is still linked to a student. Unlink them "
        "first.\" A direct database check showed zero matching "
        "StudentParents rows for that parent, yet the guard kept "
        "firing.")
    add_para(doc,
        "The root cause is an EF Core relationship fix-up quirk that "
        "is easy to walk straight past until it bites. The fix replaces "
        "an in-memory navigation count with a fresh database query, "
        "and — while we are in the same method — also retires the "
        "ApplicationUser that sprint 9 auto-provisioned, so the "
        "deleted parent cannot keep signing in to the portal.")

    add_heading(doc, "1.1 Reproduction", 2)
    add_numbered(doc, [
        "Sign in as superadmin@naijaprimeschool.ng / Admin@12345.",
        "Open Family → Students, edit a pupil who has at least one "
        "parent linked.",
        "Switch to the Parents tab and click Unlink on each row. The "
        "notification says \"Unlinked. '<parent>' removed.\" — the "
        "soft-delete on StudentParent is committed.",
        "Without refreshing or signing out, navigate to Family → "
        "Parents and click Delete on the parent you just unlinked.",
        "Observe the misleading error: \"Cannot delete a parent who is "
        "still linked to a student. Unlink them first.\"",
        "Open the database and run SELECT COUNT(*) FROM StudentParents "
        "WHERE ParentId = '<id>' AND IsDeleted = 0 — the answer is 0.",
    ])

    add_heading(doc, "1.2 Acceptance criteria", 2)
    add_numbered(doc, [
        "After unlinking every pupil from a parent (via Edit student → "
        "Parents → Unlink), pressing Delete on the parent in Parents "
        "succeeds.",
        "If at least one active StudentParent row still points at the "
        "parent, the delete is refused with an error that names the "
        "actual remaining count (e.g. \"Cannot delete this parent "
        "because 2 active student link(s) remain. Unlink them first.\").",
        "When the parent does delete, the linked ApplicationUser (if "
        "any) is also soft-deleted. The parent can no longer sign in "
        "to the portal.",
        "No new tables, no migration, no UI change — the fix is "
        "entirely behind the IParentService boundary.",
    ])

    add_heading(doc, "1.3 What this sprint does not do", 2)
    add_bullets(doc, [
        "It does not auto-unlink pupils on delete. The office still "
        "has to remove the StudentParent rows on the Edit student → "
        "Parents tab first — the new count-aware error message just "
        "tells them whether they actually have anything left to do.",
        "It does not change StudentService.SoftDeleteAsync, which has "
        "the same in-memory-count vs. database-count exposure on "
        "student.Enrolments. The bug has not been reported there yet "
        "and the predicate (WithdrawnOn == null) tends to be more "
        "forgiving, but a symmetric hotfix is a sensible follow-up.",
        "It does not wrap the parent-delete-then-user-delete pair in "
        "an explicit IDbContextTransaction. Today both writes share "
        "the same EF Core context, so a failure between them is "
        "observable as \"parent gone but user remains\" and reversible "
        "by hand — see the Follow-ups chapter.",
    ])

    add_heading(doc, "1.4 Files touched", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs\n"
        "README.md\n"
        "Sprint 10 - Implementation Guide.md (markdown companion)\n"
        "tools/generate_sprint10_guide.py (this generator)"
    )

    add_page_break(doc)


# ---------- Chapter 2 -------------------------------------------------------

def chapter2_pre_sprint(doc):
    add_heading(doc, "2. The trap: EF Core relationship fix-up", 1)
    add_para(doc,
        "EF Core does two related things that interact badly here. "
        "First, it ships a global query filter mechanism that you can "
        "register with HasQueryFilter, and which it applies to every "
        "top-level SELECT it generates for that entity (including "
        "Include navigations). Second, it ships a relationship fix-up "
        "mechanism that, after materialising the result of a query, "
        "walks the change tracker and stitches navigation properties "
        "from whatever else it already knows. The first respects the "
        "filter. The second does not. This sprint is what happens when "
        "you rely on the navigation collection for a guard that should "
        "reflect the live database state.")

    add_heading(doc, "2.1 The query filter on StudentParent", 2)
    add_para(doc,
        "Every entity in the system soft-deletes through ApplyAuditAndSoftDelete "
        "in ApplicationDbContext: a Delete becomes a Modified write that "
        "stamps IsDeleted = true, DeletedOn, and DeletedBy. The matching "
        "global filter on StudentParent hides those rows from every "
        "subsequent SELECT.")

    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs")
    add_code(doc,
        "builder.Entity<StudentParent>(b =>\n"
        "{\n"
        "    b.ToTable(\"StudentParents\");\n"
        "    // ... HasOne / HasMany / unique index ...\n"
        "    b.HasIndex(sp => sp.IsDeleted);\n"
        "    b.HasQueryFilter(sp => !sp.IsDeleted);\n"
        "});")

    add_caption(doc, "Excerpt — ApplyAuditAndSoftDelete (same file)")
    add_code(doc,
        "if (entry.Entity is ISoftDelete softDelete && entry.State == EntityState.Deleted)\n"
        "{\n"
        "    entry.State = EntityState.Modified;\n"
        "    softDelete.IsDeleted = true;\n"
        "    softDelete.DeletedOn = now;\n"
        "    softDelete.DeletedBy = userName;\n"
        "}")

    add_heading(doc, "2.2 The unlink path", 2)
    add_para(doc,
        "StudentService.UnlinkParentAsync looks innocent: load the link, "
        "call db.StudentParents.Remove(link), save. ApplyAuditAndSoftDelete "
        "converts the Delete to a Modified write that sets IsDeleted = "
        "true. The catch is that the StudentParent entity stays "
        "tracked in the DbContext after SaveChanges — its EntityState "
        "is Unchanged (because its modified values were persisted) and "
        "its IsDeleted property is true.")

    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Infrastructure/Services/StudentService.cs")
    add_code(doc,
        "public async Task<OperationResult> UnlinkParentAsync(Guid linkId, CancellationToken ct = default)\n"
        "{\n"
        "    var link = await db.StudentParents.FirstOrDefaultAsync(l => l.Id == linkId, ct);\n"
        "    if (link is null) return OperationResult.Failure(\"Link not found.\");\n"
        "\n"
        "    db.StudentParents.Remove(link);\n"
        "    await db.SaveChangesAsync(ct);\n"
        "    return OperationResult.Success();\n"
        "}")

    add_heading(doc, "2.3 Why the navigation count went wrong", 2)
    add_para(doc,
        "The original ParentService.SoftDeleteAsync loaded the parent "
        "with Include(p => p.StudentLinks). The generated SQL applied "
        "the StudentParent filter, so the result set was empty for a "
        "parent whose links had all been unlinked. But EF Core then "
        "ran relationship fix-up against the change tracker and added "
        "the already-tracked, soft-deleted StudentParent rows to "
        "parent.StudentLinks. The guard then tested parent.StudentLinks.Count "
        "> 0, saw the pre-unlink count, and refused the delete.")

    add_caption(doc, "Pre-sprint-10 ParentService.SoftDeleteAsync (now replaced)")
    add_code(doc,
        "public async Task<OperationResult> SoftDeleteAsync(Guid id, CancellationToken ct = default)\n"
        "{\n"
        "    var parent = await db.Parents\n"
        "        .Include(p => p.StudentLinks)\n"
        "        .FirstOrDefaultAsync(p => p.Id == id, ct);\n"
        "\n"
        "    if (parent is null) return OperationResult.Failure(\"Parent not found.\");\n"
        "\n"
        "    if (parent.StudentLinks.Count > 0)\n"
        "        return OperationResult.Failure(\n"
        "            \"Cannot delete a parent who is still linked to a student. Unlink them first.\");\n"
        "\n"
        "    db.Parents.Remove(parent);\n"
        "    await db.SaveChangesAsync(ct);\n"
        "    return OperationResult.Success();\n"
        "}")

    add_para(doc,
        "The asymmetry is the bug: the database says zero, the change "
        "tracker says N, the guard reads the change tracker. In a "
        "Blazor Server circuit the ApplicationDbContext is scoped per "
        "request and the navigation can repopulate from prior request "
        "state in long-lived circuits, which is exactly the workflow "
        "the office uses (open Edit student, unlink several pupils, "
        "then jump straight to Parents → Delete).")

    add_page_break(doc)


# ---------- Chapter 3 -------------------------------------------------------

def chapter3_fix(doc):
    add_heading(doc, "3. The fix", 1)
    add_para(doc,
        "ParentService.SoftDeleteAsync no longer Includes the "
        "StudentLinks navigation. Instead it loads the parent on its "
        "own and counts active links with a fresh top-level query "
        "against db.StudentParents. That query respects the global "
        "filter at the database level and does not consult the change "
        "tracker, so soft-deleted rows are correctly invisible. The "
        "error message also grew a count so the office can see at a "
        "glance whether something is genuinely still linked.")

    add_heading(doc, "3.1 New ParentService.SoftDeleteAsync", 2)
    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs")
    add_code(doc,
        "public async Task<OperationResult> SoftDeleteAsync(Guid id, CancellationToken ct = default)\n"
        "{\n"
        "    var parent = await db.Parents.FirstOrDefaultAsync(p => p.Id == id, ct);\n"
        "    if (parent is null) return OperationResult.Failure(\"Parent not found.\");\n"
        "\n"
        "    // Count active links via a fresh database query rather than the\n"
        "    // Parent.StudentLinks navigation. EF Core relationship fix-up\n"
        "    // populates that navigation with every StudentParent currently in\n"
        "    // the change tracker that matches the FK — including soft-deleted\n"
        "    // rows that an earlier unlink in the same circuit marked\n"
        "    // IsDeleted = true. The global query filter only affects new\n"
        "    // SELECTs, not the in-memory graph, so the navigation count would\n"
        "    // stay > 0 after a successful unlink and block this delete.\n"
        "    var activeLinks = await db.StudentParents\n"
        "        .CountAsync(l => l.ParentId == id, ct);\n"
        "\n"
        "    if (activeLinks > 0)\n"
        "        return OperationResult.Failure(\n"
        "            $\"Cannot delete this parent because {activeLinks} active student link(s) remain. Unlink them first.\");\n"
        "\n"
        "    db.Parents.Remove(parent);\n"
        "    await db.SaveChangesAsync(ct);\n"
        "\n"
        "    // Sprint 9 auto-provisions an ApplicationUser when a parent is\n"
        "    // created. Mirror that on the way out so a deleted parent cannot\n"
        "    // sign in to the portal and hit the \"we can't find your record\"\n"
        "    // fallback card forever.\n"
        "    if (parent.UserId is { } userId)\n"
        "    {\n"
        "        var user = await userManager.FindByIdAsync(userId.ToString());\n"
        "        if (user is not null)\n"
        "        {\n"
        "            await userManager.DeleteAsync(user);\n"
        "        }\n"
        "    }\n"
        "\n"
        "    return OperationResult.Success();\n"
        "}")

    add_heading(doc, "3.2 Why retire the ApplicationUser?", 2)
    add_para(doc,
        "Sprint 9 made every new parent come with a sign-in (a fresh "
        "ApplicationUser in the Parent role, with Parent.UserId stamped "
        "on the way through CreateAsync). If we soft-delete the parent "
        "but leave the user alive, two unhappy paths open up:")
    add_bullets(doc, [
        "The (now-orphaned) parent can keep signing in. The portal's "
        "Parent dashboard runs PortalService.ResolveParentIdForCurrentUserAsync, "
        "which honours the IsDeleted query filter and therefore returns "
        "null — so the parent lands on the friendly \"We can't find "
        "your record\" fallback card forever. That is more confusing "
        "than \"your account has been disabled\".",
        "The unique filtered index on Parent.UserId (\"[UserId] IS NOT "
        "NULL\") would refuse to let the office create a fresh parent "
        "linked to the same user without manually unlinking first.",
    ])
    add_para(doc,
        "Calling UserManager.DeleteAsync flows through the same "
        "scoped ApplicationDbContext and the same ApplyAuditAndSoftDelete "
        "interceptor that handles every other entity, so the user "
        "row is preserved with IsDeleted = true rather than physically "
        "removed. The audit trail (CreatedOn/By, ModifiedOn/By, "
        "DeletedOn/By, the Identity stamps) survives.")

    add_heading(doc, "3.3 Why no explicit transaction?", 2)
    add_para(doc,
        "The same reasoning as sprint 9: UserManager and our service "
        "share the same scoped DbContext, so the two writes execute "
        "sequentially against the same SQL Server connection. A "
        "failure between SaveChangesAsync (parent committed) and "
        "UserManager.DeleteAsync (user delete) is observable as "
        "\"parent gone but user remains\" — easy to spot in the Users "
        "page and easy to fix by hand. Wrapping the pair in "
        "IDbContextTransaction is a sensible future hardening; it is "
        "deliberately left out of this hotfix so the diff stays small.")

    add_page_break(doc)


# ---------- Chapter 4 -------------------------------------------------------

def chapter4_listing(doc):
    add_heading(doc, "4. Full updated listing", 1)
    add_para(doc,
        "For reference, here is the complete ParentService.cs after "
        "the hotfix. Only SoftDeleteAsync changed — everything else is "
        "the sprint 9 baseline.")
    add_file(doc,
             "src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs")

    add_page_break(doc)


# ---------- Chapter 5 -------------------------------------------------------

def chapter5_test(doc):
    add_heading(doc, "5. How to test end-to-end", 1)

    add_heading(doc, "5.1 Happy path: unlink then delete", 2)
    add_numbered(doc, [
        "Sign in as superadmin@naijaprimeschool.ng / Admin@12345.",
        "Pick a pupil with at least one parent linked.",
        "Open Family → Students → Edit, switch to the Parents tab, and "
        "click Unlink on every link. The notification confirms each "
        "unlink.",
        "Without refreshing the page or signing out, navigate to "
        "Family → Parents and click Delete on the parent that was just "
        "unlinked.",
        "Confirm the dialog. The delete completes; the parent "
        "disappears from the Parents list.",
        "Sign out, then attempt to sign in with that parent's username "
        "and password (the credentials sprint 9 captured on create). "
        "The sign-in is refused — the linked ApplicationUser is now "
        "soft-deleted.",
    ])

    add_heading(doc, "5.2 Negative path: links still active", 2)
    add_numbered(doc, [
        "Sign in as the SuperAdmin.",
        "Pick a parent who is linked to at least one pupil. Do not "
        "unlink.",
        "Open Family → Parents and click Delete on that parent.",
        "Observe the count-aware error: \"Cannot delete this parent "
        "because N active student link(s) remain. Unlink them first.\" "
        "— where N is the actual number of active StudentParent rows.",
        "Open Family → Students → Edit on the linked pupil, switch to "
        "the Parents tab, and click Unlink on each row.",
        "Return to Family → Parents and click Delete again. The "
        "delete now succeeds.",
    ])

    add_heading(doc, "5.3 Database verification", 2)
    add_para(doc,
        "If you want to confirm the data shape from the database "
        "directly, three SQL queries are useful.")

    add_caption(doc, "Confirm the parent row is now soft-deleted")
    add_code(doc,
        "SELECT Id, FirstName, LastName, IsDeleted, DeletedOn, DeletedBy\n"
        "FROM Parents\n"
        "WHERE Id = '<parent-id>';")

    add_caption(doc, "Confirm the linked ApplicationUser is retired")
    add_code(doc,
        "SELECT u.Id, u.UserName, u.Email, u.IsDeleted, u.DeletedOn, u.DeletedBy\n"
        "FROM Users u\n"
        "WHERE u.Id = '<parent.UserId>';")

    add_caption(doc, "Confirm all StudentParents links are soft-deleted")
    add_code(doc,
        "SELECT Id, StudentId, ParentId, IsDeleted, DeletedOn, DeletedBy\n"
        "FROM StudentParents\n"
        "WHERE ParentId = '<parent-id>';")

    add_page_break(doc)


# ---------- Chapter 6 -------------------------------------------------------

def chapter6_followups(doc):
    add_heading(doc, "6. Follow-ups & known limitations", 1)
    add_bullets(doc, [
        "StudentService.SoftDeleteAsync has the same in-memory-count "
        "vs. database-count exposure on student.Enrolments (it runs "
        ".Include(s => s.Enrolments).Any(e => e.WithdrawnOn == null) "
        "after a similar Include). The bug has not been reported "
        "there yet and the WithdrawnOn-based predicate is more "
        "forgiving, but the symmetric hotfix is a sensible follow-up.",
        "Other services across the codebase that guard on "
        "Include + .Count / .Any should be audited; the recommended "
        "pattern is a fresh, top-level CountAsync / AnyAsync against "
        "the relevant DbSet so the global query filter is honoured and "
        "the change tracker is bypassed.",
        "The parent-delete-then-user-delete pair would benefit from "
        "an explicit IDbContextTransaction so a failure between the "
        "two writes leaves no daylight. Today both share the same "
        "scoped DbContext and a failure is observable as \"parent "
        "gone but user remains\" — easy to spot in the Users page and "
        "reversible by hand.",
        "There is no UI to undo a soft-delete on a parent or its "
        "linked user. Use IgnoreQueryFilters() in the relevant "
        "service or run a SQL UPDATE clearing IsDeleted = 0 if a "
        "deletion needs to be reversed.",
    ])

    add_para(doc,
        "With sprint 10 shipped, the parent lifecycle (create → link → "
        "unlink → delete → user retired) is now coherent end to end. "
        "The next sprint can move on to the items on the roadmap "
        "(notifications, two-way messaging, online fee payment, the "
        "audit-log viewer) without worrying that the office will "
        "stack up undeletable parent rows in the directory.")


# ---------- Build ------------------------------------------------------------

def build():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_pre_sprint(doc)
    chapter3_fix(doc)
    chapter4_listing(doc)
    chapter5_test(doc)
    chapter6_followups(doc)
    out_path = ROOT / OUTPUT
    doc.save(str(out_path))
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    build()
