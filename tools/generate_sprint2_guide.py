"""Generates 'Sprint 2 - Implementation Guide.docx' covering ONLY the
academic-domain work delivered in sprint 2 (Sessions, Terms, Classes,
Subjects, Periods, Timetable).

This is the long-form (40+ page) edition. Code blocks embed actual source
files from the repository so the guide stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint2_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 2 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    """Embed a source file with a caption labelling its path."""
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    """Embed a slice of a file between markers (or N lines after start)."""
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 2 — Academic Domain")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Sessions · Terms · Classes · Subjects · Timetables")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/2-academic-domain")
    meta.add_run("\nBuilt on: Sprint 1 identity foundation")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 2 in context", 1)
    add_para(doc,
        "Sprint 2 layers the academic domain on top of the identity foundation "
        "shipped in sprint 1. It delivers the structural skeleton on which every "
        "future feature will hang — students attend Classes, attendance and "
        "results are recorded against Terms, fees are tied to Sessions, and "
        "every classroom interaction takes place inside a slot of the Timetable. "
        "Without this sprint there is nowhere for the rest of the application to "
        "sit. The work is therefore foundational rather than user-facing in the "
        "obvious sense, and it tries hard to keep the data model honest, so that "
        "later sprints do not have to re-shape the schema to accommodate "
        "real-world Nigerian schools.")
    add_para(doc,
        "This document is a long-form implementation guide. It is written so "
        "that an engineer who has read the sprint 1 guide and has the codebase "
        "checked out can recreate every change in this sprint without referring "
        "to the diff. The document is organised in roughly the order I built "
        "the code in: design decisions first, Domain entities next, Application "
        "contracts after that, Infrastructure (DbContext, services, seeder, "
        "migration) in the middle, and finally the Razor UI and navigation. "
        "There is also a smoke-test chapter near the end that walks through the "
        "happy-path you can use to confirm a fresh checkout works.")

    add_heading(doc, "1.1 Where this sits relative to sprint 1", 2)
    add_para(doc,
        "Sprint 1 delivered authentication, role-based authorization, the "
        "BaseEntity pattern, the SaveChanges audit and soft-delete rewriter, "
        "the OperationResult shape, and the Radzen-themed admin shell. Every "
        "single one of those is reused in sprint 2 unchanged. In particular:")
    add_bullets(doc, [
        "BaseEntity — every academic entity inherits from it, picking up Guid "
        "Id, IAuditable, and ISoftDelete with no boilerplate.",
        "ApplicationDbContext.SaveChanges — the override stamps "
        "CreatedOn/By and ModifiedOn/By and rewrites Delete to "
        "IsDeleted = true. Every academic write therefore inherits auditing "
        "and soft delete with zero changes to the override.",
        "Global query filters — every academic entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries without service code having to remember to filter.",
        "ICurrentUser / CurrentUserAccessor — the academic services do not "
        "consume ICurrentUser directly, but the DbContext does, and that is "
        "what stamps audit columns when they save.",
        "OperationResult / OperationResult<T> — every academic service uses "
        "this for predictable success/failure responses without throwing on "
        "expected validation errors.",
        "ILookupService — already existed for titles, genders, and roles. "
        "Sprint 2 adds four more methods (term types, class levels, week days, "
        "teachers) and rewires the service to inject UserManager so the "
        "teachers list can use Identity correctly.",
        "Radzen Blazor + the green/gold app.css — the academic pages adopt the "
        "same .nps-page-header / .nps-card / .nps-form-grid primitives so "
        "they read as part of the same product, not a bolt-on.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_para(doc,
        "Concretely, after this sprint a SuperAdmin or HeadTeacher signing in "
        "to the application can:")
    add_numbered(doc, [
        "Create the academic year (Session), set its date range, mark it as "
        "the current session, and edit or delete it later.",
        "Define the three (or however many) Terms within each session, mark "
        "one of them current, and view or filter terms by session.",
        "Set up class arms (Primary 1A, JSS 2B …) tied to a session and a "
        "class level, with an optional class teacher pulled from active "
        "users in the Teacher role.",
        "Maintain the catalogue of taught Subjects, including a unique short "
        "code that flows through to timetables and result sheets later.",
        "Configure the school day's TimetablePeriods — start time, end time, "
        "display order, and whether the period is a break.",
        "Lay out the weekly timetable for any (term, class) combination by "
        "clicking each period × weekday cell to assign a Subject and "
        "(optionally) a Teacher, Room, and Notes.",
        "Soft-delete any of the above, with friendly errors if the row is "
        "still referenced elsewhere (e.g. you cannot delete a Session that "
        "still has Terms, or a Class that still has timetable entries).",
    ])
    add_para(doc,
        "Teachers — already a seeded role from sprint 1 — get read access to "
        "the timetable so they can look up where they should be teaching. "
        "Other roles (Parent, Student, Bursar, Storekeeper) do not see the "
        "Academics navigation panel; their menu items remain placeholder "
        "items reserved for later sprints.")

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_para(doc,
        "It is just as important to be explicit about what sprint 2 does NOT "
        "do, because every one of these has been weighed and consciously "
        "deferred. Trying to land them now would have pushed scope past what "
        "is comfortably reviewable in one branch, and several of them depend "
        "on data we have not modelled yet:")
    add_bullets(doc, [
        "Student enrolment into a class. Without students there can be no "
        "attendance records, no result sheets, and no parent portal — those "
        "are sprint 3 territory.",
        "Attendance capture. Once enrolment is in place, attendance is just "
        "a per-day, per-class sheet of attendance status rows; the timetable "
        "structure laid down here makes that trivial.",
        "Assessment definition and result computation. Sprint 4 territory; "
        "scores will hang off Subject × Term × Student.",
        "Fees and invoices. Tied to Session and Class, but driven by the "
        "bursar's own workflow which is sprint 5 territory.",
        "Concurrency safeguards. The current upsert on a timetable cell uses "
        "the natural key (term, class, weekday, period) but does not yet "
        "detect a teacher being booked into two classes at the same time. "
        "That validation should live alongside teacher workload reporting in "
        "a later sprint.",
        "Export to PDF / Excel. The timetable grid is rendered with plain "
        "HTML and CSS, so a PDF export is mostly a CSS print-stylesheet job; "
        "we will add it the first time a user genuinely asks for it.",
        "Bulk operations. There is no 'copy this term's timetable to next "
        "term' button yet, no 'bulk-create classes from a template'. We can "
        "build that on top of the existing services without changing the "
        "schema, so it is a feature-add, not a redesign.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers, this sprint adds:")
    add_bullets(doc, [
        "9 new domain entities under src/NaijaPrimeSchool.Domain/Academics/.",
        "5 DTO files (one per feature area) under src/NaijaPrimeSchool.Application/Academics/Dtos/.",
        "5 service contracts under src/NaijaPrimeSchool.Application/Academics/.",
        "5 service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "1 ILookupService extension covering 4 new lookups.",
        "1 EF Core migration introducing 9 new tables and 14 indexes.",
        "1 DatabaseInitializer extension seeding TermTypes, ClassLevels, "
        "WeekDays, and a default set of TimetablePeriods.",
        "6 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Academics/.",
        "1 navigation menu rewrite to surface the new pages.",
        "1 set of CSS additions in wwwroot/app.css for the timetable grid.",
    ])
    add_para(doc,
        "Everything compiles with zero warnings on .NET 10 (the team-wide "
        "warning bar). The code follows the patterns already accepted in "
        "sprint 1, so the diff is low-friction to review.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)
    add_para(doc,
        "Before any code was written I made a small number of architectural "
        "calls that shaped everything that followed. Re-reading these in "
        "isolation makes the code easier to navigate later, and it gives "
        "future maintainers permission to revisit choices when the trade-offs "
        "change.")

    add_heading(doc, "2.1 No enums — lookups are first-class tables", 2)
    add_para(doc,
        "Sprint 1 set the rule that domain concepts which would normally be "
        "C# enums are stored as tables instead. Sprint 2 honours that without "
        "exception. TermType, ClassLevel, and WeekDay are all proper entities "
        "that derive from BaseEntity and live in the database. The reasons "
        "remain the same as sprint 1:")
    add_bullets(doc, [
        "Translatability. A future Hausa or Yoruba translation only needs to "
        "edit Name in a row, not recompile a deployment.",
        "Extensibility by data. If a school chooses to run a fourth term, the "
        "TermType table accepts a new row; no code change is needed.",
        "Reportability. Joins to lookup tables surface readable labels in "
        "every reporting tool; magic enum integers do not.",
        "Display order. Each lookup carries a DisplayOrder column which the "
        "UI uses to sort dropdowns. Enums would have to lean on the integer "
        "value or impose an attribute-based ordering.",
    ])
    add_para(doc,
        "The cost is a few extra joins on read paths, which EF Core handles "
        "efficiently because the lookup tables are tiny and effectively "
        "always cached in the buffer pool.")

    add_heading(doc, "2.2 DateOnly and TimeOnly", 2)
    add_para(doc,
        "Session.StartDate and TimetablePeriod.StartTime are deliberately "
        "modelled as DateOnly and TimeOnly, the .NET 6+ types that map cleanly "
        "to SQL Server's date and time columns without dragging along the time "
        "or date component nobody asked for.")
    add_para(doc,
        "Using DateTime for a Session start date would have invited a string "
        "of subtle bugs — local-time vs UTC interpretation, midnight vs 23:59 "
        "boundaries, and Radzen pickers showing time controls users do not "
        "want. The Razor pages still bind Radzen pickers to DateTime? in the "
        "form model (because Radzen has not yet adopted DateOnly/TimeOnly) "
        "and then convert at the boundary using DateOnly.FromDateTime and "
        "TimeOnly.FromDateTime. The conversion is the cheapest piece of "
        "ceremony in the entire sprint.")

    add_heading(doc, "2.3 IsCurrent flags vs separate 'active' lookup", 2)
    add_para(doc,
        "Both Session and Term carry an IsCurrent boolean, even though only "
        "one row of each is supposed to be current at a time. I considered "
        "making 'current session' a separate single-row table or a "
        "configuration setting; I rejected those for two reasons.")
    add_bullets(doc, [
        "Discoverability. A column called IsCurrent on the Session row is "
        "obvious to anyone reading the schema. A separate CurrentSession table "
        "is not.",
        "Reportability. Reports that filter to historic sessions can do so "
        "with a simple WHERE clause; with a separate table they need an outer "
        "join.",
    ])
    add_para(doc,
        "The trade-off is that the service has to enforce 'only one current "
        "row at a time' explicitly. SessionService and TermService both do "
        "this with a single ExecuteUpdateAsync that flips IsCurrent = false "
        "on every other row before saving the new current row, all inside the "
        "same SaveChanges and therefore inside the same implicit transaction.")

    add_heading(doc, "2.4 Composite unique indexes", 2)
    add_para(doc,
        "Three composite unique indexes do most of the structural integrity "
        "heavy lifting:")
    add_bullets(doc, [
        "(SessionId, TermTypeId) on Term — a session cannot have two First "
        "Terms.",
        "(SessionId, Name) on SchoolClass — class names are unique within a "
        "session, but Primary 1A in 2024/2025 and Primary 1A in 2025/2026 "
        "are legitimately different rows.",
        "(TermId, SchoolClassId, WeekDayId, TimetablePeriodId) on "
        "TimetableEntry — at most one lesson per slot.",
    ])
    add_para(doc,
        "The services also pre-check these constraints with explicit "
        "AnyAsync queries so the user gets a friendly OperationResult.Failure "
        "message rather than a database exception. The unique index is the "
        "ultimate authority; the application-level check is the friendly UI.")

    add_heading(doc, "2.5 Foreign-key delete behaviour", 2)
    add_para(doc,
        "Required navigations on TimetableEntry use OnDelete(Restrict) — "
        "you cannot delete a Subject or a Term while it is referenced by "
        "any timetable row. Optional navigations (TimetableEntry.Teacher, "
        "SchoolClass.ClassTeacher) use OnDelete(SetNull) so deactivating or "
        "removing a teacher does not cascade through the schedule. This "
        "matches the soft-delete model: when an ApplicationUser is "
        "soft-deleted the foreign keys remain valid in storage but the global "
        "query filter hides the user, and our DTO projections surface "
        "ClassTeacherName as null in those cases.")

    add_heading(doc, "2.6 Service per feature area, not per entity", 2)
    add_para(doc,
        "Five services split the work along feature lines: ISessionService, "
        "ITermService, ISchoolClassService, ISubjectService, and "
        "ITimetableService. The last one bundles two closely related concepts "
        "(timetable periods and timetable entries) because callers almost "
        "always need both at once and a fresh interface for each would have "
        "fragmented the API.")
    add_para(doc,
        "I rejected the alternatives of one giant IAcademicService (would "
        "become unwieldy) and one tiny per-entity service (would force the "
        "Razor pages to coordinate two services for any cross-aggregate "
        "operation). The five-service split is what the UI naturally needs.")

    add_heading(doc, "2.7 Inline forms vs Radzen dialogs", 2)
    add_para(doc,
        "Every CRUD page uses an inline form pattern — clicking New or Edit "
        "reveals a RadzenCard below the data grid containing the form, rather "
        "than opening a Radzen dialog. The reasons:")
    add_bullets(doc, [
        "Validation messages are easier to read inline; modal dialogs often "
        "obscure feedback or push it under the keyboard on small screens.",
        "Inline forms are trivially testable in Playwright/Selenium because "
        "they share the same DOM as the list — no dialog handle to chase.",
        "Server-side rendering of the form is straightforward; dialog "
        "components require extra ceremony around DialogService and child "
        "component lifetimes.",
        "The visual rhythm of the page (header → filter → grid → form) "
        "matches what an admin already does mentally: filter, scan, edit.",
    ])
    add_para(doc,
        "The one exception is the timetable grid, where the cell-level "
        "editor still uses an inline RadzenCard rather than a dialog, "
        "consistent with the rest of the sprint.")

    add_heading(doc, "2.8 Soft delete revisited for the academic domain", 2)
    add_para(doc,
        "The soft-delete pattern from sprint 1 is preserved here, but the "
        "academic services add a new wrinkle: refusing to delete a row that "
        "still has dependents. For example, SessionService.SoftDeleteAsync "
        "loads the session with its Terms and Classes collections and "
        "returns OperationResult.Failure if either is non-empty. The same "
        "pattern guards Class deletion against existing TimetableEntries, "
        "Subject deletion against TimetableEntries, and Period deletion "
        "against TimetableEntries.")
    add_para(doc,
        "An alternative would have been to cascade-soft-delete dependents, "
        "so deleting a session also deletes its terms. I chose the explicit "
        "refusal because cascade soft delete is hard to undo accurately — "
        "if a row is restored later you have to remember which rows were "
        "deleted because of it, versus which were already deleted on their "
        "own. Refusing the delete keeps the audit trail readable.")

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)
    add_para(doc,
        "Every academic entity lives in a single new folder, "
        "src/NaijaPrimeSchool.Domain/Academics/. There are no abstract base "
        "classes other than BaseEntity, and the entities are deliberately "
        "anaemic — no domain methods, no validation logic. Validation lives "
        "in the Application DTOs (DataAnnotations) and in the Infrastructure "
        "services (cross-aggregate checks). The Domain layer is a typed "
        "vocabulary; it is not the place for behaviour in this codebase.")

    add_heading(doc, "3.1 Folder layout and namespacing", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "├── Academics/                ← (new in sprint 2)\n"
        "│   ├── ClassLevel.cs         ← lookup\n"
        "│   ├── SchoolClass.cs        ← class arm\n"
        "│   ├── Session.cs            ← academic year\n"
        "│   ├── Subject.cs\n"
        "│   ├── Term.cs               ← session × term-type\n"
        "│   ├── TermType.cs           ← lookup\n"
        "│   ├── TimetableEntry.cs     ← (term, class, day, period) → subject + teacher\n"
        "│   ├── TimetablePeriod.cs    ← school-day slot\n"
        "│   └── WeekDay.cs            ← lookup, Mon–Fri\n"
        "├── Common/                   ← from sprint 1\n"
        "└── Identity/                 ← from sprint 1\n")
    add_para(doc,
        "Namespace is NaijaPrimeSchool.Domain.Academics. The four lookup "
        "entities (TermType, ClassLevel, WeekDay) share that namespace because "
        "they are conceptually part of the academic domain rather than the "
        "Identity domain — keeping them together makes the dependency graph "
        "obvious.")

    add_heading(doc, "3.2 Entity-by-entity tour", 2)

    add_heading(doc, "3.2.1 Session.cs", 3)
    add_para(doc,
        "Session is the academic year root. Its Name follows the Nigerian "
        "convention of two-year strings (e.g. '2025/2026'), constrained to "
        "40 characters at the database level. StartDate and EndDate are "
        "DateOnly, so the column type is SQL Server's date rather than "
        "datetime2. IsCurrent is the at-most-one-row flag enforced by the "
        "service layer.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/Session.cs")
    add_para(doc,
        "The two collection navigations (Terms and Classes) are convenience "
        "back-references the service layer uses for blocking-dependent "
        "checks during soft-delete. EF Core does not require them; they are "
        "purely for readability.")

    add_heading(doc, "3.2.2 TermType.cs", 3)
    add_para(doc,
        "A pure lookup with Name (e.g. 'First Term'), DisplayOrder for the "
        "dropdown ordering, and the back-reference to its Terms.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/TermType.cs")

    add_heading(doc, "3.2.3 Term.cs", 3)
    add_para(doc,
        "Term carries its Session and TermType as required navigations, plus "
        "its own date range and IsCurrent flag. The back-reference to "
        "TimetableEntry helps the service block deletion of a term that "
        "still has lessons scheduled.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/Term.cs")

    add_heading(doc, "3.2.4 ClassLevel.cs", 3)
    add_para(doc,
        "ClassLevel is the academic level: Creche, Pre-Nursery, Nursery 1, "
        "KG 1, Primary 1 through Primary 6, and so on. The seeded list "
        "currently stops at Primary 6 because junior-secondary support "
        "(JSS 1–3) is not on the sprint 2 board, but adding rows is a data "
        "change, not a code change.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/ClassLevel.cs")

    add_heading(doc, "3.2.5 SchoolClass.cs", 3)
    add_para(doc,
        "SchoolClass is the actual class arm. Note the optional ClassTeacher "
        "navigation pointing at ApplicationUser — that crosses into the "
        "Identity namespace and demonstrates that the academic domain is not "
        "isolated from the user model.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs")

    add_heading(doc, "3.2.6 Subject.cs", 3)
    add_para(doc,
        "Subject has both a Name (Mathematics) and a Code (MTH). Both are "
        "unique-indexed at the schema level. The service uppercases Code on "
        "save so 'MTH' and 'mth' cannot be created as separate subjects.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/Subject.cs")

    add_heading(doc, "3.2.7 WeekDay.cs", 3)
    add_para(doc,
        "WeekDay is the lookup of working days. ShortName powers the compact "
        "headers in the timetable grid; only Monday through Friday are "
        "seeded today, but the schema allows Saturday school days to be "
        "added by inserting a row.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/WeekDay.cs")

    add_heading(doc, "3.2.8 TimetablePeriod.cs", 3)
    add_para(doc,
        "TimetablePeriod is the most data-intensive small entity: it carries "
        "TimeOnly start/end, a DisplayOrder for grid rendering, and an "
        "IsBreak flag the UI uses to grey out break rows. The seeded period "
        "set is realistic for a Nigerian primary school day: seven 40-minute "
        "lesson periods bracketing a Short Break and a Lunch Break.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/TimetablePeriod.cs")

    add_heading(doc, "3.2.9 TimetableEntry.cs", 3)
    add_para(doc,
        "The keystone of the sprint. Five required Guid foreign keys form "
        "the natural identity of the row: Term, SchoolClass, Subject, "
        "WeekDay, TimetablePeriod. TeacherId is optional — for "
        "self-directed work or library periods — and Room and Notes are "
        "free-text annotations.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/TimetableEntry.cs")
    add_para(doc,
        "There is no domain logic on TimetableEntry; the upsert in "
        "TimetableService treats (TermId, SchoolClassId, WeekDayId, "
        "TimetablePeriodId) as the natural key, which is why "
        "TimetableEntry.cs has no methods at all.")

    add_heading(doc, "3.3 Relationships at a glance", 2)
    add_para(doc,
        "An ASCII relationship diagram captures everything the schema "
        "currently knows about the academic side of the world:")
    add_code(doc,
        "                       ┌─────────────┐\n"
        "                       │   Session   │\n"
        "                       │  IsCurrent  │\n"
        "                       └──┬───────┬──┘\n"
        "                          │ 1     │ 1\n"
        "                          │       │\n"
        "                       N  │       │ N\n"
        "                ┌─────────▼──┐ ┌──▼────────────┐\n"
        "                │    Term    │ │  SchoolClass  │\n"
        "                │  IsCurrent │ │ ClassTeacher? │\n"
        "                └──┬─────────┘ └────┬──────────┘\n"
        "                   │ 1               │ 1\n"
        "                   │                 │\n"
        "         ┌─────────┼─────────────────┘\n"
        "         │         │   N\n"
        "         │  ┌──────▼──────────────┐\n"
        "         │  │   TimetableEntry    │\n"
        "         │  │ (term, class, day,  │\n"
        "         │  │  period)→subject+   │\n"
        "         │  │  teacher? + room?   │\n"
        "         │  └─┬───────┬───────┬──┘\n"
        "         │    │ N     │ N     │ N\n"
        "         │    │       │       │\n"
        "         │ ┌──▼──┐ ┌──▼───┐ ┌─▼────────────┐\n"
        "         │ │Subj.│ │Week- │ │ Timetable    │\n"
        "         │ │     │ │Day   │ │ Period       │\n"
        "         │ └─────┘ └──────┘ └──────────────┘\n"
        "         │\n"
        "    Lookups (BaseEntity, soft-deletable):\n"
        "      TermType ↔ Term\n"
        "      ClassLevel ↔ SchoolClass\n")
    add_para(doc,
        "Reading downward in this diagram corresponds to reading 'has many' "
        "in EF terminology. Solid lookups (TermType, ClassLevel, WeekDay) "
        "are referenced from their domain owners with required Guid foreign "
        "keys; only the optional teacher links are nullable.")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)
    add_para(doc,
        "Application is the boundary between the Razor UI and the persistence "
        "layer. It is deliberately thin: just DTOs and service interfaces. "
        "There is no domain logic and no EF Core. Sprint 2 adds two new "
        "subfolders: Academics/Dtos and Academics. The DTOs live in the "
        "first; the contracts in the second.")

    add_heading(doc, "4.1 DTO design rules", 2)
    add_bullets(doc, [
        "Every list view returns a flat read-DTO with denormalised display "
        "fields (e.g. SessionName on TermDto, ClassTeacherName on "
        "SchoolClassDto). This means the UI can render a row without "
        "lazy-loading a navigation.",
        "Create and Update requests are separate types so DataAnnotations "
        "can express the differences (Update has an Id; Create does not). "
        "It also means the API surface is honest about what is required.",
        "DTOs that pair (Create, Update) live in the same .cs file with the "
        "Dto suffix, so adding a new domain concept means one new file, not "
        "three.",
        "DTO properties for ids on optional dropdowns are Guid? rather than "
        "Guid, so the UI can model 'no selection' cleanly.",
    ])

    add_heading(doc, "4.2 SessionDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/Dtos/SessionDtos.cs")
    add_para(doc,
        "TermCount and ClassCount are populated by the service via "
        ".Select(s => new SessionDto { ... TermCount = s.Terms.Count }) so "
        "they are computed by the database, not in memory.")

    add_heading(doc, "4.3 TermDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/Dtos/TermDtos.cs")
    add_para(doc,
        "TermDto carries the denormalised SessionName and TermTypeName so "
        "the Terms grid can show 'Primary 1A · Second Term' without a "
        "second round-trip to the database.")

    add_heading(doc, "4.4 SchoolClassDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/Dtos/SchoolClassDtos.cs")
    add_para(doc,
        "ClassTeacherName is computed by EF Core's projection — see the "
        "service implementation in chapter 6 for the exact projection "
        "expression. Notably, when the teacher is soft-deleted the global "
        "query filter on ApplicationUser hides the row, the projection sees "
        "ClassTeacher == null, and the DTO's ClassTeacherName comes back "
        "null. The UI renders an empty cell. The class itself is not "
        "affected.")

    add_heading(doc, "4.5 SubjectDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/Dtos/SubjectDtos.cs")
    add_para(doc,
        "Code has a hard length cap of 10 characters because anything longer "
        "looks bad in the timetable grid. This is enforced by both "
        "DataAnnotations (so the form blocks submission) and a HasMaxLength "
        "in the DbContext (so the database enforces it on direct SQL).")

    add_heading(doc, "4.6 TimetableDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/Dtos/TimetableDtos.cs")
    add_para(doc,
        "TimetableEntryDto is intentionally fat — it carries every label the "
        "grid needs to render a cell, including the period's start and end "
        "TimeOnly values and the weekday's display order. This makes the "
        "Razor render loop a pure read; no extra service calls are made "
        "while looping.")
    add_para(doc,
        "UpsertTimetableEntryRequest is the one exception to the "
        "Create/Update separation rule, and it is intentional: the timetable "
        "grid's natural key is (TermId, SchoolClassId, WeekDayId, "
        "TimetablePeriodId), not Id, so a single endpoint that inserts or "
        "updates by natural key is the simplest model. The service ignores "
        "any Id supplied by the client.")
    add_para(doc,
        "TimetableQuery is the input to the list endpoint. Sprint 2's only "
        "list scenario is 'load every entry for this term × class', so the "
        "query type holds exactly those two ids. Future scenarios — 'list "
        "every lesson taught by this teacher this week', 'list every lesson "
        "in this room' — will get their own request types.")

    add_heading(doc, "4.7 Service contracts", 2)
    add_para(doc,
        "Five interface files. Read methods return immutable IReadOnlyList; "
        "writes return OperationResult or OperationResult<Guid>. Every "
        "method takes an optional CancellationToken so Blazor's component "
        "cancellation propagates all the way to the database query.")

    add_heading(doc, "4.7.1 ISessionService", 3)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/ISessionService.cs")

    add_heading(doc, "4.7.2 ITermService", 3)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/ITermService.cs")

    add_heading(doc, "4.7.3 ISchoolClassService", 3)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/ISchoolClassService.cs")

    add_heading(doc, "4.7.4 ISubjectService", 3)
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/ISubjectService.cs")

    add_heading(doc, "4.7.5 ITimetableService", 3)
    add_para(doc,
        "ITimetableService bundles the period CRUD and the entry CRUD "
        "together because callers usually need both at once when rendering "
        "the grid. The split inside the file is by section comment.")
    add_file(doc, "src/NaijaPrimeSchool.Application/Academics/ITimetableService.cs")

    add_heading(doc, "4.8 ILookupService extensions", 2)
    add_para(doc,
        "Sprint 1 introduced ILookupService for titles, genders, and roles. "
        "Sprint 2 adds four more methods. The interface as it now stands:")
    add_file(doc, "src/NaijaPrimeSchool.Application/Users/ILookupService.cs")
    add_para(doc,
        "GetTeachersAsync is the only method that returns a list with a "
        "non-static membership: it filters ApplicationUser by role 'Teacher' "
        "and IsActive = true. The implementation lives in the Infrastructure "
        "layer (chapter 6.6) because it depends on UserManager, which is an "
        "ASP.NET Core Identity type.")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)
    add_para(doc,
        "ApplicationDbContext gained nine DbSet properties and a private "
        "ConfigureAcademics(builder) helper invoked from OnModelCreating. "
        "Nothing about the class shape changed — it is still a "
        "primary-constructor IdentityDbContext that takes ICurrentUser and "
        "the same SaveChanges override does the audit and soft-delete work.")

    add_heading(doc, "5.1 New DbSet properties", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="public DbSet<Session> Sessions",
        end_marker="public DbSet<TimetableEntry>")
    add_para(doc,
        "DbSet uses the expression-bodied member style introduced in "
        "EF Core 6 — it is just a thin alias for Set<T>(). I list the "
        "lookups before the entities they describe (TermTypes before Terms, "
        "ClassLevels before SchoolClasses, WeekDays before TimetableEntries) "
        "to read top-to-bottom by causation.")

    add_heading(doc, "5.2 ConfigureAcademics — the model builder block", 2)
    add_para(doc,
        "OnModelCreating now ends with one extra call:")
    add_code(doc,
        "ConfigureLookup<Title>(builder, \"Titles\", extra: b =>\n"
        "{\n"
        "    b.Property(t => t.Name).HasMaxLength(50).IsRequired();\n"
        "    b.HasIndex(t => t.Name).IsUnique();\n"
        "});\n"
        "\n"
        "ConfigureAcademics(builder);   // ← new in sprint 2\n")
    add_para(doc, "ConfigureAcademics is itself a single private static method:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="private static void ConfigureAcademics",
        end_marker="// One subject per (term, class, day, period) slot.",
        caption="Excerpt — beginning of ConfigureAcademics")
    add_para(doc,
        "Each lookup is configured with the existing ConfigureLookup<T> "
        "helper that was introduced in sprint 1. Each non-lookup entity gets "
        "its own builder.Entity<T>(...) block. The full body is too long to "
        "embed here in one piece; the rest of this section breaks down the "
        "decisions piece by piece, with the relevant excerpts.")

    add_heading(doc, "5.3 Session — uniqueness and indexing", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<Session>(b =>",
        end_marker="});",
        caption="Excerpt — Session entity configuration")
    add_para(doc,
        "Three indexes: a unique index on Name to stop two sessions sharing "
        "the same year, and supporting indexes on IsCurrent and IsDeleted to "
        "make 'find current session' and the global query filter cheap.")

    add_heading(doc, "5.4 Term — composite uniqueness", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<Term>(b =>",
        end_marker="});",
        caption="Excerpt — Term entity configuration")
    add_para(doc,
        "OnDelete(Restrict) on both required navigations forces the service "
        "layer to handle deletion order explicitly. If a Session is "
        "soft-deleted while it still has Terms, EF Core will not try to "
        "cascade: instead, the SessionService.SoftDeleteAsync method "
        "pre-checks for blocking dependents and refuses with a friendly "
        "error.")

    add_heading(doc, "5.5 SchoolClass — uniqueness scoped by session", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<SchoolClass>(b =>",
        end_marker="});",
        caption="Excerpt — SchoolClass entity configuration")
    add_para(doc,
        "Note OnDelete(SetNull) on ClassTeacher — a teacher can leave the "
        "school without orphaning their classes. The unique index "
        "(SessionId, Name) lets two different sessions both have a Primary "
        "1A row, which is the right semantics for an academic-year-keyed "
        "schema.")

    add_heading(doc, "5.6 Subject — name and code unique", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<Subject>(b =>",
        end_marker="});",
        caption="Excerpt — Subject entity configuration")
    add_para(doc,
        "Both Name and Code are unique-indexed independently. The service "
        "uppercases Code on save (request.Code.Trim().ToUpperInvariant()) "
        "so the constraint catches case-only collisions.")

    add_heading(doc, "5.7 TimetablePeriod — display order", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<TimetablePeriod>(b =>",
        end_marker="});",
        caption="Excerpt — TimetablePeriod entity configuration")
    add_para(doc,
        "There is no unique constraint on (DisplayOrder) — two periods can "
        "share an order and the grid will sort them by start time as a "
        "secondary key. That is a deliberate easing of the schema in case "
        "two periods genuinely begin in the same minute (which would be "
        "very unusual but not theoretically forbidden).")

    add_heading(doc, "5.8 TimetableEntry — the most-constrained table", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="builder.Entity<TimetableEntry>(b =>",
        end_marker="});",
        caption="Excerpt — TimetableEntry entity configuration")
    add_para(doc,
        "The composite unique index on (TermId, SchoolClassId, WeekDayId, "
        "TimetablePeriodId) is the most important constraint in the entire "
        "sprint. It guarantees that for any (term, class) pair, the timetable "
        "grid never has two rows in the same cell. Even if the upsert "
        "service is bypassed by direct SQL, the row will not insert.")
    add_para(doc,
        "Restrict on the required navigations + SetNull on the optional "
        "Teacher navigation gives the right behaviour for soft delete: a "
        "soft-deleted teacher's lessons remain on the timetable with the "
        "teacher cell blank, while a soft-deleted Subject blocks deletion "
        "until the timetable rows are first cleared.")

    add_heading(doc, "5.9 Soft delete — same plumbing, more entities", 2)
    add_para(doc,
        "Every academic entity inherits IsDeleted/DeletedOn/DeletedBy from "
        "BaseEntity, declares HasQueryFilter(x => !x.IsDeleted) in its "
        "model configuration, and is therefore filtered out of every "
        "ordinary query automatically. The SaveChanges override that does "
        "the actual work was written in sprint 1; sprint 2 did not touch it.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        start_marker="if (entry.Entity is ISoftDelete s",
        end_marker="}",
        lines_after_start=8,
        caption="Excerpt — SaveChanges soft-delete rewrite, unchanged from sprint 1")
    add_para(doc,
        "This is the line that converts every db.Set.Remove(entity) call "
        "in the academic services into an UPDATE that flips IsDeleted to "
        "true and stamps the deleter's name. No service code in sprint 2 "
        "manipulates IsDeleted directly.")

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)
    add_para(doc,
        "Five new service classes live under "
        "src/NaijaPrimeSchool.Infrastructure/Services/. Each is a primary-"
        "constructor class that takes ApplicationDbContext (and "
        "occasionally UserManager) and implements its corresponding "
        "Application interface. None of them throws for an expected "
        "validation failure; everything that the user could reasonably "
        "cause is reported via OperationResult.Failure.")

    add_heading(doc, "6.1 SessionService — full listing", 2)
    add_para(doc,
        "Sessions are simple — there is no parent entity to validate "
        "against — but the IsCurrent flag introduces an interesting "
        "constraint. The service uses ExecuteUpdateAsync to flip every "
        "other current row to false in a single SQL UPDATE, all inside the "
        "same SaveChanges and therefore the same implicit transaction.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/SessionService.cs")
    add_para(doc, "Worth pausing on:")
    add_bullets(doc, [
        "ListAsync orders by IsCurrent descending, then StartDate "
        "descending, so the current session is always at the top of the "
        "grid.",
        "GetCurrentAsync returns the SessionDto for the row flagged "
        "current — this is what other services and pages will use to "
        "default their own session selectors.",
        "Validation is double-layered: SessionDto validation comes from "
        "DataAnnotations on CreateSessionRequest; cross-row validation "
        "(date sanity, name uniqueness) is checked in the service before "
        "any DB write.",
        "SoftDeleteAsync loads the session with its Terms and Classes "
        "collections, refuses if either is non-empty, and then calls "
        "Remove(...) — which the SaveChanges override turns into a soft "
        "delete.",
    ])

    add_heading(doc, "6.2 TermService — full listing", 2)
    add_para(doc,
        "TermService is structurally identical to SessionService but "
        "validates against two parents (Session and TermType) and enforces "
        "the (SessionId, TermTypeId) uniqueness explicitly so the user "
        "gets a friendly message before the database constraint fires.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/TermService.cs")
    add_para(doc,
        "The static Project helper is a recurring pattern in this sprint: "
        "the read methods all share the same mapping from entity "
        "navigations to DTO display fields, so wrapping it in a function "
        "keeps the three list/get/get-current methods three lines each.")

    add_heading(doc, "6.3 SchoolClassService — full listing", 2)
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/SchoolClassService.cs")
    add_para(doc,
        "Two things stand out. First, the projection's ClassTeacherName "
        "uses the EF Core null-propagation pattern "
        "(c.ClassTeacher == null ? null : ...) which translates into a "
        "CASE WHEN in SQL — a global query filter that hides a "
        "soft-deleted teacher will return null, and the DTO will surface a "
        "blank cell. Second, the unique-name check is scoped to "
        "(SessionId, Name) so 2024/2025's Primary 1A and 2025/2026's "
        "Primary 1A are legitimately different rows.")

    add_heading(doc, "6.4 SubjectService — full listing", 2)
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/SubjectService.cs")
    add_para(doc,
        "The .Trim() on Name and the .Trim().ToUpperInvariant() on Code "
        "are deliberate normalisations — admins typing 'Mathematics ' or "
        "'mth' should not be able to create near-duplicates of an existing "
        "subject. The unique indexes are the safety net; the service "
        "normalisation is the friendly UI.")

    add_heading(doc, "6.5 TimetableService — full listing", 2)
    add_para(doc,
        "TimetableService is the most complex service in the sprint. It "
        "owns two unrelated CRUD surfaces (TimetablePeriod and "
        "TimetableEntry) and a non-trivial upsert by natural key. The full "
        "listing follows.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/TimetableService.cs")
    add_para(doc,
        "The interesting half is UpsertEntryAsync. It validates each "
        "supplied foreign key with an AnyAsync probe — turning a database "
        "error into a friendly OperationResult.Failure — then looks up "
        "the existing row by natural key, copies the request fields onto "
        "either the loaded row or a fresh new TimetableEntry, and lets "
        "SaveChanges do the rest. The natural-key lookup is the "
        "behavioural difference between this and a generic CreateOrUpdate "
        "by Id.")
    add_para(doc, "ProjectEntries is shared by ListEntriesAsync and "
        "GetEntryByIdAsync. Notably, it computes the teacher display name "
        "with EF Core's null-safe concatenation — the resulting SQL uses "
        "ISNULL or COALESCE rather than throwing if the teacher is null.")

    add_heading(doc, "6.6 LookupService — full listing", 2)
    add_para(doc,
        "LookupService gained UserManager as a constructor dependency in "
        "sprint 2 so it can implement GetTeachersAsync. The other three "
        "new methods are simple ordered DbSet projections.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs")
    add_para(doc,
        "GetTeachersAsync uses UserManager.GetUsersInRoleAsync — a "
        "hand-written join would be possible but Identity provides the "
        "method already, and using it keeps the code resilient if the "
        "Identity schema changes shape later.")
    add_para(doc,
        "The Where clause filters out IsActive == false in memory after "
        "the role lookup, which is acceptable for the small set of "
        "teachers a primary school typically has. If the school grows past "
        "a few hundred teachers we will rewrite this as a database-side "
        "join, but until then the simple form is more readable.")

    add_heading(doc, "6.7 DatabaseInitializer — seeding the new lookups", 2)
    add_para(doc,
        "DatabaseInitializer.InitializeAsync still does what sprint 1 "
        "wired it to do: apply migrations, seed lookups, seed roles, seed "
        "the SuperAdmin. Sprint 2 added one new step in the middle:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        start_marker="await SeedLookupsAsync(db, ct);",
        end_marker="await SeedSuperAdminAsync",
        caption="Excerpt — InitializeAsync now also seeds academic lookups")
    add_para(doc, "The new method itself:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        start_marker="private static async Task SeedAcademicLookupsAsync",
        end_marker="await db.SaveChangesAsync(ct);\n    }",
        caption="Excerpt — SeedAcademicLookupsAsync")
    add_para(doc,
        "Every block is wrapped in an AnyAsync().IgnoreQueryFilters() "
        "guard so the seeder does not duplicate rows on subsequent "
        "starts, and IgnoreQueryFilters means even rows that have been "
        "soft-deleted count as 'present' — preventing the seeder from "
        "happily re-creating data that an admin had specifically deleted.")

    add_heading(doc, "6.8 Dependency injection registration", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        start_marker="services.AddScoped<IUserService",
        end_marker="return services;",
        caption="Excerpt — DI registrations")
    add_para(doc,
        "All five new services are scoped, matching the lifetime of the "
        "DbContext they depend on. A scoped service is created once per "
        "Blazor circuit, which is the expected unit of work for a single "
        "user's session.")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. EF Core migration", 1)
    add_para(doc,
        "After the entity configuration changes were in place I generated a "
        "single migration covering the entire academic domain. The command "
        "is identical to the sprint 1 incantation:")
    add_code(doc,
        "dotnet ef migrations add AcademicDomain \\\n"
        "    --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "    --startup-project src/NaijaPrimeSchool.Web \\\n"
        "    --output-dir Persistence/Migrations\n")
    add_para(doc,
        "The output is a pair of files under "
        "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/ named "
        "<timestamp>_AcademicDomain.cs and <timestamp>_AcademicDomain.Designer.cs, "
        "plus an updated ApplicationDbContextModelSnapshot.cs.")

    add_heading(doc, "7.1 What the migration creates", 2)
    add_para(doc, "Tables added (in dependency order):")
    add_numbered(doc, [
        "TermTypes — lookup, parent of Term.",
        "ClassLevels — lookup, parent of SchoolClass.",
        "WeekDays — lookup, parent of TimetableEntry.",
        "TimetablePeriods — parent of TimetableEntry.",
        "Sessions — parent of Term and SchoolClass.",
        "Terms — depends on Sessions and TermTypes.",
        "SchoolClasses — depends on Sessions, ClassLevels, and Users (FK is Set Null).",
        "Subjects — referenced by TimetableEntry.",
        "TimetableEntries — depends on the previous seven plus Users (Set Null).",
    ])
    add_para(doc,
        "Indexes created — fourteen in total, every one of them mentioned "
        "explicitly in chapter 5. The migration also adds the IsDeleted "
        "indexes that the global query filters need to remain cheap, even "
        "though the query filter itself does not require a specific index "
        "(it simply takes advantage of one when it exists).")

    add_heading(doc, "7.2 Verifying the migration", 2)
    add_para(doc,
        "Two quick checks make sure the migration is shaped the way you "
        "expect before you run it. First, scan the generated Up() method "
        "for any unintended drops or alters — adding nine new tables should "
        "produce exactly nine CreateTable calls. Second, look at the "
        "ApplicationDbContextModelSnapshot diff to confirm every new entity "
        "is listed there exactly once.")
    add_code(doc,
        "# Confirm CreateTable count matches the new entity count\n"
        "grep -c \"CreateTable\" src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/*_AcademicDomain.cs\n"
        "# Should print 9")

    add_heading(doc, "7.3 Applying the migration", 2)
    add_para(doc,
        "DatabaseInitializer.InitializeAsync calls db.Database.MigrateAsync "
        "on every startup, so simply running the app applies the "
        "migration. If you need to apply it out-of-band — for instance to "
        "stage a deployment without booting the app — the dotnet-ef "
        "equivalent is:")
    add_code(doc,
        "dotnet ef database update \\\n"
        "    --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "    --startup-project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "Rolling back is symmetric. dotnet ef database update InitialCreate "
        "drops all the academic tables and indexes; dotnet ef migrations "
        "remove deletes the AcademicDomain migration files. Useful during "
        "development but not something you want to run on production data — "
        "the soft-deleted rows would be permanently lost.")

    add_heading(doc, "7.4 What the migration does NOT do", 2)
    add_para(doc,
        "The migration only creates schema. It does not populate the four "
        "new lookup tables — that is DatabaseInitializer.SeedAcademicLookupsAsync's "
        "job, which runs immediately after the migration on every boot. "
        "If you apply migrations with the dotnet ef CLI in isolation and "
        "then connect a SQL client to the database, you will see empty "
        "TermTypes, ClassLevels, WeekDays, and TimetablePeriods tables — "
        "they only fill on first application start.")

    add_page_break(doc)


def chapter8_web_setup(doc):
    add_heading(doc, "8. Web layer — Razor structure", 1)
    add_para(doc,
        "Sprint 2 adds six Razor pages, all of them routed under top-level "
        "URLs (/sessions, /terms, /classes, /subjects, /timetable-periods, "
        "/timetable). They live in a new folder, "
        "src/NaijaPrimeSchool.Web/Components/Pages/Academics/, mirroring the "
        "Application/Academics folder structure. Three small infrastructure "
        "changes outside that folder finished the wiring.")

    add_heading(doc, "8.1 Updated _Imports.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/_Imports.razor")
    add_para(doc,
        "The two new @using directives mean Razor pages can reference "
        "ISessionService, SchoolClassDto, and friends without per-page "
        "@using statements. This is the same pattern sprint 1 used for the "
        "Users namespaces.")

    add_heading(doc, "8.2 Updated NavMenu.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor")
    add_para(doc,
        "The single biggest visible change to the shell is the Academics "
        "panel item, which used to be Disabled='true' in sprint 1. It is "
        "now an expanded panel with six children, gated behind "
        "Roles.SuperAdmin / HeadTeacher / Teacher. Teachers see the "
        "expanded panel because the timetable view is read-only useful to "
        "them; the management pages are still gated at the page level "
        "(SuperAdmin / HeadTeacher only) by the @attribute "
        "[Authorize(Roles = …)] directive on each page.")

    add_heading(doc, "8.3 Page layout pattern", 2)
    add_para(doc,
        "Every CRUD page follows the same five-block structure, which "
        "makes them quick to read and quick to maintain:")
    add_numbered(doc, [
        "Page header (.nps-page-header) — H1, lede, and a 'New …' button.",
        "(Optional) filter bar in a RadzenCard — for example, the Terms "
        "page filters by Session, the Classes page filters by Session.",
        "List grid in a RadzenCard — RadzenDataGrid with columns and a "
        "right-aligned action column.",
        "Inline form in a RadzenCard, conditionally rendered when "
        "showForm == true — the form's fields and validation messages live "
        "inside a RadzenTemplateForm with a DataAnnotationsValidator.",
        "@code block — typed list state, form state, OnInitializedAsync, "
        "LoadAsync, OpenForm, SaveAsync, DeleteAsync.",
    ])
    add_para(doc,
        "OpenForm is the bridge between Edit / New and the form. When "
        "called with a DTO, it pre-populates the form. When called with "
        "null, it resets to defaults. Either way it sets showForm = true "
        "and lets the inline form render below the grid.")
    add_para(doc,
        "SaveAsync is uniformly straightforward: validate the not-DataAnnotations "
        "constraints (e.g. that a date is set), call the right service, "
        "show a Radzen notification with the OperationResult message, and "
        "either close the form on success or leave it open on failure so "
        "the user can fix and retry.")

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. Page-by-page walk-through", 1)
    add_para(doc,
        "What follows is the full source of every academic page, with "
        "running commentary on the design decisions, the bindings, and the "
        "Radzen component choices. Reading these in order is the fastest "
        "way to build a mental model of how all the layers come together.")

    add_heading(doc, "9.1 Sessions.razor (/sessions)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/Sessions.razor")
    add_para(doc, "Notable bits:")
    add_bullets(doc, [
        "@attribute [Authorize(Roles = $\"{Roles.SuperAdmin},{Roles.HeadTeacher}\")] "
        "uses an interpolated string so the Roles constants drive the value, "
        "not a magic string.",
        "@rendermode InteractiveServer keeps this as a server-rendered "
        "interactive page; the WebAssembly runtime is not loaded for it.",
        "The grid's row-level Mark current button only renders if "
        "!s.IsCurrent, so the current session has only Edit and Delete "
        "buttons, avoiding self-targeting state changes.",
        "The inline form's RadzenSwitch for IsCurrent is bound directly to "
        "form.IsCurrent. The service then handles the at-most-one-current "
        "guarantee with ExecuteUpdateAsync.",
        "DateOnly is converted to/from DateTime? at the boundary because "
        "RadzenDatePicker still works in DateTime; this is a deliberate "
        "boundary, not a leak.",
    ])

    add_heading(doc, "9.2 Terms.razor (/terms)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/Terms.razor")
    add_para(doc, "Things to highlight:")
    add_bullets(doc, [
        "OnInitializedAsync loads three reference lists in parallel: "
        "sessions, term types, and the initial term list. They feed the "
        "two dropdowns and the grid respectively.",
        "The session filter dropdown calls LoadAsync on Change, so the "
        "grid re-queries whenever the user picks a session. AllowClear "
        "lets them go back to 'all sessions'.",
        "The TermType selector is a dropdown rather than a radio list "
        "because a school might add a fourth term type later — keeping "
        "the UI extensible to data.",
    ])

    add_heading(doc, "9.3 Classes.razor (/classes)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/Classes.razor")
    add_para(doc,
        "The class form has three reference dropdowns (level, session, "
        "teacher) and a name + description pair. The teacher dropdown "
        "AllowClear=\"true\" so admins can leave a class teacherless "
        "during the planning period before assignments are finalised.")
    add_para(doc,
        "OpenForm defaults the SessionId to the current session "
        "(sessions.FirstOrDefault(s => s.IsCurrent)?.Id) so creating a new "
        "class on a typical day requires the absolute minimum number of "
        "clicks.")

    add_heading(doc, "9.4 Subjects.razor (/subjects)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/Subjects.razor")
    add_para(doc,
        "Subjects.razor is the simplest CRUD in the sprint and a useful "
        "template for any future flat lookup. The Code column has a "
        "120-pixel width so it does not stretch arbitrarily; the "
        "Description column takes whatever is left.")

    add_heading(doc, "9.5 TimetablePeriods.razor (/timetable-periods)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/TimetablePeriods.razor")
    add_para(doc,
        "Two small subtleties. The display order column sits at width "
        "100px and is the first column, so admins can scan the day's "
        "structure at a glance. The Type column renders a Break / Lesson "
        "badge with different colours so break rows stand out without "
        "needing an extra column.")
    add_para(doc,
        "The form's StartTime and EndTime use RadzenDatePicker with "
        "TimeOnly=\"true\" and DateFormat=\"HH:mm\"/HourFormat=\"24\" so "
        "admins enter 24-hour times. The conversion between DateTime? and "
        "TimeOnly happens at the boundary in OpenForm and SaveAsync.")

    add_heading(doc, "9.6 Timetable.razor (/timetable)", 2)
    add_para(doc,
        "The flagship page of the sprint. Two dropdowns at the top let "
        "the user choose a term and a class. Once both are selected, the "
        "page loads every TimetableEntry for that pair and renders a "
        "weekly grid of period rows × weekday columns. Each cell shows "
        "either an empty 'assign' hint, the assigned subject + teacher, "
        "or a greyed-out break row.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Academics/Timetable.razor")
    add_para(doc, "The interesting parts:")
    add_bullets(doc, [
        "The grid is plain HTML <table> rather than a RadzenDataGrid, "
        "because the layout is two-dimensional (rows = periods, cols = "
        "days) and pagination is meaningless.",
        "Click handlers are wired with @onclick on the <td> elements "
        "themselves, capturing the loop-local 'period' and 'day' variables "
        "into local copies (var p = period; var d = day;) to avoid the "
        "classic closure-over-loop-variable pitfall.",
        "OpenCellAsync skips break rows so they cannot be assigned a "
        "subject — UI affordance plus defensive coding.",
        "The cell editor is rendered as a RadzenCard below the grid (not "
        "as a dialog) for the same reasons the rest of the sprint avoids "
        "dialogs.",
        "SaveEntryAsync dispatches a single UpsertTimetableEntryRequest. "
        "The service treats (term, class, day, period) as the natural key, "
        "so a click into an already-filled cell that just changes the "
        "subject simply UPDATEs the existing row rather than inserting "
        "a duplicate.",
        "RemoveEntryAsync calls SoftDeleteEntryAsync — and because the "
        "global query filter hides the soft-deleted row, the cell becomes "
        "empty on the next refresh.",
    ])

    add_page_break(doc)


def chapter10_theming(doc):
    add_heading(doc, "10. Theming additions", 1)
    add_para(doc,
        "The shell from sprint 1 already has the green/gold Nigerian "
        "palette, the page header rhythm, and the form grid. Sprint 2 only "
        "needed to add CSS for the timetable grid itself. The new rules "
        "live at the bottom of wwwroot/app.css, namespaced under "
        ".nps-timetable / .nps-tt-* so they cannot collide with the rest "
        "of the design system.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/wwwroot/app.css",
        start_marker="===================== Timetable grid =====================",
        end_marker=".nps-tt-break {",
        caption="Excerpt — wwwroot/app.css timetable grid rules")

    add_heading(doc, "10.1 Visual decisions", 2)
    add_bullets(doc, [
        "The header row uses the same #f0fdf4 / #05613c green pair the "
        "sidebar uses, so the grid feels like part of the shell rather "
        "than a foreign component.",
        "Cells get a 64-pixel minimum height so an empty cell still feels "
        "clickable and gives visual room for two-line lessons.",
        "Hover state lights the cell with the same #f0fdf4 green so the "
        "user gets a clear 'click me' hint before they actually click.",
        "Break rows are tinted #fff8e1 (a soft gold tint) and the period "
        "name renders in #b8860b (gold), which signals 'not editable' "
        "without using grey, which would read as 'disabled'.",
        "The whole grid is wrapped in .nps-timetable { overflow-x: auto } "
        "so on a narrow screen the grid can scroll horizontally rather "
        "than wrap or shrink to illegibility.",
    ])

    add_heading(doc, "10.2 Filter bar reuse", 2)
    add_para(doc,
        ".nps-filter-bar from sprint 1 is reused as-is on Terms.razor and "
        "Classes.razor, where the session filter dropdown sits inside it. "
        "I deliberately did not introduce variants — keeping the same "
        "primitive everywhere makes the design system small and learnable.")

    add_page_break(doc)


def chapter11_authz(doc):
    add_heading(doc, "11. Authorization matrix", 1)
    add_para(doc,
        "Each academic page declares its allowed roles with the "
        "[Authorize(Roles = …)] attribute. The navigation menu wraps each "
        "panel in <AuthorizeView Roles=\"…\"> with a matching role list, "
        "so users see only the items they can use. The matrix:")

    matrix = [
        ("Page",                  "URL",                  "Allowed roles"),
        ("Sessions",              "/sessions",            "SuperAdmin, HeadTeacher"),
        ("Terms",                 "/terms",               "SuperAdmin, HeadTeacher"),
        ("Classes",               "/classes",             "SuperAdmin, HeadTeacher"),
        ("Subjects",              "/subjects",            "SuperAdmin, HeadTeacher"),
        ("Timetable periods",     "/timetable-periods",   "SuperAdmin, HeadTeacher"),
        ("Timetable grid",        "/timetable",           "SuperAdmin, HeadTeacher, Teacher"),
        ("Academics nav panel",   "(menu)",               "SuperAdmin, HeadTeacher, Teacher"),
    ]

    table = doc.add_table(rows=len(matrix), cols=3)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(matrix):
        cells = table.rows[i].cells
        for j, value in enumerate(row):
            cells[j].text = value
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True

    add_para(doc, " ")
    add_para(doc,
        "Notice that the Teacher role can see the Academics panel and the "
        "Timetable page but none of the configuration pages. The page "
        "attributes are the source of truth; the menu is just an "
        "experience optimisation. If a teacher manually navigates to "
        "/sessions, the Authorize attribute redirects them through the "
        "AccessDenied flow.")

    add_heading(doc, "11.1 Compile-time safety", 2)
    add_para(doc,
        "Roles are referenced via the Roles static class — Roles.SuperAdmin, "
        "Roles.HeadTeacher, Roles.Teacher — not as string literals. "
        "Renaming a role value is therefore a single search-and-replace, "
        "and a typo on a role name fails to compile rather than silently "
        "denying access.")

    add_heading(doc, "11.2 Defence in depth", 2)
    add_para(doc,
        "The page attributes are not the only line of defence. The menu's "
        "AuthorizeView ensures Teachers never see configuration links to "
        "click. The services do not check roles directly; they trust the "
        "authorize-attribute boundary, which is the right layering. If a "
        "future feature exposes a service over an HTTP API rather than "
        "Razor, the authorization decision will move to that endpoint, "
        "and the service will keep doing what it has always done.")

    add_page_break(doc)


def chapter12_lifecycle(doc):
    add_heading(doc, "12. Soft-delete and audit lifecycle", 1)
    add_para(doc,
        "It is worth tracing a single Save through the layers in detail, "
        "because the fact that the academic services do not need explicit "
        "audit or soft-delete code is a direct consequence of the "
        "infrastructure built in sprint 1.")

    add_heading(doc, "12.1 A successful Term creation", 2)
    add_numbered(doc, [
        "User clicks 'Save' on the Terms create form (Terms.razor).",
        "Blazor binds the form model into the page's TermFormModel; "
        "DataAnnotationsValidator confirms required fields are present.",
        "SaveAsync constructs a CreateTermRequest and calls "
        "TermService.CreateAsync.",
        "The service validates: dates make sense, Session exists, "
        "TermType exists, no duplicate (SessionId, TermTypeId).",
        "If IsCurrent is true, ExecuteUpdateAsync flips IsCurrent = false "
        "on every other Term row in a single SQL UPDATE.",
        "A new Term is added to the DbContext.",
        "SaveChangesAsync runs. The override iterates ChangeTracker.Entries "
        "and stamps CreatedOn = now, CreatedBy = current user. Then it "
        "delegates to base.SaveChangesAsync which writes the INSERT.",
        "The service returns OperationResult<Guid>.Success(term.Id).",
        "The page closes the form, reloads the grid, and shows a "
        "RadzenNotification with the success message.",
    ])

    add_heading(doc, "12.2 A successful Term soft-deletion", 2)
    add_numbered(doc, [
        "User clicks the Delete button on a Terms grid row.",
        "DialogService.Confirm opens a confirmation dialog; user clicks "
        "Delete.",
        "TermService.SoftDeleteAsync loads the Term with its "
        "TimetableEntries collection and refuses with a friendly error if "
        "any entries exist.",
        "If the term is clean, db.Terms.Remove(term) is called.",
        "SaveChangesAsync runs. The override sees EntityState.Deleted on "
        "an ISoftDelete row and rewrites it to EntityState.Modified, then "
        "stamps IsDeleted = true, DeletedOn = now, DeletedBy = current "
        "user. Then it delegates to base which writes the UPDATE.",
        "The service returns OperationResult.Success.",
        "On the next LoadAsync the global query filter (t => !t.IsDeleted) "
        "hides the row, so the grid no longer shows it.",
    ])

    add_heading(doc, "12.3 A failed save (validation error)", 2)
    add_numbered(doc, [
        "User clicks Save with EndDate before StartDate.",
        "Service returns OperationResult.Failure(\"End date must be after start date.\").",
        "Page renders the error in a Radzen notification; the form stays "
        "open with current input intact so the user can correct.",
        "Nothing was sent to the database; ChangeTracker has no pending "
        "changes; the audit override does not fire.",
    ])
    add_para(doc,
        "Notice that 'I made a mistake' and 'I successfully changed my "
        "mind' both leave the system in a consistent state. The "
        "service-level pre-checks save us from depending entirely on "
        "database constraints to surface user-friendly errors.")

    add_page_break(doc)


def chapter13_smoketest(doc):
    add_heading(doc, "13. Smoke test — end-to-end happy path", 1)
    add_para(doc,
        "After cloning the sprint/2-academic-domain branch and configuring "
        "appsettings.json with a working SQL connection string, the smoke "
        "test below confirms every layer is wired correctly. Run the app "
        "with dotnet run --project src/NaijaPrimeSchool.Web — first-run "
        "behaviour applies the AcademicDomain migration and seeds the "
        "TermTypes, ClassLevels, WeekDays, and TimetablePeriods tables.")

    add_heading(doc, "13.1 Sign in", 2)
    add_numbered(doc, [
        "Navigate to the URL printed by dotnet run.",
        "Sign in with superadmin@naijaprimeschool.ng / Admin@12345.",
        "You are redirected to the dashboard. The Academics panel in the "
        "sidebar is now expanded with six children.",
    ])

    add_heading(doc, "13.2 Create a Session", 2)
    add_numbered(doc, [
        "Click Sessions in the sidebar.",
        "Click 'New session'.",
        "Fill: Name = 2025/2026, StartDate = 1 Sep 2025, EndDate = 31 Jul 2026, "
        "Mark current = on.",
        "Click Save.",
        "Expected: a green success toast appears; the row is in the grid "
        "with a 'Current' badge.",
    ])

    add_heading(doc, "13.3 Create the three Terms", 2)
    add_numbered(doc, [
        "Click Terms in the sidebar.",
        "Click 'New term', pick Session = 2025/2026, Term = First Term, "
        "supply dates, and Mark current = on. Save.",
        "Repeat for Second Term and Third Term, IsCurrent = off.",
        "Expected: three rows in the Terms grid; First Term carries the "
        "'Current' badge.",
    ])

    add_heading(doc, "13.4 Create a Class", 2)
    add_numbered(doc, [
        "Click Classes.",
        "Click 'New class'. Name = Primary 1A, Class level = Primary 1, "
        "Session defaults to 2025/2026, Class teacher left blank for now.",
        "Save.",
        "Expected: the row shows under 2025/2026 with the level Primary 1 "
        "and an empty teacher column.",
    ])

    add_heading(doc, "13.5 Create a Subject", 2)
    add_numbered(doc, [
        "Click Subjects.",
        "Click 'New subject'. Name = Mathematics, Code = mth (lowercase). "
        "Save.",
        "Expected: the row shows with Code MTH (uppercased on save).",
    ])

    add_heading(doc, "13.6 Verify the seeded Periods", 2)
    add_para(doc,
        "Click 'Periods' under Academics. The grid should show Period 1 "
        "through Period 7 plus Short Break and Lunch, totalling nine rows. "
        "Two of them (Short Break and Lunch) carry the orange 'Break' "
        "badge.")

    add_heading(doc, "13.7 Lay out a Timetable", 2)
    add_numbered(doc, [
        "Click Timetable.",
        "Pick Term = 2025/2026 — First Term, Class = Primary 1A.",
        "Click any non-break cell on Monday's Period 1.",
        "In the inline editor below the grid, pick Subject = Mathematics. "
        "Save.",
        "Expected: the cell now shows MTH on green text and the editor "
        "closes.",
        "Click the same cell again. The editor opens with the existing "
        "values loaded.",
        "Click Remove. Confirm. The cell becomes empty.",
        "Click an empty cell on Tuesday's Period 2 and assign Mathematics "
        "again to confirm the upsert path is fine when starting fresh.",
    ])

    add_heading(doc, "13.8 Soft-delete behaviour", 2)
    add_numbered(doc, [
        "Try deleting the Mathematics subject. Expect a friendly error "
        "saying it is on a timetable.",
        "Remove the Mathematics lesson from the timetable cell that was "
        "just assigned, then try deleting Mathematics again. Expect "
        "success.",
        "Open SQL Server Management Studio (or your favourite client) and "
        "query SELECT * FROM Subjects. The Mathematics row is still "
        "present but IsDeleted = 1 and DeletedOn / DeletedBy are stamped. "
        "Soft delete confirmed.",
    ])

    add_page_break(doc)


def chapter14_troubleshooting(doc):
    add_heading(doc, "14. Troubleshooting playbook", 1)
    add_para(doc,
        "This is the short list of things that most often go wrong while "
        "bringing up the sprint 2 branch on a new machine, with the fix "
        "for each.")

    add_heading(doc, "14.1 Build fails on missing using directives", 2)
    add_para(doc,
        "Symptom: the Razor compiler complains that ISessionService or "
        "SchoolClassDto is undefined. Cause: _Imports.razor was not updated "
        "to add the Academics namespaces. Fix: confirm the file under "
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor has the lines:")
    add_code(doc,
        "@using NaijaPrimeSchool.Application.Academics\n"
        "@using NaijaPrimeSchool.Application.Academics.Dtos")

    add_heading(doc, "14.2 Migration fails with 'database already contains a table'", 2)
    add_para(doc,
        "Symptom: dotnet ef database update fails because the AcademicDomain "
        "migration finds tables it intends to create. Cause: a previous "
        "checkout left behind ad-hoc tables, or the database is being "
        "shared between branches. Fix: drop the database (in development "
        "only) or run dotnet ef database update <PreviousMigrationName> "
        "first to roll back, then forward.")

    add_heading(doc, "14.3 LookupService cannot find UserManager", 2)
    add_para(doc,
        "Symptom: 'Unable to resolve service for type "
        "Microsoft.AspNetCore.Identity.UserManager`1[ApplicationUser] "
        "while attempting to activate LookupService'. Cause: the "
        "Infrastructure DI registration runs before AddIdentityCore and "
        "AddRoles, but the LookupService scoped lifetime was changed to "
        "singleton (which would not work because UserManager is scoped). "
        "Fix: keep services.AddScoped<ILookupService, LookupService>() — "
        "do not change the lifetime.")

    add_heading(doc, "14.4 Timetable grid does not render after picking term + class", 2)
    add_para(doc,
        "Symptom: dropdowns work but the grid stays empty. Most common "
        "cause: there are no TimetablePeriod rows in the database "
        "(somebody soft-deleted them all). The grid renders nothing if "
        "there are no period rows, and the empty grid is "
        "indistinguishable from 'no entries yet'. Fix: visit "
        "/timetable-periods and confirm there are visible periods. Add "
        "one if there are none.")

    add_heading(doc, "14.5 Cannot mark a Term current", 2)
    add_para(doc,
        "Symptom: clicking 'Mark current' fails with no obvious feedback. "
        "Most likely cause: cookie expired and SignalR circuit dropped. "
        "Fix: reload the page; the action will work again.")

    add_heading(doc, "14.6 Time pickers show empty", 2)
    add_para(doc,
        "Symptom: editing an existing TimetablePeriod row, the StartTime "
        "and EndTime pickers come up blank. Cause: locale / culture "
        "settings on the user's machine prevent Radzen from parsing the "
        "DateTime back into the picker. Fix: confirm the form's "
        "OpenForm code is doing today.Add(p.StartTime.ToTimeSpan()) "
        "rather than DateTime.MinValue.Add(...) — DateTime.MinValue is "
        "below SqlDateTime.MinValue and Radzen sometimes refuses to "
        "render it.")

    add_page_break(doc)


def chapter15_forward(doc):
    add_heading(doc, "15. Sprint 3 outlook", 1)
    add_para(doc,
        "Now that the academic domain is in place, sprint 3 will plug "
        "students and parents into it. The work falls broadly into four "
        "tracks:")

    add_heading(doc, "15.1 Students", 2)
    add_bullets(doc, [
        "Student entity in NaijaPrimeSchool.Domain.Students, deriving "
        "from BaseEntity. Probably an aggregate root rather than tied to "
        "ApplicationUser one-to-one — many students never log in and "
        "should not consume an Identity row.",
        "Enrolment as a join table StudentClassEnrolment with FromDate / "
        "ToDate so historic enrolments are preserved when a student moves "
        "class mid-year.",
        "A 'class roster' page that lists every student in a class for a "
        "term — natural extension of the existing Classes management.",
    ])

    add_heading(doc, "15.2 Parents and guardians", 2)
    add_bullets(doc, [
        "Parent entity (or maybe a Guardian generalisation) with "
        "many-to-many to Student through a ParentStudent join with "
        "Relationship and IsPrimaryContact flags.",
        "Optional ApplicationUser link so parents who want online access "
        "can sign in; parents who do not have access still exist in the "
        "system as data.",
        "A 'parent contact' panel inside the student profile, "
        "showing every contact for a child plus their relationship.",
    ])

    add_heading(doc, "15.3 Forward-compatibility, today", 2)
    add_para(doc,
        "Sprint 2 already left a few breadcrumbs that make the sprint 3 "
        "work cleaner.")
    add_bullets(doc, [
        "TimetableEntry has a TeacherId but no StudentId — students do "
        "not appear on the schedule directly. Their participation will "
        "come from enrolment + class.",
        "Term has IsCurrent — sprint 3 will key 'current student "
        "enrolment' on (term IsCurrent + enrolment FromDate ≤ today ≤ "
        "ToDate or null), reusing the existing flag.",
        "Subjects already exist independent of any class. Sprint 4 will "
        "add ClassSubjectAssignment so a class can declare its subject "
        "list, and the timetable picker can be narrowed to those.",
    ])

    add_heading(doc, "15.4 What might need a small refactor", 2)
    add_bullets(doc, [
        "TimetableService.UpsertEntryAsync currently ignores any Id the "
        "client sends. Once timetable entries can be referenced from "
        "elsewhere (e.g. attendance markings keyed off them), a true Id "
        "round-trip will matter and the upsert will need to choose a "
        "strategy more carefully.",
        "ILookupService.GetTeachersAsync filters in memory; once the "
        "teacher count is non-trivial we will rewrite as a "
        "database-side join.",
        "The 'no enums' rule will be tested by attendance status (Present, "
        "Absent, Late, Excused). I will treat it as a lookup table again, "
        "consistent with sprint 1 and 2, even though the temptation to "
        "use an enum is genuine.",
    ])

    add_page_break(doc)


def chapter16_appendix(doc):
    add_heading(doc, "16. Appendix — files added or changed in sprint 2", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Academics/Session.cs",          "Academic year root."),
        ("src/NaijaPrimeSchool.Domain/Academics/TermType.cs",         "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Academics/Term.cs",             "Session × TermType, with date range."),
        ("src/NaijaPrimeSchool.Domain/Academics/ClassLevel.cs",       "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs",      "Class arm with optional teacher."),
        ("src/NaijaPrimeSchool.Domain/Academics/Subject.cs",          "Subject + unique code."),
        ("src/NaijaPrimeSchool.Domain/Academics/WeekDay.cs",          "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Academics/TimetablePeriod.cs",  "School-day slot."),
        ("src/NaijaPrimeSchool.Domain/Academics/TimetableEntry.cs",   "(term, class, day, period) → subject + teacher."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Academics/Dtos/SessionDtos.cs",     "Session DTOs."),
        ("src/NaijaPrimeSchool.Application/Academics/Dtos/TermDtos.cs",        "Term DTOs."),
        ("src/NaijaPrimeSchool.Application/Academics/Dtos/SchoolClassDtos.cs", "Class DTOs."),
        ("src/NaijaPrimeSchool.Application/Academics/Dtos/SubjectDtos.cs",     "Subject DTOs."),
        ("src/NaijaPrimeSchool.Application/Academics/Dtos/TimetableDtos.cs",   "Period + entry DTOs."),
        ("src/NaijaPrimeSchool.Application/Academics/ISessionService.cs",      "Session service contract."),
        ("src/NaijaPrimeSchool.Application/Academics/ITermService.cs",         "Term service contract."),
        ("src/NaijaPrimeSchool.Application/Academics/ISchoolClassService.cs",  "Class service contract."),
        ("src/NaijaPrimeSchool.Application/Academics/ISubjectService.cs",      "Subject service contract."),
        ("src/NaijaPrimeSchool.Application/Academics/ITimetableService.cs",    "Timetable service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",           "Added GetTermTypes, GetClassLevels, GetWeekDays, GetTeachers."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SessionService.cs",     "Session CRUD + SetCurrent."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/TermService.cs",        "Term CRUD + SetCurrent."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SchoolClassService.cs", "Class CRUD."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SubjectService.cs",     "Subject CRUD."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/TimetableService.cs",   "Period CRUD + entry list/upsert/delete."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/<ts>_AcademicDomain.cs", "EF migration adding 9 tables and 14 indexes."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",         "Registered the 5 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs", "Added 9 DbSets, ConfigureAcademics."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs", "Seeded TermTypes, ClassLevels, WeekDays, TimetablePeriods."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",      "Added the 4 new lookup methods + UserManager DI."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/Sessions.razor",         "Sessions list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/Terms.razor",            "Terms list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/Classes.razor",          "Classes list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/Subjects.razor",         "Subjects list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/TimetablePeriods.razor", "Periods list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Academics/Timetable.razor",        "Weekly grid with cell editor."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                          "Added Academics + Academics.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                    "Replaced disabled Academics item with full panel."),
        ("src/NaijaPrimeSchool.Web/wwwroot/app.css",                                    "Added timetable grid styles."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint2_guide.py",                                             "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 2 implementation guide. The next sprint plugs "
        "students and parents into the academic structure laid down here.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_web_setup(doc)
    chapter9_pages(doc)
    chapter10_theming(doc)
    chapter11_authz(doc)
    chapter12_lifecycle(doc)
    chapter13_smoketest(doc)
    chapter14_troubleshooting(doc)
    chapter15_forward(doc)
    chapter16_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
