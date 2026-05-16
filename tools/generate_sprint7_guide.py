"""Generates 'Sprint 7 - Implementation Guide.docx' covering the store
and inventory management domain delivered in sprint 7 (ItemCategory,
UnitOfMeasure, StockMovementType lookups; Supplier, StoreItem,
StockMovement entities; the matching services and the Razor pages
that drive the storekeeper's workspace).

Long-form edition. Code blocks embed actual source files so the guide
stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint7_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 7 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260516054730_StoreAndInventory.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 7 — Store & Inventory Management")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Suppliers · Catalog · Stock movements · Storekeeper dashboard")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/7-store-inventory")
    meta.add_run("\nBuilt on: Sprints 1–6 (identity, academic domain, students & parents, attendance, results & report cards, pupil photos, fees & bursar workflows)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 7 in context", 1)
    add_para(doc,
        "Sprint 7 builds the storekeeper's workspace. The bursar's side "
        "of the financial coin closed in sprint 6 — fee schedules, "
        "invoices, receipts, and the dashboard that tells the bursar how "
        "much money is coming in and how much is still owed. Sprint 7 "
        "turns its attention to the other side of the same building: the "
        "store room. Every textbook the school issues to a pupil, every "
        "uniform handed across the counter, every bag of rice the kitchen "
        "draws down, every reams of A4 the office consumes, every "
        "football the games master collects — those movements all flow "
        "through a single audit table that this sprint introduces.")
    add_para(doc,
        "Functionally the sprint introduces three layered concepts. The "
        "Supplier directory captures the vendors the school buys from. "
        "The StoreItem catalog lists every distinct article the "
        "storekeeper tracks, with category, unit of measure, reorder "
        "level, and a running on-hand quantity. The StockMovement table "
        "is the audit trail: every receipt, issuance, return, write-off, "
        "or adjustment is one row, and the on-hand quantity on the parent "
        "StoreItem is brought into line with the stream of movements via "
        "a service-layer helper.")
    add_para(doc,
        "Once this sprint ships, the SchoolStoreKeeper role — seeded "
        "back in sprint 1 but unused since — finally has a workspace. "
        "The previously-disabled 'Store & Inventory' navigation "
        "placeholder lights up with six items: a dashboard summarising "
        "stock value and low-stock alerts, the catalog, item detail, "
        "movement log, record-new-movement form, and supplier directory.")
    add_para(doc,
        "This document is a long-form implementation guide written in the "
        "tone of the sprint 6 guide. An engineer who has read the earlier "
        "guides and has the codebase checked out can recreate every "
        "change here without referring to the diff.")

    add_heading(doc, "1.1 Where this sits relative to sprint 6", 2)
    add_para(doc, "Every load-bearing piece of the earlier sprints is reused:")
    add_bullets(doc, [
        "BaseEntity — every new entity inherits Guid Id, IAuditable, "
        "and ISoftDelete from it.",
        "ApplicationDbContext.SaveChanges — the override stamps audit "
        "columns and rewrites Delete to IsDeleted = true. Every Supplier, "
        "StoreItem and StockMovement therefore inherits the same audit "
        "and soft-delete machinery as everything else.",
        "Global query filters — every new entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries automatically.",
        "OperationResult / OperationResult<T> — every new service uses "
        "this for predictable success/failure responses.",
        "ILookupService — already had twenty-four methods. Sprint 7 adds "
        "five more (item categories, units of measure, stock movement "
        "types, active suppliers, and a typeahead for store items) "
        "without rewriting the existing ones.",
        "Student, SchoolClass — pick up StockIssuances back-references "
        "so the storekeeper can see, on a pupil's profile (in a future "
        "sprint), every item ever issued to them. No scalar columns "
        "change on either table.",
        "StudentAvatar (sprint 5b) — already lives in Components.Shared "
        "and is reused on the movement log when an item is issued to a "
        "pupil.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_numbered(doc, [
        "Maintain a Supplier directory — name, contact, phone, email, "
        "address, notes. Soft delete refused once a purchase has been "
        "recorded against the supplier; deactivate instead.",
        "Maintain the StoreItem catalog. Categories and units of measure "
        "are seeded lookup tables. Each item carries a reorder level and "
        "a running QuantityOnHand. Optional SKU is a unique filtered "
        "index.",
        "Open a new item with an opening balance — the create form lets "
        "the storekeeper enter an opening quantity and unit cost, which "
        "the service persists as a single 'Opening Balance' stock "
        "movement and the item's on-hand value.",
        "Record every stock movement: Purchase, Return, Adjustment In, "
        "Opening Balance (inbound); Issue, Write-off, Adjustment Out "
        "(outbound). Each movement carries quantity, unit cost, total "
        "cost, optional reference (PO number, requisition number) and "
        "free-text notes.",
        "Pair an issuance with its recipient: a pupil, a class, or a "
        "staff member. A purchase pairs with a supplier. The service "
        "enforces at most one recipient per movement.",
        "Watch the dashboard: items in catalog, items below reorder, "
        "stock value, this-month inbound / outbound counts, low-stock "
        "list, recent movements, value by category.",
        "Soft-delete items, suppliers, and movements with guards: "
        "cannot delete an item with movement history, cannot delete a "
        "supplier with purchase history, cannot record an outbound "
        "movement larger than the on-hand quantity. Deleting a "
        "movement reverses its effect on the running balance.",
    ])

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_bullets(doc, [
        "Multi-location inventory. The school is single-site; a future "
        "Store table per physical location would slot in without "
        "redesigning anything else.",
        "Batch / lot tracking. Some items (food, medicines) would "
        "benefit from per-batch expiry dates. The current model holds "
        "one running balance per item; a future StoreItemBatch table "
        "would be additive.",
        "Purchase orders. The Reference column on a StockMovement "
        "carries a PO number today; a dedicated PurchaseOrder table "
        "with approvals and partial receipts is its own sprint.",
        "Pupil-facing returns (the family bringing the textbook back at "
        "year end). The Return movement type captures it on the "
        "storekeeper side; the parent portal sprint will add the "
        "pupil-facing surface.",
        "Tying issuances to the bursar's invoices. A school that "
        "charges separately for uniforms could link an Issue movement "
        "to a Sprint-6 InvoiceLine. The FK is not in the schema today "
        "because we have not yet seen a school that wants that link "
        "automated.",
        "Reorder workflow. Today the dashboard surfaces items below "
        "reorder; clicking through to 'create a purchase from this "
        "list' is a future workflow.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers:")
    add_bullets(doc, [
        "6 new domain entities under src/NaijaPrimeSchool.Domain/Inventory/.",
        "2 collection navigations on existing entities (Student, SchoolClass).",
        "4 DTO files under src/NaijaPrimeSchool.Application/Inventory/Dtos/.",
        "3 new service contracts under src/NaijaPrimeSchool.Application/Inventory/.",
        "3 service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "5 new methods on ILookupService (and the matching LookupService).",
        "1 EF Core migration introducing 6 new tables and the indexes "
        "that go with them.",
        "1 DatabaseInitializer extension seeding ItemCategories, "
        "UnitsOfMeasure, and StockMovementTypes.",
        "6 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Inventory/.",
        "1 navigation menu rewrite: the previously-disabled 'Store & "
        "Inventory' placeholder is replaced with a six-item panel gated "
        "to SuperAdmin + HeadTeacher + SchoolStoreKeeper.",
    ])
    add_para(doc,
        "The code follows the patterns already accepted in sprints 1–6.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)

    add_heading(doc, "2.1 Movement-log as the source of truth", 2)
    add_para(doc,
        "Every change to stock — in or out — is one row in StockMovement. "
        "StoreItem.QuantityOnHand is a cached running total that the "
        "service keeps in agreement with the movement stream. The "
        "alternative — calculate on-hand from a SUM(Direction * "
        "Quantity) every time the catalog renders — is correct but "
        "slow. Persisting the running balance and updating it inside the "
        "same SaveChanges that creates the movement gives us:")
    add_bullets(doc, [
        "Constant-time list queries: the catalog reads QuantityOnHand "
        "directly without aggregating thousands of movements.",
        "Atomic correctness: a movement insert and the balance update "
        "are in the same transaction, so the two cannot diverge.",
        "A reversible audit log: soft-deleting a movement reverses its "
        "contribution to the running balance, which makes 'undo a "
        "wrongly-keyed receipt' a one-click operation.",
    ])

    add_heading(doc, "2.2 No enums — three lookup tables", 2)
    add_para(doc,
        "The rule from earlier sprints holds. ItemCategory, "
        "UnitOfMeasure, and StockMovementType are all proper entities "
        "derived from BaseEntity, seeded on first run, and editable "
        "from the database without a redeploy. StockMovementType "
        "carries a Direction column (+1 = inbound, -1 = outbound) that "
        "the service multiplies Quantity by to roll up the on-hand "
        "figure — so a future 'Loan Out' or 'Loan Back' row could be "
        "added without changing any service code.")

    add_heading(doc, "2.3 Decimal precision", 2)
    add_para(doc,
        "Quantities are decimal(14,3) — three decimal places of "
        "resolution lets the school track partial kilograms of rice or "
        "litres of cleaning fluid without rounding. Unit cost is "
        "decimal(12,2); total cost is decimal(14,2). These match the "
        "decimal precisions used in the finance domain in sprint 6.")

    add_heading(doc, "2.4 Movement numbering", 2)
    add_para(doc,
        "MovementNumber follows the same year-prefixed pattern as "
        "Sprint-6 invoices and receipts: NPS/STK/<year>/<4-digit-"
        "sequence>. The service computes the next sequence by reading "
        "the largest existing number for the year and adding one. "
        "Numbers are guaranteed unique by a database index, and "
        "auditable across years.")

    add_heading(doc, "2.5 At most one recipient per movement", 2)
    add_para(doc,
        "An Issue movement can be paired with a pupil, a class, or a "
        "staff member — but not two of them. Three nullable FKs "
        "(IssuedToStudentId, IssuedToSchoolClassId, IssuedToUserId) "
        "give us flexibility without the alternative shapes "
        "(polymorphic foreign keys, JSON columns) that would make "
        "the audit trail brittle. The service rejects a request with "
        "more than one set. Leaving all three null is allowed and "
        "represents 'general consumption' — e.g. the cleaning team "
        "drew a bag of soap powder for the school.")

    add_heading(doc, "2.6 Soft delete plus operation guards", 2)
    add_bullets(doc, [
        "Supplier cannot be deleted if it has purchase history. "
        "Deactivate instead.",
        "StoreItem cannot be deleted if it has movement history. "
        "Deactivate instead.",
        "Outbound StockMovement is refused if Quantity exceeds the "
        "item's on-hand balance.",
        "Soft-deleting a StockMovement reverses the running balance "
        "on its item, so the audit history can be hidden without "
        "leaving a phantom quantity behind.",
    ])

    add_heading(doc, "2.7 Unique indexes", 2)
    add_bullets(doc, [
        "Unique on each lookup's Name and Code (4 unique indexes "
        "across ItemCategory, UnitOfMeasure, StockMovementType).",
        "Filtered unique on StoreItem.Sku WHERE [Sku] IS NOT NULL — "
        "items don't have to carry a SKU but if they do, it must be "
        "unique.",
        "Unique on StockMovement.MovementNumber — auditable.",
    ])

    add_heading(doc, "2.8 Inline forms, the same UI rhythm", 2)
    add_para(doc,
        "Every Razor page in this sprint follows the same shape as the "
        "earlier sprints: a filter card on top, a paged grid in the "
        "middle, an inline RadzenCard form below the grid for new-or-"
        "edit, and a confirm-then-act dialog for destructive operations. "
        "Two pages — CreateStoreItem and RecordStockMovement — are "
        "dedicated full-page forms because their workflows have enough "
        "fields and side-effects that an inline form would have been "
        "cramped. The record-movement page also flips its lower half "
        "between an inbound 'received from supplier' panel and an "
        "outbound 'issued to recipient' panel based on the chosen "
        "movement type's Direction.")

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)

    add_heading(doc, "3.1 Folder layout", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "└── Inventory/                     <- (new in sprint 7)\n"
        "    ├── ItemCategory.cs            <- lookup\n"
        "    ├── UnitOfMeasure.cs           <- lookup\n"
        "    ├── StockMovementType.cs       <- lookup with Direction\n"
        "    ├── Supplier.cs                <- vendor entity\n"
        "    ├── StoreItem.cs               <- catalog entry\n"
        "    └── StockMovement.cs           <- audit row\n")

    add_heading(doc, "3.2 The three lookup entities", 2)
    add_heading(doc, "3.2.1 ItemCategory.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/ItemCategory.cs")
    add_heading(doc, "3.2.2 UnitOfMeasure.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/UnitOfMeasure.cs")
    add_heading(doc, "3.2.3 StockMovementType.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/StockMovementType.cs")

    add_heading(doc, "3.3 The three core entities", 2)
    add_heading(doc, "3.3.1 Supplier.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/Supplier.cs")
    add_heading(doc, "3.3.2 StoreItem.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/StoreItem.cs")
    add_para(doc,
        "QuantityOnHand and LastUnitCost are persisted convenience "
        "columns kept in agreement with the movement stream by the "
        "service. The catalog list page reads these directly so it can "
        "render thousands of items without aggregating millions of "
        "movements.")
    add_heading(doc, "3.3.3 StockMovement.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Inventory/StockMovement.cs")
    add_para(doc,
        "Five optional FKs: ReceivedFromSupplierId for purchases, plus "
        "the three IssuedTo* counter-parties for issuances, plus "
        "PerformedBy linking back to the storekeeper who keyed the row. "
        "All five OnDelete to SetNull so the soft-delete of any of "
        "those entities does not vaporise the audit history.")

    add_heading(doc, "3.4 Relationships at a glance", 2)
    add_code(doc,
        "      ItemCategory      UnitOfMeasure       StockMovementType\n"
        "          \\                |                       |\n"
        "           \\               |                       |\n"
        "            v 1..N         v 1..N                  v 1..N\n"
        "                StoreItem  --------+               |\n"
        "                                   |               |\n"
        "                                   v N             v\n"
        "                              StockMovement -------+\n"
        "                                   |\n"
        "                                   +--- ReceivedFromSupplier (Supplier)\n"
        "                                   +--- IssuedToStudent      (Student)\n"
        "                                   +--- IssuedToSchoolClass  (SchoolClass)\n"
        "                                   +--- IssuedToUser         (ApplicationUser)\n"
        "                                   +--- PerformedBy          (ApplicationUser)\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)

    add_heading(doc, "4.1 SupplierDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/Dtos/SupplierDtos.cs")

    add_heading(doc, "4.2 StoreItemDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/Dtos/StoreItemDtos.cs")
    add_para(doc,
        "StoreItemDto carries the IsBelowReorder computed property so "
        "the catalog list can colour its badge without the consumer "
        "redoing the comparison. CreateStoreItemRequest carries an "
        "OpeningQuantity and OpeningUnitCost so a brand-new item can "
        "be seeded with stock-on-hand in one call.")

    add_heading(doc, "4.3 StockMovementDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/Dtos/StockMovementDtos.cs")
    add_para(doc,
        "StockMovementDto carries the sprint-5b photo plumbing for "
        "issuances to pupils so the <StudentAvatar /> component "
        "renders a pupil's face on the movement log without a second "
        "round-trip.")

    add_heading(doc, "4.4 StoreSummaryDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/Dtos/StoreSummaryDtos.cs")

    add_heading(doc, "4.5 Service contracts", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/ISupplierService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/IStoreItemService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Inventory/IStockMovementService.cs")

    add_heading(doc, "4.6 ILookupService extension", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Application/Users/ILookupService.cs",
        "Task<IReadOnlyList<LookupDto>> GetItemCategoriesAsync",
        end_marker="Task<IReadOnlyList<LookupDto>> GetStoreItemsAsync",
        caption="Excerpt — ILookupService.cs (sprint 7 additions)")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)

    add_heading(doc, "5.1 Six new DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<ItemCategory>",
        end_marker="public DbSet<StockMovement>",
        lines_after_start=7,
        caption="Excerpt — the six inventory DbSets")

    add_heading(doc, "5.2 ConfigureInventory", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureInventory",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureInventory")
    add_para(doc, "Highlights:")
    add_bullets(doc, [
        "HasPrecision(14, 3) on Quantity and ReorderLevel; "
        "HasPrecision(12, 2) on UnitCost; HasPrecision(14, 2) on TotalCost.",
        "Filtered unique index on StoreItem.Sku — schools that don't "
        "use SKUs can leave the column null without colliding on the "
        "unique constraint.",
        "Restrict on StockMovement.StoreItemId / StockMovementTypeId — "
        "the schema-level constraint reinforces the service-layer guard "
        "that refuses to delete an item or type that has movements.",
        "SetNull on every optional counter-party (Supplier, Student, "
        "SchoolClass, IssuedToUser, PerformedBy) so soft-deleting any "
        "of those entities does not vaporise the audit history.",
        "Composite index on (StoreItemId, MovedOn) — speeds the per-"
        "item movement-history page.",
        "b.Ignore(i => i.StockValue) and b.Ignore(i => i.IsBelowReorder) "
        "on the DTO side — but the entity itself has no computed "
        "properties to ignore. Those are DTO conveniences.",
    ])

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)

    add_heading(doc, "6.1 SupplierService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/SupplierService.cs")

    add_heading(doc, "6.2 StoreItemService.cs", 2)
    add_para(doc,
        "StoreItemService owns catalog CRUD plus the create-with-"
        "opening-balance shortcut. When OpeningQuantity > 0 it stages a "
        "fresh StoreItem and an 'OPENING' StockMovement in the same "
        "SaveChanges, so a new item never lives in the system without "
        "its origin movement.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/StoreItemService.cs")

    add_heading(doc, "6.3 StockMovementService.cs", 2)
    add_para(doc,
        "The audit-trail service. RecordAsync validates the type, "
        "checks the at-most-one-recipient rule, refuses outbound "
        "movements that would drive the on-hand quantity negative, "
        "computes the next sequential movement number, persists the "
        "row and updates StoreItem.QuantityOnHand and (for inbound "
        "with a unit cost) LastUnitCost in the same SaveChanges. "
        "SoftDeleteAsync reverses the contribution before removing "
        "the row.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/StockMovementService.cs")

    add_heading(doc, "6.4 DI registration", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<ISupplierService",
        end_marker="return services;",
        caption="Excerpt — DependencyInjection.cs (sprint 7 additions)")

    add_heading(doc, "6.5 LookupService — five new methods", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",
        "public async Task<IReadOnlyList<LookupDto>> GetItemCategoriesAsync",
        lines_after_start=50,
        caption="Excerpt — LookupService.cs (sprint 7 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_code(doc,
        "dotnet ef migrations add StoreAndInventory \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the inventory lookups", 1)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedInventoryLookupsAsync",
        end_marker="private static async Task SeedFinanceLookupsAsync",
        caption="Excerpt — SeedInventoryLookupsAsync")
    add_para(doc, "What gets seeded:")
    add_bullets(doc, [
        "ItemCategories — Books (BOOK), Uniforms (UNIF), Stationery "
        "(STAT), Sports Equipment (SPRT), Cleaning Supplies (CLN), "
        "Food (FOOD), Furniture (FURN), ICT Equipment (ICT), "
        "Medical Supplies (MED), Other (OTH).",
        "UnitsOfMeasure — Each (EA), Piece (PC), Pack (PK), Box (BOX), "
        "Carton (CTN), Set (SET), Bag (BAG), Kilogram (KG), Litre (L), "
        "Metre (M).",
        "StockMovementTypes — Opening Balance (OPENING, +1), Purchase "
        "(PURCHASE, +1), Return (RETURN, +1), Adjustment In (ADJ_IN, "
        "+1), Issue (ISSUE, -1), Write-off (WRITEOFF, -1), Adjustment "
        "Out (ADJ_OUT, -1).",
    ])

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/\n"
        "├── StoreDashboard.razor          <- /store\n"
        "├── StoreItems.razor              <- /store/items\n"
        "├── CreateStoreItem.razor         <- /store/items/new\n"
        "├── StoreItemDetail.razor         <- /store/items/{id}\n"
        "├── StockMovements.razor          <- /store/movements\n"
        "├── RecordStockMovement.razor     <- /store/movements/new\n"
        "└── Suppliers.razor               <- /store/suppliers\n")
    add_para(doc,
        "Every page is gated to SuperAdmin + HeadTeacher + "
        "SchoolStoreKeeper. Sprint 7 is the first sprint where the "
        "SchoolStoreKeeper role moves from placeholder navigation to "
        "a fully realised workspace.")

    add_heading(doc, "9.1 StoreDashboard.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreDashboard.razor")

    add_heading(doc, "9.2 StoreItems.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreItems.razor")

    add_heading(doc, "9.3 CreateStoreItem.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/CreateStoreItem.razor")

    add_heading(doc, "9.4 StoreItemDetail.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreItemDetail.razor")

    add_heading(doc, "9.5 StockMovements.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StockMovements.razor")

    add_heading(doc, "9.6 RecordStockMovement.razor", 2)
    add_para(doc,
        "The most interesting page in the sprint. Picking a movement "
        "type flips the lower half of the form between an inbound "
        "panel (asks for a Supplier) and an outbound panel (asks for "
        "at most one of pupil / class / staff member). Both sides "
        "share the same upper half: item, type, date, quantity, unit "
        "cost, reference, notes.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/RecordStockMovement.razor")

    add_heading(doc, "9.7 Suppliers.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Inventory/Suppliers.razor")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and authorization", 1)
    add_para(doc,
        "The previously-disabled 'Store & Inventory' nav placeholder "
        "is replaced with a six-item panel. The panel is wrapped in "
        "an AuthorizeView that grants access to SuperAdmin, "
        "HeadTeacher, and SchoolStoreKeeper.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Store & Inventory\"",
        end_marker="</AuthorizeView>",
        caption="Excerpt — NavMenu.razor")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Inventory",
        lines_after_start=2,
        caption="Excerpt — _Imports.razor")

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of an inventory flow", 1)

    add_heading(doc, "11.1 Seeding the store", 2)
    add_bullets(doc, [
        "Storekeeper opens /store/suppliers, clicks New supplier. "
        "Enters 'Lagos Books Ltd', a contact, a phone, an email.",
        "Opens /store/items, clicks New item. Enters 'English "
        "Textbook — Primary 4', picks category Books, unit Each, "
        "reorder level 5, opening quantity 60, opening unit cost "
        "₦2,000. Saves.",
        "StoreItemService stages the StoreItem with QuantityOnHand = "
        "60 plus an OPENING StockMovement for 60 × ₦2,000 in the "
        "same SaveChanges.",
        "Storekeeper lands on the item detail page, which now shows "
        "the opening movement and a stock value of ₦120,000.",
    ])

    add_heading(doc, "11.2 Recording a purchase", 2)
    add_bullets(doc, [
        "Lagos Books delivers 50 more copies at ₦2,100.",
        "Storekeeper opens /store/movements/new, picks the textbook "
        "item, picks Purchase as the type. The lower half flips to the "
        "inbound 'received from' panel.",
        "Picks Lagos Books Ltd as the supplier, enters quantity 50, "
        "unit cost 2,100, reference 'PO-2026-009'.",
        "Saves. StockMovementService writes the row, increments "
        "QuantityOnHand by 50 (now 110), and updates LastUnitCost to "
        "2,100. Movement number is auto-generated as "
        "NPS/STK/<year>/<seq>.",
    ])

    add_heading(doc, "11.3 Issuing to a class", 2)
    add_bullets(doc, [
        "The Primary 4A teacher requisitions 38 textbooks (one per "
        "pupil).",
        "Storekeeper records a movement, picks Issue. The lower half "
        "flips to the outbound 'issued to recipient' panel.",
        "Picks Primary 4A as the class. Enters quantity 38, "
        "leaves unit cost blank (cost only matters for inbound "
        "movements).",
        "Saves. QuantityOnHand drops to 72.",
    ])

    add_heading(doc, "11.4 Catching a low-stock alert", 2)
    add_bullets(doc, [
        "Over the term, 65 more textbooks are issued. QuantityOnHand "
        "now sits at 7.",
        "Dashboard's 'Below reorder level' tile flips to 1.",
        "The Low-stock panel surfaces the textbook with a Low badge "
        "and a click-through to the item detail.",
        "Storekeeper raises a fresh purchase, runs through the same "
        "Purchase movement flow, and the badge clears.",
    ])

    add_heading(doc, "11.5 Reversing a mistake", 2)
    add_bullets(doc, [
        "Storekeeper realises a Purchase was keyed against the wrong "
        "item.",
        "Opens the movement log, finds the row, soft-deletes it.",
        "StockMovementService reads the row, multiplies Direction by "
        "Quantity, subtracts that from QuantityOnHand on the affected "
        "item, then soft-deletes the row.",
        "The audit history shows the row as hidden (only "
        "IgnoreQueryFilters returns it); the item's on-hand quantity "
        "is back to what it was before the mistake.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "First run applies the StoreAndInventory migration and seeds "
        "the three new lookup tables. Sign in as the SuperAdmin (or a "
        "user in the SchoolStoreKeeper role).")

    add_heading(doc, "12.2 Verify navigation and lookups", 2)
    add_bullets(doc, [
        "The Store & Inventory nav panel is now an active six-item "
        "dropdown rather than a disabled placeholder.",
        "SELECT Name, Code FROM ItemCategories ORDER BY DisplayOrder; "
        "shows ten rows.",
        "SELECT Name, Code FROM UnitsOfMeasure ORDER BY DisplayOrder; "
        "shows ten rows.",
        "SELECT Name, Code, Direction FROM StockMovementTypes ORDER "
        "BY DisplayOrder; shows seven rows with the right "
        "+1 / -1 values.",
    ])

    add_heading(doc, "12.3 End-to-end happy path", 2)
    add_numbered(doc, [
        "Store → Suppliers → New supplier. Save 'Lagos Books Ltd'.",
        "Store → New item. Save 'English Textbook — Primary 4', "
        "Books, Each, reorder level 5, opening quantity 60 at "
        "₦2,000.",
        "Open the item detail. Confirm one Opening Balance movement "
        "and QuantityOnHand = 60.",
        "Store → Record movement. Pick the textbook, Purchase, "
        "supplier Lagos Books, quantity 50 at ₦2,100. Save.",
        "Item detail now shows two movements and on-hand 110.",
        "Record an Issue movement to a class for quantity 38. "
        "On-hand drops to 72.",
        "Store dashboard shows the textbook in the 'Recent movements' "
        "list, the stock value reflects (72 × 2,100), and items below "
        "reorder is zero.",
    ])

    add_heading(doc, "12.4 Error paths", 2)
    add_numbered(doc, [
        "Try recording an Issue larger than the on-hand quantity. "
        "Refused.",
        "Try recording an Issue with both a pupil and a class picked. "
        "Refused by the form.",
        "Try deleting a supplier that has a purchase. Refused.",
        "Try deleting an item that has movement history. Refused.",
        "Soft-delete a movement and watch the parent item's on-hand "
        "quantity rewind by the right amount.",
    ])

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)

    add_heading(doc, "13.1 'Cannot remove N — only M are on hand'", 2)
    add_para(doc,
        "Sprint 7 refuses to let an outbound movement drive the "
        "running balance negative. Either record a fresh Purchase or "
        "Adjustment In first, or split the movement so the requested "
        "quantity matches what is on hand. The friendly error message "
        "names the on-hand figure.")

    add_heading(doc, "13.2 'A movement can be issued to at most one recipient'", 2)
    add_para(doc,
        "Both the UI and the service-layer guard against this. If the "
        "Save button silently refuses, double-check the recipient "
        "panel — only one of pupil / class / staff member should be "
        "selected.")

    add_heading(doc, "13.3 'Cannot delete a supplier with purchase history'", 2)
    add_para(doc,
        "Suppliers that have ever been linked to a purchase movement "
        "stay on file for the audit trail. Use Deactivate instead — "
        "deactivated suppliers no longer appear in the active-suppliers "
        "dropdown on the record-movement page, but their history is "
        "preserved.")

    add_heading(doc, "13.4 'SKU already exists'", 2)
    add_para(doc,
        "The filtered unique index on StoreItem.Sku ignores NULLs, so "
        "you can have any number of items without an SKU. If you "
        "insist on filling it in, the value has to be globally unique "
        "across active items.")

    add_heading(doc, "13.5 Dashboard stock value disagrees with my expectations", 2)
    add_para(doc,
        "StockValue is QuantityOnHand × LastUnitCost — a "
        "single-point estimate. If the school buys at varying unit "
        "costs (the 60 textbooks at ₦2,000 plus 50 at ₦2,100), the "
        "valuation uses the most recent unit cost (2,100), not a "
        "weighted average. This matches how most small schools "
        "value stock; a future enhancement could expose a weighted-"
        "average column.")

    add_heading(doc, "13.6 The IdentityRole migration warning", 2)
    add_para(doc,
        "Same pre-existing warning that has accompanied every sprint "
        "since sprint 1. Harmless at runtime.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_bullets(doc, [
        "Parent and student portals (later sprint) will read "
        "StudentId-tagged issuance movements to show 'items received "
        "this term' on a pupil's profile.",
        "Purchase Orders can grow into their own table with approvals; "
        "the Reference column on a StockMovement captures a PO number "
        "today and can later become a foreign key.",
        "Per-batch tracking (expiry, lot number) is additive — a "
        "future StoreItemBatch table sitting between StoreItem and "
        "StockMovement does not require changing existing rows.",
        "Linking issuances to bursar invoices (a uniform issued is "
        "also a uniform billed) is a one-column add to StockMovement.",
        "Stock-take audits — periodic full counts that produce a "
        "stream of Adjustment In / Out movements — already work with "
        "the seeded types.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 7", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Inventory/ItemCategory.cs",     "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Inventory/UnitOfMeasure.cs",    "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Inventory/StockMovementType.cs","Lookup with Direction."),
        ("src/NaijaPrimeSchool.Domain/Inventory/Supplier.cs",         "Vendor entity."),
        ("src/NaijaPrimeSchool.Domain/Inventory/StoreItem.cs",        "Catalog entry."),
        ("src/NaijaPrimeSchool.Domain/Inventory/StockMovement.cs",    "Audit row."),
        ("Domain layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Domain/Family/Student.cs",             "Added StockIssuances."),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs",      "Added StockIssuances."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Inventory/Dtos/SupplierDtos.cs",       "Supplier DTOs."),
        ("src/NaijaPrimeSchool.Application/Inventory/Dtos/StoreItemDtos.cs",      "Catalog DTOs."),
        ("src/NaijaPrimeSchool.Application/Inventory/Dtos/StockMovementDtos.cs",  "Movement DTOs."),
        ("src/NaijaPrimeSchool.Application/Inventory/Dtos/StoreSummaryDtos.cs",   "Dashboard DTOs."),
        ("src/NaijaPrimeSchool.Application/Inventory/ISupplierService.cs",        "Supplier service contract."),
        ("src/NaijaPrimeSchool.Application/Inventory/IStoreItemService.cs",       "Catalog service contract."),
        ("src/NaijaPrimeSchool.Application/Inventory/IStockMovementService.cs",   "Movement service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",              "Added 5 new lookup methods."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SupplierService.cs",       "Supplier CRUD."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/StoreItemService.cs",      "Catalog CRUD + opening balance."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/StockMovementService.cs",  "Audit log + dashboard summary."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 6 tables."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",            "Registered 3 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs", "Added 6 DbSets, ConfigureInventory."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs", "Seeded the 3 new lookup tables."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",         "Added 5 new lookup methods."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreDashboard.razor",      "Storekeeper summary."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreItems.razor",          "Catalog list + filters."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/CreateStoreItem.razor",     "New item with opening balance."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StoreItemDetail.razor",     "Edit + movement history."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/StockMovements.razor",      "Movement log."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/RecordStockMovement.razor", "Record movement with direction-aware panel."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Inventory/Suppliers.razor",           "Supplier directory + inline form."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                            "Added Inventory + Inventory.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                      "Replaced disabled Store & Inventory placeholder with full panel."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint7_guide.py",                                               "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 7 implementation guide. With the "
        "storekeeper's workspace alive, the next sprint can turn to "
        "the parent and student portals — the last role groups still "
        "looking at disabled navigation.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
