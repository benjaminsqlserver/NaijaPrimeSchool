"""Generates 'Sprint 4 - Implementation Guide.docx' covering the
attendance work delivered in sprint 4 (DailyAttendanceRegister,
DailyAttendanceEntry, SubjectAttendanceSession, SubjectAttendanceEntry,
the AttendanceStatus lookup, and the supporting Razor pages).

This is the long-form edition. Code blocks embed actual source files
from the repository so the guide stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint4_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 4 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260430212205_Attendance.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 4 — Attendance")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Daily registers · Per-subject sessions · Submission lifecycle · Class summaries")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/4-attendance")
    meta.add_run("\nBuilt on: Sprint 1 identity + Sprint 2 academic domain + Sprint 3 students & parents")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 4 in context", 1)
    add_para(doc,
        "Sprint 4 plugs attendance into the structure that sprints 1–3 laid "
        "down. It is the first sprint in which the application starts "
        "carrying actual day-to-day classroom data — every previous sprint "
        "set up the scaffolding, but nothing changed from one school day to "
        "the next. After sprint 4 ships, every classroom day generates a "
        "register; every lesson on the timetable can generate a per-subject "
        "session; and the school can pull a per-pupil attendance percentage "
        "for any term or session.")
    add_para(doc,
        "There are two parallel attendance models, deliberately built side "
        "by side. The daily register is what most Nigerian primary schools "
        "actually run on: at homeroom each morning, the class teacher marks "
        "every pupil Present, Absent, Late, Excused, Sick or Suspended in a "
        "single sheet. The subject register is finer-grained: it sits on "
        "top of the timetable from sprint 2 and lets a subject teacher take "
        "attendance at any specific (class × subject × period × day) slot. "
        "Both models share the same AttendanceStatus lookup, the same "
        "submit/reopen lifecycle, and the same auditing and soft-delete "
        "primitives the previous sprints established.")
    add_para(doc,
        "This document is a long-form implementation guide. It is written "
        "so that an engineer who has read the sprint 3 guide and has the "
        "codebase checked out can recreate every change in this sprint "
        "without referring to the diff. The document is organised in "
        "roughly the order I built the code in: design decisions first, "
        "Domain entities next, Application contracts after that, "
        "Infrastructure (DbContext, services, seeder, migration) in the "
        "middle, and finally the Razor UI and navigation. There is a "
        "smoke-test chapter near the end that walks through the happy-path "
        "you can use to confirm a fresh checkout works.")

    add_heading(doc, "1.1 Where this sits relative to sprint 3", 2)
    add_para(doc,
        "Sprint 3 delivered Student, Parent, StudentParent linkage, and "
        "Enrolment. Sprint 4 leans on every one of those:")
    add_bullets(doc, [
        "BaseEntity — every new attendance entity inherits from it and "
        "picks up Guid Id, IAuditable, and ISoftDelete with no boilerplate.",
        "ApplicationDbContext.SaveChanges — the override stamps "
        "CreatedOn/By and ModifiedOn/By and rewrites Delete to "
        "IsDeleted = true. Every attendance write therefore inherits "
        "auditing and soft delete with zero changes to the override.",
        "Global query filters — every new entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries without service code having to remember to filter.",
        "OperationResult / OperationResult<T> — every attendance service "
        "uses this for predictable success/failure responses.",
        "Enrolment — DailyAttendanceService.OpenAsync and "
        "SubjectAttendanceService.OpenAsync read open enrolments "
        "(WithdrawnOn IS NULL) to pre-populate the register with every "
        "pupil who is currently in the class on the date being marked. "
        "This is exactly the load-bearing definition the sprint 3 guide "
        "promised.",
        "TimetableEntry — sprint 2's (term, class, weekday, period, "
        "subject, teacher) tuple is the natural anchor for a per-subject "
        "session. SubjectAttendanceSession.TimetableEntryId is the only "
        "academic foreign key that session needs.",
        "SchoolClass / Term — the daily register is keyed on (class × "
        "date) but also carries TermId so 'all registers for first term' "
        "is a single index lookup.",
        "ILookupService — sprint 4 adds one method "
        "(GetAttendanceStatusesAsync). Every other lookup call the "
        "attendance pages need (sessions, terms, classes, students) "
        "already exists.",
        "Radzen Blazor + the green/gold app.css — the attendance pages "
        "adopt the same .nps-page-header / .nps-card / .nps-form-grid "
        "primitives so they read as part of the same product.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_para(doc,
        "Concretely, after this sprint a SuperAdmin, HeadTeacher, or "
        "Teacher signing in to the application can:")
    add_numbered(doc, [
        "Open a daily attendance register for any (class × date) combination. "
        "Every pupil currently enrolled in that class is pre-loaded with "
        "the default status (Present), and the term that covers the date "
        "is resolved automatically.",
        "Mark or change each pupil's status, set an arrival time when the "
        "status is Late, and capture free-text remarks.",
        "Submit the register, which locks every entry against further "
        "edits until an admin reopens it.",
        "Take per-subject attendance: pick a term, class, and date; the "
        "page lists every lesson on the timetable for that weekday; click "
        "any lesson to open its register; mark each pupil; submit.",
        "View a per-class attendance summary across a term — days "
        "counted, days present, days late, days absent, days excused, "
        "and a colour-coded percentage band per pupil.",
        "Soft-delete a register or session that should not have been "
        "opened, with a friendly refusal if the register is still in "
        "the submitted state.",
    ])
    add_para(doc,
        "Parents, Students, Bursars and Storekeepers do not see the "
        "Attendance navigation panel; their menu items remain placeholders "
        "reserved for later sprints (parent portal in particular will read "
        "a child's attendance summary).")

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_para(doc,
        "It is just as important to be explicit about what sprint 4 does "
        "NOT do, because every one of these has been weighed and "
        "consciously deferred:")
    add_bullets(doc, [
        "Bell-schedule auto-roll. The current subject-attendance flow "
        "still requires a teacher to open each lesson register manually. "
        "Auto-creating a session at the start of every period via a "
        "background job is a feature-add, not a redesign — the schema "
        "supports it today.",
        "Attendance-driven SMS alerts to parents (\"your ward was "
        "absent today\"). Once the notifications backbone lands "
        "(separate sprint), wiring up an event off Submit is a few "
        "lines.",
        "Public holidays / non-school days. The system happily lets you "
        "open a register on a Sunday because that is sometimes correct "
        "(boarding schools, religious schools); a school-calendar "
        "feature deserves its own sprint.",
        "Excused-absence workflow. Today an admin marks a pupil "
        "Excused. A follow-on workflow that captures the parent's "
        "request (note from home, doctor's note URL) belongs to the "
        "parent portal sprint.",
        "Bulk import of historical attendance from spreadsheets. Some "
        "schools want to backfill the year-to-date register on day one. "
        "The schema already supports this; the lift is a service "
        "method that loops over rows and calls BulkSetAsync, plus a "
        "Razor page with file upload.",
        "Subject-attendance summary reporting. We have a daily "
        "summary; per-subject summaries (\"Adaeze missed three Maths "
        "lessons this term\") are useful but the existing data already "
        "supports the query — only the UI is missing.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers, this sprint adds:")
    add_bullets(doc, [
        "5 new domain entities under src/NaijaPrimeSchool.Domain/Attendance/.",
        "4 collection navigations on existing entities (SchoolClass, "
        "Term, TimetableEntry, Student).",
        "3 DTO files under src/NaijaPrimeSchool.Application/Attendance/Dtos/.",
        "2 service contracts under src/NaijaPrimeSchool.Application/Attendance/.",
        "2 service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "1 method on ILookupService for the new attendance-status lookup.",
        "1 EF Core migration introducing 5 new tables and 13 indexes.",
        "1 DatabaseInitializer extension seeding the AttendanceStatus "
        "lookup with six rows.",
        "3 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Attendance/.",
        "1 navigation menu addition to surface the new pages.",
    ])
    add_para(doc,
        "Everything compiles with zero warnings on .NET 10 (the team-wide "
        "warning bar). The code follows the patterns already accepted in "
        "sprints 1–3, so the diff is low-friction to review.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)
    add_para(doc,
        "Before any code was written I made a small number of architectural "
        "calls that shaped everything that followed. Re-reading these in "
        "isolation makes the code easier to navigate later, and it gives "
        "future maintainers permission to revisit choices when the "
        "trade-offs change.")

    add_heading(doc, "2.1 Two parallel attendance models, one lookup", 2)
    add_para(doc,
        "The most consequential decision in this sprint was modelling "
        "daily attendance and per-subject attendance as two separate "
        "aggregate roots — DailyAttendanceRegister and "
        "SubjectAttendanceSession — instead of forcing them into one "
        "shape. The trade-offs are real:")
    add_bullets(doc, [
        "A single 'AttendanceRecord' table with a nullable "
        "TimetableEntryId would be tempting on day one, but it would "
        "muddle two different keys: (class, date) for daily and "
        "(timetable-entry, date) for subject. Indexes for the two "
        "shapes would conflict, and the UI would have to filter every "
        "list query by 'is daily? is subject?'.",
        "Two roots keep each table narrow. The unique index on "
        "(SchoolClassId, Date) for daily and (TimetableEntryId, Date) "
        "for subject is the natural identity for each row, and the "
        "indexes are non-overlapping.",
        "AttendanceStatus is shared between the two. The same "
        "Present/Absent/Late/Excused/Sick/Suspended values apply in "
        "both models — there is no conceptual reason to split the "
        "lookup.",
        "If a future sprint needs a unified 'show me everything that "
        "happened to this pupil today' view, joining the two tables "
        "with UNION ALL is cheap. The opposite refactor — splitting "
        "one table back into two — would have been much harder.",
    ])

    add_heading(doc, "2.2 No enums — AttendanceStatus is a table", 2)
    add_para(doc,
        "The Present/Absent/Late/Excused/Sick/Suspended set is the "
        "textbook argument for a C# enum: it is genuinely closed, "
        "semantically clear, and rarely changes. I still modelled it "
        "as a table, consistent with sprints 1–3:")
    add_bullets(doc, [
        "Translatability. A future Hausa or Yoruba translation only "
        "needs to edit Name in a row, not recompile.",
        "Extensibility by data. A school that wants to distinguish "
        "between 'Sick (with note)' and 'Sick (no note)' adds one "
        "row; no code change.",
        "Reportability. Joins to the lookup surface readable labels in "
        "every reporting tool; magic enum integers do not.",
        "Display order and rendering hints. AttendanceStatus carries a "
        "DisplayOrder for the dropdown ordering and a CountsAsPresent "
        "flag the summary code uses to compute percentages without "
        "hard-coding which statuses count.",
        "Code column. Each status carries a short Code (P, L, E, S, "
        "A, SP) that the UI can use as a compact visual marker. "
        "DailyAttendanceService keys 'late arrival time' off the code "
        "value 'L', not the integer index of an enum.",
    ])
    add_para(doc,
        "The cost is one extra join on every read query, which EF Core "
        "handles efficiently because the lookup table is six rows and "
        "effectively always cached.")

    add_heading(doc, "2.3 Submit / reopen lifecycle, not 'finalised forever'", 2)
    add_para(doc,
        "Both DailyAttendanceRegister and SubjectAttendanceSession "
        "carry an IsSubmitted boolean and a nullable SubmittedOn. "
        "Once submitted, the service layer refuses every edit until "
        "an admin explicitly reopens it. Three considerations went "
        "into this:")
    add_bullets(doc, [
        "Reality. Class teachers submit a register at the end of "
        "homeroom; thirty minutes later a child's parent rings to "
        "say their car broke down on the school run. The teacher "
        "needs to flip the child from Absent to Late. We need to "
        "support that without losing the audit trail of who changed "
        "what and when.",
        "Audit trail. The audit columns (CreatedOn/By, "
        "ModifiedOn/By) on the register and each entry capture every "
        "change. SubmittedOn captures the moment the register was "
        "first considered final. ModifiedOn after a reopen makes the "
        "post-submission edit visible.",
        "Soft delete still applies. A submitted register cannot be "
        "soft-deleted directly — the service refuses with a friendly "
        "message — because deleting submitted attendance is "
        "different from amending it. Reopen first, then delete if "
        "really required.",
    ])

    add_heading(doc, "2.4 Pre-populate registers from open enrolments", 2)
    add_para(doc,
        "OpenAsync on both services loads every Enrolment for the "
        "class with EnrolledOn ≤ Date AND (WithdrawnOn IS NULL OR "
        "WithdrawnOn ≥ Date), then creates a default-status entry "
        "for each one. The reasoning:")
    add_bullets(doc, [
        "Teachers should not have to add pupils to the register "
        "themselves. The class roster on the date in question is "
        "the source of truth.",
        "A pupil who was withdrawn before the date does not appear. "
        "A pupil who is enrolled but was admitted three days after "
        "the date does not appear either. The query is the "
        "load-bearing definition of 'who is in this class on this "
        "day'.",
        "If a pupil is added or withdrawn after the register is "
        "opened, the existing entries are not touched. The "
        "register is a snapshot of what the teacher saw that day, "
        "not a recomputation of today's roster.",
    ])

    add_heading(doc, "2.5 Daily register carries TermId; subject session does not", 2)
    add_para(doc,
        "There is a deliberate asymmetry between the two roots. "
        "DailyAttendanceRegister carries an explicit TermId column. "
        "SubjectAttendanceSession does not, because TermId is "
        "reachable through TimetableEntry.Term and there is no "
        "direct query on session that does not also need to join "
        "to the timetable entry anyway.")
    add_para(doc,
        "OpenAsync on DailyAttendanceService resolves the term at "
        "creation time using the date and the class's session: "
        "first it looks for a term whose start/end range covers "
        "the date, and falls back to the current term in the "
        "session if no term covers the date (typical for "
        "edge-of-holiday days). The chosen term is stamped on the "
        "row so subsequent reads do not have to recompute it.")

    add_heading(doc, "2.6 SubjectAttendanceSession validates date vs weekday", 2)
    add_para(doc,
        "OpenAsync on SubjectAttendanceService refuses to create a "
        "session when the calendar weekday of the requested date "
        "does not match the WeekDay of the timetable entry. A "
        "Monday-period maths lesson on a Wednesday is almost "
        "always a typo, and the service catches it before the "
        "register is created. The exact check:")
    add_code(doc,
        "var weekDayName = entry.WeekDay!.Name;\n"
        "var actualDayOfWeek = request.Date.DayOfWeek.ToString();\n"
        "if (!string.Equals(weekDayName, actualDayOfWeek,\n"
        "                   StringComparison.OrdinalIgnoreCase))\n"
        "    return OperationResult<Guid>.Failure(\n"
        "        $\"Selected date is a {actualDayOfWeek} but \"\n"
        "      + $\"this lesson runs on {weekDayName}.\");\n")
    add_para(doc,
        "This relies on the WeekDay.Name (\"Monday\") matching "
        "DayOfWeek.ToString() (\"Monday\") exactly, which is true "
        "for the seeded Monday–Friday rows. If a school adds a "
        "Saturday WeekDay row, it just has to be named "
        "\"Saturday\" and the check still works.")

    add_heading(doc, "2.7 Composite unique indexes", 2)
    add_para(doc,
        "Four composite unique indexes do most of the structural "
        "integrity heavy lifting in sprint 4:")
    add_bullets(doc, [
        "(SchoolClassId, Date) on DailyAttendanceRegister — at most "
        "one daily register per class per day.",
        "(RegisterId, StudentId) on DailyAttendanceEntry — at most "
        "one entry per pupil per register.",
        "(TimetableEntryId, Date) on SubjectAttendanceSession — at "
        "most one session per timetable entry per day.",
        "(SessionId, StudentId) on SubjectAttendanceEntry — at most "
        "one entry per pupil per session.",
    ])
    add_para(doc,
        "OpenAsync on both services pre-checks the parent-row "
        "constraint with .FirstOrDefaultAsync; if a register/session "
        "already exists for that key, it returns the existing Id "
        "rather than failing. This is what lets a teacher click "
        "'Open register' twice without seeing a database error.")

    add_heading(doc, "2.8 Foreign-key delete behaviour", 2)
    add_bullets(doc, [
        "DailyAttendanceEntry.RegisterId / SubjectAttendanceEntry.SessionId "
        "→ Cascade. If the parent row is hard-deleted, child rows go "
        "with it. We cannot reach this state from the UI — the soft-"
        "delete flow refuses submitted registers — but the schema is "
        "honest about it.",
        "DailyAttendanceRegister.SchoolClassId / TermId → Restrict. A "
        "class or term cannot be hard-deleted while any register "
        "points at it. This is consistent with sprint 2's stance on "
        "academic structural rows.",
        "AttendanceStatusId on both entry tables → Restrict. You "
        "cannot accidentally lose 'Present' from the lookup while "
        "millions of rows point at it.",
        "TakenById on both parent tables → SetNull. Removing a "
        "teacher's account should not delete the registers they "
        "took.",
        "StudentId on both entry tables → Restrict. We never delete "
        "a Student row, but if some future maintenance script does, "
        "it should fail loudly while attendance entries still point "
        "at the pupil.",
    ])

    add_heading(doc, "2.9 Service per attendance model, plus shared lookup", 2)
    add_para(doc,
        "Two services split the work along the natural seams: "
        "IDailyAttendanceService and ISubjectAttendanceService. "
        "Each one owns OpenAsync, SetEntryAsync, BulkSetAsync, "
        "SubmitAsync, ReopenAsync, SoftDeleteAsync, plus the "
        "appropriate read methods. IDailyAttendanceService also "
        "owns the per-pupil and per-class summary queries because "
        "those are sourced from daily attendance — per-subject "
        "summaries are deferred to a later sprint.")
    add_para(doc,
        "I considered a single combined IAttendanceService and "
        "rejected it: each service is already 200+ lines, and "
        "merging them would have made the file unreviewable.")

    add_heading(doc, "2.10 Inline editing in the grid, not modal forms", 2)
    add_para(doc,
        "Both attendance pages render the entries as a "
        "RadzenDataGrid where each cell is the editor. Status is "
        "a dropdown bound directly to the row; arrival time and "
        "remarks are inline text inputs. Save is a single button "
        "that BulkSetAsync's every row in one call. Reasons:")
    add_bullets(doc, [
        "A class of thirty pupils marked one-by-one through a "
        "modal would be cripplingly slow. Inline grid editing is "
        "what teachers actually do today on paper.",
        "BulkSetAsync diffs against the existing entries and "
        "updates only the rows that changed, but the API "
        "deliberately accepts the full set so the UI does not "
        "have to track dirty rows.",
        "A submit button next to save lets the teacher mark and "
        "submit in two clicks: change a few statuses, click "
        "'Save changes' to write the diff, then click 'Submit "
        "register' to lock it.",
    ])

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)
    add_para(doc,
        "Every attendance entity lives in a single new folder, "
        "src/NaijaPrimeSchool.Domain/Attendance/. There are no abstract "
        "base classes other than BaseEntity, and the entities are "
        "deliberately anaemic — no domain methods, no validation logic. "
        "Validation lives in the Application DTOs (DataAnnotations) and "
        "in the Infrastructure services (cross-aggregate checks). The "
        "Domain layer is a typed vocabulary; it is not the place for "
        "behaviour in this codebase.")

    add_heading(doc, "3.1 Folder layout and namespacing", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "├── Attendance/                       ← (new in sprint 4)\n"
        "│   ├── AttendanceStatus.cs           ← lookup\n"
        "│   ├── DailyAttendanceRegister.cs    ← class × date root\n"
        "│   ├── DailyAttendanceEntry.cs       ← register × pupil\n"
        "│   ├── SubjectAttendanceSession.cs   ← timetable entry × date\n"
        "│   └── SubjectAttendanceEntry.cs     ← session × pupil\n"
        "├── Family/                           ← from sprint 3\n"
        "├── Academics/                        ← from sprint 2 (3 small additions)\n"
        "├── Common/                           ← from sprint 1\n"
        "└── Identity/                         ← from sprint 1\n")

    add_heading(doc, "3.2 AttendanceStatus.cs — the lookup", 2)
    add_para(doc,
        "The shared lookup. Code is the short two-character marker "
        "(P, L, E, S, A, SP); CountsAsPresent gates the percentage "
        "calculations.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Attendance/AttendanceStatus.cs")

    add_heading(doc, "3.3 DailyAttendanceRegister.cs — the daily root", 2)
    add_para(doc,
        "DailyAttendanceRegister is keyed naturally by "
        "(SchoolClassId, Date). The TermId is denormalised at "
        "creation time to make 'all registers for first term' a "
        "single index seek. TakenBy is optional — if a "
        "head-teacher takes a class teacher's register on their "
        "behalf, that is captured. IsSubmitted + SubmittedOn make "
        "up the lifecycle flag.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Attendance/DailyAttendanceRegister.cs")

    add_heading(doc, "3.4 DailyAttendanceEntry.cs — one row per pupil", 2)
    add_para(doc,
        "Each entry points at a Student and an AttendanceStatus. "
        "ArrivalTime is optional and only meaningful when the "
        "status is Late; the UI hides the field for other "
        "statuses. Remarks is free-text for any unusual case.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Attendance/DailyAttendanceEntry.cs")

    add_heading(doc, "3.5 SubjectAttendanceSession.cs — the per-lesson root", 2)
    add_para(doc,
        "SubjectAttendanceSession hangs off TimetableEntry — every "
        "concept the academic side already knows about (term, "
        "class, subject, weekday, period, teacher, room) is reached "
        "through that single navigation. The session itself only "
        "adds a Date, a TakenBy, the lifecycle flags, and an "
        "optional Notes column.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Attendance/SubjectAttendanceSession.cs")

    add_heading(doc, "3.6 SubjectAttendanceEntry.cs", 2)
    add_para(doc,
        "Mirrors DailyAttendanceEntry but without the ArrivalTime "
        "column — for a 40-minute lesson, 'arrival time' is not a "
        "meaningful concept. Status + remarks is enough.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Attendance/SubjectAttendanceEntry.cs")

    add_heading(doc, "3.7 Existing entities that grew a navigation", 2)
    add_para(doc,
        "Four sprint 1–3 entities pick up exactly one collection "
        "navigation each. Existing code that ignored these "
        "properties continues to compile and behave identically; "
        "the migration emits no schema change to the host tables.")
    add_bullets(doc, [
        "SchoolClass.DailyAttendanceRegisters",
        "Term.DailyAttendanceRegisters",
        "TimetableEntry.AttendanceSessions",
        "Student.DailyAttendanceEntries and Student.SubjectAttendanceEntries",
    ])

    add_heading(doc, "3.8 Relationships at a glance", 2)
    add_code(doc,
        "                          ┌──────────────────┐\n"
        "                          │ AttendanceStatus │\n"
        "                          │  P L E S A SP    │\n"
        "                          └────────┬─────────┘\n"
        "                                   │ N\n"
        "                  ┌────────────────┴────────────────────┐\n"
        "                  │ N                                   │ N\n"
        "         ┌────────▼──────────────┐         ┌────────────▼─────────────┐\n"
        "         │ DailyAttendanceEntry  │         │ SubjectAttendanceEntry   │\n"
        "         │ register × pupil      │         │ session × pupil          │\n"
        "         └──┬───────────┬────────┘         └──┬─────────────┬─────────┘\n"
        "            │ N         │ N                  │ N           │ N\n"
        "            │           │                    │             │\n"
        " ┌──────────▼─────┐  ┌──▼──────┐   ┌─────────▼──────────┐  │\n"
        " │ DailyAttendance│  │ Student │   │ SubjectAttendance- │  │\n"
        " │   Register     │  └─────────┘   │   Session          │  │\n"
        " │ class × date   │                │ timetableEntry×day │  │\n"
        " └─┬─────┬────────┘                └────────┬───────────┘  │\n"
        "   │ N   │ N                                │ N            │\n"
        "   │     │                                  │              │\n"
        " ┌─▼─┐ ┌─▼──┐                       ┌───────▼────────┐     │\n"
        " │Cls│ │Term│                       │ TimetableEntry │     │\n"
        " └───┘ └────┘                       │ (sprint 2)     │ ────┘\n"
        "                                    └────────────────┘\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)
    add_para(doc,
        "Application is the boundary between the Razor UI and the "
        "persistence layer. It is deliberately thin: just DTOs and "
        "service interfaces. Sprint 4 adds two new subfolders: "
        "Attendance/Dtos and Attendance.")

    add_heading(doc, "4.1 DTO design rules", 2)
    add_bullets(doc, [
        "Each list-view DTO carries denormalised display fields "
        "(SchoolClassName, SubjectName, TermName) so the UI can "
        "render a row without lazy-loading any navigations.",
        "Detail DTOs (DailyAttendanceRegisterDetailDto, "
        "SubjectAttendanceSessionDetailDto) bundle the parent row "
        "and its entries into one payload, so the editor page is "
        "a single round-trip.",
        "Bulk-set requests carry the full set of entries; the "
        "service diffs against existing rows. The UI does not "
        "have to track dirty rows.",
        "Filter types (DailyRegisterListFilter, "
        "SubjectSessionListFilter, StudentAttendanceSummaryFilter) "
        "carry the search/scope parameters.",
    ])

    add_heading(doc, "4.2 DailyAttendanceDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Attendance/Dtos/DailyAttendanceDtos.cs")
    add_para(doc,
        "PresentCount, AbsentCount, LateCount and TotalCount are "
        "populated by the service projection so the list view can "
        "show 'P:28 A:1 L:1 / 30' for each register without a "
        "follow-up query.")

    add_heading(doc, "4.3 SubjectAttendanceDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Attendance/Dtos/SubjectAttendanceDtos.cs")

    add_heading(doc, "4.4 AttendanceSummaryDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Attendance/Dtos/AttendanceSummaryDtos.cs")
    add_para(doc,
        "PresentRate is a computed property in the DTO itself — "
        "rounded to one decimal place, returns 0 when DaysCounted "
        "is zero. The summary grid uses this directly to colour "
        "the badge (≥ 90% green, ≥ 75% amber, otherwise red).")

    add_heading(doc, "4.5 IDailyAttendanceService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Attendance/IDailyAttendanceService.cs")

    add_heading(doc, "4.6 ISubjectAttendanceService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Attendance/ISubjectAttendanceService.cs")

    add_heading(doc, "4.7 ILookupService extension", 2)
    add_para(doc,
        "ILookupService grew by one method, GetAttendanceStatusesAsync, "
        "which surfaces the AttendanceStatus lookup as a list of "
        "LookupDto items (Id, Name, Code). The Code is what the UI "
        "uses to spot 'Late' specifically when deciding whether to "
        "show the arrival-time field.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Application/Users/ILookupService.cs",
        "Task<IReadOnlyList<LookupDto>> GetAttendanceStatusesAsync",
        lines_after_start=1,
        caption="Excerpt — ILookupService.cs (sprint 4 addition)")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)
    add_para(doc,
        "ApplicationDbContext picks up five new DbSets and a single "
        "new ConfigureAttendance method. Everything else in the file "
        "is unchanged from sprint 3. The pattern of a private "
        "ConfigureXxx method per feature folder keeps OnModelCreating "
        "lean.")

    add_heading(doc, "5.1 New DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<AttendanceStatus>",
        end_marker="protected override void OnModelCreating",
        caption="Excerpt — the five attendance DbSets")

    add_heading(doc, "5.2 ConfigureAttendance", 2)
    add_para(doc,
        "ConfigureAttendance is invoked from OnModelCreating after "
        "ConfigureFamily. It uses the same ConfigureLookup helper "
        "from sprint 1 for AttendanceStatus, then configures the "
        "two register/entry pairs in turn.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureAttendance",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureAttendance")

    add_heading(doc, "5.3 SaveChanges still does all the work", 2)
    add_para(doc,
        "The SaveChanges override from sprint 1 needs no changes "
        "for sprint 4. Every new entity inherits BaseEntity, so "
        "the audit stamping and soft-delete rewriting apply "
        "automatically. The pattern continues to pay off.")

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)
    add_para(doc,
        "Two new services land in src/NaijaPrimeSchool.Infrastructure/"
        "Services/. Each one is a straightforward implementation of "
        "the matching Application interface. LookupService grows by "
        "one method.")

    add_heading(doc, "6.1 DailyAttendanceService.cs", 2)
    add_para(doc,
        "DailyAttendanceService is the heavier of the two. It owns "
        "the daily-register CRUD, the bulk update path, the "
        "submit/reopen lifecycle, and the per-student / per-class "
        "summary queries.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/DailyAttendanceService.cs")
    add_para(doc, "Three patterns are worth calling out:")
    add_bullets(doc, [
        "OpenAsync resolves the term automatically. It first tries "
        "to find a Term whose StartDate ≤ Date ≤ EndDate; if no "
        "term covers the date (typical for a date that lands in a "
        "vacation), it falls back to the current term in the "
        "session. If neither resolves, it returns a friendly error "
        "asking the admin to set up a term first.",
        "OpenAsync is idempotent. Calling it twice for the same "
        "(class, date) returns the existing register's Id without "
        "duplicating rows. The unique index makes this the safest "
        "approach.",
        "BulkSetAsync diffs against existing rows. It loads the "
        "existing entries into a dictionary keyed by StudentId, "
        "then for each request entry it either updates the matching "
        "row or inserts a new one. The pattern is simple and "
        "correct under concurrent edits because the UI only ever "
        "submits the full set.",
    ])

    add_heading(doc, "6.2 SubjectAttendanceService.cs", 2)
    add_para(doc,
        "SubjectAttendanceService follows the same shape as the "
        "daily service, with the weekday-vs-date validation in "
        "OpenAsync that section 2.6 described. There is no summary "
        "query because per-subject summaries are deferred.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/SubjectAttendanceService.cs")

    add_heading(doc, "6.3 LookupService — the new method", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",
        "public async Task<IReadOnlyList<LookupDto>> GetAttendanceStatusesAsync",
        lines_after_start=5,
        caption="Excerpt — LookupService.GetAttendanceStatusesAsync")

    add_heading(doc, "6.4 DI registration", 2)
    add_para(doc,
        "DependencyInjection.cs picks up two new lines.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IDailyAttendanceService",
        end_marker="return services;",
        caption="Excerpt — DependencyInjection.cs (sprint 4 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_para(doc,
        "A single migration named Attendance adds five new tables "
        "(AttendanceStatuses, DailyAttendanceRegisters, "
        "DailyAttendanceEntries, SubjectAttendanceSessions, "
        "SubjectAttendanceEntries) and the indexes that go with "
        "them. It was generated with:")
    add_code(doc,
        "dotnet ef migrations add Attendance \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_para(doc,
        "On a fresh checkout the migration runs at startup via "
        "DatabaseInitializer.MigrateAsync. The full Up() method "
        "is embedded below for reference.")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")
    add_para(doc,
        "Notice the sequence: AttendanceStatuses first, then the "
        "two parent tables (DailyAttendanceRegisters and "
        "SubjectAttendanceSessions), then the two entry tables. "
        "EF Core computes the ordering automatically from the "
        "foreign-key graph; we did not have to specify it.")

    add_heading(doc, "7.1 Indexes created", 2)
    add_bullets(doc, [
        "Unique on AttendanceStatuses.Name and "
        "AttendanceStatuses.Code (2 unique).",
        "Unique on DailyAttendanceRegisters.(SchoolClassId, Date) "
        "(1 unique).",
        "Unique on DailyAttendanceEntries.(RegisterId, StudentId) "
        "(1 unique).",
        "Unique on SubjectAttendanceSessions.(TimetableEntryId, "
        "Date) (1 unique).",
        "Unique on SubjectAttendanceEntries.(SessionId, StudentId) "
        "(1 unique).",
        "Plus IsDeleted indexes on every soft-deletable table (5).",
        "Plus a Date index on each of the two parent tables for "
        "date-range queries (2).",
    ])

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the AttendanceStatus lookup", 1)
    add_para(doc,
        "DatabaseInitializer picks up a new SeedAttendanceLookupsAsync "
        "method that is invoked between SeedFamilyLookupsAsync and "
        "SeedRolesAsync. It seeds six rows: Present, Late, Excused, "
        "Sick, Absent, Suspended. CountsAsPresent is true for "
        "Present and Late, false for the rest.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedAttendanceLookupsAsync",
        end_marker="private static async Task SeedLookupsAsync",
        caption="Excerpt — SeedAttendanceLookupsAsync")
    add_para(doc,
        "The seeder uses the same .IgnoreQueryFilters().AnyAsync() "
        "guard as the previous sprints' seeders — it only inserts "
        "if the table is empty (counting soft-deleted rows), so "
        "running the app a second time after a soft delete does "
        "not resurrect the missing row.")

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_para(doc,
        "Three new pages land in src/NaijaPrimeSchool.Web/Components/"
        "Pages/Attendance/. Two are register-taking pages (one daily, "
        "one per-subject), and the third is a class-level summary.")

    add_heading(doc, "9.1 Page roster", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/\n"
        "├── DailyAttendance.razor      ← /attendance/daily   (open + mark + submit)\n"
        "├── SubjectAttendance.razor    ← /attendance/subject (per-lesson register)\n"
        "└── AttendanceSummary.razor    ← /attendance/summary (class % view)\n")
    add_para(doc,
        "All three pages are gated to SuperAdmin + HeadTeacher + "
        "Teacher. Teachers genuinely need write access to the daily "
        "and subject pages — they are the people taking attendance. "
        "Bursars, Storekeepers, Parents, and Students do not see the "
        "panel.")

    add_heading(doc, "9.2 DailyAttendance.razor — the workhorse", 2)
    add_para(doc,
        "DailyAttendance.razor is the page teachers will use every "
        "morning. The top filter bar narrows down the (session, "
        "class, date) tuple; the 'Open register' button creates "
        "(or opens an existing) register for that combination.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/DailyAttendance.razor")
    add_para(doc,
        "Three details deserve a moment:")
    add_bullets(doc, [
        "Status is rendered as a RadzenDropDown bound directly to "
        "the entry's AttendanceStatusId. The OnStatusChanged "
        "handler also updates the entry's Code locally so that "
        "the arrival-time column can re-render based on the new "
        "code without a server round-trip.",
        "Arrival time appears only when the status code is 'L'. "
        "The TimeOnly value is parsed from a free-text 'HH:mm' "
        "field; bad input simply leaves the value alone, which "
        "is friendly enough for a primary-school workflow.",
        "Submit calls Save first, then SubmitAsync. This makes "
        "the two-button workflow simple: edit in the grid, click "
        "Submit, and the page does the right thing without "
        "asking the teacher to remember to click Save.",
    ])

    add_heading(doc, "9.3 SubjectAttendance.razor — the per-lesson page", 2)
    add_para(doc,
        "SubjectAttendance.razor has a two-stage flow. First the "
        "teacher picks (term, class, date); the page lists every "
        "lesson on the timetable for that class on that weekday. "
        "Clicking 'Take attendance' on a lesson opens (or "
        "creates) a SubjectAttendanceSession for that "
        "(timetable entry, date) pair and renders the register.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/SubjectAttendance.razor")
    add_para(doc,
        "The page leans on sprint 2's TimetableService to find "
        "lessons. There is a deliberate filter on WeekDayName "
        "matching DateTime.DayOfWeek.ToString() — Saturday "
        "lessons would only appear if a Saturday WeekDay row "
        "existed and the date being marked is a Saturday.")

    add_heading(doc, "9.4 AttendanceSummary.razor — the percentage view", 2)
    add_para(doc,
        "AttendanceSummary.razor renders the per-class summary "
        "the IDailyAttendanceService.GetClassSummaryAsync method "
        "returns. Each pupil's percentage is bucketed into "
        "green (≥ 90%), amber (≥ 75%), or red (< 75%) — a "
        "common rule of thumb in Nigerian primary schools.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Attendance/AttendanceSummary.razor")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and authorization", 1)
    add_heading(doc, "10.1 NavMenu — a new Attendance panel", 2)
    add_para(doc,
        "NavMenu.razor picks up a third role-gated panel between "
        "Family and the Finance/Inventory placeholders. Three "
        "entries live inside it: Daily attendance, Subject "
        "attendance, Summary. The whole panel is wrapped in an "
        "AuthorizeView that requires SuperAdmin, HeadTeacher, "
        "or Teacher.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Attendance\"",
        end_marker="</AuthorizeView>",
        caption="Excerpt — NavMenu.razor")

    add_heading(doc, "10.2 _Imports.razor", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Attendance",
        lines_after_start=2,
        caption="Excerpt — _Imports.razor")

    add_heading(doc, "10.3 The page-class / inject-field collision", 2)
    add_para(doc,
        "There is one Blazor-specific gotcha worth flagging. The "
        "Razor page's class name is derived from the file name, "
        "so DailyAttendance.razor compiles into a class called "
        "DailyAttendance. If you @inject IDailyAttendanceService "
        "with the field name DailyAttendance, the C# compiler "
        "rejects it with CS0542 because a member cannot share a "
        "name with its enclosing type.")
    add_para(doc,
        "Both attendance pages therefore inject their service "
        "with the field name AttendanceService. It is a small "
        "thing, but it is the kind of error that takes ten "
        "minutes to recognise the first time.")

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of a daily register", 1)
    add_para(doc,
        "Walking a single register from creation to submission "
        "is the clearest way to see how all the layers cooperate. "
        "Imagine the Primary 1A class teacher taking morning "
        "attendance on Monday 4 May 2026.")

    add_heading(doc, "11.1 Teacher opens /attendance/daily", 2)
    add_bullets(doc, [
        "Authorization fires: the user must be SuperAdmin, "
        "HeadTeacher, or Teacher. If they are not, ASP.NET Core "
        "returns 403.",
        "OnInitializedAsync calls SessionService.ListAsync, "
        "LookupService.GetClassesForSessionAsync, and "
        "GetAttendanceStatusesAsync — three small queries, "
        "all on tiny tables.",
        "The session dropdown defaults to the current session, "
        "and the date picker defaults to today.",
    ])

    add_heading(doc, "11.2 Teacher picks Primary 1A and clicks Open register", 2)
    add_bullets(doc, [
        "DailyAttendanceService.OpenAsync looks up the class to "
        "verify it exists and to read its SessionId.",
        "It looks for a Term where StartDate ≤ 2026-05-04 ≤ "
        "EndDate in that session. Suppose it finds Second Term.",
        "It checks for an existing register on (class, "
        "2026-05-04). None exists, so it creates one with "
        "TermId pointing at Second Term.",
        "It loads the lowest-DisplayOrder AttendanceStatus, "
        "which is 'Present'.",
        "It loads every Enrolment in the class with EnrolledOn "
        "≤ 2026-05-04 AND (WithdrawnOn IS NULL OR WithdrawnOn "
        "≥ 2026-05-04). Suppose it returns 28 pupils.",
        "It creates 28 DailyAttendanceEntry rows, each pointing "
        "at the new register and the Present status.",
        "SaveChangesAsync commits register + 28 entries in one "
        "transaction. The audit trail captures the teacher's "
        "username on every CreatedBy column.",
    ])

    add_heading(doc, "11.3 Teacher marks two pupils Late and one Absent", 2)
    add_bullets(doc, [
        "Each dropdown change fires OnStatusChanged in the "
        "Razor file. The change is local to the page model — "
        "no service call yet.",
        "For the two Late pupils, the arrival-time text box "
        "appears next to the dropdown. The teacher types "
        "08:15 and 08:22 respectively.",
        "Teacher clicks 'Save changes'. BulkSetAsync sends "
        "all 28 entries to the server.",
        "The service loads the existing 28 entries into a "
        "dictionary keyed by StudentId. For each request "
        "entry, it finds the matching existing row and updates "
        "the AttendanceStatusId, ArrivalTime, and Remarks. "
        "The two Late rows and the one Absent row are "
        "modified; the other 25 rows show no change. EF "
        "Core's change tracker emits an UPDATE statement only "
        "for the dirty rows.",
        "ApplicationDbContext.SaveChanges stamps ModifiedOn/By "
        "on each modified row.",
    ])

    add_heading(doc, "11.4 Teacher clicks Submit", 2)
    add_bullets(doc, [
        "SubmitAsync (after a final SaveAsync that catches any "
        "edits made between the last Save and the Submit click) "
        "sets IsSubmitted = true and SubmittedOn = now on the "
        "register row.",
        "The grid re-renders. Every dropdown is now disabled. "
        "The 'Save' and 'Submit' buttons are replaced with a "
        "'Submitted just now' badge and a 'Reopen' button.",
    ])

    add_heading(doc, "11.5 Half an hour later — a parent calls", 2)
    add_bullets(doc, [
        "An admin (SuperAdmin or HeadTeacher) clicks Reopen. "
        "ReopenAsync flips IsSubmitted = false and clears "
        "SubmittedOn.",
        "The grid re-renders with editable dropdowns. The "
        "admin changes one pupil from Absent to Late, "
        "captures the arrival time, and clicks Submit again.",
        "ModifiedOn/By on the affected row records the second "
        "edit. The audit trail makes the post-submission edit "
        "visible.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)
    add_para(doc,
        "Once the build is green and the migration has applied, "
        "here is the end-to-end smoke test you can run against a "
        "fresh checkout. It covers the happy path and a couple "
        "of important error paths.")

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "The first run applies migrations and seeds the "
        "AttendanceStatus lookup. Sign in as "
        "superadmin@naijaprimeschool.ng / Admin@12345.")

    add_heading(doc, "12.2 Set up the prerequisites", 2)
    add_numbered(doc, [
        "Academics → Sessions: create a session for the "
        "current academic year if one does not exist; mark it "
        "current.",
        "Academics → Terms: create at least one term covering "
        "today's date.",
        "Academics → Classes: create Primary 1A.",
        "Academics → Subjects: create Mathematics.",
        "Academics → Periods: keep the seeded periods or "
        "tweak them.",
        "Academics → Timetable: pick the term and Primary "
        "1A; assign Mathematics to the Monday Period 1 cell.",
        "Family → Add Student: create at least one pupil and "
        "enrol them in Primary 1A.",
    ])

    add_heading(doc, "12.3 Take a daily register", 2)
    add_numbered(doc, [
        "Attendance → Daily attendance.",
        "Pick the current session, Primary 1A, and today's "
        "date.",
        "Click Open register. The pre-loaded grid shows the "
        "pupil(s) you enrolled, all set to Present.",
        "Change one pupil to Late and capture an arrival time.",
        "Click Save changes — confirmation toast appears.",
        "Click Submit register. The grid disables; the badge "
        "flips to 'Submitted just now'.",
        "Click Reopen. The grid re-enables. Click Submit "
        "again to lock it back.",
    ])

    add_heading(doc, "12.4 Take a subject register", 2)
    add_numbered(doc, [
        "Attendance → Subject attendance.",
        "Pick the term, Primary 1A, and a Monday date that "
        "covers the timetable entry you set up.",
        "Click Take attendance on the Mathematics row.",
        "Mark the pupil(s) and Submit.",
    ])

    add_heading(doc, "12.5 Verify error paths", 2)
    add_numbered(doc, [
        "From Subject attendance, pick a Tuesday. The "
        "Mathematics row should not appear (it runs on "
        "Monday only).",
        "Manually call OpenAsync via the daily page on a "
        "Saturday in a session that has no term covering "
        "Saturdays. The register cannot be opened — the "
        "service returns 'No term covers this date and no "
        "current term is set'.",
        "Submit a register, then try to delete it from the "
        "page. The delete is refused with 'Submitted "
        "registers cannot be deleted. Reopen first.'",
    ])

    add_heading(doc, "12.6 Verify the summary view", 2)
    add_para(doc,
        "Attendance → Summary. Pick the session, term, and "
        "Primary 1A. Each pupil's row shows days-counted, "
        "days-present, days-absent, days-late, days-excused, "
        "and a green/amber/red percentage badge. If a pupil "
        "was absent on the day you marked, their percentage "
        "drops accordingly.")

    add_heading(doc, "12.7 Verify soft-delete and audit", 2)
    add_para(doc, "Connect to SQL Server and run:")
    add_code(doc,
        "SELECT Id, SchoolClassId, [Date], IsSubmitted,\n"
        "       SubmittedOn, IsDeleted, DeletedOn\n"
        "FROM DailyAttendanceRegisters;\n"
        "\n"
        "SELECT Id, RegisterId, StudentId,\n"
        "       AttendanceStatusId, ArrivalTime,\n"
        "       CreatedOn, ModifiedOn\n"
        "FROM DailyAttendanceEntries;\n")
    add_para(doc,
        "If you reopened a submitted register and changed a "
        "row, ModifiedOn on that DailyAttendanceEntry row is "
        "stamped. The register's SubmittedOn is null after "
        "the reopen and re-stamped after the second submit.")

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)
    add_heading(doc, "13.1 'No term covers this date and no current term is set'", 2)
    add_para(doc,
        "OpenAsync on the daily service refuses to create a "
        "register if the date has no term and the session has "
        "no current term. Set up a term first via "
        "/terms (sprint 2) or mark an existing term current.")

    add_heading(doc, "13.2 'Selected date is a Wednesday but this lesson runs on Monday'", 2)
    add_para(doc,
        "OpenAsync on the subject service refuses a date whose "
        "DayOfWeek does not match the timetable entry's WeekDay. "
        "Check the calendar; the timetable entry is correct.")

    add_heading(doc, "13.3 'Submitted registers cannot be deleted'", 2)
    add_para(doc,
        "Deleting a submitted register is intentionally a "
        "two-step: reopen it first, then delete. This forces "
        "the action to leave a clear audit trail.")

    add_heading(doc, "13.4 The pupil list on a register is stale", 2)
    add_para(doc,
        "The register pre-loads pupils at OpenAsync time. If "
        "you enrol a new pupil after opening the register, "
        "they will not appear automatically — that is by "
        "design, because the register is a snapshot of who "
        "the teacher saw that day. SetEntryAsync can still "
        "be called for the new pupil if you want to add them, "
        "but the page does not surface that flow today; "
        "deleting the register and reopening it is the "
        "simplest path.")

    add_heading(doc, "13.5 'member names cannot be the same as their enclosing type' (CS0542)", 2)
    add_para(doc,
        "If you copy the existing pages and rename the file "
        "without renaming the @inject field, the C# compiler "
        "will reject the build. The Razor page class name is "
        "derived from the file name; an inject field cannot "
        "share that name. Both attendance pages therefore use "
        "AttendanceService as the inject field name.")

    add_heading(doc, "13.6 The summary page shows 0 days counted", 2)
    add_para(doc,
        "GetClassSummaryAsync reads from DailyAttendanceEntry, "
        "not from SubjectAttendanceEntry. If you have only "
        "taken subject attendance for the class, the summary "
        "will show zeros until at least one daily register is "
        "submitted.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_para(doc,
        "Sprint 4 has deliberately left a few breadcrumbs that "
        "make the next sprints' work cleaner.")
    add_bullets(doc, [
        "AttendanceStatus.CountsAsPresent is the load-bearing "
        "flag for percentage calculations. Sprint 5 (results) "
        "will likely lean on attendance percentages as a "
        "feeder into the term report card; the column is "
        "ready.",
        "DailyAttendanceRegister.TermId means 'all attendance "
        "for first term' is a single index seek. Reporting "
        "queries do not need to span term boundaries by "
        "joining on date ranges.",
        "SubjectAttendanceSession.TimetableEntryId means "
        "'all attendance for Mathematics in second term' is "
        "a join through TimetableEntry — fast and friendly.",
        "Reopening + the audit trail mean a parent portal "
        "(later sprint) can show 'attendance updated this "
        "morning at 08:42' without losing the original "
        "register's submission timestamp.",
        "TakenById is nullable, so SchoolClass.ClassTeacher "
        "is not load-bearing here. If a class has no class "
        "teacher assigned, attendance still works — TakenBy "
        "just stays null.",
    ])

    add_heading(doc, "14.1 What might need a small refactor later", 2)
    add_bullets(doc, [
        "GetClassSummaryAsync issues one query per pupil. "
        "Once classes have 30+ pupils and the school has 2000+ "
        "pupils total, this will need to become a single "
        "GROUP BY query in SQL. The lift is small (LINQ "
        "GroupBy on StudentId) but I deliberately kept the "
        "implementation simple here.",
        "The arrival-time field is a free-text 'HH:mm' input "
        "rather than a Radzen TimePicker. Radzen's TimePicker "
        "expects DateTime?, and converting through the binding "
        "boundary is fiddly enough to defer until the design "
        "system has a TimeOnly story.",
        "Per-subject summary reporting (\"Adaeze missed three "
        "Maths lessons this term\") is straightforward off "
        "the existing schema but the UI is missing. A "
        "/attendance/subject-summary page is on the backlog.",
        "Bulk editor for a whole week. Currently a teacher "
        "edits one date at a time. A 'mark this pupil "
        "absent for the whole week' workflow would be a "
        "useful add-on.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 4", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Attendance/AttendanceStatus.cs",          "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Attendance/DailyAttendanceRegister.cs",   "Class × date root."),
        ("src/NaijaPrimeSchool.Domain/Attendance/DailyAttendanceEntry.cs",      "Register × pupil."),
        ("src/NaijaPrimeSchool.Domain/Attendance/SubjectAttendanceSession.cs",  "Timetable entry × date root."),
        ("src/NaijaPrimeSchool.Domain/Attendance/SubjectAttendanceEntry.cs",    "Session × pupil."),
        ("Domain layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs",                "Added DailyAttendanceRegisters navigation."),
        ("src/NaijaPrimeSchool.Domain/Academics/Term.cs",                       "Added DailyAttendanceRegisters navigation."),
        ("src/NaijaPrimeSchool.Domain/Academics/TimetableEntry.cs",             "Added AttendanceSessions navigation."),
        ("src/NaijaPrimeSchool.Domain/Family/Student.cs",                       "Added DailyAttendanceEntries + SubjectAttendanceEntries navigations."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Attendance/Dtos/DailyAttendanceDtos.cs",     "Daily register DTOs."),
        ("src/NaijaPrimeSchool.Application/Attendance/Dtos/SubjectAttendanceDtos.cs",   "Subject session DTOs."),
        ("src/NaijaPrimeSchool.Application/Attendance/Dtos/AttendanceSummaryDtos.cs",   "Per-pupil/class summary DTOs."),
        ("src/NaijaPrimeSchool.Application/Attendance/IDailyAttendanceService.cs",      "Daily register service contract."),
        ("src/NaijaPrimeSchool.Application/Attendance/ISubjectAttendanceService.cs",    "Subject session service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",                    "Added GetAttendanceStatusesAsync."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/DailyAttendanceService.cs",      "Daily register CRUD + summaries."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/SubjectAttendanceService.cs",    "Subject session CRUD."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 5 tables and 13 indexes."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",                  "Registered the 2 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",     "Added 5 DbSets, ConfigureAttendance."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",      "Seeded AttendanceStatus."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",               "Added GetAttendanceStatusesAsync."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Attendance/DailyAttendance.razor",      "Daily register editor."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Attendance/SubjectAttendance.razor",    "Per-lesson register editor."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Attendance/AttendanceSummary.razor",    "Class % view."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                              "Added Attendance + Attendance.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                        "Added the Attendance panel."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint4_guide.py",                                                 "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 4 implementation guide. The next sprint "
        "lands assessments and report cards on top of the data primitives "
        "established here.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
