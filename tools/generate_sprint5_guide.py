"""Generates 'Sprint 5 - Implementation Guide.docx' covering the
assessments, results and report-card work delivered in sprint 5
(TermAssessment, AssessmentScore, SubjectResult, ReportCard, the
five new lookup tables, the supporting services, and the Razor
pages that drive the gradebook -> result -> report-card pipeline).

This is the long-form edition. Code blocks embed actual source files
from the repository so the guide stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint5_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 5 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260504172028_AssessmentsAndResults.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 5 — Assessments, Results & Report Cards")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Continuous assessment · Exam scores · Result computation · Term report cards")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/5-results-reports")
    meta.add_run("\nBuilt on: Sprint 1 identity + Sprint 2 academic domain + Sprint 3 students & parents + Sprint 4 attendance")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 5 in context", 1)
    add_para(doc,
        "Sprint 5 closes the academic loop. The first four sprints set up "
        "identities, the academic calendar, the families that depend on the "
        "school, and the daily attendance register that says who turned up. "
        "Sprint 5 takes the work the pupils actually do — quizzes, tests, "
        "projects, exams — translates it into subject-level percentages and "
        "grades, and compiles those grades into the per-pupil end-of-term "
        "report card that goes home in the school bag.")
    add_para(doc,
        "Functionally, the sprint introduces a three-stage pipeline. Stage "
        "one is the gradebook: a SuperAdmin or HeadTeacher (or Teacher with "
        "the right scope) creates one TermAssessment row per piece of work, "
        "specifying max score and weight, then keys the per-pupil scores. "
        "Stage two is computation: a single button on the results page sums "
        "every weighted score for a (term, class), produces a "
        "SubjectResult row per pupil per subject, looks up the GradeBand, "
        "and ranks the class. Stage three is the report card: another "
        "button rolls every subject result for a pupil into a single "
        "ReportCard row, snaps in the attendance summary from sprint 4, and "
        "leaves room for the class teacher and head teacher's comments and "
        "for the affective and psychomotor ratings the Nigerian curriculum "
        "expects.")
    add_para(doc,
        "Once this sprint ships, every other module the school needs has a "
        "complete data path through the system. Fees can hang off pupils "
        "with active enrolments and now-finalised report cards. Parent "
        "portals can show the same report cards. Promotion can read the "
        "Average Percentage column. None of those features needed sprint 5 "
        "to start, but each of them now has a load-bearing dataset to read.")
    add_para(doc,
        "This document is a long-form implementation guide written in the "
        "tone of the sprint 4 guide. An engineer who has read sprints 1–4 "
        "and has the codebase checked out can recreate every change in "
        "this sprint without referring to the diff. The structure mirrors "
        "the build order: design decisions first, Domain entities next, "
        "Application contracts after that, Infrastructure (DbContext, "
        "services, seeder, migration) in the middle, then the Razor UI and "
        "navigation. Smoke-test, troubleshooting, and forward-compatibility "
        "chapters round it off.")

    add_heading(doc, "1.1 Where this sits relative to sprint 4", 2)
    add_para(doc,
        "Sprint 4 delivered the (Class × Date) and (TimetableEntry × Date) "
        "primitives for attendance plus the AttendanceStatus lookup. "
        "Sprint 5 reuses several of those load-bearing pieces unchanged "
        "and bolts new tables next to them. In particular:")
    add_bullets(doc, [
        "BaseEntity — every new entity in sprint 5 inherits it and picks up "
        "Guid Id, IAuditable, and ISoftDelete with no boilerplate.",
        "ApplicationDbContext.SaveChanges — the override stamps "
        "CreatedOn/By and ModifiedOn/By and rewrites Delete to "
        "IsDeleted = true. Every assessment, score, result, and report "
        "card therefore inherits auditing and soft delete with no extra "
        "code.",
        "Global query filters — every new entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries automatically.",
        "OperationResult / OperationResult<T> — every new service returns "
        "this for predictable success/failure responses.",
        "ILookupService — already had fifteen methods. Sprint 5 adds five "
        "(assessment types, grade bands, affective traits, psychomotor "
        "skills, trait ratings) without rewriting the existing ones.",
        "Term, SchoolClass, Subject, Student — the four big sprint 1–3 "
        "entities pick up new collection navigations only. None of their "
        "scalar columns change, so existing code that ignored these "
        "navigations continues to work.",
        "DailyAttendanceEntries / DailyAttendanceRegisters — the report "
        "card generator joins through these to compute days present, "
        "absent, and late at the time of generation.",
        "Radzen Blazor + the green/gold app.css — the new pages adopt the "
        "same .nps-page-header / .nps-card / .nps-form-grid primitives, "
        "plus a small new pair of grid styles (.nps-score-grid and "
        ".nps-trait-grid) so they read as part of the same product.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_para(doc,
        "Concretely, after this sprint a SuperAdmin or HeadTeacher (or, for "
        "the assessment pages, a Teacher) signing in to the application "
        "can:")
    add_numbered(doc, [
        "Create a TermAssessment for any (Term, Class, Subject) tuple, "
        "picking an AssessmentType (CA1, CA2, Mid-Term, Assignment, "
        "Project, Exam) and setting a max score and a multiplier weight.",
        "Open the score sheet for an assessment, see every actively-"
        "enrolled pupil pre-listed, key in scores or mark a pupil absent, "
        "save in bulk, and publish the assessment when ready.",
        "From the Results page, pick a (Term, Class) and recompute the "
        "weighted subject totals across every assessment in scope. The "
        "service produces a SubjectResult per pupil per subject, looks up "
        "the GradeBand from the percentage, and ranks the class with "
        "dense ordering on ties.",
        "Finalise individual SubjectResults so further recomputes leave "
        "them alone, or reopen them for a correction.",
        "From the Report Cards page, generate or refresh the per-pupil "
        "term card for a (Term, Class). The generator reads every "
        "SubjectResult, joins to the day-level attendance counts, and "
        "ranks pupils by average percentage.",
        "Open a single report card and key in the class teacher's and "
        "head teacher's comments, the next-term-begins date, and the "
        "five-point ratings for each affective trait (Punctuality, "
        "Honesty, …) and each psychomotor skill (Handwriting, Music, …).",
        "Publish a report card to lock it from further edits, unpublish "
        "to amend something the head teacher caught, soft-delete a card "
        "that should not have existed at all.",
    ])

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_para(doc,
        "Sprint 5 deliberately stops short of several adjacent ideas that "
        "all sit on top of the schema it lays down. Each was weighed and "
        "consciously deferred:")
    add_bullets(doc, [
        "PDF / printed report cards. The data model is complete and the "
        "ReportCardDetailDto carries everything a printable layout needs. "
        "We will add a Razor-based print stylesheet (and likely a "
        "QuestPDF-driven server endpoint) in a follow-up sprint, since "
        "the design choices there are partly visual.",
        "End-of-session aggregation. A pupil's annual position and "
        "average involve summing across three Term cards, which is a "
        "second computation pass. The data is available; the rule "
        "(promote at >=50%, repeat below) is school-policy-specific and "
        "deferred until we are talking to a real school.",
        "Subject-by-class subject lists. Today any assessment can be "
        "created for any (class, subject) pair, even if the class does "
        "not officially offer that subject. The 'official subject list "
        "per class' will land alongside curriculum/lesson-note tooling "
        "in a later sprint.",
        "Per-class assessment-scheme templates. Every assessment is "
        "created individually. A 'standard scheme: CA1 + CA2 + Exam = 100' "
        "template per class level would speed setup but is a feature on "
        "top of the existing schema, not a redesign.",
        "Per-subject teacher comments at scale. The SubjectResult row "
        "carries a TeacherComment column; the UI exposes editing on the "
        "list page. A bulk-edit form for a (class, subject) is on the "
        "wishlist but not yet built.",
        "Result analytics dashboards (subject pass rates, class trends "
        "across terms). Every input is in place; the visualisations "
        "themselves are a separate piece of work.",
        "Parent/student portal access to results. The Parent and Student "
        "roles still see placeholder navigation; the portal sprint will "
        "expose the existing IsPublished flag as the gate for what "
        "external users can see.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers, this sprint adds:")
    add_bullets(doc, [
        "11 new domain entities under src/NaijaPrimeSchool.Domain/Results/.",
        "4 collection navigations on existing entities (Subject, Term, "
        "SchoolClass, Student) so EF can navigate in the other direction; "
        "no scalar columns change on existing tables.",
        "4 DTO files under src/NaijaPrimeSchool.Application/Results/Dtos/.",
        "3 new service contracts under src/NaijaPrimeSchool.Application/Results/.",
        "3 service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "5 new methods on ILookupService (and the matching LookupService).",
        "1 EF Core migration introducing 11 new tables and the indexes "
        "that go with them.",
        "1 DatabaseInitializer extension seeding AssessmentTypes, "
        "GradeBands, AffectiveTraits, PsychomotorSkills, and TraitRatings.",
        "4 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Results/.",
        "1 navigation menu addition (Results & Reports panel) and 1 set "
        "of CSS additions (.nps-score-grid, .nps-trait-grid).",
    ])
    add_para(doc,
        "Everything compiles with zero warnings on .NET 10 (the team-wide "
        "warning bar). The code follows the patterns already accepted in "
        "sprints 1–4, so the diff is low-friction to review.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)
    add_para(doc,
        "Before any code was written I pinned down the shape of the "
        "pipeline. Two of the calls below shape the whole feature; the "
        "rest are the kind of medium decisions that pile up when you ship "
        "a sprint and that future maintainers will thank you for "
        "documenting.")

    add_heading(doc, "2.1 Three layers, three tables, no enums", 2)
    add_para(doc,
        "The single biggest decision in this sprint was to model the data "
        "in three layers — assessment, subject result, report card — "
        "rather than computing report cards directly from raw scores at "
        "render time. The reasons:")
    add_bullets(doc, [
        "Performance. A class of 40 pupils with 8 subjects and 4 "
        "assessments per subject is 1,280 score rows. Re-aggregating that "
        "graph every time the head teacher opens a report card would be "
        "wasteful; persisting the SubjectResult is essentially a cached "
        "projection.",
        "Auditability. SubjectResult rows record TotalScore, Percentage, "
        "GradeBandId, Position and a FinalisedOn timestamp. If a parent "
        "queries a card three months later, the row that was persisted "
        "is the row the school stands behind — even if a teacher edits a "
        "score afterwards.",
        "Composability. ReportCard then hangs off SubjectResults rather "
        "than scores. The ReportCard generator's loop is simple, and the "
        "service can refuse to refresh a card whose results no longer "
        "agree.",
        "Narrow services. With three tables come three services "
        "(AssessmentService, ResultService, ReportCardService). Each one "
        "is easy to read and easy to change. A combined IResultsService "
        "would have grown into a kitchen sink.",
    ])
    add_para(doc,
        "The rule from earlier sprints — that domain concepts which would "
        "normally be C# enums are stored as tables — survives this sprint "
        "without exception. AssessmentType, GradeBand, AffectiveTrait, "
        "PsychomotorSkill, and TraitRating are all proper entities that "
        "derive from BaseEntity and live in the database. Schools that "
        "want a fourth assessment type, an extra grade band, or a "
        "different rating ladder change a row in a table; no recompile, "
        "no deployment.")

    add_heading(doc, "2.2 Weighted percentage, dense ranking, no surprises", 2)
    add_para(doc,
        "The heart of the sprint is one method on ResultService: "
        "ComputeAsync. It walks every (term, class, subject) triple in "
        "scope, sums weighted raw scores per pupil, and stamps a "
        "percentage out of the per-subject total possible. A worked "
        "example, with three assessments on Mathematics:")
    add_code(doc,
        "  CA1   max=20  weight=1   ->  contributes  0..20\n"
        "  CA2   max=20  weight=1   ->  contributes  0..20\n"
        "  Exam  max=60  weight=2   ->  contributes  0..120\n"
        "  -----------------------------------------\n"
        "  Total possible = 20*1 + 20*1 + 60*2 = 160\n"
        "  Pupil scored CA1=18, CA2=15, Exam=40\n"
        "  Weighted    = 18 + 15 + 80 = 113\n"
        "  Percentage  = 113 / 160 = 70.625%  -> rounded to 70.63\n")
    add_para(doc,
        "The denominator is the per-subject 'total possible weighted', so "
        "weights act as multipliers exactly the way teachers expect. "
        "Setting every weight to 1 reduces to a simple sum-out-of-100.")
    add_para(doc,
        "Position is computed with dense ranking on the percentage column "
        "(ties share a place, the next pupil gets the next number, not a "
        "skipped one). Pupils with no enrolment in the class are filtered "
        "out before ranking begins. Students with all-zero scores still "
        "receive a percentage of 0 and a position; we do not silently drop "
        "them. The class size in the StudentsInClass column is the "
        "denominator the position is taken out of, so '5 of 38' tells the "
        "reader exactly what they think it does.")

    add_heading(doc, "2.3 Soft-delete plus operation guards, again", 2)
    add_para(doc,
        "Every entity in this sprint implements ISoftDelete via "
        "BaseEntity. The pattern matches sprint 4 exactly: SaveChanges "
        "intercepts EntityState.Deleted, flips IsDeleted, and stamps "
        "DeletedOn/By. Global query filters then hide the row from "
        "every subsequent query that does not call IgnoreQueryFilters.")
    add_para(doc, "What is genuinely new in sprint 5 is a layered set of "
        "delete and edit guards, all enforced in the services:")
    add_bullets(doc, [
        "TermAssessment cannot be edited or deleted while it is published. "
        "Unpublish first.",
        "TermAssessment cannot be deleted while it has scores attached. "
        "The service refuses with a friendly OperationResult.Failure. "
        "Clear the scores first or accept that the assessment exists in "
        "the audit trail.",
        "AssessmentScore is editable only while the parent assessment is "
        "in draft. Publishing locks the gradebook page to read-only.",
        "SubjectResult cannot be deleted while it is finalised. Reopen "
        "it first; this leaves an obvious audit trail.",
        "ReportCard cannot be edited or deleted while it is published. "
        "Comments, ratings, and the next-term-begins date are all "
        "rejected by the service.",
    ])

    add_heading(doc, "2.4 Unique indexes do the heavy lifting", 2)
    add_para(doc,
        "Three composite unique indexes are the integrity backbone of the "
        "sprint:")
    add_bullets(doc, [
        "(TermAssessmentId, StudentId) on AssessmentScore — a pupil's "
        "score for a given assessment is exactly one row.",
        "(StudentId, TermId, SubjectId) on SubjectResult — at most one "
        "computed result per (pupil, term, subject).",
        "(StudentId, TermId) on ReportCard — at most one card per pupil "
        "per term.",
    ])
    add_para(doc,
        "Plus (ReportCardId, AffectiveTraitId) and (ReportCardId, "
        "PsychomotorSkillId) on the rating tables, so each rating "
        "category is only ever stamped once per card. The services pre-"
        "check these unique constraints and surface friendly messages, "
        "but the database remains the authority.")

    add_heading(doc, "2.5 Decimal precision, not floats", 2)
    add_para(doc,
        "Every score, total, and percentage column uses HasPrecision so "
        "EF Core emits decimal columns, not the SqlServer default of "
        "decimal(18,2). Specifically:")
    add_bullets(doc, [
        "AssessmentScore.Score is decimal(7,2) — five digits before the "
        "point is plenty for any reasonable assessment, two after "
        "preserves half-mark resolution.",
        "TermAssessment.Weight is decimal(5,2) — schools rarely want a "
        "weight greater than 100, and two decimals lets you write 1.5 "
        "or 2.25.",
        "SubjectResult.TotalScore is decimal(7,2); SubjectResult."
        "Percentage is decimal(5,2). Percentages can be 100.00 exactly.",
        "GradeBand.LowerBound and UpperBound are decimal(5,2) so a band "
        "can be specified as 70.00 to 79.99 with no round-off ambiguity.",
        "ReportCard.TotalScore is decimal(7,2), AveragePercentage is "
        "decimal(5,2).",
    ])

    add_heading(doc, "2.6 Foreign-key delete behaviour", 2)
    add_para(doc, "The mix is deliberate:")
    add_bullets(doc, [
        "AssessmentScore.TermAssessmentId -> Cascade. Wiping a draft "
        "assessment that should never have been created is allowed (in "
        "the data sense; the service still requires the assessment to "
        "have no scores in the soft-delete pre-check).",
        "AffectiveRating.ReportCardId / PsychomotorRating.ReportCardId "
        "-> Cascade. Removing a draft report card cleans up its trait "
        "rows. The service still refuses if the card is published.",
        "Lookup foreign keys (AssessmentTypeId, GradeBandId, AffectiveTraitId, "
        "PsychomotorSkillId, TraitRatingId) -> Restrict where required, "
        "SetNull where optional. You cannot accidentally lose a grade "
        "band whose name appears on existing finalised results.",
        "Cross-aggregate keys (StudentId, TermId, SubjectId, "
        "SchoolClassId) -> Restrict. The schema refuses to lose a "
        "subject while results still reference it. Soft-delete from the "
        "service layer is the supported path.",
    ])

    add_heading(doc, "2.7 Inline forms again, no dialogs", 2)
    add_para(doc,
        "The sprint 2 / 3 / 4 pattern is preserved. Every CRUD page "
        "reveals an inline RadzenCard form below its data grid rather "
        "than opening a dialog. Three reasons:")
    add_bullets(doc, [
        "Validation messages stay close to fields and are easy to read.",
        "The score sheet and trait sheet are HTML tables (.nps-score-grid "
        "and .nps-trait-grid) — flat tables hosted in a card, not "
        "Radzen-data-grid editing — because every row needs three or "
        "four widgets next to each other and Radzen's grid editor "
        "chokes on that pattern.",
        "Server-side rendering of a Razor page is straightforward; "
        "dialog components require extra ceremony and life-time "
        "management we do not need.",
    ])

    add_heading(doc, "2.8 Tabs for the report card detail", 2)
    add_para(doc,
        "ReportCardDetail.razor uses a Radzen tab strip with four tabs: "
        "Subjects (the read-only table of computed results), Affective "
        "traits (the trait × rating grid), Psychomotor skills (same "
        "shape), and Comments (free-text plus the next-term date). "
        "Splitting these prevents the page from becoming an unreadable "
        "wall of fields and keeps the user's mental task — 'I am picking "
        "ratings now' — aligned with what is on the screen.")

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)
    add_para(doc,
        "Every sprint-5 entity lives in a single new folder, "
        "src/NaijaPrimeSchool.Domain/Results/. There are no abstract base "
        "classes other than BaseEntity, no domain methods, no validation "
        "logic. Validation lives in the Application DTOs (DataAnnotations) "
        "and in the Infrastructure services (cross-aggregate checks). "
        "The Domain layer remains a typed vocabulary.")

    add_heading(doc, "3.1 Folder layout and namespacing", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "├── Results/                       <- (new in sprint 5)\n"
        "│   ├── AssessmentType.cs          <- lookup\n"
        "│   ├── GradeBand.cs               <- lookup\n"
        "│   ├── AffectiveTrait.cs          <- lookup\n"
        "│   ├── PsychomotorSkill.cs        <- lookup\n"
        "│   ├── TraitRating.cs             <- lookup\n"
        "│   ├── TermAssessment.cs          <- gradebook entry\n"
        "│   ├── AssessmentScore.cs         <- per-pupil score\n"
        "│   ├── SubjectResult.cs           <- per (pupil, term, subject) total\n"
        "│   ├── ReportCard.cs              <- per (pupil, term) summary\n"
        "│   ├── AffectiveRating.cs         <- card x trait rating\n"
        "│   └── PsychomotorRating.cs       <- card x skill rating\n"
        "├── Attendance/                    <- from sprint 4\n"
        "├── Family/                        <- from sprint 3\n"
        "├── Academics/                     <- from sprint 2\n"
        "├── Common/                        <- from sprint 1\n"
        "└── Identity/                      <- from sprint 1\n")

    add_heading(doc, "3.2 The five lookup entities", 2)

    add_heading(doc, "3.2.1 AssessmentType.cs", 3)
    add_para(doc,
        "The lookup of what kind of assessment a row represents. Each row "
        "carries a default max score (so the form can pre-fill it for new "
        "assessments) and an IsExam flag (so the UI can stamp 'Exam' "
        "badges on summative assessments).")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/AssessmentType.cs")

    add_heading(doc, "3.2.2 GradeBand.cs", 3)
    add_para(doc,
        "GradeBand is the lookup of grade letters with bounds. The "
        "computation finds the band whose [LowerBound, UpperBound] range "
        "contains the percentage. Description is the human-friendly label "
        "(Excellent, Very Good, …); Remark is the boilerplate one-line "
        "comment that appears on the report card.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/GradeBand.cs")

    add_heading(doc, "3.2.3 AffectiveTrait.cs and PsychomotorSkill.cs", 3)
    add_para(doc,
        "Two parallel lookup tables for the soft-skills section of the "
        "Nigerian primary report card. AffectiveTraits cover behaviour "
        "and character (Punctuality, Honesty, Cooperation); "
        "PsychomotorSkills cover physical and creative work "
        "(Handwriting, Music, Sports). The schema is identical, just the "
        "owning collection differs.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/AffectiveTrait.cs")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/PsychomotorSkill.cs")

    add_heading(doc, "3.2.4 TraitRating.cs", 3)
    add_para(doc,
        "The shared 1–5 rating ladder used for both affective and "
        "psychomotor entries. Both navigation collections live here so "
        "the schema stays consistent.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/TraitRating.cs")

    add_heading(doc, "3.3 The six core entities", 2)

    add_heading(doc, "3.3.1 TermAssessment.cs", 3)
    add_para(doc,
        "TermAssessment is the gradebook entry: a single quiz, project, "
        "or exam tied to a (Term, Class, Subject) tuple. MaxScore and "
        "Weight together govern how the assessment contributes to the "
        "subject total during computation. IsPublished and PublishedOn "
        "drive the gradebook lock — once published, the score sheet is "
        "read-only.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/TermAssessment.cs")

    add_heading(doc, "3.3.2 AssessmentScore.cs", 3)
    add_para(doc,
        "AssessmentScore is the per-pupil score on a given assessment. "
        "Score is nullable so an absent pupil can be recorded with "
        "IsAbsent = true and Score = null without polluting the "
        "computation. Remarks lets a teacher annotate a particular "
        "score (e.g. 'sat the make-up exam', 'partial credit only').")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/AssessmentScore.cs")

    add_heading(doc, "3.3.3 SubjectResult.cs", 3)
    add_para(doc,
        "SubjectResult is the persisted output of ResultService."
        "ComputeAsync. TotalScore is the weighted total; Percentage is "
        "TotalScore / TotalPossibleWeighted * 100; GradeBandId is the "
        "lookup match for that percentage; Position is the dense rank in "
        "the (subject, class, term) cohort. IsFinalised + FinalisedOn "
        "lock the row from being recomputed without an explicit reopen.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/SubjectResult.cs")

    add_heading(doc, "3.3.4 ReportCard.cs", 3)
    add_para(doc,
        "ReportCard is the per-(pupil, term) roll-up. Carries the "
        "average across the pupil's subjects, the position in class, "
        "the attendance counts (read from sprint 4 at generation time), "
        "the two free-text comments, the next-term date, and the "
        "publishing flag. Two collection navigations dangle off it: the "
        "affective and psychomotor ratings.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/ReportCard.cs")

    add_heading(doc, "3.3.5 AffectiveRating.cs and PsychomotorRating.cs", 3)
    add_para(doc,
        "Two parallel join entities. Each one ties a specific report "
        "card to a specific trait/skill, with the chosen TraitRating. "
        "The schema enforces uniqueness on (ReportCardId, TraitId) and "
        "(ReportCardId, SkillId) respectively, so each row in the "
        "report-card UI grid maps to exactly one row.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/AffectiveRating.cs")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Results/PsychomotorRating.cs")

    add_heading(doc, "3.4 Back-references on existing entities", 2)
    add_para(doc,
        "Four sprint-1-to-4 entities pick up new collection navigations "
        "so EF can navigate the new graph in both directions. None of "
        "their scalar columns change, so existing code that ignored "
        "these properties continues to compile unchanged.")
    add_bullets(doc, [
        "Subject — TermAssessments and SubjectResults.",
        "Term — TermAssessments, SubjectResults, and ReportCards.",
        "SchoolClass — TermAssessments, SubjectResults, and ReportCards.",
        "Student — AssessmentScores, SubjectResults, and ReportCards.",
    ])
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/Subject.cs")

    add_heading(doc, "3.5 Relationships at a glance", 2)
    add_code(doc,
        "                       AssessmentType (lookup)\n"
        "                                |\n"
        "                                v 1..N\n"
        "    Term --------+------ TermAssessment ------+\n"
        "                 |              | 1            |\n"
        "    SchoolClass -+              |              |\n"
        "                 |              v N            |\n"
        "    Subject -----+------ AssessmentScore ------+--- Student\n"
        "                                                 N\n"
        "\n"
        "    Term + SchoolClass + Subject + Student --> SubjectResult\n"
        "                                          |\n"
        "                                          v\n"
        "                                       GradeBand (lookup)\n"
        "\n"
        "    Term + SchoolClass + Student --> ReportCard ---+\n"
        "                                                   |\n"
        "                            +----------------------+----------------+\n"
        "                            |                                       |\n"
        "                       AffectiveRating                       PsychomotorRating\n"
        "                            |                                       |\n"
        "                  AffectiveTrait + TraitRating          PsychomotorSkill + TraitRating\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)
    add_para(doc,
        "Application stays thin: just DTOs and service interfaces. There "
        "is no domain logic and no EF Core. Sprint 5 adds two new "
        "subfolders: Results/Dtos and Results.")

    add_heading(doc, "4.1 DTO design rules", 2)
    add_bullets(doc, [
        "Every list view returns a flat read-DTO with denormalised "
        "display fields. TermAssessmentDto carries the term/class/subject "
        "names alongside the assessment-type code so a row renders "
        "without lazy-loading any navigations.",
        "Computed counts (ScoredCount, ExpectedCount) are computed in "
        "the projection so the gradebook list shows progress at a "
        "glance.",
        "Aggregate operations (compute, generate) carry their own "
        "request/response shapes so the service signatures don't leak "
        "tuples or anonymous types.",
        "GradeBand is denormalised onto SubjectResultDto via "
        "GradeBandName and GradeBandRemark — the consuming Razor page "
        "is exclusively read-only for these fields, so denormalising "
        "saves one join per row.",
    ])

    add_heading(doc, "4.2 TermAssessmentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/Dtos/TermAssessmentDtos.cs")

    add_heading(doc, "4.3 AssessmentScoreDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/Dtos/AssessmentScoreDtos.cs")
    add_para(doc,
        "AssessmentScoreSheetDto is the shape the score-entry page "
        "consumes. It bundles the assessment metadata with the per-pupil "
        "rows so a single round-trip populates the whole UI.")

    add_heading(doc, "4.4 SubjectResultDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/Dtos/SubjectResultDtos.cs")
    add_para(doc,
        "Note ComputeResultsRequest with its Finalise flag — the same "
        "API supports 'recompute everything' (Finalise=false, idempotent "
        "and safe to run after every score change) and 'compute and "
        "lock' (Finalise=true, used at end of term). Already-finalised "
        "rows are left alone unless the user explicitly reopens them "
        "first; ComputeResultsResponse.Warnings carries any rows that "
        "were skipped for that reason.")

    add_heading(doc, "4.5 ReportCardDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/Dtos/ReportCardDtos.cs")
    add_para(doc,
        "ReportCardDetailDto is the read-shape the detail page asks for: "
        "the card itself, the list of subject results, plus the "
        "affective/psychomotor rating grids. Each grid row carries the "
        "trait/skill name and the rating name and value so the UI can "
        "render without any extra lookups.")

    add_heading(doc, "4.6 Service contracts", 2)
    add_para(doc,
        "Three new contracts. Each one is small and reads top-to-bottom "
        "as the lifecycle of one entity.")
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/IAssessmentService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/IResultService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Results/IReportCardService.cs")

    add_heading(doc, "4.7 ILookupService extension", 2)
    add_para(doc,
        "ILookupService grew by five methods, one per new lookup table.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Application/Users/ILookupService.cs",
        "Task<IReadOnlyList<LookupDto>> GetAssessmentTypesAsync",
        end_marker="Task<IReadOnlyList<LookupDto>> GetTraitRatingsAsync",
        caption="Excerpt — ILookupService.cs (sprint 5 additions)")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)
    add_para(doc,
        "ApplicationDbContext picks up eleven new DbSets, one new "
        "ConfigureResults method, and one extra invocation in "
        "OnModelCreating. Everything else in the file is unchanged from "
        "sprint 4.")

    add_heading(doc, "5.1 New DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<AssessmentType>", end_marker="public DbSet<PsychomotorRating>",
        lines_after_start=12,
        caption="Excerpt — the eleven results DbSets")

    add_heading(doc, "5.2 ConfigureResults", 2)
    add_para(doc,
        "ConfigureResults is invoked from OnModelCreating after "
        "ConfigureAttendance. Like its sprint 3 and 4 siblings, it uses "
        "the ConfigureLookup<T> helper for the five lookup tables, then "
        "configures each big entity one block at a time.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureResults",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureResults")
    add_para(doc, "Things worth dwelling on:")
    add_bullets(doc, [
        "HasPrecision on every score, total, and percentage column so "
        "the schema uses decimal(p,s) — see chapter 2.5.",
        "Cascade delete on AssessmentScore.TermAssessmentId, "
        "AffectiveRating.ReportCardId, and PsychomotorRating."
        "ReportCardId so wiping a draft assessment or draft card cleans "
        "up the children. The service still pre-checks for the right "
        "draft status before letting that happen.",
        "Restrict on every cross-aggregate FK (StudentId, TermId, "
        "SubjectId, SchoolClassId) — the schema refuses to drop a "
        "subject while results exist for it. Soft-delete from the "
        "service is the supported path.",
        "Composite unique indexes on (TermAssessmentId, StudentId), "
        "(StudentId, TermId, SubjectId), (StudentId, TermId), "
        "(ReportCardId, AffectiveTraitId), (ReportCardId, "
        "PsychomotorSkillId) — these are the five rules the schema "
        "considers inviolable.",
    ])

    add_heading(doc, "5.3 The ConfigureLookup helper, reused", 2)
    add_para(doc,
        "Same generic helper from sprints 1–4. Each new lookup picks up "
        "table name, primary key, audit columns, and the global "
        "soft-delete query filter — the only customisation per lookup "
        "is the Name column constraint and any extra unique indexes.")

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)
    add_para(doc,
        "Three new services land in src/NaijaPrimeSchool.Infrastructure/"
        "Services/. Each one implements the matching interface, leans "
        "on db.SaveChangesAsync for auditing and soft delete, and uses "
        "OperationResult to surface friendly errors.")

    add_heading(doc, "6.1 AssessmentService.cs", 2)
    add_para(doc,
        "The gradebook service. CRUD over TermAssessment plus the score-"
        "sheet operations. Two design notes:")
    add_bullets(doc, [
        "GetScoreSheetAsync pre-loads every actively-enrolled pupil for "
        "the assessment's class and matches them with any existing "
        "scores. Pupils with no score yet still appear in the response, "
        "with Id = Guid.Empty so the UI knows the row is new.",
        "BulkSetScoresAsync skips rows with no score and no IsAbsent and "
        "no remarks — keying nothing into a row should not create an "
        "empty AssessmentScore.",
    ])
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/AssessmentService.cs")

    add_heading(doc, "6.2 ResultService.cs", 2)
    add_para(doc,
        "The compute service. ComputeAsync is the longest method in the "
        "sprint and the one worth reading slowly:")
    add_numbered(doc, [
        "Validate the term and class exist.",
        "Pull the assessments in scope (everything in the term/class, or "
        "narrowed by subject if the request specifies one).",
        "Pull the score rows for those assessments in a single query.",
        "Pull the active pupil set for the class (an enrolment with the "
        "right SchoolClassId).",
        "Pull all grade bands ordered by display.",
        "For each subject in scope: compute the per-pupil weighted "
        "total, divide by the per-subject total possible weighted, "
        "round to two decimals, dense-rank by percentage.",
        "Look up the matching grade band per pupil and write a fresh "
        "SubjectResult or update the existing one. If the existing row "
        "is finalised and the request is not asking to finalise, log a "
        "warning and skip; never silently overwrite a finalised row.",
        "SaveChangesAsync once at the end. Auditing and soft-delete "
        "stamping run automatically.",
    ])
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/ResultService.cs")

    add_heading(doc, "6.3 ReportCardService.cs", 2)
    add_para(doc,
        "The composer service. GenerateAsync rolls the SubjectResult set "
        "for a (term, class) into one ReportCard per pupil, joins to the "
        "DailyAttendanceEntries from sprint 4 to compute days present / "
        "absent / late, and ranks pupils by average percentage. Like "
        "ComputeAsync it never overwrites a published card; the service "
        "log entry will say 'updated' or 'skipped' and the UI surfaces "
        "the count.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/ReportCardService.cs")
    add_para(doc,
        "GetByIdAsync and GetForStudentTermAsync share a private "
        "BuildDetailAsync that loads the SubjectResults, AffectiveRatings, "
        "and PsychomotorRatings as flat DTOs ready for the detail page.")

    add_heading(doc, "6.4 LookupService — five new methods", 2)
    add_para(doc,
        "All five new methods follow the same pattern: order by "
        "DisplayOrder, project to LookupDto, return the list. "
        "GetGradeBandsAsync stuffs the description into the Code slot "
        "of LookupDto so the UI can show 'A — Excellent' without an "
        "extra round-trip.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",
        "GetAssessmentTypesAsync(CancellationToken",
        end_marker="GetTraitRatingsAsync(CancellationToken",
        lines_after_start=40,
        caption="Excerpt — LookupService.cs (sprint 5 additions)")

    add_heading(doc, "6.5 DI registration", 2)
    add_para(doc, "DependencyInjection.cs picks up three new lines.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IAssessmentService",
        end_marker="return services;",
        caption="Excerpt — DependencyInjection.cs (sprint 5 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_para(doc,
        "A single migration named AssessmentsAndResults adds eleven new "
        "tables (AssessmentTypes, GradeBands, AffectiveTraits, "
        "PsychomotorSkills, TraitRatings, TermAssessments, "
        "AssessmentScores, SubjectResults, ReportCards, AffectiveRatings, "
        "PsychomotorRatings) and the indexes that go with them. It was "
        "generated with:")
    add_code(doc,
        "dotnet ef migrations add AssessmentsAndResults \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_para(doc,
        "On a fresh checkout the migration runs at startup via "
        "DatabaseInitializer.MigrateAsync. The full Up() method is "
        "embedded below for reference.")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")
    add_para(doc,
        "Sequence: lookup tables first, then TermAssessments (which "
        "references AssessmentTypes plus the sprint-2 Term/SchoolClass/"
        "Subject), then AssessmentScores (which references TermAssessment "
        "and the sprint-3 Student), then SubjectResults and ReportCards "
        "(which reference everything else), then the two rating tables. "
        "EF Core computes this from the foreign-key graph automatically.")

    add_heading(doc, "7.1 The schema warning", 2)
    add_para(doc,
        "The migration emits the same long-standing IdentityRole "
        "warning every previous sprint has emitted:")
    add_code(doc,
        "warn: Microsoft.EntityFrameworkCore.Model.Validation[10622]\n"
        "      Entity 'ApplicationRole' has a global query filter\n"
        "      defined and is the required end of a relationship with\n"
        "      the entity 'ApplicationUserRole'. ...")
    add_para(doc,
        "Unrelated to sprint 5. Runtime is fine because we never "
        "soft-delete a role from the UI. The role-management UI will "
        "revisit this in a later sprint.")

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the results lookups", 1)
    add_para(doc,
        "DatabaseInitializer picks up a new SeedResultsLookupsAsync "
        "method invoked between SeedAttendanceLookupsAsync and "
        "SeedRolesAsync. It seeds the five new lookup tables with "
        "sensible defaults a Nigerian primary school would want on "
        "day one.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedResultsLookupsAsync",
        end_marker="private static async Task SeedAcademicLookupsAsync",
        caption="Excerpt — SeedResultsLookupsAsync")
    add_para(doc, "What gets seeded:")
    add_bullets(doc, [
        "AssessmentTypes — First CA, Second CA, Mid-Term Test, "
        "Assignment, Project, Examination (6 rows). Examination is the "
        "one with IsExam = true so the UI can flag it.",
        "GradeBands — A (80–100), B (70–79.99), C (60–69.99), D "
        "(50–59.99), E (40–49.99), F (0–39.99). Each band carries a "
        "Description and a Remark, and the bounds use decimal(5,2) for "
        "precision.",
        "AffectiveTraits — Punctuality, Attentiveness, Honesty, "
        "Neatness, Politeness, Cooperation, Class Participation, Self-"
        "control (8 rows).",
        "PsychomotorSkills — Handwriting, Drawing & Painting, Music, "
        "Sports & Games, Public Speaking, Crafts, Verbal Fluency (7 "
        "rows).",
        "TraitRatings — Excellent (5), Very Good (4), Good (3), Fair "
        "(2), Poor (1).",
    ])
    add_para(doc,
        "Same .IgnoreQueryFilters().AnyAsync() guard pattern as earlier "
        "sprints — the seeder only inserts if the table is empty "
        "(counting soft-deleted rows), so running the app a second time "
        "after a soft delete does not resurrect the row.")

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_para(doc,
        "Four new pages land in src/NaijaPrimeSchool.Web/Components/"
        "Pages/Results/. The pattern mirrors earlier sprints: a list "
        "page (with optional inline form), a detail page (for the score "
        "sheet or the report card detail), a per-domain action page "
        "(Results page, ReportCards page).")

    add_heading(doc, "9.1 Page roster", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Results/\n"
        "├── Assessments.razor          <- /assessments\n"
        "├── AssessmentScores.razor     <- /assessments/{id}/scores\n"
        "├── Results.razor              <- /results\n"
        "├── ReportCards.razor          <- /reports\n"
        "└── ReportCardDetail.razor     <- /reports/{id}\n")
    add_para(doc,
        "Assessment pages are gated to SuperAdmin + HeadTeacher + "
        "Teacher (a class teacher needs to enter scores). Results and "
        "ReportCards are SuperAdmin + HeadTeacher only — those are "
        "school-wide computations that we do not want a single class "
        "teacher kicking off accidentally.")

    add_heading(doc, "9.2 Assessments.razor — the gradebook list", 2)
    add_para(doc,
        "Search by session/term/class/subject, paged grid showing each "
        "assessment with its scored count, status badge, and a row of "
        "actions (open score sheet, edit, publish/unpublish, delete). "
        "The new-assessment form is inline below the grid; once an "
        "assessment exists, its term/class/subject are locked because "
        "moving an assessment to a different (class, subject) would "
        "invalidate scores.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Results/Assessments.razor")

    add_heading(doc, "9.3 AssessmentScores.razor — the score sheet", 2)
    add_para(doc,
        "The score-entry workhorse. Renders an HTML table "
        "(.nps-score-grid) — one row per actively-enrolled pupil, "
        "columns for score, absent, and remarks. Score numerics are "
        "bounded by the assessment's MaxScore. Marking absent clears "
        "the score; entering a score clears absent. A single Save "
        "Scores button posts the whole sheet via BulkSetScoresAsync. "
        "Once the assessment is published, every input on the page is "
        "disabled.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Results/AssessmentScores.razor")

    add_heading(doc, "9.4 Results.razor — compute and view", 2)
    add_para(doc,
        "Pick a term and class, optionally narrow to a subject, then "
        "either Recompute (idempotent and safe) or Compute & finalise "
        "(locks the rows). The grid renders the SubjectResults with "
        "grade badges, position counters, and per-row "
        "finalise/reopen/delete actions. The colour-coded grade badges "
        "use the Radzen BadgeStyle palette so A/B render green/blue, "
        "C/D primary/yellow, E/F warning/red.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Results/Results.razor")

    add_heading(doc, "9.5 ReportCards.razor — generate and list", 2)
    add_para(doc,
        "Pick a (term, class), see every report card already generated "
        "for that pair (with the average, position, attendance summary, "
        "and publishing status), or click Generate / refresh to roll a "
        "fresh batch. The generate dialog is inline below the grid and "
        "lets the head teacher set 'next term begins' once for the "
        "whole batch.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCards.razor")

    add_heading(doc, "9.6 ReportCardDetail.razor — the long form", 2)
    add_para(doc,
        "The most substantial page in the sprint. Header summary cards "
        "show subjects taken, average, position, days present. Below "
        "that, a Radzen tab strip splits the page into Subjects "
        "(read-only result table), Affective traits, Psychomotor "
        "skills, and Comments. The two ratings tabs render an HTML "
        "table per category; choosing a rating from the dropdown auto-"
        "saves via UpsertAffectiveRatingAsync / "
        "UpsertPsychomotorRatingAsync. The Comments tab is a small "
        "form with class-teacher comment, head-teacher comment, and "
        "the next-term date.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCardDetail.razor")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and authorization", 1)
    add_heading(doc, "10.1 NavMenu — a new Results & Reports panel", 2)
    add_para(doc,
        "NavMenu.razor picks up a fifth role-gated panel between "
        "Attendance and the Finance/Inventory placeholders. Three "
        "entries live inside it: Assessments, Subject results, Report "
        "cards. The panel is wrapped in an AuthorizeView that requires "
        "SuperAdmin, HeadTeacher, or Teacher.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Results & Reports\"",
        end_marker="</AuthorizeView>",
        caption="Excerpt — NavMenu.razor")

    add_heading(doc, "10.2 _Imports.razor", 2)
    add_para(doc,
        "_Imports.razor picks up two new @using lines so the results "
        "DTOs and service interfaces are visible to every Razor file in "
        "the Web project.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Results",
        lines_after_start=2,
        caption="Excerpt — _Imports.razor")

    add_heading(doc, "10.3 Authorization at the page level", 2)
    add_para(doc,
        "Every sprint-5 page declares an [Authorize] attribute. "
        "Assessment pages require SuperAdmin + HeadTeacher + Teacher; "
        "result and report-card pages require SuperAdmin + HeadTeacher. "
        "There is no policy added — the existing role attribute is "
        "enough. Future sprints (parent portal, student portal) will "
        "introduce a Pastoral or Family policy that exposes published "
        "report cards to parents.")

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of a results pipeline run", 1)
    add_para(doc,
        "Walking a single (term, class) from blank gradebook to "
        "published report cards is the clearest way to see how all the "
        "layers cooperate.")

    add_heading(doc, "11.1 Setting up the gradebook", 2)
    add_bullets(doc, [
        "HeadTeacher opens /assessments. Filter to the current term and "
        "the target class.",
        "Click New assessment. Pick the subject, pick CA1 from the "
        "type dropdown, set max score 20 and weight 1, click Save.",
        "Repeat for CA2 (max 20, weight 1) and Examination (max 60, "
        "weight 2). The grid now lists three draft assessments for the "
        "subject.",
    ])

    add_heading(doc, "11.2 Entering scores", 2)
    add_bullets(doc, [
        "Open the score sheet for CA1. Every actively-enrolled pupil is "
        "pre-listed. Key in scores; mark a pupil absent if they did "
        "not sit the test.",
        "Click Save scores. BulkSetScoresAsync writes the rows in one "
        "round-trip; ApplicationDbContext.SaveChanges stamps "
        "CreatedOn/By on each new row.",
        "Repeat for CA2 and Examination.",
        "Optionally Publish each assessment when scores are final. "
        "Published assessments are read-only.",
    ])

    add_heading(doc, "11.3 Computing subject results", 2)
    add_bullets(doc, [
        "Open /results, pick the same term and class.",
        "Click Recompute. ResultService.ComputeAsync produces one "
        "SubjectResult per (pupil, subject) and one warning per "
        "skipped finalised row, if any.",
        "The grid now shows percentages, grade badges, and positions "
        "out of the cohort size.",
        "When the head teacher is satisfied, click Compute & finalise. "
        "The same compute runs and rows are stamped FinalisedOn = now. "
        "Further recomputes will not touch them unless reopened.",
    ])

    add_heading(doc, "11.4 Composing report cards", 2)
    add_bullets(doc, [
        "Open /reports, click Generate / refresh, pick the same (term, "
        "class), set the next-term-begins date, click Generate.",
        "ReportCardService.GenerateAsync rolls every SubjectResult for "
        "the (term, class), joins to the sprint-4 attendance counts, "
        "and ranks pupils by average percentage.",
        "Click any row. The detail page tabs split profile / subjects "
        "/ ratings / comments.",
        "Pick affective and psychomotor ratings; type the class "
        "teacher's comment; type the head teacher's comment.",
        "Publish. Card is now read-only and ready for the parent "
        "portal sprint to surface.",
    ])

    add_heading(doc, "11.5 Correcting an error after publishing", 2)
    add_bullets(doc, [
        "A teacher notices a score was wrong. Open the report card, "
        "click Unpublish.",
        "Open /results, find the affected SubjectResult, click Reopen.",
        "Open the underlying assessment, unpublish it, fix the score, "
        "save.",
        "Re-publish the assessment, run Recompute on the results page. "
        "The single row updates.",
        "Run Generate / refresh on /reports — the existing card row "
        "refreshes (it was unpublished, so the generator does write to "
        "it).",
        "Re-publish the card. Audit columns (ModifiedOn/By) reflect "
        "the chain of edits.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)
    add_para(doc,
        "Once the build is green and the migration has applied, this is "
        "the end-to-end smoke test for a fresh checkout.")

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "First run applies migrations and seeds the five new lookup "
        "tables. Sign in as superadmin@naijaprimeschool.ng / Admin@12345.")

    add_heading(doc, "12.2 Verify navigation", 2)
    add_bullets(doc, [
        "The Results & Reports panel appears in the sidebar, between "
        "Attendance and the Finance placeholder.",
        "It contains three items: Assessments, Subject results, Report "
        "cards.",
        "Each one renders without error.",
    ])

    add_heading(doc, "12.3 Verify the seeded lookups", 2)
    add_para(doc, "Connect to SQL Server and run:")
    add_code(doc,
        "SELECT Name, Code, DefaultMaxScore, IsExam FROM AssessmentTypes ORDER BY DisplayOrder;\n"
        "SELECT Name, LowerBound, UpperBound, Description FROM GradeBands ORDER BY DisplayOrder;\n"
        "SELECT Name FROM AffectiveTraits ORDER BY DisplayOrder;\n"
        "SELECT Name FROM PsychomotorSkills ORDER BY DisplayOrder;\n"
        "SELECT Name, Value FROM TraitRatings ORDER BY DisplayOrder;\n")

    add_heading(doc, "12.4 Create assessments, score them, compute results", 2)
    add_numbered(doc, [
        "Pre-requisites — at least one Term, one SchoolClass with "
        "actively enrolled pupils, and at least one Subject. (Sprints "
        "2–3 cover the setup.)",
        "Create three assessments (CA1, CA2, Exam) for one subject in "
        "the class.",
        "Open each score sheet, key scores for every pupil, save.",
        "Open /results, pick the term and class, click Recompute. "
        "Verify the percentages and positions look right.",
        "Click Compute & finalise on the same row. Verify the rows "
        "flip to the Finalised badge.",
    ])

    add_heading(doc, "12.5 Generate and publish report cards", 2)
    add_numbered(doc, [
        "Open /reports, click Generate / refresh, pick the same (term, "
        "class), generate.",
        "Open one card. Verify the subject totals match what /results "
        "showed. Verify the attendance counts match a quick spot-check "
        "against the sprint-4 register.",
        "Pick affective and psychomotor ratings; type teacher and head-"
        "teacher comments.",
        "Click Publish. Verify the action buttons disable and the "
        "status badge flips to Published.",
    ])

    add_heading(doc, "12.6 Verify error paths", 2)
    add_numbered(doc, [
        "Try to delete a published assessment. The OperationResult "
        "comes back with 'Unpublish the assessment before deleting.'",
        "Try to delete a finalised SubjectResult. The OperationResult "
        "comes back with 'Cannot delete a finalised result. Reopen it "
        "first.'",
        "Try to delete a published ReportCard. The OperationResult "
        "comes back with 'Unpublish the card before deleting.'",
        "Try to enter a score larger than the assessment's MaxScore. "
        "The score endpoint rejects it with the bounded-range message.",
    ])

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)
    add_heading(doc, "13.1 'No assessments exist for this term/class'", 2)
    add_para(doc,
        "Recompute fails with this message when the (term, class) has "
        "zero TermAssessment rows. Create assessments under "
        "/assessments first.")

    add_heading(doc, "13.2 'No subject results exist for this term/class'", 2)
    add_para(doc,
        "Generate report cards fails when results have not been "
        "computed yet. Open /results, pick the same (term, class), "
        "click Recompute, then come back.")

    add_heading(doc, "13.3 The score sheet shows pupils I did not enrol", 2)
    add_para(doc,
        "GetScoreSheetAsync pre-lists every pupil with an active "
        "enrolment in the assessment's class. If a pupil who left "
        "still appears, their enrolment row is still open. Withdraw "
        "the enrolment under Family → Students → pupil → Enrolment "
        "history, then reload the score sheet.")

    add_heading(doc, "13.4 Position is the same for two pupils", 2)
    add_para(doc,
        "Dense ranking gives ties the same position number. If two "
        "pupils have 78.50% they are both '5 of 38'. The next pupil is "
        "'6 of 38' (not 7). This is intentional and matches the way "
        "most schools rank. If you need 1224-style ranking, the "
        "ResultService.ComputeAsync method is the place to change it.")

    add_heading(doc, "13.5 'Score must be between 0 and N'", 2)
    add_para(doc,
        "Every score is bounded by the assessment's MaxScore. If "
        "MaxScore is 20, valid scores are 0..20 inclusive. Rounding "
        "(20.0001 = 20) is not done; either edit the MaxScore or "
        "lower the entered score.")

    add_heading(doc, "13.6 'Already finalised — skipped recompute'", 2)
    add_para(doc,
        "ResultService.ComputeAsync warns rather than fails when it "
        "encounters a finalised SubjectResult during a non-finalising "
        "recompute. The warning identifies the (student, subject) so "
        "the user can reopen the right row before recomputing.")

    add_heading(doc, "13.7 The IdentityRole migration warning", 2)
    add_para(doc,
        "The EF Core warning about ApplicationRole's query filter and "
        "IdentityUserRole's required end is harmless at runtime and "
        "predates sprint 5. The build is still green.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_para(doc,
        "Sprint 5 has left a few breadcrumbs that make later sprints "
        "easier.")
    add_bullets(doc, [
        "ReportCard.IsPublished is the load-bearing flag the parent "
        "and student portals will check. Until that sprint, only "
        "internal staff can see cards.",
        "SubjectResult.IsFinalised distinguishes draft from authoritative "
        "rows, so a future analytics dashboard can choose to show only "
        "finalised data.",
        "GradeBand.LowerBound/UpperBound use decimal(5,2). Adding a new "
        "band (e.g. A+ for 90+) is a one-row insert plus a re-seed "
        "guard, no schema change.",
        "AssessmentType.IsExam is reserved for end-of-session "
        "promotion logic — 'pupil passes the term if their exam mean "
        "is at least 50%' will read this column.",
        "TermAssessment.Weight defaults to 1m, so existing schemes "
        "remain valid as the school evolves its grading policy.",
        "ReportCard.NextTermBegins is a date the parent portal will "
        "show next to the published card so families know when the "
        "child returns.",
    ])

    add_heading(doc, "14.1 What might need a small refactor later", 2)
    add_bullets(doc, [
        "ResultService.ComputeAsync issues four queries (assessments, "
        "scores, enrolments, grade bands) before its in-memory loop. "
        "For a school of 5,000 pupils this stays fine; if a school "
        "scales to 20,000+ we may want to push the loop into a SQL "
        "stored procedure or a CTE-driven single query.",
        "BulkSetScoresAsync currently writes one row per Add. EF Core "
        "10's bulk operations would be faster but the API surface "
        "changes. Worth doing the day someone reports a slow save.",
        "Report cards are persisted with the 'next term begins' date "
        "as a flat column. If that policy ever varies per pupil "
        "(unlikely) we'd promote it to its own table; for now the "
        "denormalisation matches every Nigerian school I have built "
        "for.",
        "The HTML score-grid is plain HTML rather than RadzenDataGrid. "
        "If a future feature wants column sorting on this grid, we "
        "will rewrite it as a Radzen grid with a custom edit template; "
        "the back-end DTOs are already in the right shape.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 5", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Results/AssessmentType.cs",     "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Results/GradeBand.cs",          "Lookup with bounds + remarks."),
        ("src/NaijaPrimeSchool.Domain/Results/AffectiveTrait.cs",     "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Results/PsychomotorSkill.cs",   "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Results/TraitRating.cs",        "Lookup, 1–5."),
        ("src/NaijaPrimeSchool.Domain/Results/TermAssessment.cs",     "Gradebook entry."),
        ("src/NaijaPrimeSchool.Domain/Results/AssessmentScore.cs",    "Per-pupil score."),
        ("src/NaijaPrimeSchool.Domain/Results/SubjectResult.cs",      "Per (pupil, term, subject) total."),
        ("src/NaijaPrimeSchool.Domain/Results/ReportCard.cs",         "Per (pupil, term) summary."),
        ("src/NaijaPrimeSchool.Domain/Results/AffectiveRating.cs",    "Card x trait rating."),
        ("src/NaijaPrimeSchool.Domain/Results/PsychomotorRating.cs",  "Card x skill rating."),
        ("Domain layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Domain/Academics/Subject.cs",     "Added TermAssessments + SubjectResults."),
        ("src/NaijaPrimeSchool.Domain/Academics/Term.cs",        "Added TermAssessments + SubjectResults + ReportCards."),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs", "Added TermAssessments + SubjectResults + ReportCards."),
        ("src/NaijaPrimeSchool.Domain/Family/Student.cs",        "Added AssessmentScores + SubjectResults + ReportCards."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/TermAssessmentDtos.cs",   "TermAssessment DTOs."),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/AssessmentScoreDtos.cs",  "Score-sheet DTOs."),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/SubjectResultDtos.cs",    "Subject result DTOs + compute request/response."),
        ("src/NaijaPrimeSchool.Application/Results/Dtos/ReportCardDtos.cs",       "Report card DTOs + ratings + generate request."),
        ("src/NaijaPrimeSchool.Application/Results/IAssessmentService.cs",        "Assessment + score service contract."),
        ("src/NaijaPrimeSchool.Application/Results/IResultService.cs",            "Result service contract."),
        ("src/NaijaPrimeSchool.Application/Results/IReportCardService.cs",        "Report card service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",              "Added 5 new lookup methods."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/AssessmentService.cs",     "Assessment CRUD + score-sheet methods."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/ResultService.cs",         "Compute, finalise, reopen."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/ReportCardService.cs",     "Generate, comments, ratings, publish."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 11 tables and indexes."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",            "Registered the 3 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs", "Added 11 DbSets, ConfigureResults."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs", "Seeded the 5 new lookup tables."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",         "Added 5 new lookup methods."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/Assessments.razor",        "Gradebook list + inline form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/AssessmentScores.razor",   "Score sheet."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/Results.razor",            "Compute, view, finalise."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCards.razor",        "Generate + list."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Results/ReportCardDetail.razor",   "Tabs: subjects / traits / skills / comments."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                          "Added Results + Results.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                    "Added the Results & Reports panel."),
        ("src/NaijaPrimeSchool.Web/wwwroot/app.css",                                    "Added .nps-score-grid and .nps-trait-grid styles."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint5_guide.py",                                             "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 5 implementation guide. The next sprint "
        "lands fees and bursar workflows on top of the (Pupil x Term) "
        "primitives established here.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
