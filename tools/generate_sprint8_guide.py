"""Generates 'Sprint 8 - Implementation Guide.docx' covering the parent
and student portals, the announcements / communications module, and
the role-based plumbing that ties them to the existing identity, finance
and academic domains.

Long-form edition. Code blocks embed actual source files so the guide
stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint8_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 8 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260518164808_Communications.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 8 — Parent & Student Portals + Announcements")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Parent dashboard · Student dashboard · Announcements · Read tracking")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/8-portals")
    meta.add_run("\nBuilt on: Sprints 1–7 (identity, academic domain, students & parents, "
                 "attendance, results & report cards, pupil photos, fees & bursar workflows, "
                 "store & inventory)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 8 in context", 1)
    add_para(doc,
        "Sprint 8 is the sprint where the school finally talks to its "
        "families. The seven preceding sprints built a complete back "
        "office: identity and roles, the academic calendar, students "
        "and parents and their linkage, daily and per-subject "
        "attendance, a results pipeline that produces published "
        "report cards, pupil photos, a bursar workspace for fees and "
        "receipts, and a storekeeper workspace for inventory. Every "
        "one of those workflows ends at the school office. Sprint 8 "
        "turns the system outward.")
    add_para(doc,
        "Two role groups have been waiting their turn since sprint 1: "
        "Parent and Student. Both have lived as seeded ApplicationRoles "
        "with the navigation menu showing a single disabled "
        "'My Children' placeholder. This sprint replaces that "
        "placeholder with a complete parent portal — a card per ward "
        "linking to the pupil's results, attendance, and fee history — "
        "and a complete student portal — today's timetable, latest "
        "results, attendance summary, and outstanding fees, all read-"
        "only and scoped to the signed-in pupil.")
    add_para(doc,
        "A second feature ships in the same sprint because it is what "
        "actually closes the loop. Announcements give the head teacher "
        "and the super admin a way to broadcast notices — 'mid-term "
        "break starts next Friday', 'school fees due before the third "
        "Monday', 'PTA meeting on Saturday' — to the parent and "
        "student portals. Announcements are scoped (Everyone, Parents, "
        "Students, or a specific Class) and they track who has read "
        "what so the dashboards can surface a meaningful 'unread' "
        "count.")
    add_para(doc,
        "Once this sprint ships, every seeded role in the system has "
        "a workspace: SuperAdmin and HeadTeacher run the school, "
        "Teachers mark attendance and capture results, SchoolBursar "
        "issues invoices and receipts, SchoolStoreKeeper manages the "
        "store, Parents and Students view their data. The naive 'open "
        "the seven dashboards in turn' tour now covers everyone.")
    add_para(doc,
        "This document is a long-form implementation guide written in "
        "the tone of the sprint 7 guide. An engineer who has read the "
        "earlier guides and has the codebase checked out can recreate "
        "every change here without referring to the diff.")

    add_heading(doc, "1.1 Where this sits relative to sprint 7", 2)
    add_para(doc, "Every load-bearing piece of the earlier sprints is reused:")
    add_bullets(doc, [
        "BaseEntity — Announcement and AnnouncementRead inherit Guid Id, "
        "IAuditable, ISoftDelete from it.",
        "ApplicationDbContext.SaveChanges — the override stamps audit "
        "columns and rewrites Delete to IsDeleted = true. Soft deletion "
        "of an announcement is the only delete path the UI exposes.",
        "Global query filters — !IsDeleted on Announcement and on "
        "AnnouncementRead means hidden notices disappear from every "
        "ordinary query.",
        "OperationResult / OperationResult<T> — the announcement "
        "service uses both shapes for predictable success / failure "
        "responses.",
        "ILookupService — already had twenty-nine methods. Sprint 8 "
        "adds two more (announcement categories, audiences) without "
        "rewriting the existing ones.",
        "IInvoiceService.GetStudentLedgerAsync (sprint 6) — the "
        "portal services lean on this for the outstanding balance, "
        "lifetime invoiced and lifetime paid figures.",
        "IReportCardService.ListAsync (sprint 5) — the student and "
        "ward views ask for published report cards filtered by pupil.",
        "ApplicationUser.IsActive / role membership — the portal "
        "services resolve a Parent or Student row by following "
        "Parent.UserId / Student.UserId and the access guard cross-"
        "checks role and StudentParent linkage.",
        "ICurrentUser — gains a Roles list and an IsInRole helper so "
        "the announcement service can decide which audiences apply "
        "to the currently-signed-in user without a second round-trip "
        "through the auth state.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_numbered(doc, [
        "Maintain an announcement category lookup table — General, "
        "Academic, Finance, Events, Holiday, Health, Emergency — "
        "seeded on first run with stable Codes.",
        "Maintain an announcement audience lookup table — Everyone, "
        "Parents, Students, Specific Class — also seeded. Specific "
        "Class is the only audience that requires the Announcement "
        "row to carry a TargetSchoolClassId; the service enforces "
        "the rule.",
        "Compose an announcement: title, body, category, audience, "
        "optional target class, optional expiry date, optional "
        "pinned-to-top flag. Save as draft or publish immediately.",
        "Publish / unpublish / soft-delete an announcement. Draft "
        "rows are visible only to the admin who composed them. "
        "Published rows that have expired are still visible to the "
        "admin under the 'Include expired' filter but disappear from "
        "the parent and student feeds.",
        "Read tracking — one AnnouncementRead row per (announcement, "
        "user). The composite unique index prevents duplicates if a "
        "user clicks Mark as read twice. The dashboards count unread "
        "announcements per signed-in user and surface the figure on a "
        "stat tile.",
        "Parent dashboard — one card per ward, showing the linkage "
        "(relationship + primary-contact flag), the current school "
        "class, the outstanding balance, the latest attendance "
        "percentage, and the count of published report cards.",
        "Ward detail view — a tabbed view per pupil with Overview, "
        "Report cards (linking to the existing PDF download), and "
        "Fees & invoices (with the same StudentLedger the bursar "
        "uses, filtered to the pupil).",
        "Student dashboard — today's timetable resolved from the "
        "pupil's current class and the current term's weekly grid, "
        "plus stat tiles for class, outstanding fees, attendance "
        "percentage, published report cards, and unread "
        "announcements. Shortcuts to the four sub-pages.",
        "Student sub-pages — read-only profile, results list, "
        "attendance summary, and fees ledger.",
        "Parent / student announcements feed — only published, "
        "non-expired notices targeted at the user's audience (or at "
        "a class one of their wards is enrolled in). Mark-as-read "
        "is one click.",
    ])

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_bullets(doc, [
        "Comments on announcements. The current model is one-way; a "
        "future AnnouncementComment table would slot in without "
        "redesigning anything else.",
        "Push notifications and emails. The portal surfaces notices "
        "but does not push them. A future job would scan unread "
        "AnnouncementRead gaps and emit emails or SMS via a separate "
        "notification service.",
        "Attachments on announcements. Body is plain text today. A "
        "future AnnouncementAttachment table sitting alongside "
        "Announcement would be additive.",
        "Parent-to-school messaging. The portal is read-only on "
        "purpose for sprint 8 — every existing workflow assumes the "
        "school office is the source of truth. A two-way messaging "
        "sprint would introduce a Conversation / Message pair.",
        "Online fee payment. The student fees page shows the "
        "outstanding balance and invoice list but does not collect "
        "money. A payment gateway sprint would graft onto the "
        "existing PaymentService.",
        "Mobile apps. The portal is a Blazor Web App and renders "
        "responsively in mobile browsers; a native wrapper or a "
        "dedicated MAUI front-end is its own track of work.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers:")
    add_bullets(doc, [
        "4 new domain entities under src/NaijaPrimeSchool.Domain/Communications/.",
        "2 DTO files under src/NaijaPrimeSchool.Application/Communications/Dtos/ "
        "and src/NaijaPrimeSchool.Application/Portals/Dtos/.",
        "2 new service contracts (IAnnouncementService, IPortalService).",
        "2 service implementations (AnnouncementService, PortalService).",
        "2 new methods on ILookupService and the matching LookupService.",
        "1 EF Core migration introducing 4 new tables and the indexes "
        "that go with them.",
        "1 DatabaseInitializer extension seeding AnnouncementCategories "
        "and AnnouncementAudiences.",
        "2 Razor pages for the head-teacher / super-admin under "
        "src/NaijaPrimeSchool.Web/Components/Pages/Communications/.",
        "8 Razor pages for the portals under "
        "src/NaijaPrimeSchool.Web/Components/Pages/Portals/.",
        "1 NavMenu rewrite: the disabled 'My Children' placeholder is "
        "replaced with a Communications panel (admin) and dedicated "
        "Parent and Student panels.",
        "1 CSS appendix in wwwroot/app.css for the new ward grid, "
        "portal row, read-only grid, and announcement cards.",
    ])
    add_para(doc,
        "The code follows the patterns already accepted in sprints 1–7.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)

    add_heading(doc, "2.1 Two lookup tables for announcements", 2)
    add_para(doc,
        "The 'don't use enums' rule from earlier sprints holds. "
        "AnnouncementCategory and AnnouncementAudience are proper "
        "entities derived from BaseEntity, seeded on first run, and "
        "editable from the database without a redeploy. AudienceAudience "
        "carries a RequiresTargetClass column that the service consults "
        "when validating an announcement payload — the only audience "
        "with that flag set is 'Specific Class', so a future 'Specific "
        "Year Group' could be added without changing service code.")

    add_heading(doc, "2.2 Read tracking is one row per (announcement, user)", 2)
    add_para(doc,
        "AnnouncementRead has a composite unique index on "
        "(AnnouncementId, UserId). The MarkAsReadAsync method checks "
        "for an existing row first and short-circuits if one is "
        "present, so a parent who clicks Mark as read twice in quick "
        "succession does not crash on the unique constraint. The read "
        "row carries its own audit columns inherited from BaseEntity, "
        "so a future 'who has not read this notice?' report can lean "
        "on CreatedOn.")

    add_heading(doc, "2.3 Audience resolution lives in the service, not the page", 2)
    add_para(doc,
        "The parent and student announcements feed asks "
        "AnnouncementService.ListForCurrentUserAsync. The service "
        "resolves which audience codes apply (Everyone always; "
        "Parent for users in the Parent role; Student for users in "
        "the Student role; SuperAdmin and HeadTeacher see every "
        "code so the preview matches reality) and which class ids "
        "are relevant (a parent sees their wards' classes; a student "
        "sees their own current class). The page just renders the "
        "result. That keeps the UI Razor file thin and lets a future "
        "API consumer call the same method.")

    add_heading(doc, "2.4 PortalService is a façade over existing services", 2)
    add_para(doc,
        "Neither the parent nor the student dashboard introduces "
        "new persistence. The two dashboards re-use "
        "IInvoiceService.GetStudentLedgerAsync (sprint 6) for the "
        "balance figures, the report card table (sprint 5) for the "
        "published-card count, the daily attendance entries (sprint "
        "4) for the term attendance percentage, and the timetable "
        "(sprint 2) for today's lessons. PortalService composes "
        "these into the two dashboard DTOs and applies the access "
        "guard. That gives sprint 8 a small surface area: when one "
        "of the underlying services changes, the portal "
        "automatically reflects the change.")

    add_heading(doc, "2.5 Resolution from a login to a Parent or Student row", 2)
    add_para(doc,
        "The school office links a parent profile (Family.Parent) "
        "to an Identity user (ApplicationUser) by setting "
        "Parent.UserId at admission. The same is true for students. "
        "PortalService.ResolveParentIdForCurrentUserAsync simply "
        "looks for the Parent row whose UserId matches the signed-in "
        "user's Id. If no link exists, the dashboard renders a "
        "friendly 'we can't find your parent record — ask the school "
        "office to link your account' card. This is intentionally a "
        "soft failure: a Parent role user without a linked Parent "
        "row is not a security incident, just an incomplete admission.")

    add_heading(doc, "2.6 Access guard on the ward detail page", 2)
    add_para(doc,
        "The /portal/parent/wards/{StudentId} route is reachable by "
        "any user in the Parent role, but the guard "
        "(CurrentUserCanViewStudentAsync) confirms that the signed-"
        "in user is linked via Family.StudentParent to the pupil "
        "before any data is loaded. SuperAdmin and HeadTeacher pass "
        "the guard unconditionally so the same page can be used as "
        "an admin preview during smoke testing.")

    add_heading(doc, "2.7 Authorisation by role, not by claim", 2)
    add_para(doc,
        "Every Razor page in sprint 8 uses @attribute "
        "[Authorize(Roles = ...)] — the same pattern as the rest of "
        "the project. The role names come from the Roles static "
        "class seeded in sprint 1, which means a typo in a role "
        "string is a compile error. The portals are gated to Parent "
        "or Student plus the two admin roles; the announcement "
        "admin pages are gated to SuperAdmin or HeadTeacher; the "
        "portal announcements feed is gated to Parent + Student + "
        "the two admin roles so the same page works as an admin "
        "preview.")

    add_heading(doc, "2.8 Soft delete plus reversible publish", 2)
    add_bullets(doc, [
        "Announcement soft-delete is the only delete path the UI "
        "exposes. The audit trail is preserved.",
        "Unpublish flips IsPublished = false without clearing "
        "PublishedOn, so 'when was this last visible to families?' "
        "is still answerable.",
        "Publishing a draft (IsPublished = true with PublishedOn = "
        "UtcNow) is what makes a row visible to parents and students. "
        "A drafted-but-never-published row never reaches the portal.",
        "Expiry is a soft cut-off: ExpiresOn < today drops the row "
        "from parent and student feeds but the row is still visible "
        "to admins under Include expired.",
    ])

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)

    add_heading(doc, "3.1 Folder layout", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "└── Communications/                <- (new in sprint 8)\n"
        "    ├── AnnouncementCategory.cs    <- lookup\n"
        "    ├── AnnouncementAudience.cs    <- lookup with RequiresTargetClass\n"
        "    ├── Announcement.cs            <- the notice itself\n"
        "    └── AnnouncementRead.cs        <- per-user read row\n")

    add_heading(doc, "3.2 The two lookup entities", 2)
    add_heading(doc, "3.2.1 AnnouncementCategory.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Communications/AnnouncementCategory.cs")
    add_heading(doc, "3.2.2 AnnouncementAudience.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Communications/AnnouncementAudience.cs")
    add_para(doc,
        "RequiresTargetClass is the only column that splits one audience "
        "from another at the service level. 'Specific Class' carries "
        "true; the three broadcast audiences carry false.")

    add_heading(doc, "3.3 The two core entities", 2)
    add_heading(doc, "3.3.1 Announcement.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Communications/Announcement.cs")
    add_para(doc,
        "Title is capped at 200 characters; Body at 4,000. "
        "TargetSchoolClassId is nullable and only used by the "
        "Specific Class audience. PostedById links back to the "
        "ApplicationUser who composed the row so the parent and "
        "student feeds can show 'posted by'. PublishedOn is captured "
        "the first time IsPublished flips to true; it survives a "
        "later unpublish.")
    add_heading(doc, "3.3.2 AnnouncementRead.cs", 3)
    add_file(doc, "src/NaijaPrimeSchool.Domain/Communications/AnnouncementRead.cs")
    add_para(doc,
        "ReadOn is captured at insert time. The composite unique on "
        "(AnnouncementId, UserId) lives in the DbContext configuration "
        "(see chapter 5).")

    add_heading(doc, "3.4 Relationships at a glance", 2)
    add_code(doc,
        "  AnnouncementCategory    AnnouncementAudience\n"
        "         \\                  /\n"
        "          \\                /\n"
        "           v 1..N         v 1..N\n"
        "             Announcement  ----+----> SchoolClass (TargetSchoolClassId)\n"
        "                  |            +----> ApplicationUser (PostedById)\n"
        "                  v 1..N\n"
        "             AnnouncementRead\n"
        "                  |\n"
        "                  +----> ApplicationUser (UserId)\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)

    add_heading(doc, "4.1 AnnouncementDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Communications/Dtos/AnnouncementDtos.cs")
    add_para(doc,
        "AnnouncementDto carries enough denormalised metadata "
        "(category code, audience code, target class name, posted-by "
        "name) for the admin list and the portal feed to render "
        "without a second round-trip. The IsExpired computed property "
        "lets the admin grid badge expired rows without redoing the "
        "comparison. ReadCount and ReadByCurrentUser support both "
        "the admin 'how many have read this?' column and the portal "
        "'New' badge.")

    add_heading(doc, "4.2 PortalDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Portals/Dtos/PortalDtos.cs")
    add_para(doc,
        "ParentDashboardDto carries one WardSummaryDto per linked "
        "pupil. StudentDashboardDto carries the term-scoped attendance "
        "tally (DaysPresent / DaysAbsent / DaysLate / DaysCounted) "
        "plus the percentage, so the attendance sub-page can render "
        "five stat tiles without a second call. UpcomingLessonDto is "
        "the shape the student dashboard's 'today's timetable' grid "
        "binds to.")

    add_heading(doc, "4.3 Service contracts", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Communications/IAnnouncementService.cs")
    add_file(doc, "src/NaijaPrimeSchool.Application/Portals/IPortalService.cs")

    add_heading(doc, "4.4 ICurrentUser gains a role API", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Common/ICurrentUser.cs")
    add_para(doc,
        "The announcement service needs to know which roles the "
        "signed-in user belongs to so it can compute the right "
        "audience codes. Adding Roles and IsInRole keeps that "
        "knowledge out of the service-layer (no direct dependency "
        "on ClaimsPrincipal). The web-layer implementation is in "
        "CurrentUserAccessor (chapter 6).")

    add_heading(doc, "4.5 ILookupService extension", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Application/Users/ILookupService.cs",
        "Task<IReadOnlyList<LookupDto>> GetAnnouncementCategoriesAsync",
        lines_after_start=3,
        caption="Excerpt — ILookupService.cs (sprint 8 additions)")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)

    add_heading(doc, "5.1 Four new DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<AnnouncementCategory>",
        end_marker="public DbSet<AnnouncementRead>",
        lines_after_start=5,
        caption="Excerpt — the four communications DbSets")

    add_heading(doc, "5.2 ConfigureCommunications", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureCommunications",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureCommunications")
    add_para(doc, "Highlights:")
    add_bullets(doc, [
        "ConfigureLookup<AnnouncementCategory> and "
        "ConfigureLookup<AnnouncementAudience> reuse the same helper "
        "that has set up every other lookup table since sprint 1 — "
        "BaseEntity columns, DisplayOrder, soft-delete query filter.",
        "Unique on each lookup's Name and Code — four unique indexes "
        "in total.",
        "Restrict on Announcement.AnnouncementCategoryId and "
        "AnnouncementAudienceId — a category or audience cannot be "
        "deleted while it has announcements pointing at it.",
        "SetNull on Announcement.TargetSchoolClassId and "
        "Announcement.PostedById — a class or a user can be soft-"
        "deleted without vaporising the historical announcement.",
        "Cascade on AnnouncementRead.AnnouncementId and "
        "AnnouncementRead.UserId — soft-deleting an announcement or "
        "removing a user should not leave dangling read rows.",
        "Composite unique on AnnouncementRead.(AnnouncementId, UserId) "
        "— guarantees one read row per user per notice.",
        "Indexes on PublishedOn, ExpiresOn, IsPublished — the portal "
        "feed query filters by all three.",
        "HasQueryFilter(!IsDeleted) on both Announcement and "
        "AnnouncementRead — soft-deleted rows disappear from every "
        "ordinary query.",
    ])

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)

    add_heading(doc, "6.1 AnnouncementService.cs", 2)
    add_para(doc,
        "AnnouncementService owns CRUD plus the two portal-facing "
        "helpers. ListForCurrentUserAsync resolves the right audience "
        "codes and the right class ids and filters published, non-"
        "expired rows accordingly. CountUnreadForCurrentUserAsync "
        "runs the same visibility query but counts rows that are not "
        "in AnnouncementReads for the current user.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/AnnouncementService.cs")

    add_heading(doc, "6.2 PortalService.cs", 2)
    add_para(doc,
        "PortalService is the thin façade described in chapter 2. "
        "GetParentDashboardAsync walks the parent's StudentLinks and "
        "calls a private BuildWardSummaryAsync for each. "
        "GetStudentDashboardAsync resolves the current enrolment, "
        "the current term, the term-scoped attendance tally and "
        "today's timetable, then assembles the StudentDashboardDto.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/PortalService.cs")

    add_heading(doc, "6.3 CurrentUserAccessor — role plumbing", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Services/CurrentUserAccessor.cs")

    add_heading(doc, "6.4 DI registration", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IAnnouncementService",
        lines_after_start=3,
        caption="Excerpt — DependencyInjection.cs (sprint 8 additions)")

    add_heading(doc, "6.5 LookupService — two new methods", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",
        "public async Task<IReadOnlyList<LookupDto>> GetAnnouncementCategoriesAsync",
        lines_after_start=14,
        caption="Excerpt — LookupService.cs (sprint 8 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_code(doc,
        "dotnet ef migrations add Communications \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the communications lookups", 1)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedCommunicationsLookupsAsync",
        end_marker="private static async Task SeedFinanceLookupsAsync",
        caption="Excerpt — SeedCommunicationsLookupsAsync")
    add_para(doc, "What gets seeded:")
    add_bullets(doc, [
        "AnnouncementCategories — General (GEN), Academic (ACAD), "
        "Finance (FIN), Events (EVENT), Holiday (HOL), Health "
        "(HEALTH), Emergency (EMERG).",
        "AnnouncementAudiences — Everyone (ALL, RequiresTargetClass "
        "false), Parents (PARENT, false), Students (STUDENT, false), "
        "Specific Class (CLASS, true).",
    ])
    add_para(doc,
        "Codes are stable and short because the announcement service "
        "branches on them (the ResolveAudienceCodesForCurrentUserAsync "
        "method compares the seeded codes against the roles the "
        "current user belongs to). Adding a new audience is a row "
        "insert plus, if RequiresTargetClass is true, a one-line "
        "addition to the resolver.")

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/\n"
        "├── Communications/\n"
        "│   ├── Announcements.razor          <- /announcements         (admin)\n"
        "│   └── AnnouncementEditor.razor     <- /announcements/new and /{id}\n"
        "└── Portals/\n"
        "    ├── ParentDashboard.razor        <- /portal/parent\n"
        "    ├── WardDetail.razor             <- /portal/parent/wards/{id}\n"
        "    ├── StudentDashboard.razor       <- /portal/student\n"
        "    ├── StudentProfile.razor         <- /portal/student/profile\n"
        "    ├── StudentResults.razor         <- /portal/student/results\n"
        "    ├── StudentAttendance.razor      <- /portal/student/attendance\n"
        "    ├── StudentFees.razor            <- /portal/student/fees\n"
        "    └── PortalAnnouncements.razor    <- /portal/announcements\n")

    add_heading(doc, "9.1 Announcements.razor (admin list)", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Communications/Announcements.razor")

    add_heading(doc, "9.2 AnnouncementEditor.razor", 2)
    add_para(doc,
        "Composite editor: serves both /announcements/new and "
        "/announcements/{id}. The audience dropdown's Change handler "
        "flips a flag (selectedAudienceRequiresClass) that drives "
        "the visibility of the target-class picker — picking "
        "'Specific Class' surfaces a class dropdown, picking any "
        "other audience hides it and clears the value.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Communications/AnnouncementEditor.razor")

    add_heading(doc, "9.3 ParentDashboard.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/ParentDashboard.razor")

    add_heading(doc, "9.4 WardDetail.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/WardDetail.razor")

    add_heading(doc, "9.5 StudentDashboard.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentDashboard.razor")

    add_heading(doc, "9.6 StudentProfile.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentProfile.razor")

    add_heading(doc, "9.7 StudentResults.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentResults.razor")

    add_heading(doc, "9.8 StudentAttendance.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentAttendance.razor")

    add_heading(doc, "9.9 StudentFees.razor", 2)
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentFees.razor")

    add_heading(doc, "9.10 PortalAnnouncements.razor", 2)
    add_para(doc,
        "Shared between parents and students — the same page renders "
        "for both audiences because the announcement service decides "
        "what is visible. The Back button routes to /portal/parent "
        "for parents and /portal/student for students based on the "
        "signed-in user's role.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Portals/PortalAnnouncements.razor")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and styling", 1)
    add_para(doc,
        "The previously-disabled 'My Children' nav placeholder is "
        "removed. In its place, the menu now grows three new "
        "AuthorizeView panels:")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Communications\"",
        end_marker="</RadzenPanelMenu>",
        caption="Excerpt — NavMenu.razor (sprint 8 additions)")
    add_bullets(doc, [
        "Communications — gated to SuperAdmin / HeadTeacher. "
        "Surfaces /announcements and /announcements/new.",
        "Parent portal — gated to Parent. Surfaces /portal/parent "
        "and /portal/announcements.",
        "Student portal — gated to Student. Surfaces the dashboard "
        "and the four sub-pages plus the shared announcements feed.",
    ])

    add_heading(doc, "10.1 _Imports.razor", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Communications",
        lines_after_start=4,
        caption="Excerpt — _Imports.razor (sprint 8 additions)")

    add_heading(doc, "10.2 Styling additions in app.css", 2)
    add_para(doc,
        "The portal pages introduce a handful of new utility classes. "
        "They live in a 'Sprint 8 — Portals & announcements' section "
        "appended to wwwroot/app.css and reuse the existing colour "
        "variables (--nps-green-700, --nps-ink-200, --nps-ink-500, "
        "--nps-surface-alt) so the look matches the admin pages.")
    add_bullets(doc, [
        "nps-portal-row — the flex row that lays out the student "
        "dashboard's timetable card next to the shortcuts card.",
        "nps-shortcuts — the vertical button stack on the student "
        "dashboard.",
        "nps-readonly-grid — the auto-fit grid behind the student "
        "profile page; one card per labelled field.",
        "nps-ward-grid + nps-ward-card (and the BEM modifiers "
        "__head, __stats, __actions) — the ward gallery on the "
        "parent dashboard.",
        "nps-announcement-list, nps-announcement plus the is-read "
        "/ is-unread state modifiers — the announcement feed.",
    ])

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of the portal experience", 1)

    add_heading(doc, "11.1 The school office links the accounts", 2)
    add_bullets(doc, [
        "SuperAdmin (or HeadTeacher) opens /users/new and creates a "
        "Parent user with the family's primary email and the Parent "
        "role.",
        "Opens /parents, finds the parent profile (created when the "
        "pupil was admitted), edits it, sets the UserId to the new "
        "login. Saves.",
        "Repeats for any student who is old enough to log in directly "
        "— creates a Student-role user, sets Student.UserId on the "
        "pupil's record.",
        "From now on the parent / student can sign in at the normal "
        "login page and land on their portal dashboard automatically "
        "(via the AuthorizeView in NavMenu).",
    ])

    add_heading(doc, "11.2 The head teacher posts an announcement", 2)
    add_bullets(doc, [
        "Opens /announcements, clicks New announcement.",
        "Enters a title and a body. Picks the Holiday category and "
        "the Parents audience.",
        "Sets ExpiresOn to the day the holiday ends so the notice "
        "auto-drops from the parent feed once the term resumes.",
        "Toggles Pin to top so the notice sits above other rows "
        "until the auto-expiry.",
        "Clicks Publish now. AnnouncementService stamps PublishedOn "
        "= UtcNow and IsPublished = true in the same SaveChanges.",
    ])

    add_heading(doc, "11.3 A parent signs in", 2)
    add_bullets(doc, [
        "Lands on /portal/parent. The dashboard shows one card per "
        "ward, the total outstanding balance across all wards, and "
        "an unread-announcement count.",
        "Opens /portal/announcements. The Holiday notice is at the "
        "top (pinned) with a New badge.",
        "Clicks Mark as read. AnnouncementService inserts one "
        "AnnouncementRead row. The unread count drops by one on the "
        "next dashboard load.",
        "Clicks Back. The route goes to /portal/parent (the page "
        "detects the Parent role and routes accordingly).",
        "Clicks Open on the eldest ward's card. Lands on the ward "
        "detail page with three tabs.",
        "Switches to Report cards. Sees the last published card. "
        "Clicks the eye icon to open the PDF download from sprint 5.",
        "Switches to Fees & invoices. Sees the same StudentLedger "
        "the bursar uses, scoped to the pupil.",
    ])

    add_heading(doc, "11.4 A pupil signs in", 2)
    add_bullets(doc, [
        "Lands on /portal/student. Five stat tiles: class, "
        "outstanding fees, attendance percentage, published report "
        "cards, unread announcements.",
        "The 'Today's timetable' grid lists every TimetableEntry "
        "row whose WeekDay matches today's DayOfWeek, ordered by "
        "period start time.",
        "Clicks the My results shortcut. Sees every published "
        "report card. Clicks the eye to open the PDF.",
        "Clicks My fees. Sees the outstanding balance, the invoice "
        "history and the payment history (read-only).",
    ])

    add_heading(doc, "11.5 An admin previews the portal", 2)
    add_bullets(doc, [
        "All portal routes are reachable by SuperAdmin and "
        "HeadTeacher (the [Authorize(Roles = ...)] attribute "
        "includes them). PortalService.CurrentUserCanViewStudentAsync "
        "returns true for these roles regardless of StudentParent "
        "linkage.",
        "ResolveParentIdForCurrentUserAsync / "
        "ResolveStudentIdForCurrentUserAsync return null for admin "
        "users (no matching Parent.UserId or Student.UserId), so the "
        "two dashboards render the 'we can't find your record' "
        "friendly fallback. The admin can still reach the ward and "
        "student sub-pages directly by URL — useful for support "
        "calls.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "First run applies the Communications migration and seeds "
        "the two new lookup tables. Sign in as the SuperAdmin and "
        "open /announcements to confirm the lookups loaded.")

    add_heading(doc, "12.2 Verify navigation and lookups", 2)
    add_bullets(doc, [
        "Signed in as SuperAdmin: the Communications nav panel is "
        "now an active two-item dropdown.",
        "Signed in as a Parent role user: the Parent portal panel "
        "is visible; the Communications panel is not.",
        "Signed in as a Student role user: the Student portal "
        "panel is visible.",
        "SELECT Name, Code FROM AnnouncementCategories ORDER BY "
        "DisplayOrder; shows seven rows.",
        "SELECT Name, Code, RequiresTargetClass FROM "
        "AnnouncementAudiences ORDER BY DisplayOrder; shows four "
        "rows with RequiresTargetClass set only on 'Specific Class'.",
    ])

    add_heading(doc, "12.3 End-to-end happy path", 2)
    add_numbered(doc, [
        "As SuperAdmin: /users/new — create a Parent role user "
        "linked to an existing Family.Parent row by UserId.",
        "Sign out, sign in as the parent. Land on /portal/parent. "
        "Confirm one card per linked ward.",
        "Sign in (in a private window) as SuperAdmin and post an "
        "announcement to the Parents audience. Publish.",
        "Back in the parent's window, refresh /portal/parent. "
        "Unread count = 1.",
        "Open /portal/announcements. Click Mark as read. Refresh "
        "the dashboard — unread count = 0.",
        "Click Open on a ward card. Switch through Overview / "
        "Report cards / Fees & invoices.",
        "Repeat as a Student role user against /portal/student "
        "and its four sub-pages.",
    ])

    add_heading(doc, "12.4 Error paths", 2)
    add_numbered(doc, [
        "Sign in as a Parent-role user whose Parent row has no "
        "UserId link. Confirm the dashboard renders the 'we can't "
        "find your parent record' card rather than crashing.",
        "Try the URL /portal/parent/wards/{id} with a pupil id that "
        "is not linked to the current parent. Confirm the 'Not "
        "authorised' card renders.",
        "As admin, create an announcement targeted at Specific "
        "Class without picking a class. Confirm the service "
        "returns a friendly error.",
        "Try clicking Mark as read twice on the same announcement. "
        "Confirm no exception (the service short-circuits on the "
        "existing read row).",
        "Set ExpiresOn on a published announcement to yesterday's "
        "date. Confirm it disappears from the parent feed but "
        "still shows in /announcements with the Include expired "
        "filter on.",
    ])

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)

    add_heading(doc, "13.1 'We can't find your parent / student record'", 2)
    add_para(doc,
        "The portal dashboards resolve the signed-in user back to a "
        "Family.Parent or Family.Student row through Parent.UserId / "
        "Student.UserId. If the school office has not yet linked the "
        "two, the dashboard renders a friendly card asking the user "
        "to contact the office. The fix is on the office side: open "
        "/parents (or /students), edit the row, set the UserId to "
        "the relevant login, save.")

    add_heading(doc, "13.2 'Audience X requires a target class'", 2)
    add_para(doc,
        "Only the seeded 'Specific Class' audience carries "
        "RequiresTargetClass = true. If you add a custom audience "
        "with that flag set, every announcement using it must pick "
        "a class. The editor surfaces the dropdown automatically "
        "whenever the audience code is 'CLASS'; if you renamed the "
        "code, update OnAudienceChanged in AnnouncementEditor.razor "
        "to match.")

    add_heading(doc, "13.3 The unread count looks wrong", 2)
    add_para(doc,
        "Unread is defined as 'visible to the current user AND no "
        "AnnouncementRead row for that user'. Visibility filters by "
        "IsPublished, by ExpiresOn, and by the audience codes that "
        "apply to the user. If the count seems off, check the row's "
        "AnnouncementAudience code against ResolveAudienceCodesFor"
        "CurrentUserAsync — broad audiences like ALL are always "
        "visible, but a Specific-Class row only counts if one of the "
        "user's wards (or the user themselves) is enrolled in that "
        "class with WithdrawnOn = NULL.")

    add_heading(doc, "13.4 Component attribute build error after editing the announcement card class", 2)
    add_para(doc,
        "Blazor refuses mixed literal-and-expression content on a "
        "component attribute (RZ9986). The pattern in "
        "PortalAnnouncements.razor wraps the whole class value in an "
        "interpolated string: class=\"@($\"nps-card nps-announcement "
        "{(...)}\")\". If you change the class composition, keep the "
        "whole value inside a single @($\"...\") expression.")

    add_heading(doc, "13.5 PortalService dependencies", 2)
    add_para(doc,
        "PortalService takes IInvoiceService and IAnnouncementService "
        "as constructor parameters. Both are scoped services "
        "registered earlier in DependencyInjection.cs. If you add a "
        "new portal feature that needs a service that has not yet "
        "been registered, register it above the PortalService "
        "registration or the container will reject the resolution.")

    add_heading(doc, "13.6 IdentityRole migration warning", 2)
    add_para(doc,
        "Same pre-existing warning that has accompanied every sprint "
        "since sprint 1. Harmless at runtime.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_bullets(doc, [
        "Email / SMS notifications: a future job can scan unread "
        "AnnouncementRead gaps and emit messages via a notification "
        "service. The visibility logic is already in "
        "AnnouncementService.CountUnreadForCurrentUserAsync — re-use "
        "the same predicates.",
        "Attachments: a future AnnouncementAttachment table "
        "(AnnouncementId, FileName, BlobUrl, MimeType) slots in "
        "without changing Announcement. The editor can grow a file "
        "picker; the feed can render thumbnails.",
        "Two-way messaging: the StudentParent linkage already gives "
        "the school a way to route messages to a pupil's family. A "
        "Conversation / Message table pair would close the loop.",
        "Mobile push: the portal pages render responsively. Wrapping "
        "them in a MAUI shell, or building a native app against the "
        "same endpoints, is additive.",
        "Online fee payment: the student fees page already shows "
        "the ledger; bolting a payment gateway onto PaymentService "
        "would make the page transactional.",
        "Per-ward custom announcements: today 'Specific Class' is "
        "the narrowest audience. A future 'Specific Pupil' audience "
        "would slot in by adding a TargetStudentId column to "
        "Announcement and one more branch to the audience resolver.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 8", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Communications/AnnouncementCategory.cs", "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Communications/AnnouncementAudience.cs", "Lookup with RequiresTargetClass."),
        ("src/NaijaPrimeSchool.Domain/Communications/Announcement.cs",         "The notice itself."),
        ("src/NaijaPrimeSchool.Domain/Communications/AnnouncementRead.cs",     "Per-user read row."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Communications/Dtos/AnnouncementDtos.cs", "Announcement DTOs."),
        ("src/NaijaPrimeSchool.Application/Communications/IAnnouncementService.cs",  "Announcement service contract."),
        ("src/NaijaPrimeSchool.Application/Portals/Dtos/PortalDtos.cs",              "Portal DTOs."),
        ("src/NaijaPrimeSchool.Application/Portals/IPortalService.cs",               "Portal service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Common/ICurrentUser.cs",                  "Added Roles + IsInRole."),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",                 "Added 2 new lookup methods."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/AnnouncementService.cs",      "CRUD + portal feed + read tracking."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/PortalService.cs",            "Parent + student dashboard façade."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 4 tables."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",               "Registered 2 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",  "Added 4 DbSets, ConfigureCommunications."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",   "Seeded the 2 new lookup tables."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",            "Added 2 new lookup methods."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Communications/Announcements.razor",       "Admin list + filters."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Communications/AnnouncementEditor.razor",  "Composer with audience-aware fields."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/ParentDashboard.razor",            "Parent home."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/WardDetail.razor",                 "Per-ward tabs."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentDashboard.razor",           "Student home."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentProfile.razor",             "Read-only demographic profile."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentResults.razor",             "Published report card list."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentAttendance.razor",          "Term attendance summary."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/StudentFees.razor",                "Pupil ledger."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Portals/PortalAnnouncements.razor",        "Shared feed for parents + students."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                                 "Added Communications + Portals usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",                           "Replaced disabled 'My Children' placeholder with three new panels."),
        ("src/NaijaPrimeSchool.Web/Services/CurrentUserAccessor.cs",                           "Exposed Roles + IsInRole from claims."),
        ("src/NaijaPrimeSchool.Web/wwwroot/app.css",                                           "Appended portal + announcement styles."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint8_guide.py",                                                    "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 8 implementation guide. With the parent "
        "and student portals alive and announcements in flight, every "
        "seeded role in the system now has a workspace. The next sprint "
        "can pick up whichever forward-compatibility lever the school "
        "needs most — email notifications, attachments, two-way "
        "messaging, or online fee payment.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
