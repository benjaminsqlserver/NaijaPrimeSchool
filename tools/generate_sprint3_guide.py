"""Generates 'Sprint 3 - Implementation Guide.docx' covering the
students-and-parents work delivered in sprint 3 (Students, Parents,
StudentParent linkage, Enrolments, and the supporting lookup tables).

This is the long-form (40+ page) edition. Code blocks embed actual source
files from the repository so the guide stays in lock-step with the code.

Run from the repo root:  python tools/generate_sprint3_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 3 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent
MIGRATION_FILE = "src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/20260429211125_StudentsAndParents.cs"


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    """Embed a source file with a caption labelling its path."""
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_excerpt(doc, rel_path, start_marker, end_marker=None,
                lines_after_start=None, caption=None):
    """Embed a slice of a file between markers (or N lines after start)."""
    text = read_text(rel_path)
    lines = text.splitlines()
    start_idx = next((i for i, l in enumerate(lines) if start_marker in l), None)
    if start_idx is None:
        add_caption(doc, f"(could not find start marker in {rel_path})")
        return
    if end_marker is not None:
        end_idx = next((i for i, l in enumerate(lines[start_idx + 1:],
                                                start=start_idx + 1)
                        if end_marker in l), None)
        end_idx = end_idx + 1 if end_idx is not None else len(lines)
    else:
        end_idx = start_idx + (lines_after_start or 25)
    add_caption(doc, caption or f"Excerpt — {rel_path}")
    add_code(doc, "\n".join(lines[start_idx:end_idx]).rstrip())


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 3 — Students & Parents")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Pupils · Parents/Guardians · Linkage · Enrolment")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/3-students-parents")
    meta.add_run("\nBuilt on: Sprint 1 identity foundation + Sprint 2 academic domain")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 3 in context", 1)
    add_para(doc,
        "Sprint 3 plugs students and their parents into the academic structure "
        "laid down in sprint 2. The work in this sprint is the first time the "
        "system has anyone in it who is not a member of staff. Pupils get an "
        "admission record, a profile, a medical note where it matters, and an "
        "enrolment in a class arm; parents and guardians get their own "
        "directory; the two sides of the family are stitched together with a "
        "many-to-many linkage table that captures relationship, primary "
        "contact, and pickup authorisation.")
    add_para(doc,
        "Once this sprint ships, every other module the school is going to "
        "need has somewhere to attach to. Attendance hangs off (Student × "
        "Class × Date). Result sheets hang off (Student × Subject × Term). "
        "Fee invoices hang off (Student × Term × Class). Parent portals hang "
        "off the StudentParent linkage. None of those features could land "
        "before sprint 3, because there were no students in the system to "
        "attach them to.")
    add_para(doc,
        "This document is a long-form implementation guide. It is written so "
        "that an engineer who has read the sprint 2 guide and has the "
        "codebase checked out can recreate every change in this sprint "
        "without referring to the diff. The document is organised in roughly "
        "the order I built the code in: design decisions first, Domain "
        "entities next, Application contracts after that, Infrastructure "
        "(DbContext, services, seeder, migration) in the middle, and finally "
        "the Razor UI and navigation. There is a smoke-test chapter near the "
        "end that walks through the happy-path you can use to confirm a "
        "fresh checkout works.")

    add_heading(doc, "1.1 Where this sits relative to sprint 2", 2)
    add_para(doc,
        "Sprint 1 delivered authentication, role-based authorization, the "
        "BaseEntity pattern, the SaveChanges audit and soft-delete rewriter, "
        "the OperationResult shape, and the Radzen-themed admin shell. "
        "Sprint 2 layered the academic skeleton on top: Sessions, Terms, "
        "Class arms, Subjects, Periods, and the weekly timetable. Both of "
        "those sets of foundations are reused unchanged here. In particular:")
    add_bullets(doc, [
        "BaseEntity — every new entity in sprint 3 inherits from it and "
        "picks up Guid Id, IAuditable, and ISoftDelete with no boilerplate.",
        "ApplicationDbContext.SaveChanges — the override stamps "
        "CreatedOn/By and ModifiedOn/By and rewrites Delete to "
        "IsDeleted = true. Every family-domain write therefore inherits "
        "auditing and soft delete with zero changes to the override.",
        "Global query filters — every new entity declares "
        "HasQueryFilter(x => !x.IsDeleted), so deleted rows vanish from "
        "ordinary queries without service code having to remember to filter.",
        "ICurrentUser / CurrentUserAccessor — the family services do not "
        "consume ICurrentUser directly, but the DbContext does, and that is "
        "what stamps audit columns when they save.",
        "OperationResult / OperationResult<T> — every family service uses "
        "this for predictable success/failure responses.",
        "ILookupService — already had eight methods. Sprint 3 adds seven "
        "more (relationships, enrolment statuses, blood groups, marital "
        "statuses, classes-for-session, students search, parents search) "
        "without rewriting the existing ones.",
        "SchoolClass — sprint 2's class arm picks up one new collection "
        "navigation (Enrolments) so the schema can join in the other "
        "direction without breaking anything. Existing code that ignored "
        "the new property continues to compile and behave identically.",
        "Radzen Blazor + the green/gold app.css — the family pages adopt "
        "the same .nps-page-header / .nps-card / .nps-form-grid primitives "
        "so they read as part of the same product, not a bolt-on.",
    ])

    add_heading(doc, "1.2 Functional scope delivered", 2)
    add_para(doc,
        "Concretely, after this sprint a SuperAdmin or HeadTeacher signing "
        "in to the application can:")
    add_numbered(doc, [
        "Add a new pupil with their identity, demographic profile, blood "
        "group, allergies, and medical notes; optionally enrol them into a "
        "class straight from the create form.",
        "Browse the student directory with full-text search, filter by "
        "session and class, and toggle active/inactive pupils.",
        "Open a pupil's profile to edit the demographic and medical "
        "fields, or to manage their parent linkages and enrolment history.",
        "Add a new parent or guardian to the directory; capture title, "
        "phone, alternate phone, email, occupation, employer, and marital "
        "status.",
        "Browse the parents directory with full-text search.",
        "From a pupil's profile, link an existing parent to them with a "
        "Relationship row (Father, Mother, Uncle, Guardian, Other, …), "
        "mark that parent as the primary contact, and toggle whether they "
        "are authorised to collect the child.",
        "From a parent's profile, see every pupil they are linked to and "
        "click straight through to the pupil.",
        "Enrol a pupil into a class for a given session, then later "
        "withdraw them or move them by enrolling into a different class.",
        "Browse all enrolments across the school with session, class, and "
        "status filters; withdraw or delete enrolments inline.",
        "Soft-delete any of the above, with friendly errors if the row is "
        "still referenced elsewhere — you cannot delete a student who has "
        "an open enrolment, and you cannot delete a parent who is still "
        "linked to a pupil.",
    ])
    add_para(doc,
        "Teachers, Bursars, Storekeepers, Parents and Students do not see "
        "the Family navigation panel; their menu items remain placeholders "
        "reserved for later sprints (parent portal, attendance capture, "
        "report cards). Adding a new pupil does not yet create an "
        "ApplicationUser for them — there is a deliberate UserId hook on "
        "the Student entity that the portal sprint will populate.")

    add_heading(doc, "1.3 Non-goals deliberately deferred", 2)
    add_para(doc,
        "It is just as important to be explicit about what sprint 3 does "
        "NOT do, because every one of these has been weighed and consciously "
        "deferred. Trying to land them now would have pushed scope past "
        "what is comfortably reviewable in one branch:")
    add_bullets(doc, [
        "Self-service parent and student logins. The Parent and Student "
        "roles are seeded; a UserId column is reserved on each entity. "
        "Wiring those up is a separate, opinionated piece of work that "
        "belongs to the portal sprint.",
        "Bulk import (CSV/Excel). Schools generally have a spreadsheet of "
        "all current pupils and would like to upload it. The schema will "
        "support this without changes; it is a feature-add, not a redesign.",
        "Photographs. There is a PhotoUrl column on Student but no upload "
        "flow yet — schools using the system today can host photos in the "
        "wwwroot folder or a CDN and paste a URL. A real upload pipeline "
        "(probably backed by a /uploads volume or Azure Blob Storage) will "
        "land alongside report cards.",
        "Sibling discovery. We can find pupils who share a parent through "
        "the StudentParent linkage, but there is no dedicated 'siblings' "
        "panel on the pupil profile yet. The query is one join; the UI is "
        "the lift.",
        "Attendance. Attendance is sprint 4; the (Student × Class × Date) "
        "key it will use is now expressible.",
        "Promotion / end-of-session bulk roll-over. Moving every pupil "
        "from Primary 1A to Primary 2A at the start of a new session is a "
        "useful workflow, but it is a one-shot operation against the "
        "schema we now have, not a change to the schema.",
    ])

    add_heading(doc, "1.4 Scale of the sprint", 2)
    add_para(doc, "By the numbers, this sprint adds:")
    add_bullets(doc, [
        "8 new domain entities under src/NaijaPrimeSchool.Domain/Family/.",
        "1 new collection navigation on SchoolClass (Enrolments).",
        "4 DTO files under src/NaijaPrimeSchool.Application/Family/Dtos/.",
        "3 new service contracts under src/NaijaPrimeSchool.Application/Family/.",
        "3 new service implementations under src/NaijaPrimeSchool.Infrastructure/Services/.",
        "7 new methods on ILookupService (and the matching LookupService).",
        "1 EF Core migration introducing 8 new tables and 16 indexes.",
        "1 DatabaseInitializer extension seeding Relationships, "
        "EnrolmentStatuses, BloodGroups, and MaritalStatuses.",
        "6 Razor pages under src/NaijaPrimeSchool.Web/Components/Pages/Family/.",
        "1 navigation menu addition to surface the new pages.",
    ])
    add_para(doc,
        "Everything compiles with zero warnings on .NET 10 (the team-wide "
        "warning bar). The code follows the patterns already accepted in "
        "sprints 1 and 2, so the diff is low-friction to review.")

    add_page_break(doc)


def chapter2_decisions(doc):
    add_heading(doc, "2. Design decisions and trade-offs", 1)
    add_para(doc,
        "Before any code was written I made a small number of architectural "
        "calls that shaped everything that followed. Re-reading these in "
        "isolation makes the code easier to navigate later, and it gives "
        "future maintainers permission to revisit choices when the "
        "trade-offs change.")

    add_heading(doc, "2.1 Student is its own entity, not an ApplicationUser", 2)
    add_para(doc,
        "The most consequential decision in this sprint was modelling "
        "Student as a first-class entity that is independent of "
        "ApplicationUser, rather than making Student a profile attached to "
        "an ApplicationUser carrying the Student role. The trade-off is "
        "real, so it is worth being explicit:")
    add_bullets(doc, [
        "Schools admit pupils long before those pupils log in — and many "
        "primary-school pupils never will. Forcing a Student to have an "
        "Identity row would mean creating throw-away credentials at "
        "admission, and asking the bursar's clerk to think about passwords "
        "and email addresses while they are filling in a guarantor's "
        "telephone number.",
        "ApplicationUser is shaped for staff: it has UserName, "
        "EmailConfirmed, LockoutEnd, PasswordHash, and a host of other "
        "Identity columns that are noise for a five-year-old. Subclassing "
        "or extending ApplicationUser further would have made the Identity "
        "table a junk drawer.",
        "The Student → ApplicationUser link is still expressible: there is "
        "a nullable UserId on Student. The portal sprint will populate it "
        "for older pupils who actually need to sign in. Until then, the "
        "column is null and the UI does not surface it.",
        "Several denormalised columns (FirstName, LastName, GenderId, "
        "DateOfBirth) live on Student rather than relying on a join. This "
        "is intentional: the student record is the source of truth for "
        "their identity in the school, even if they later get an "
        "ApplicationUser and that login carries different display names.",
    ])
    add_para(doc,
        "Parent follows the same shape for the same reasons: one parent in "
        "ten gets an actual portal login, the other nine in ten just want "
        "phone and email captured so the front office can call them.")

    add_heading(doc, "2.2 No enums — every lookup is a table", 2)
    add_para(doc,
        "Sprint 1 set the rule that domain concepts which would normally be "
        "C# enums are stored as tables instead. Sprint 2 honoured it. "
        "Sprint 3 honours it without exception. Relationship, "
        "EnrolmentStatus, BloodGroup, and MaritalStatus are all proper "
        "entities that derive from BaseEntity and live in the database. "
        "The reasons remain the same as the earlier sprints:")
    add_bullets(doc, [
        "Translatability. A future Hausa or Yoruba translation only needs "
        "to edit Name in a row, not recompile a deployment.",
        "Extensibility by data. If a school wants to add a 'Foster Parent' "
        "or 'Step-Grandparent' relationship, the Relationships table "
        "accepts a new row; no code change.",
        "Reportability. Joins to lookup tables surface readable labels in "
        "every reporting tool; magic enum integers do not.",
        "Display order. Each lookup carries a DisplayOrder column which "
        "the UI uses to sort dropdowns. Enums would have to lean on the "
        "integer value or impose an attribute-based ordering.",
    ])
    add_para(doc,
        "BloodGroup is a particularly interesting case here. The set of "
        "blood groups is genuinely closed (A+ A- B+ B- AB+ AB- O+ O-, plus "
        "Unknown), which is the textbook argument for an enum. I still "
        "modelled it as a table — partly for consistency with the rest of "
        "the codebase, partly so that the seed file is the single place a "
        "developer reads to discover what the value set is, and partly "
        "because '+' and '-' would have made for awkward C# enum members.")

    add_heading(doc, "2.3 Soft delete, again", 2)
    add_para(doc,
        "Every entity in sprint 3 implements ISoftDelete via BaseEntity. "
        "The pattern matches sprint 2 exactly: SaveChanges intercepts "
        "EntityState.Deleted, flips IsDeleted to true, and stamps "
        "DeletedOn/By. Global query filters then hide the row from every "
        "subsequent query that does not explicitly call IgnoreQueryFilters.")
    add_para(doc,
        "What is genuinely new in sprint 3 is the layered pre-delete "
        "guards. Three of the family entities refuse a delete that would "
        "leave the schema in an awkward state:")
    add_bullets(doc, [
        "StudentService.SoftDeleteAsync refuses to delete a pupil who "
        "still has an open enrolment (WithdrawnOn IS NULL). Withdrawing "
        "first forces an audit trail, then the deletion is allowed.",
        "ParentService.SoftDeleteAsync refuses to delete a parent who is "
        "still linked to one or more pupils.",
        "EnrolmentService allows delete unconditionally — the StudentParent "
        "and Student rows that point at it survive because of cascade-on-"
        "soft-delete semantics, but historically deleted enrolment rows are "
        "rare enough that we do not bother gating them.",
    ])
    add_para(doc,
        "An alternative would have been to cascade-soft-delete dependents, "
        "so deleting a student also deletes every enrolment and parent "
        "link. I chose the explicit refusal because cascade soft delete is "
        "hard to undo accurately. Refusing the delete keeps the audit "
        "trail readable.")

    add_heading(doc, "2.4 The StudentParent link is its own entity", 2)
    add_para(doc,
        "EF Core 10 handles many-to-many through a join entity, and we "
        "could have used the implicit pattern (just navigation collections "
        "and let EF magic up the join table). I made StudentParent an "
        "explicit entity for two reasons:")
    add_bullets(doc, [
        "It carries data in its own right — RelationshipId, "
        "IsPrimaryContact, CanPickUp, Notes. Implicit join tables cannot "
        "do that without ceremony.",
        "Soft delete needs an explicit type. The implicit join entity does "
        "not implement ISoftDelete; making it explicit lets us reuse the "
        "BaseEntity pattern unchanged.",
    ])
    add_para(doc,
        "The unique index (StudentId, ParentId) ensures a given parent "
        "appears at most once per pupil. The service layer also "
        "pre-checks this in LinkParentAsync so the user gets a friendly "
        "OperationResult.Failure rather than a database exception.")

    add_heading(doc, "2.5 Enrolment is per-class, not per-session", 2)
    add_para(doc,
        "An enrolment row points at a SchoolClass. Because every "
        "SchoolClass already belongs to a Session (sprint 2), the session "
        "is implicit — there is no SessionId column on Enrolment. This "
        "matters because:")
    add_bullets(doc, [
        "Reading 'who is in Primary 1A in 2025/2026' is a simple WHERE "
        "SchoolClassId = …, no extra join.",
        "Moving a pupil from Primary 1A to Primary 1B in the middle of a "
        "term is just two enrolment rows: one withdrawn, one new.",
        "There can never be a contradiction between Enrolment.SessionId "
        "and SchoolClass.SessionId, because the former does not exist.",
    ])
    add_para(doc,
        "The cost is that 'show me a pupil's history across sessions' "
        "joins through SchoolClass to Session. EF Core handles that "
        "elegantly in the projection (see EnrolmentService.Project), and "
        "the cost is a single extra index lookup on a tiny Sessions table.")
    add_para(doc,
        "The service layer enforces 'one open enrolment per session per "
        "pupil' explicitly: CreateAsync rejects an enrolment if the pupil "
        "already has an enrolment with WithdrawnOn IS NULL in the same "
        "session. That keeps a pupil from being in Primary 1A and "
        "Primary 1B simultaneously.")

    add_heading(doc, "2.6 Composite unique indexes", 2)
    add_para(doc,
        "Two composite unique indexes do most of the structural integrity "
        "heavy lifting in sprint 3:")
    add_bullets(doc, [
        "(StudentId, ParentId) on StudentParent — a parent cannot be "
        "linked twice to the same pupil.",
        "(StudentId, SchoolClassId) on Enrolment — a pupil cannot be "
        "enrolled twice in the same class arm. Re-enrolment after a "
        "withdrawal would have to be a new row in a new class.",
    ])
    add_para(doc,
        "Plus a filtered unique index on Student.UserId (and Parent.UserId) "
        "with [UserId] IS NOT NULL — sprint 8 will use this to ensure each "
        "pupil/parent maps to at most one Identity user, while still "
        "allowing many rows to have a null UserId today.")

    add_heading(doc, "2.7 Foreign-key delete behaviour", 2)
    add_para(doc,
        "Family relationships use a deliberately mixed delete behaviour:")
    add_bullets(doc, [
        "StudentParent.StudentId / ParentId → Cascade. If a student row "
        "ever gets hard-deleted (e.g. via a script that bypasses the soft-"
        "delete pattern), their links go with them. We cannot reach this "
        "state from the UI but the schema is honest about it.",
        "Enrolment.StudentId → Cascade for the same reason.",
        "Enrolment.SchoolClassId → Restrict. A class arm cannot be hard-"
        "deleted while it still has enrolments. This matches the "
        "soft-delete UI flow that already refuses to delete a class with "
        "timetable entries.",
        "Lookup foreign keys (RelationshipId, EnrolmentStatusId, "
        "BloodGroupId, MaritalStatusId) → SetNull where the field is "
        "optional, Restrict where it is required. You cannot accidentally "
        "lose 'Father' from the Relationships table while there are "
        "linkages pointing at it.",
        "UserId on Student and Parent → SetNull. Removing an Identity user "
        "should not delete the underlying student or parent.",
    ])

    add_heading(doc, "2.8 Service per feature area", 2)
    add_para(doc,
        "Three new services split the work along feature lines: "
        "IStudentService, IParentService, and IEnrolmentService. The "
        "StudentService is the heaviest, because it owns the "
        "StudentParent linkage operations as well — the linkage is "
        "conceptually 'a thing the student has', and a separate "
        "IStudentParentService felt like over-engineering for a half-"
        "dozen methods.")
    add_para(doc,
        "Likewise, ILookupService grew by seven methods rather than "
        "spawning a separate IFamilyLookupService. The lookup interface "
        "is the UI-facing dropdown source for the entire app; "
        "fragmenting it would have just meant Razor pages injecting two "
        "services where they used to inject one.")

    add_heading(doc, "2.9 DateOnly for biographical dates", 2)
    add_para(doc,
        "Student.DateOfBirth, Student.AdmissionDate, Enrolment.EnrolledOn, "
        "and Enrolment.WithdrawnOn are all DateOnly. The reasoning is the "
        "same as in sprint 2: SQL Server's date column is the right "
        "storage type, and DateOnly avoids the time-zone landmines that "
        "DateTime would introduce. The Razor pages still bind Radzen "
        "DatePickers to DateTime? in the form model and convert at the "
        "boundary using DateOnly.FromDateTime — Radzen has not yet "
        "adopted DateOnly natively.")

    add_heading(doc, "2.10 Inline forms with tabs for related data", 2)
    add_para(doc,
        "EditStudent.razor and EditParent.razor introduce a pattern that "
        "did not exist in sprint 2: a Radzen tab strip with the profile "
        "form on the first tab, related data on subsequent tabs. The "
        "Students profile has Profile / Parents / Enrolment history; the "
        "Parent profile has Profile / Linked pupils. The reasoning:")
    add_bullets(doc, [
        "Editing demographics and managing parent links are conceptually "
        "different tasks — separating them prevents a 'wall of fields' "
        "form.",
        "Each tab body is itself an inline form (the parent-link tab "
        "shows a small RadzenCard form when you click 'Link a parent'), "
        "so the sprint 2 inline-form pattern is preserved on the inside.",
        "Tabs avoid Radzen dialogs, which the sprint 2 design notes "
        "rejected for accessibility and testability.",
    ])

    add_page_break(doc)


def chapter3_domain(doc):
    add_heading(doc, "3. The Domain layer in full", 1)
    add_para(doc,
        "Every family-domain entity lives in a single new folder, "
        "src/NaijaPrimeSchool.Domain/Family/. There are no abstract base "
        "classes other than BaseEntity, and the entities are deliberately "
        "anaemic — no domain methods, no validation logic. Validation lives "
        "in the Application DTOs (DataAnnotations) and in the "
        "Infrastructure services (cross-aggregate checks). The Domain "
        "layer is a typed vocabulary; it is not the place for behaviour "
        "in this codebase.")

    add_heading(doc, "3.1 Folder layout and namespacing", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Domain/\n"
        "├── Family/                      ← (new in sprint 3)\n"
        "│   ├── Relationship.cs          ← lookup\n"
        "│   ├── EnrolmentStatus.cs       ← lookup\n"
        "│   ├── BloodGroup.cs            ← lookup\n"
        "│   ├── MaritalStatus.cs         ← lookup\n"
        "│   ├── Student.cs               ← pupil\n"
        "│   ├── Parent.cs                ← parent / guardian\n"
        "│   ├── StudentParent.cs         ← linkage with relationship\n"
        "│   └── Enrolment.cs             ← (student × class) over time\n"
        "├── Academics/                   ← from sprint 2 (SchoolClass updated)\n"
        "├── Common/                      ← from sprint 1\n"
        "└── Identity/                    ← from sprint 1\n")
    add_para(doc,
        "Namespace is NaijaPrimeSchool.Domain.Family. Lookup entities "
        "live alongside their owners in Family/ rather than in a separate "
        "Lookups folder, mirroring the sprint 2 layout where TermType and "
        "ClassLevel sit beside Term and SchoolClass.")

    add_heading(doc, "3.2 Lookup entities", 2)

    add_heading(doc, "3.2.1 Relationship.cs", 3)
    add_para(doc,
        "Relationship is the lookup of how a parent stands to the child: "
        "Father, Mother, Stepfather, Stepmother, Grandfather, Grandmother, "
        "Uncle, Aunt, Guardian, Other. The seed list is in "
        "DatabaseInitializer; adding Foster Parent or Sibling is a row "
        "insertion, not a code change.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/Relationship.cs")

    add_heading(doc, "3.2.2 EnrolmentStatus.cs", 3)
    add_para(doc,
        "EnrolmentStatus is the lookup of what state an enrolment is in: "
        "Active, Suspended, Withdrawn, Transferred, Graduated. The "
        "service layer treats 'Withdrawn' specially — the WithdrawAsync "
        "convenience method looks up that row by name and uses it. If a "
        "school renames the row, WithdrawAsync still works because "
        "WithdrawnOn is the load-bearing column, not the status name.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/EnrolmentStatus.cs")

    add_heading(doc, "3.2.3 BloodGroup.cs", 3)
    add_para(doc,
        "BloodGroup carries the eight blood types plus an explicit "
        "'Unknown' so that captured-but-not-recorded is distinguishable "
        "from never-asked. The Name column is 10 characters wide, which "
        "is plenty for 'AB+' and 'Unknown'.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/BloodGroup.cs")

    add_heading(doc, "3.2.4 MaritalStatus.cs", 3)
    add_para(doc,
        "MaritalStatus covers Single, Married, Divorced, Widowed, "
        "Separated. It feeds the Parent profile — which rarely matters "
        "to the school but matters for the records that the school "
        "regulator wants once a year.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/MaritalStatus.cs")

    add_heading(doc, "3.3 Core entities", 2)

    add_heading(doc, "3.3.1 Student.cs", 3)
    add_para(doc,
        "Student is the pupil. AdmissionNumber is the unique admin handle "
        "(NPS/2026/0001). FirstName / LastName / MiddleName are "
        "denormalised so that the student record is the source of truth "
        "for their identity in the school. The optional UserId hooks the "
        "pupil to an ApplicationUser if and when a portal login is "
        "needed; the unique filtered index ensures one ApplicationUser "
        "maps to at most one pupil.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/Student.cs")
    add_para(doc,
        "FullName is a computed property that the DbContext explicitly "
        "ignores in the model configuration (see chapter 5). Allergies "
        "and MedicalNotes are deliberately free-text — schools have a lot "
        "of variation in how they record them, and an enum or table here "
        "would be brittle.")

    add_heading(doc, "3.3.2 Parent.cs", 3)
    add_para(doc,
        "Parent mirrors Student in shape but is shaped for adults: "
        "Title, MaritalStatus, PrimaryPhone, AlternatePhone, Email, "
        "Occupation, Employer. There is no admission number — a parent "
        "is not 'admitted' to the school. There is no DateOfBirth either, "
        "because no school I have built for has ever asked the parent's "
        "DOB and a few have politely declined to ask.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/Parent.cs")

    add_heading(doc, "3.3.3 StudentParent.cs", 3)
    add_para(doc,
        "StudentParent is the explicit join entity between Student and "
        "Parent. RelationshipId tells you what hat the parent wears for "
        "this child; IsPrimaryContact picks one parent for school-day "
        "phone calls; CanPickUp gates pickup at the gate. Notes is a "
        "free-text annotation for special instructions ('only collected "
        "by mother on Tuesdays').")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/StudentParent.cs")
    add_para(doc,
        "Note that there is no StudentName / ParentName denormalisation. "
        "StudentParent reaches across navigations to surface the names in "
        "the DTOs; the join entity itself stays narrow.")

    add_heading(doc, "3.3.4 Enrolment.cs", 3)
    add_para(doc,
        "Enrolment links Student to SchoolClass with a date range and a "
        "status. EnrolledOn is required; WithdrawnOn is null while the "
        "enrolment is active. Notes is a free-text annotation for the "
        "registrar — typical entries are 'transferred from St George's' "
        "or 'parents requested mid-year change of class'.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Family/Enrolment.cs")

    add_heading(doc, "3.4 The one change to sprint 2 — SchoolClass", 2)
    add_para(doc,
        "SchoolClass picks up exactly one new line: a collection "
        "navigation called Enrolments. Existing code that ignored this "
        "property continues to compile and behave identically; the "
        "migration emits a foreign-key index but no schema change to "
        "SchoolClasses itself. This is the smallest possible change to "
        "an existing domain entity.")
    add_file(doc, "src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs")

    add_heading(doc, "3.5 Relationships at a glance", 2)
    add_para(doc,
        "An ASCII relationship diagram captures the family-side of the "
        "schema and how it joins back to academics:")
    add_code(doc,
        "                    ┌──────────────┐                  ┌──────────────┐\n"
        "                    │   Student    │                  │    Parent    │\n"
        "                    │  Admission#  │                  │   Phone(s)   │\n"
        "                    │  IsActive    │                  │   IsActive   │\n"
        "                    └──┬───────┬───┘                  └──┬───────────┘\n"
        "                       │ N     │ N                       │ N\n"
        "                       │       │                         │\n"
        "                       │  ┌────▼─────────┐               │\n"
        "                       │  │ StudentParent│ N─────────────┘\n"
        "                       │  │ Relationship │\n"
        "                       │  │ IsPrimary    │\n"
        "                       │  │ CanPickUp    │\n"
        "                       │  └──────────────┘\n"
        "                       │ 1\n"
        "                ┌──────▼────────┐\n"
        "                │   Enrolment   │ N──── EnrolmentStatus (lookup)\n"
        "                │   EnrolledOn  │\n"
        "                │   WithdrawnOn │\n"
        "                └──────┬────────┘\n"
        "                       │ N\n"
        "                ┌──────▼────────┐\n"
        "                │  SchoolClass  │  ←── (sprint 2)\n"
        "                │   Session     │\n"
        "                └───────────────┘\n"
        "\n"
        "    Lookups (BaseEntity, soft-deletable):\n"
        "      Relationship, EnrolmentStatus, BloodGroup, MaritalStatus\n"
        "    Cross-domain re-uses:\n"
        "      Student.GenderId       → Identity.Gender   (sprint 1 lookup)\n"
        "      Student.BloodGroupId   → Family.BloodGroup\n"
        "      Parent.TitleId         → Identity.Title    (sprint 1 lookup)\n"
        "      Parent.MaritalStatusId → Family.MaritalStatus\n"
        "      Student.UserId / Parent.UserId → Identity.ApplicationUser  (nullable)\n")

    add_page_break(doc)


def chapter4_application(doc):
    add_heading(doc, "4. Application layer — DTOs and contracts", 1)
    add_para(doc,
        "Application is the boundary between the Razor UI and the "
        "persistence layer. It is deliberately thin: just DTOs and "
        "service interfaces. There is no domain logic and no EF Core. "
        "Sprint 3 adds two new subfolders: Family/Dtos and Family. The "
        "DTOs live in the first; the contracts in the second.")

    add_heading(doc, "4.1 DTO design rules", 2)
    add_bullets(doc, [
        "Every list view returns a flat read-DTO with denormalised "
        "display fields. StudentDto carries CurrentClassName, "
        "CurrentSessionName, PrimaryContactName, and PrimaryContactPhone "
        "so the list view can render a row without lazy-loading any "
        "navigations.",
        "Create and Update requests are separate types with "
        "DataAnnotations validation, mirroring the sprint 2 pattern.",
        "Filter types (StudentListFilter, ParentListFilter, "
        "EnrolmentListFilter) carry the search/sort/page parameters and "
        "let the service layer translate into LINQ.",
        "Linkage DTOs (StudentParentDto, LinkStudentParentRequest, "
        "UpdateStudentParentLinkRequest) live in their own file because "
        "they are conceptually a separate concept from either the "
        "student or parent.",
    ])

    add_heading(doc, "4.2 StudentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/Dtos/StudentDtos.cs")
    add_para(doc,
        "CurrentClassId, CurrentClassName, CurrentSessionId, "
        "CurrentSessionName, ParentCount, PrimaryContactName, and "
        "PrimaryContactPhone are populated by the service via a "
        "subquery in the projection — see chapter 6. The list view "
        "therefore renders the most useful information about a pupil "
        "without needing follow-up reads.")

    add_heading(doc, "4.3 ParentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/Dtos/ParentDtos.cs")
    add_para(doc,
        "StudentCount on ParentDto is the count of linked pupils, "
        "computed via .StudentLinks.Count in the projection. It lets the "
        "parents grid show, at a glance, which entries in the directory "
        "are still tied to active children.")

    add_heading(doc, "4.4 StudentParentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/Dtos/StudentParentDtos.cs")
    add_para(doc,
        "There is a deliberate asymmetry between LinkStudentParentRequest "
        "(creates a new link) and UpdateStudentParentLinkRequest (edits "
        "an existing link). The create form needs StudentId and ParentId; "
        "the update form does not, because the link's identity is its Id "
        "and the underlying pupil and parent should not change.")

    add_heading(doc, "4.5 EnrolmentDtos.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/Dtos/EnrolmentDtos.cs")
    add_para(doc,
        "Note that CreateEnrolmentRequest has Guid? EnrolmentStatusId — "
        "callers who do not specify it pick up the lowest-DisplayOrder "
        "status (which is 'Active'), which matches the typical UX where "
        "an admin creates a new enrolment and only later realises they "
        "wanted a non-default status.")

    add_heading(doc, "4.6 IStudentService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/IStudentService.cs")
    add_para(doc,
        "The student service owns the parent-linkage methods because "
        "the linkage is conceptually 'a thing the student has'. There "
        "is no IStudentParentService.")

    add_heading(doc, "4.7 IParentService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/IParentService.cs")

    add_heading(doc, "4.8 IEnrolmentService.cs", 2)
    add_file(doc, "src/NaijaPrimeSchool.Application/Family/IEnrolmentService.cs")
    add_para(doc,
        "WithdrawAsync is a convenience: it sets WithdrawnOn, flips the "
        "status to 'Withdrawn' if that lookup row exists, and appends "
        "the supplied notes to the existing notes. The fall-through to "
        "UpdateAsync (which can also set WithdrawnOn) gives admins a "
        "more flexible path when they need to backdate or correct a "
        "record.")

    add_heading(doc, "4.9 ILookupService extension", 2)
    add_para(doc,
        "ILookupService grew by seven methods. Five new lookups expose "
        "the new tables; two are search helpers for the existing "
        "Students and Parents tables that the dropdowns on EditStudent "
        "lean on:")
    add_file(doc, "src/NaijaPrimeSchool.Application/Users/ILookupService.cs")
    add_para(doc,
        "GetClassesForSessionAsync was strictly speaking a sprint 2-era "
        "convenience that never got built — Sprint 3 needs it for the "
        "student create form, the enrolment grid, and the family pages, "
        "so it lands here. It is technically about Academics but it "
        "lives on the same lookup service for UI convenience.")

    add_page_break(doc)


def chapter5_dbcontext(doc):
    add_heading(doc, "5. Infrastructure — DbContext changes", 1)
    add_para(doc,
        "ApplicationDbContext picks up eight new DbSets and a single new "
        "ConfigureFamily method. Everything else in the file is "
        "unchanged from sprint 2. The pattern of a private "
        "ConfigureXxx method per feature folder keeps OnModelCreating "
        "from becoming a thousand-line monster.")

    add_heading(doc, "5.1 New DbSets", 2)
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "public DbSet<Relationship>", end_marker="public DbSet<Enrolment>",
        caption="Excerpt — the eight family DbSets")

    add_heading(doc, "5.2 ConfigureFamily", 2)
    add_para(doc,
        "ConfigureFamily is invoked from OnModelCreating after "
        "ConfigureAcademics. It uses the same ConfigureLookup helper "
        "from sprint 1 for the four lookup tables, then configures "
        "Student, Parent, StudentParent, and Enrolment one by one.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureFamily",
        end_marker="private static void ConfigureLookup",
        caption="Excerpt — ConfigureFamily")
    add_para(doc,
        "Three things are worth dwelling on here:")
    add_bullets(doc, [
        "b.Ignore(s => s.FullName) on Student and Parent — the "
        "computed property is a string, so EF Core would otherwise try "
        "to map it as a column.",
        "b.HasIndex(s => s.UserId).IsUnique().HasFilter(\"[UserId] IS "
        "NOT NULL\") — a filtered unique index that allows many rows to "
        "have a null UserId today, but at most one row to point at a "
        "given ApplicationUser when the portal sprint hooks them up.",
        "DeleteBehavior.Cascade on the StudentParent and Enrolment "
        "foreign keys back to Student. Re-read the design notes in 2.7 "
        "for why; the short version is that hard-delete via a script "
        "should be self-cleaning.",
    ])

    add_heading(doc, "5.3 The lookup helper, reused", 2)
    add_para(doc,
        "The ConfigureLookup<T> generic helper from sprint 1 is "
        "reused unchanged. Each new lookup picks up table name, primary "
        "key, audit columns, and the global query filter on IsDeleted — "
        "the only customisation per lookup is the Name column "
        "constraint and a unique index.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs",
        "private static void ConfigureLookup",
        end_marker="public override int SaveChanges",
        caption="Excerpt — ConfigureLookup<T> (unchanged from sprint 1)")

    add_heading(doc, "5.4 SaveChanges still does all the work", 2)
    add_para(doc,
        "The SaveChanges override from sprint 1 needs no changes for "
        "sprint 3. Every new entity inherits BaseEntity, so the audit "
        "stamping and soft-delete rewriting apply automatically. This "
        "is the single biggest payoff of the architectural decisions in "
        "sprint 1: every sprint after it gets auditing and soft delete "
        "for free.")

    add_page_break(doc)


def chapter6_services(doc):
    add_heading(doc, "6. Infrastructure — service implementations", 1)
    add_para(doc,
        "Three new services land in src/NaijaPrimeSchool.Infrastructure/"
        "Services/, each one a straightforward implementation of the "
        "matching Application interface. Plus the existing LookupService "
        "grows seven methods to surface the new lookup tables.")

    add_heading(doc, "6.1 StudentService.cs", 2)
    add_para(doc,
        "StudentService is the heaviest of the three. It owns the "
        "pupil CRUD, the parent-linkage CRUD, and the student-level "
        "primitives (search, list, set-active, soft-delete).")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/StudentService.cs")
    add_para(doc,
        "Three patterns are worth calling out from this file:")
    add_bullets(doc, [
        "ListAsync uses a single big projection that pulls "
        "CurrentClass, CurrentSession, ParentCount, and "
        "PrimaryContactName from subqueries. The whole list page is one "
        "round-trip to the database, with no N+1 queries even when the "
        "filters narrow it to a thousand pupils.",
        "CreateAsync optionally creates an enrolment row in the same "
        "SaveChanges. The enrolment uses the student's AdmissionDate as "
        "the EnrolledOn — admins almost always enrol on the day of "
        "admission, and the form gives them the AdmissionDate field "
        "anyway.",
        "LinkParentAsync flips IsPrimaryContact = false on every other "
        "link for the same student before saving — there can be at most "
        "one primary contact per pupil, and the service enforces that "
        "with a single ExecuteUpdateAsync rather than reading and "
        "writing in a loop.",
    ])

    add_heading(doc, "6.2 ParentService.cs", 2)
    add_para(doc,
        "ParentService is parallel to StudentService but simpler — "
        "there is no enrolment side. The grouping pattern (search, "
        "list, get-by-id, create, update, set-active, soft-delete) is "
        "the same.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs")
    add_para(doc,
        "GetStudentLinksAsync mirrors GetParentLinksAsync on "
        "StudentService — both project from the StudentParent table to "
        "StudentParentDto, just with a different filter. The two "
        "pieces of code are similar enough that they could share a "
        "helper, but inlining each one keeps the read-flow explicit and "
        "the LINQ short.")

    add_heading(doc, "6.3 EnrolmentService.cs", 2)
    add_para(doc,
        "EnrolmentService is the smallest, and it deliberately stays "
        "narrow. There is no ListAsync paging — enrolments are usually "
        "filtered down to a single class or a single pupil before "
        "rendering, so 25-row paging on the grid is enough.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/EnrolmentService.cs")
    add_para(doc,
        "Three details deserve a sentence each:")
    add_bullets(doc, [
        "CreateAsync rejects an enrolment whose pupil already has an "
        "open enrolment in the same session. This is the rule that "
        "prevents 'pupil is in Primary 1A AND Primary 1B in 2025/2026'.",
        "WithdrawAsync looks up 'Withdrawn' by name in the "
        "EnrolmentStatuses table. If a school has renamed that row, the "
        "service still sets WithdrawnOn — it just leaves the status at "
        "whatever it was. WithdrawnOn is the load-bearing column; the "
        "status is a label.",
        "UpdateAsync re-checks the (StudentId, SchoolClassId) "
        "uniqueness constraint with .Id != request.Id, mirroring the "
        "sprint 2 services' pattern of pre-checking unique indexes for "
        "friendly error messages.",
    ])

    add_heading(doc, "6.4 LookupService — the seven new methods", 2)
    add_para(doc,
        "LookupService picks up seven new methods. Four expose the new "
        "lookup tables (relationships, enrolment statuses, blood "
        "groups, marital statuses); one is a session-scoped class "
        "lookup; and two are search helpers for students and parents.")
    add_file(doc, "src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs")
    add_para(doc,
        "GetStudentsAsync and GetParentsAsync take an optional search "
        "term and cap the result at 50 rows. The cap matters because "
        "these methods feed the AllowFiltering dropdown on the parent-"
        "link form — Radzen renders every option in the DOM, so a list "
        "of 5,000 pupils would crater the page.")

    add_heading(doc, "6.5 DI registration", 2)
    add_para(doc,
        "DependencyInjection.cs picks up three new lines.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",
        "services.AddScoped<IStudentService",
        end_marker="return services;",
        caption="Excerpt — DependencyInjection.cs (sprint 3 additions)")

    add_page_break(doc)


def chapter7_migration(doc):
    add_heading(doc, "7. The EF Core migration", 1)
    add_para(doc,
        "A single migration named StudentsAndParents adds eight new "
        "tables (Relationships, EnrolmentStatuses, BloodGroups, "
        "MaritalStatuses, Students, Parents, StudentParents, "
        "Enrolments) and the indexes that go with them. It was "
        "generated with:")
    add_code(doc,
        "dotnet ef migrations add StudentsAndParents \\\n"
        "  --project src/NaijaPrimeSchool.Infrastructure \\\n"
        "  --startup-project src/NaijaPrimeSchool.Web \\\n"
        "  --output-dir Persistence/Migrations\n")
    add_para(doc,
        "On a fresh checkout the migration runs at startup via "
        "DatabaseInitializer.MigrateAsync. The full Up() method is "
        "embedded below for reference.")
    add_excerpt(doc, MIGRATION_FILE, "protected override void Up",
                end_marker="protected override void Down",
                caption=f"Excerpt — Up() of {Path(MIGRATION_FILE).name}")
    add_para(doc,
        "Notice the sequence: lookup tables first (Relationships, "
        "EnrolmentStatuses, BloodGroups, MaritalStatuses), then the "
        "two big entity tables (Students, Parents) which depend on "
        "those lookups, then the join tables (StudentParents, "
        "Enrolments) which depend on the entity tables. EF Core "
        "computes this ordering automatically from the foreign-key "
        "graph; we did not have to specify it.")

    add_heading(doc, "7.1 Indexes created", 2)
    add_para(doc, "The migration creates 16 indexes on the new tables:")
    add_bullets(doc, [
        "Unique on Relationships.Name, EnrolmentStatuses.Name, "
        "BloodGroups.Name, MaritalStatuses.Name (4 unique).",
        "Unique on Students.AdmissionNumber (1 unique).",
        "Filtered unique on Students.UserId WHERE [UserId] IS NOT NULL "
        "(1 unique).",
        "Filtered unique on Parents.UserId WHERE [UserId] IS NOT NULL "
        "(1 unique).",
        "Unique on StudentParents.(StudentId, ParentId) (1 unique).",
        "Unique on Enrolments.(StudentId, SchoolClassId) (1 unique).",
        "Plus IsDeleted indexes on every soft-deletable table (8).",
    ])

    add_heading(doc, "7.2 The schema warning", 2)
    add_para(doc,
        "The migration emits one warning at generation time:")
    add_code(doc,
        "warn: Microsoft.EntityFrameworkCore.Model.Validation[10622]\n"
        "      Entity 'ApplicationRole' has a global query filter\n"
        "      defined and is the required end of a relationship with\n"
        "      the entity 'ApplicationUserRole'. ...")
    add_para(doc,
        "This warning is unrelated to sprint 3. It is a long-standing "
        "interaction between Identity's IdentityUserRole and our "
        "soft-delete query filter on ApplicationRole. The runtime "
        "behaviour is fine because we never soft-delete a role from the "
        "UI — roles are seeded once and stay there. Sprint 6 (or "
        "whichever sprint adds role administration) will revisit this.")

    add_page_break(doc)


def chapter8_seeding(doc):
    add_heading(doc, "8. Seeding the family lookups", 1)
    add_para(doc,
        "DatabaseInitializer picks up a new SeedFamilyLookupsAsync "
        "method that is invoked between SeedAcademicLookupsAsync and "
        "SeedRolesAsync. It seeds the four new lookup tables with "
        "sensible defaults a Nigerian primary school would want on "
        "day one.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs",
        "private static async Task SeedFamilyLookupsAsync",
        end_marker="private static async Task SeedLookupsAsync",
        caption="Excerpt — SeedFamilyLookupsAsync")
    add_para(doc, "What gets seeded:")
    add_bullets(doc, [
        "Relationships — Father, Mother, Stepfather, Stepmother, "
        "Grandfather, Grandmother, Uncle, Aunt, Guardian, Other "
        "(10 rows).",
        "EnrolmentStatuses — Active, Suspended, Withdrawn, "
        "Transferred, Graduated (5 rows). Active has DisplayOrder = 1, "
        "which is what CreateEnrolmentRequest defaults to when no "
        "status is specified.",
        "BloodGroups — A+, A-, B+, B-, AB+, AB-, O+, O-, Unknown "
        "(9 rows).",
        "MaritalStatuses — Single, Married, Divorced, Widowed, "
        "Separated (5 rows).",
    ])
    add_para(doc,
        "The seeder uses the same .IgnoreQueryFilters().AnyAsync() "
        "guard as the sprint 1 and sprint 2 seeders — it only inserts "
        "if the table is empty (counting soft-deleted rows), so "
        "running the app a second time after a soft delete does not "
        "resurrect the missing row.")

    add_page_break(doc)


def chapter9_pages(doc):
    add_heading(doc, "9. The Razor pages", 1)
    add_para(doc,
        "Six new pages land in src/NaijaPrimeSchool.Web/Components/"
        "Pages/Family/. The pattern mirrors sprint 2: a list page, a "
        "create page, an edit page, with the addition of an "
        "enrolment-overview page that does not have its own "
        "create/edit because enrolment is created and managed from "
        "within the student profile.")

    add_heading(doc, "9.1 Page roster", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/\n"
        "├── Students.razor        ← /students       (list + filters)\n"
        "├── CreateStudent.razor   ← /students/new\n"
        "├── EditStudent.razor     ← /students/{id}  (Profile / Parents / Enrolments tabs)\n"
        "├── Parents.razor         ← /parents        (list + search)\n"
        "├── CreateParent.razor    ← /parents/new\n"
        "├── EditParent.razor      ← /parents/{id}   (Profile / Linked pupils tabs)\n"
        "└── Enrolments.razor      ← /enrolments     (overview + filters)\n")
    add_para(doc,
        "Every page is gated to SuperAdmin + HeadTeacher. Teachers, "
        "Bursars, Storekeepers, Parents and Students do not see the "
        "Family panel and cannot navigate to the routes — the route "
        "itself returns 403.")

    add_heading(doc, "9.2 Students.razor — the directory", 2)
    add_para(doc,
        "Students.razor is the front door. A search box, three "
        "dropdowns (session, class, status), and a paged grid. The "
        "grid renders the pupil's name and admission number in column "
        "1, their current class and session in column 2 (or 'Not "
        "enrolled' if they have no open enrolment), the primary "
        "contact in column 3, DOB in column 4, status badge in "
        "column 5, and three action buttons in column 6.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/Students.razor")
    add_para(doc,
        "The session dropdown defaults to the current session (looked "
        "up via SessionService.GetCurrentAsync), and the class "
        "dropdown is rebuilt whenever the session changes via "
        "LookupService.GetClassesForSessionAsync. The grid's "
        "LoadData event drives paging and sorting through the LINQ "
        "filter on the server.")

    add_heading(doc, "9.3 CreateStudent.razor — the create form", 2)
    add_para(doc,
        "CreateStudent.razor is a single long inline form, organised "
        "into three labelled sections: Identity, Health, and Initial "
        "enrolment. The Initial enrolment section is optional — if "
        "the admin picks a class, the service creates an enrolment "
        "row alongside the student in the same SaveChanges; if they "
        "skip it, the pupil is created without any enrolment and "
        "the admin can enrol them later from the EditStudent page.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateStudent.razor")
    add_para(doc,
        "The form is RadzenTemplateForm with DataAnnotationsValidator, "
        "so every required-field message and length-constraint message "
        "from the DTOs renders inline next to the offending field. "
        "On success, Nav.NavigateTo($\"/students/{result.Data}\") "
        "drops the admin straight onto the new pupil's edit page so "
        "they can immediately link parents.")

    add_heading(doc, "9.4 EditStudent.razor — the heart of the sprint", 2)
    add_para(doc,
        "EditStudent.razor is the most substantial page in the sprint. "
        "It uses Radzen tabs to present three views of the same pupil:")
    add_bullets(doc, [
        "Profile — every demographic and medical field, editable.",
        "Parents — the list of linked parents, with inline link/edit/"
        "unlink workflows.",
        "Enrolment history — the full chronology with actions for "
        "withdraw and delete.",
    ])
    add_para(doc,
        "There are also two top-right buttons that flip the pupil's "
        "active status (Activate / Deactivate). The whole page is "
        "rendered server-side (InteractiveServer) so every form "
        "submission is a SignalR message rather than a client-side "
        "WebAssembly round-trip — the data is sensitive enough that "
        "we want every keystroke gated by the cookie auth.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor")
    add_para(doc,
        "Three patterns deserve a moment:")
    add_bullets(doc, [
        "Tab counts in titles. The tab labels read 'Parents (3)' and "
        "'Enrolment history (5)' so an admin can see at a glance "
        "whether anything is there before they click. The numbers "
        "are reactive; they update when the underlying lists change.",
        "Inline link form. Clicking 'Link a parent' reveals a "
        "RadzenCard inside the Parents tab. The parent dropdown uses "
        "Radzen's AllowFiltering so admins can type the parent's "
        "name and pick from the search results — backed by "
        "LookupService.GetParentsAsync, capped at 50 results.",
        "Confirm-then-act for destructive operations. Every "
        "withdraw and delete goes through DialogService.Confirm "
        "with explicit OK/Cancel labels so a misclick does not lose "
        "data.",
    ])

    add_heading(doc, "9.5 Parents.razor and CreateParent.razor", 2)
    add_para(doc,
        "Parents.razor and CreateParent.razor mirror the sprint 2 "
        "and earlier-in-sprint-3 patterns exactly: a paged grid with "
        "search and status filters; a long inline form with title, "
        "names, marital status, contact details, and occupation. The "
        "grid surfaces the count of linked pupils per parent so the "
        "admin can spot 'directory-only' parents who never got "
        "linked to a child.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/Parents.razor")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateParent.razor")

    add_heading(doc, "9.6 EditParent.razor — the parent detail view", 2)
    add_para(doc,
        "EditParent.razor uses the same tabs pattern as EditStudent. "
        "The Profile tab is the editable form; the Linked pupils tab "
        "is read-only — it shows every pupil this parent is linked "
        "to, with a click-through to the pupil. Linking is initiated "
        "from the pupil's Edit page rather than from here, because "
        "in practice an admin is almost always thinking 'this pupil "
        "needs another parent' rather than 'this parent needs more "
        "pupils'.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/EditParent.razor")

    add_heading(doc, "9.7 Enrolments.razor — the overview", 2)
    add_para(doc,
        "Enrolments.razor is the only sprint-3 page that does not "
        "drill into a single pupil. It is the registrar's view: pick "
        "a session and a class (or a status), see every pupil whose "
        "enrolment matches. Withdraw and delete actions live inline "
        "on each row; opening a pupil pops their EditStudent page "
        "in the same browser tab.")
    add_file(doc, "src/NaijaPrimeSchool.Web/Components/Pages/Family/Enrolments.razor")
    add_para(doc,
        "Note that this page does not have its own create form — "
        "creating an enrolment is done from EditStudent's Enrolment "
        "history tab, which has access to the pupil context. "
        "Enrolling a pupil from the overview page would mean adding "
        "a Student dropdown to the form and that quickly becomes "
        "fiddly with thousands of pupils. Schools that need a bulk-"
        "enrolment workflow will get a dedicated screen in a later "
        "sprint.")

    add_page_break(doc)


def chapter10_navigation(doc):
    add_heading(doc, "10. Navigation, imports, and authorization", 1)
    add_heading(doc, "10.1 NavMenu — a new Family panel", 2)
    add_para(doc,
        "NavMenu.razor picks up a fourth role-gated panel between "
        "Academics and the Finance/Inventory placeholders. Five entries "
        "live inside it: Students, Add Student, Parents, Add Parent, "
        "Enrolments. The whole panel is wrapped in an AuthorizeView "
        "that requires SuperAdmin or HeadTeacher.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",
        "<RadzenPanelMenuItem Text=\"Family\"",
        end_marker="</AuthorizeView>",
        caption="Excerpt — NavMenu.razor")

    add_heading(doc, "10.2 _Imports.razor", 2)
    add_para(doc,
        "_Imports.razor picks up two new @using lines so the family "
        "DTOs and service interfaces are visible to every Razor file "
        "in the Web project without per-file imports.")
    add_excerpt(doc,
        "src/NaijaPrimeSchool.Web/Components/_Imports.razor",
        "@using NaijaPrimeSchool.Application.Family",
        lines_after_start=2,
        caption="Excerpt — _Imports.razor")

    add_heading(doc, "10.3 Authorization at the page level", 2)
    add_para(doc,
        "Every sprint-3 page declares an [Authorize] attribute "
        "restricting access to SuperAdmin + HeadTeacher. There is no "
        "policy added — the existing ManageUsers policy is intended "
        "for SuperAdmin only, and HeadTeacher access here is a "
        "permission specific to the family pages, not a general role "
        "promotion. Future sprints (attendance, results) will likely "
        "introduce a 'Pastoral' policy that includes class teachers; "
        "this sprint deliberately stays narrow.")

    add_page_break(doc)


def chapter11_lifecycle(doc):
    add_heading(doc, "11. Lifecycle of a family-domain row", 1)
    add_para(doc,
        "Walking a single pupil from creation to withdrawal-and-"
        "deletion is the clearest way to see how all the layers cooperate. "
        "Imagine an admin admitting a six-year-old named Adaeze "
        "Okonkwo into Primary 1A for the 2025/2026 session.")

    add_heading(doc, "11.1 Admin opens /students/new", 2)
    add_bullets(doc, [
        "Authorization fires: the user must be SuperAdmin or "
        "HeadTeacher. If they are not, ASP.NET Core returns 403.",
        "OnInitializedAsync calls LookupService.GetGendersAsync, "
        "GetBloodGroupsAsync, and GetClassesForSessionAsync(currentSession.Id). "
        "Three database hits, all on tiny tables.",
        "The page renders the long inline form with three sections.",
    ])

    add_heading(doc, "11.2 Admin fills the form and clicks Save", 2)
    add_bullets(doc, [
        "RadzenTemplateForm runs DataAnnotations validation. If "
        "anything is missing, ValidationMessage components render "
        "inline.",
        "If the form is valid, SaveAsync converts the DateTime?-"
        "bound dates to DateOnly and calls "
        "StudentService.CreateAsync(request).",
        "StudentService.CreateAsync re-runs the cross-aggregate "
        "checks (DOB before admission date, AdmissionNumber unique, "
        "InitialClassId exists if provided), constructs a Student, "
        "and — if InitialClassId was provided — also constructs an "
        "Enrolment with EnrolledOn = AdmissionDate and the lowest-"
        "DisplayOrder EnrolmentStatus.",
        "db.Students.Add and db.Enrolments.Add stage both rows; "
        "SaveChangesAsync commits them in one transaction.",
        "ApplicationDbContext.SaveChanges (sprint 1) stamps "
        "CreatedOn/By on both rows.",
        "On success, the page navigates to /students/{newId} so the "
        "admin lands on Adaeze's profile and can immediately go to "
        "the Parents tab.",
    ])

    add_heading(doc, "11.3 Admin links Adaeze's mother", 2)
    add_bullets(doc, [
        "Parents tab loads via StudentService.GetParentLinksAsync — "
        "empty, since Adaeze just got created.",
        "Admin clicks Link a parent. The inline form appears.",
        "Admin types 'Okonkwo' into the parent dropdown; "
        "AllowFiltering kicks in and "
        "LookupService.GetParentsAsync(\"Okonkwo\") returns the "
        "candidates, capped at 50.",
        "Admin picks Mrs Ngozi Okonkwo, picks Mother as the "
        "relationship, leaves Primary contact ON and Authorised to "
        "collect ON, and clicks Save link.",
        "StudentService.LinkParentAsync flips IsPrimaryContact = "
        "false on every other link for this student (none, in this "
        "case) and inserts the new StudentParent row.",
        "On success the form closes and GetParentLinksAsync re-runs, "
        "so the grid shows the new link with the 'Primary' badge.",
    ])

    add_heading(doc, "11.4 Three years later, Adaeze withdraws", 2)
    add_bullets(doc, [
        "Admin opens Adaeze's profile and goes to Enrolment history. "
        "Three rows are visible: Primary 1A 2025/2026 (withdrawn at "
        "end of session), Primary 2A 2026/2027 (withdrawn similarly), "
        "Primary 3A 2027/2028 (still active).",
        "Admin clicks the Withdraw icon on the active row.",
        "DialogService.Confirm asks for confirmation. Admin clicks "
        "Withdraw.",
        "EnrolmentService.WithdrawAsync looks up the 'Withdrawn' "
        "status row, sets WithdrawnOn = today, "
        "EnrolmentStatusId = Withdrawn, appends notes.",
        "ApplicationDbContext.SaveChanges stamps "
        "ModifiedOn/By on the row.",
        "The grid re-renders; the row's status badge flips to "
        "'Withdrawn' and the Withdraw icon disappears (the action is "
        "hidden once WithdrawnOn is set).",
    ])

    add_heading(doc, "11.5 The pupil leaves the school", 2)
    add_bullets(doc, [
        "Admin clicks Delete on the pupil. "
        "DialogService.Confirm asks for confirmation.",
        "StudentService.SoftDeleteAsync loads the student with "
        ".Include(s => s.Enrolments).Include(s => s.ParentLinks). "
        "It checks for any enrolment with WithdrawnOn IS NULL — "
        "there is none, so the delete is allowed.",
        "db.Students.Remove(student); SaveChangesAsync(). The "
        "SaveChanges override intercepts the Deleted state, flips "
        "IsDeleted = true and stamps DeletedOn/By. The row is now "
        "invisible to ordinary queries because of the global query "
        "filter, but every audit column is preserved.",
    ])

    add_page_break(doc)


def chapter12_smoketest(doc):
    add_heading(doc, "12. Smoke-test walkthrough", 1)
    add_para(doc,
        "Once the build is green and the migration has applied, here "
        "is the end-to-end smoke test you can run against a fresh "
        "checkout. It covers the happy paths and a couple of "
        "important error paths.")

    add_heading(doc, "12.1 Build, migrate, run", 2)
    add_code(doc,
        "dotnet restore\n"
        "dotnet build NaijaPrimeSchool.slnx\n"
        "dotnet run --project src/NaijaPrimeSchool.Web\n")
    add_para(doc,
        "The first run applies migrations and seeds the four new "
        "lookup tables. Sign in as superadmin@naijaprimeschool.ng / "
        "Admin@12345.")

    add_heading(doc, "12.2 Verify navigation", 2)
    add_bullets(doc, [
        "The Family panel appears in the sidebar under Academics.",
        "Family contains five items: Students, Add Student, Parents, "
        "Add Parent, Enrolments.",
        "Click each one — every page renders without error.",
    ])

    add_heading(doc, "12.3 Create a parent", 2)
    add_numbered(doc, [
        "Family → Add Parent.",
        "Pick title Mrs., gender Female, marital status Married.",
        "Fill First name 'Ngozi', Last name 'Okonkwo'.",
        "Primary phone 080-xxx-xxxx, email "
        "ngozi.okonkwo@example.com.",
        "Save. You should land on the parent edit page.",
    ])

    add_heading(doc, "12.4 Create a pupil and enrol them", 2)
    add_numbered(doc, [
        "Family → Add Student.",
        "Admission number NPS/2026/0001, admission date today, "
        "first name Adaeze, last name Okonkwo, DOB any date six "
        "years ago.",
        "Pick gender Female, blood group O+, leave allergies blank.",
        "In Initial enrolment, pick a class from the current session "
        "(if you do not have one, hop to /classes and create "
        "Primary 1A first).",
        "Save. You land on Adaeze's profile.",
    ])

    add_heading(doc, "12.5 Link the parent", 2)
    add_numbered(doc, [
        "Click the Parents tab.",
        "Click Link a parent.",
        "Pick Ngozi Okonkwo from the dropdown, relationship Mother, "
        "primary contact ON, can pick up ON.",
        "Save link. The grid shows Ngozi with the 'Primary' badge.",
    ])

    add_heading(doc, "12.6 Verify the directory views", 2)
    add_bullets(doc, [
        "Family → Students. Adaeze is in the grid; the Primary "
        "contact column shows Ngozi Okonkwo and her phone number.",
        "Family → Parents. Ngozi is in the grid; the Pupils column "
        "shows 1.",
        "Family → Enrolments. Filter to current session; Adaeze's "
        "enrolment in Primary 1A is visible with status Active.",
    ])

    add_heading(doc, "12.7 Verify the error paths", 2)
    add_numbered(doc, [
        "Try to delete Ngozi from /parents. The OperationResult "
        "comes back with 'Cannot delete a parent who is still "
        "linked to a student. Unlink them first.'",
        "Try to delete Adaeze from /students without first "
        "withdrawing her enrolment. The OperationResult comes back "
        "with 'Cannot delete a student with an active enrolment.'",
        "Withdraw Adaeze from her enrolment. Now you can delete the "
        "pupil if you choose. Don't actually do it for the demo.",
    ])

    add_heading(doc, "12.8 Verify soft-delete and audit", 2)
    add_para(doc,
        "Connect to SQL Server and run:")
    add_code(doc,
        "SELECT Id, FirstName, LastName, IsDeleted, DeletedOn, DeletedBy\n"
        "FROM Students;\n"
        "\n"
        "SELECT Id, EnrolledOn, WithdrawnOn, IsDeleted\n"
        "FROM Enrolments;\n")
    add_para(doc,
        "If you withdrew and then deleted, IsDeleted is true on the "
        "Students row, DeletedOn and DeletedBy are populated, but "
        "the row is still there. The Enrolments row is also "
        "soft-deleted (because EnrolmentService.SoftDeleteAsync was "
        "called) — but only if you triggered the Delete; "
        "withdrawing alone does not delete.")

    add_page_break(doc)


def chapter13_troubleshooting(doc):
    add_heading(doc, "13. Troubleshooting and gotchas", 1)
    add_heading(doc, "13.1 'A student with admission number X already exists'", 2)
    add_para(doc,
        "Admission numbers are unique. If you copy-paste from an "
        "existing pupil and forget to bump the suffix, the service "
        "rejects it with a friendly message before the database "
        "throws. The unique index on Students.AdmissionNumber is the "
        "ultimate authority.")

    add_heading(doc, "13.2 'Student already has an active enrolment in this session'", 2)
    add_para(doc,
        "A pupil cannot be in two classes in the same session. If "
        "they need to move from Primary 1A to Primary 1B mid-term, "
        "withdraw the first enrolment, then create the second. The "
        "history view will show both rows.")

    add_heading(doc, "13.3 'A given parent is already linked to the student'", 2)
    add_para(doc,
        "The (StudentId, ParentId) unique index prevents duplicate "
        "links. If you want to change the relationship of an existing "
        "link, edit it (the pencil icon) rather than creating a new "
        "one.")

    add_heading(doc, "13.4 The parent dropdown is empty", 2)
    add_para(doc,
        "GetParentsAsync caps at 50 results without a search term. "
        "If your school has more than 50 parents and the one you "
        "want is alphabetically late, type a few letters of their "
        "name into the dropdown — Radzen's AllowFiltering re-queries "
        "with the term and the right parent will appear.")

    add_heading(doc, "13.5 The classes dropdown is empty when creating a student", 2)
    add_para(doc,
        "GetClassesForSessionAsync defaults to the current session. "
        "If no session is marked current, the call returns every "
        "class in the database, which is fine. If neither is true — "
        "no current session and no classes anywhere — you need to "
        "create a session and at least one class first via /sessions "
        "and /classes (sprint 2 pages).")

    add_heading(doc, "13.6 The 'Cannot delete' message persists after I cleared dependents", 2)
    add_para(doc,
        "EF Core's change tracker is per-DbContext-instance, and "
        "Blazor Server pages share their context across the page "
        "lifecycle. If you deleted the dependent rows in another tab, "
        "refresh the Students or Parents page so the new context "
        "reads the current state.")

    add_heading(doc, "13.7 The Identity / soft-delete migration warning", 2)
    add_para(doc,
        "The EF Core warning about ApplicationRole's query filter "
        "and IdentityUserRole's required end is harmless at runtime "
        "and predates sprint 3. We chose not to silence it because a "
        "future role-management UI will want to revisit it. The "
        "build is still green.")

    add_page_break(doc)


def chapter14_forward(doc):
    add_heading(doc, "14. Forward-compatibility, today", 1)
    add_para(doc,
        "Sprint 3 has deliberately left a few breadcrumbs that make "
        "the next sprints' work cleaner.")
    add_bullets(doc, [
        "Student.UserId and Parent.UserId are nullable hooks for "
        "future portal logins. The filtered unique indexes make "
        "wiring up an Identity user safe — at most one ApplicationUser "
        "per pupil/parent — without forcing every existing row to "
        "have one.",
        "Enrolment.WithdrawnOn distinguishes 'historical record' "
        "from 'current state'. Attendance (sprint 4) will read "
        "'open enrolments' as 'WithdrawnOn IS NULL', which is the "
        "load-bearing definition.",
        "StudentParent.IsPrimaryContact is the column the parent "
        "portal sprint will use to decide who gets the 'home' "
        "homepage when both parents have logins.",
        "BloodGroup, Allergies and MedicalNotes are captured but "
        "not yet surfaced anywhere except the pupil profile. The "
        "school nurse sprint (or the attendance sprint, depending on "
        "scheduling) will read those fields.",
        "Enrolment.SchoolClassId joining through to Session means "
        "every fee/result/attendance feature can navigate from a "
        "pupil to their session in one join — no Enrolment.SessionId "
        "duplication.",
    ])

    add_heading(doc, "14.1 What might need a small refactor later", 2)
    add_bullets(doc, [
        "GetStudentsAsync and GetParentsAsync cap at 50 rows. Once "
        "the school exceeds 5,000 pupils, server-side virtual "
        "scrolling on the dropdowns will be worth the engineering.",
        "StudentService is taking on parent-linkage methods. If a "
        "future feature (sibling discovery, family-wide messaging) "
        "warrants a dedicated IStudentParentService, those four "
        "methods can be extracted in a single PR.",
        "The Student / Parent denormalised name columns will need "
        "to stay in sync if the portal sprint introduces a 'change "
        "your name' flow on the pupil's account. The simplest "
        "approach is to update both rows in the same SaveChanges; "
        "we will add a tiny helper when that flow lands.",
        "Bulk import (CSV/Excel of pupils) is a feature-add that "
        "the schema already supports. The lift is a service method "
        "that loops over rows and calls CreateAsync, plus a Razor "
        "page with file upload.",
    ])

    add_page_break(doc)


def chapter15_appendix(doc):
    add_heading(doc, "15. Appendix — files added or changed in sprint 3", 1)
    entries = [
        ("Domain layer (new)", "—"),
        ("src/NaijaPrimeSchool.Domain/Family/Relationship.cs",     "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Family/EnrolmentStatus.cs",  "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Family/BloodGroup.cs",       "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Family/MaritalStatus.cs",    "Lookup."),
        ("src/NaijaPrimeSchool.Domain/Family/Student.cs",          "Pupil entity."),
        ("src/NaijaPrimeSchool.Domain/Family/Parent.cs",           "Parent / guardian entity."),
        ("src/NaijaPrimeSchool.Domain/Family/StudentParent.cs",    "Linkage with relationship + flags."),
        ("src/NaijaPrimeSchool.Domain/Family/Enrolment.cs",        "(student × class) over time."),
        ("Domain layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Domain/Academics/SchoolClass.cs",   "Added ICollection<Enrolment> Enrolments."),
        ("Application layer (new)", "—"),
        ("src/NaijaPrimeSchool.Application/Family/Dtos/StudentDtos.cs",       "Student DTOs."),
        ("src/NaijaPrimeSchool.Application/Family/Dtos/ParentDtos.cs",        "Parent DTOs."),
        ("src/NaijaPrimeSchool.Application/Family/Dtos/StudentParentDtos.cs", "Linkage DTOs."),
        ("src/NaijaPrimeSchool.Application/Family/Dtos/EnrolmentDtos.cs",     "Enrolment DTOs."),
        ("src/NaijaPrimeSchool.Application/Family/IStudentService.cs",        "Student service contract (CRUD + linkage)."),
        ("src/NaijaPrimeSchool.Application/Family/IParentService.cs",         "Parent service contract."),
        ("src/NaijaPrimeSchool.Application/Family/IEnrolmentService.cs",      "Enrolment service contract."),
        ("Application layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Application/Users/ILookupService.cs",          "Added 7 new lookup methods."),
        ("Infrastructure layer (new)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/Services/StudentService.cs",    "Student CRUD + linkage methods."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs",     "Parent CRUD."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/EnrolmentService.cs",  "Enrolment CRUD + Withdraw."),
        (f"src/NaijaPrimeSchool.Infrastructure/Persistence/Migrations/{Path(MIGRATION_FILE).name}", "EF migration adding 8 tables and 16 indexes."),
        ("Infrastructure layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Infrastructure/DependencyInjection.cs",        "Registered the 3 new services."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/ApplicationDbContext.cs", "Added 8 DbSets, ConfigureFamily."),
        ("src/NaijaPrimeSchool.Infrastructure/Persistence/DatabaseInitializer.cs", "Seeded Relationships, EnrolmentStatuses, BloodGroups, MaritalStatuses."),
        ("src/NaijaPrimeSchool.Infrastructure/Services/LookupService.cs",     "Added 7 new lookup methods."),
        ("Web layer (new)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/Students.razor",       "Students directory + filters."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateStudent.razor",  "Create student form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/EditStudent.razor",    "Profile / Parents / Enrolment tabs."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/Parents.razor",        "Parents directory + filters."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateParent.razor",   "Create parent form."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/EditParent.razor",     "Profile / Linked pupils tabs."),
        ("src/NaijaPrimeSchool.Web/Components/Pages/Family/Enrolments.razor",     "Enrolment overview + filters."),
        ("Web layer (modified)", "—"),
        ("src/NaijaPrimeSchool.Web/Components/_Imports.razor",                    "Added Family + Family.Dtos usings."),
        ("src/NaijaPrimeSchool.Web/Components/Layout/NavMenu.razor",              "Added the Family panel."),
        ("Tooling (new)", "—"),
        ("tools/generate_sprint3_guide.py",                                       "This document's generator."),
    ]

    table = doc.add_table(rows=len(entries), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (path, purpose) in enumerate(entries):
        row = table.rows[i].cells
        row[0].text = path
        row[1].text = purpose
        if purpose == "—":
            for run in row[0].paragraphs[0].runs:
                run.bold = True

    add_para(doc, " ")
    closing = doc.add_paragraph()
    r = closing.add_run(
        "— End of the Sprint 3 implementation guide. The next sprint "
        "lands attendance on top of the (Student × Class) primitives "
        "established here.")
    r.italic = True


# ---------- Main --------------------------------------------------------------

def main():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_decisions(doc)
    chapter3_domain(doc)
    chapter4_application(doc)
    chapter5_dbcontext(doc)
    chapter6_services(doc)
    chapter7_migration(doc)
    chapter8_seeding(doc)
    chapter9_pages(doc)
    chapter10_navigation(doc)
    chapter11_lifecycle(doc)
    chapter12_smoketest(doc)
    chapter13_troubleshooting(doc)
    chapter14_forward(doc)
    chapter15_appendix(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
