"""Generates 'Sprint 9 - Implementation Guide.docx' covering the auto-
provisioned portal accounts work: creating a parent or student now also
creates the matching ApplicationUser in the Parent / Student role and
stamps the new user's id onto Parent.UserId / Student.UserId, so the
families can sign in straight away and the portals load without an
admin follow-up.

Run from the repo root:  python tools/generate_sprint9_guide.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Sprint 9 - Implementation Guide.docx"
ROOT = Path(__file__).resolve().parent.parent


# ---------- Helpers ----------------------------------------------------------

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    cell.width = Inches(6.2)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    first = True
    for line in text.splitlines():
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        run = p.add_run(line if line else " ")
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    trailing = doc.add_paragraph()
    trailing.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def read_text(rel_path):
    path = ROOT / rel_path
    return path.read_text(encoding="utf-8")


def add_file(doc, rel_path, caption=None):
    add_caption(doc, caption or f"Listing — {rel_path}")
    add_code(doc, read_text(rel_path).rstrip("\r\n"))


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' (or press F9) to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instr, fldChar2, placeholder, fldChar3):
        run._element.append(el)


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ---------- Sections ---------------------------------------------------------

def title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(120)
    r = t.add_run("Naija Prime School")
    r.font.size = Pt(32); r.font.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Sprint 9 — Auto-Provisioned Portal Accounts")
    rs.font.size = Pt(18); rs.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)

    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Create-parent → user · Create-student → user · Portal lights up on first sign-in")
    rs2.font.size = Pt(14); rs2.italic = True

    sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs3 = sub3.add_run("Long-form implementation walk-through")
    rs3.font.size = Pt(12); rs3.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Benjamin Fadina").bold = True
    meta.add_run("\nBranch: sprint/9-portal-accounts")
    meta.add_run("\nBuilt on: Sprints 1–8 (identity, academic domain, students & parents, "
                 "attendance, results & report cards, pupil photos, fees & bursar workflows, "
                 "store & inventory, parent & student portals + announcements)")
    meta.add_run("\nStack: .NET 10, Blazor Web App (Auto), EF Core 10, SQL Server, Radzen Blazor")
    meta.add_run("\nEditor: Visual Studio Code with the C# Dev Kit")
    meta.add_run("\nRepository: https://github.com/benjaminsqlserver/NaijaPrimeSchool")
    meta.add_run("\nLicence: MIT — see LICENSE at the repo root")

    add_page_break(doc)


def toc_page(doc):
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x61, 0x3C)
    add_toc(doc)
    add_page_break(doc)


# ---------- Chapter 1 -------------------------------------------------------

def chapter1_overview(doc):
    add_heading(doc, "1. Sprint 9 in context", 1)
    add_para(doc,
        "Sprint 8 shipped the parent and student portals, but neither "
        "portal could light up on its own. Each portal resolves the "
        "signed-in user to a Family.Parent or a Family.Student row by "
        "following the optional UserId foreign key on those entities — "
        "and that key had to be filled in by hand. A SuperAdmin had to "
        "create the ApplicationUser under Users → New user, remember "
        "the new user's Guid, walk back to the Parent or Student page, "
        "and stamp the id onto the record. Until that link existed, "
        "every portal screen greeted the family with a friendly 'we "
        "can't find your record' card and stopped.")
    add_para(doc,
        "Sprint 9 closes that loop. Creating a parent at /parents/new "
        "now provisions a matching ApplicationUser in the Parent role "
        "as part of the same operation and stamps the new user's Id "
        "onto Parent.UserId. Creating a student at /students/new does "
        "the same against the Student role and Student.UserId. The "
        "very first time the family signs in with those credentials, "
        "the portal loads cleanly — no admin follow-up required.")
    add_para(doc,
        "There are no new entities, no new tables, no new lookup rows, "
        "and no new EF Core migration. The fix is a constructor change "
        "on two services (so they can ask UserManager to create a "
        "user and add a role), three new fields on two DTOs (Username, "
        "Email made required on Parent, Password), and the matching "
        "additions on the New Parent and New Student Razor forms. "
        "Portal-layer code is untouched — it already keyed off "
        "Parent.UserId / Student.UserId, and once those columns are "
        "populated the portals resolve correctly.")

    add_heading(doc, "1.1 Acceptance criteria", 2)
    add_numbered(doc, [
        "Creating a parent at /parents/new also creates an "
        "ApplicationUser in the Parent role, and Parent.UserId points "
        "at that user.",
        "Creating a student at /students/new also creates an "
        "ApplicationUser in the Student role, and Student.UserId points "
        "at that user.",
        "When that parent signs in, /portal/parent resolves their "
        "ParentId from Parent.UserId and renders their wards.",
        "When that student signs in, /portal/student resolves their "
        "StudentId from Student.UserId and renders their dashboard.",
        "Sprint 8's portal access guard (CurrentUserCanViewStudentAsync) "
        "already trusts the UserId link, so no portal-layer changes are "
        "required — the fix is upstream in the create-parent and "
        "create-student services.",
    ])

    add_heading(doc, "1.2 What this sprint does not do", 2)
    add_bullets(doc, [
        "It does not back-fill UserId on Parent / Student rows created "
        "before sprint 9. Those rows still need to be linked manually "
        "(or re-created) the way they always did.",
        "It does not change how parents and students are linked to one "
        "another. Sprint 3's StudentParent join table is the source of "
        "truth for 'who can pick up whom'. Sprint 9 only ensures the "
        "Parent and Student rows have an associated login.",
        "It does not surface a 'change password' or 'reset password' "
        "self-service flow. Those continue to be admin-driven through "
        "the Users area shipped in sprint 1.",
        "It does not relax the password policy. The 8-character "
        "lower / upper / digit / non-alphanumeric rule configured in "
        "DependencyInjection.cs still applies — the Razor forms now "
        "surface that requirement inline.",
    ])

    add_heading(doc, "1.3 Files touched", 2)
    add_code(doc,
        "src/NaijaPrimeSchool.Application/Family/Dtos/ParentDtos.cs\n"
        "src/NaijaPrimeSchool.Application/Family/Dtos/StudentDtos.cs\n"
        "src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs\n"
        "src/NaijaPrimeSchool.Infrastructure/Services/StudentService.cs\n"
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateParent.razor\n"
        "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateStudent.razor\n"
        "README.md\n"
        "Sprint 9 - Implementation Guide.md (markdown companion)\n"
        "tools/generate_sprint9_guide.py (this generator)"
    )

    add_page_break(doc)


# ---------- Chapter 2 -------------------------------------------------------

def chapter2_pre_sprint(doc):
    add_heading(doc, "2. Where sprint 8 left the wiring", 1)
    add_para(doc,
        "Before reading the new code it helps to remember exactly what "
        "sprint 8 left half-wired. The portals are perfectly capable "
        "of finding the right Parent or Student row — they just expect "
        "someone else to have stamped the UserId for them.")

    add_heading(doc, "2.1 The two foreign keys", 2)
    add_para(doc,
        "Parent and Student each carry a nullable UserId column with "
        "a filtered unique index. The column has been in the database "
        "since sprint 3, marked in the README and the sprint 3 guide "
        "as 'optional UserId hook for a future portal login'.")

    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Domain/Family/Parent.cs")
    add_code(doc,
        "public class Parent : BaseEntity\n"
        "{\n"
        "    // ... profile fields ...\n"
        "    public Guid? UserId { get; set; }\n"
        "    public ApplicationUser? User { get; set; }\n"
        "}")

    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Domain/Family/Student.cs")
    add_code(doc,
        "public class Student : BaseEntity\n"
        "{\n"
        "    // ... profile fields ...\n"
        "    public Guid? UserId { get; set; }\n"
        "    public ApplicationUser? User { get; set; }\n"
        "}")

    add_heading(doc, "2.2 How the portals resolve a row", 2)
    add_para(doc,
        "PortalService keys off ICurrentUser.UserId and walks the "
        "UserId foreign key on the appropriate table. Both methods "
        "return null when nothing matches, which is what triggers the "
        "friendly fallback card in the portal pages.")

    add_caption(doc, "Excerpt — src/NaijaPrimeSchool.Infrastructure/Services/PortalService.cs")
    add_code(doc,
        "public async Task<Guid?> ResolveParentIdForCurrentUserAsync(CancellationToken ct = default)\n"
        "{\n"
        "    if (currentUser.UserId is not { } userId) return null;\n"
        "    return await db.Parents\n"
        "        .Where(p => p.UserId == userId)\n"
        "        .Select(p => (Guid?)p.Id)\n"
        "        .FirstOrDefaultAsync(ct);\n"
        "}\n"
        "\n"
        "public async Task<Guid?> ResolveStudentIdForCurrentUserAsync(CancellationToken ct = default)\n"
        "{\n"
        "    if (currentUser.UserId is not { } userId) return null;\n"
        "    return await db.Students\n"
        "        .Where(s => s.UserId == userId)\n"
        "        .Select(s => (Guid?)s.Id)\n"
        "        .FirstOrDefaultAsync(ct);\n"
        "}")

    add_para(doc,
        "Once Parent.UserId / Student.UserId is populated these two "
        "queries land on the right row and the rest of sprint 8 "
        "(dashboards, ward detail, fees, results, attendance, "
        "today's timetable, the announcements feed) hangs off them "
        "without modification.")

    add_heading(doc, "2.3 What the create flows used to do", 2)
    add_para(doc,
        "The pre-sprint-9 ParentService.CreateAsync mapped the form "
        "straight onto a Parent entity and saved. There was no user "
        "creation, no role assignment, and no UserId. The same was "
        "true for StudentService.CreateAsync (with a couple of extra "
        "guards around the admission number, date of birth, and an "
        "optional initial enrolment).")

    add_caption(doc, "Pre-sprint-9 — ParentService.CreateAsync (now replaced)")
    add_code(doc,
        "public async Task<OperationResult<Guid>> CreateAsync(\n"
        "    CreateParentRequest request, CancellationToken ct = default)\n"
        "{\n"
        "    var parent = new Parent\n"
        "    {\n"
        "        FirstName = request.FirstName.Trim(),\n"
        "        LastName = request.LastName.Trim(),\n"
        "        // ... rest of fields ...\n"
        "        IsActive = request.IsActive,\n"
        "    };\n"
        "    db.Parents.Add(parent);\n"
        "    await db.SaveChangesAsync(ct);\n"
        "    return OperationResult<Guid>.Success(parent.Id);\n"
        "}")

    add_para(doc,
        "Until sprint 9, the admin's workflow to actually get a parent "
        "into the portal was: create the parent, then go to Users → "
        "New user and create the matching ApplicationUser in the "
        "Parent role, then edit the parent (or run a hand-written SQL "
        "update) to set Parent.UserId. The sprint 8 portal-fallback "
        "card existed precisely because this final step was so easy "
        "to forget.")

    add_page_break(doc)


# ---------- Chapter 3 -------------------------------------------------------

def chapter3_dtos(doc):
    add_heading(doc, "3. DTO changes — capturing credentials", 1)
    add_para(doc,
        "The first layer to change is the Application project: the two "
        "Create*Request DTOs grow the fields needed to provision a "
        "login. We deliberately keep them on the existing Create*Request "
        "shape rather than introducing a wrapper DTO, because the new "
        "fields are required to satisfy the acceptance criteria — there "
        "is no scenario in sprint 9 where a parent or student is created "
        "without a user.")

    add_heading(doc, "3.1 CreateParentRequest", 2)
    add_para(doc,
        "Email becomes required (Identity is configured with "
        "RequireUniqueEmail = true). Two new fields — UserName and "
        "Password — are added with the same length and password-strength "
        "validators the rest of the system uses.")

    add_file(doc,
             "src/NaijaPrimeSchool.Application/Family/Dtos/ParentDtos.cs")

    add_heading(doc, "3.2 CreateStudentRequest", 2)
    add_para(doc,
        "Student gains three new fields — UserName, Email, Password — "
        "because the previous student form was happy to capture only "
        "admission, demographic and health data. The student's email "
        "is required for the same reason as the parent's: it is the "
        "Identity user's unique email.")

    add_file(doc,
             "src/NaijaPrimeSchool.Application/Family/Dtos/StudentDtos.cs")

    add_para(doc,
        "Notice that neither DTO grows a 'roles' list the way the "
        "sprint-1 CreateUserRequest did. Sprint 9 enforces the role at "
        "the service layer because the role is implicit in the entry "
        "point — creating a parent always means 'in the Parent role'; "
        "creating a student always means 'in the Student role'. If the "
        "school needs a parent who is also a SchoolBursar (rare but "
        "possible), the SuperAdmin can extend the role list on the "
        "Users → Roles page after the fact.")

    add_page_break(doc)


# ---------- Chapter 4 -------------------------------------------------------

def chapter4_services(doc):
    add_heading(doc, "4. Service changes — provisioning the user", 1)
    add_para(doc,
        "The service implementations grow a UserManager dependency and "
        "an ICurrentUser dependency (the latter so the CreatedBy audit "
        "column reflects the SuperAdmin / HeadTeacher who pressed the "
        "button). The actual flow is the same on both sides: validate "
        "the request, check that the username / email are free, create "
        "the ApplicationUser, add the appropriate role, and only then "
        "create the Parent or Student row with UserId stamped.")

    add_heading(doc, "4.1 ParentService", 2)
    add_para(doc,
        "The new CreateAsync builds an ApplicationUser from the same "
        "form fields it would have stamped onto Parent — first name, "
        "last name, middle name, title, gender, address. Setting "
        "EmailConfirmed = true lets the new account log in immediately "
        "without going through an email confirmation token (the school "
        "office knows the addresses are real). PhoneNumber on the "
        "ApplicationUser is set from the parent's primary phone so the "
        "user's profile is not empty on day one. CreatedBy is "
        "ICurrentUser.UserName so the audit trail captures which "
        "SuperAdmin / HeadTeacher provisioned the account.")
    add_para(doc,
        "If UserManager.CreateAsync fails — most likely because the "
        "password did not meet the configured strength rules — the "
        "Identity error descriptions are returned verbatim to the "
        "caller. If the role assignment fails after the user was "
        "created (extremely unlikely, but the role could be missing in "
        "a developer database), the half-built user is deleted to "
        "leave the database clean.")
    add_para(doc,
        "Only after the user and role exist does the service add the "
        "Parent row with UserId = user.Id and call SaveChangesAsync. "
        "Identity has already committed the user in its own internal "
        "save, so the Parent insert is the last write of the operation.")

    add_file(doc,
             "src/NaijaPrimeSchool.Infrastructure/Services/ParentService.cs")

    add_heading(doc, "4.2 StudentService", 2)
    add_para(doc,
        "The student service follows the same recipe. Two things are "
        "slightly different from the parent path: the existing validation "
        "of admission number, date-of-birth-before-admission-date and "
        "initial class is preserved (and runs before any user work, so a "
        "validation failure does not leave a dangling Identity user), and "
        "the student's DateOfBirth is a DateOnly that is normalised to a "
        "DateTime when stamping the ApplicationUser.")

    add_file(doc,
             "src/NaijaPrimeSchool.Infrastructure/Services/StudentService.cs")

    add_heading(doc, "4.3 Why no explicit transaction?", 2)
    add_para(doc,
        "Identity's UserManager talks to the same ApplicationDbContext "
        "the service uses. That means the inserts happen sequentially "
        "but inside the same context. If the Parent / Student insert "
        "throws after the user is committed, the user is still there — "
        "this is the same exposure every previous sprint has had with "
        "Identity. The pre-creation duplicate-name / duplicate-email "
        "checks and the post-creation role-failure cleanup remove the "
        "most likely failure modes; a database outage between the "
        "user save and the parent / student save would be observable "
        "as 'a user exists but cannot sign in to the portal' and is "
        "rare enough to fix by hand. Wrapping the whole thing in an "
        "IDbContextTransaction is left as a future hardening exercise.")

    add_page_break(doc)


# ---------- Chapter 5 -------------------------------------------------------

def chapter5_ui(doc):
    add_heading(doc, "5. UI changes — capturing credentials on the forms", 1)
    add_para(doc,
        "The two New* Razor pages each grow a 'Portal sign-in' "
        "section beneath the existing form. The model behind the form "
        "grows three new fields with their own data annotations so "
        "client-side validation surfaces the required-and-strong "
        "rules before the request reaches the service.")

    add_heading(doc, "5.1 CreateParent.razor", 2)
    add_file(doc,
             "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateParent.razor")

    add_heading(doc, "5.2 CreateStudent.razor", 2)
    add_para(doc,
        "The student form already collected admission details, "
        "demographics, health information and an optional initial "
        "enrolment. The new 'Portal sign-in' block joins them as the "
        "final section, so the SuperAdmin / HeadTeacher fills the "
        "form top-to-bottom and ends with the credentials.")

    add_file(doc,
             "src/NaijaPrimeSchool.Web/Components/Pages/Family/CreateStudent.razor")

    add_page_break(doc)


# ---------- Chapter 6 -------------------------------------------------------

def chapter6_test(doc):
    add_heading(doc, "6. How to test end-to-end", 1)
    add_para(doc,
        "Run the app against a fresh database (or one that already has "
        "sprint 1's SuperAdmin seeded) and walk through the two new "
        "flows. The portals should resolve on first sign-in.")

    add_heading(doc, "6.1 Parent flow", 2)
    add_numbered(doc, [
        "Sign in as superadmin@naijaprimeschool.ng / Admin@12345.",
        "Navigate to Family → New parent (/parents/new).",
        "Fill in the profile (First name, Last name, Email, etc.).",
        "Scroll to the 'Portal sign-in' section and pick a Username + "
        "Initial password. The password must satisfy the rules listed "
        "under the field.",
        "Save. The notification banner reads 'Parent <name> created.' "
        "and you are redirected to /parents/<id>.",
        "Sign out and sign back in with the new parent's username "
        "and password.",
        "The sidebar shows the Parent portal panel. Click 'My wards'. "
        "If the parent is already linked to one or more pupils via "
        "the sprint-3 StudentParent join table, each ward renders as "
        "a card; otherwise the page prompts the office to link them.",
    ])

    add_heading(doc, "6.2 Student flow", 2)
    add_numbered(doc, [
        "Sign in as superadmin@naijaprimeschool.ng / Admin@12345.",
        "Navigate to Family → New student (/students/new).",
        "Fill in the admission, demographic, health and (optionally) "
        "initial-enrolment sections as before.",
        "Scroll to the 'Portal sign-in' section and pick a Username, "
        "Email and Initial password.",
        "Save. The notification banner reads 'Student <name> created.' "
        "and you are redirected to /students/<id>.",
        "Sign out and sign back in with the new student's username "
        "and password.",
        "The sidebar shows the Student portal panel. Click 'Today'. "
        "/portal/student renders the dashboard tiles (class, "
        "outstanding fees, attendance, report cards, unread "
        "announcements) and today's timetable.",
    ])

    add_heading(doc, "6.3 Database verification", 2)
    add_para(doc,
        "If you want to confirm the data shape from the database "
        "directly, three SQL queries are useful:")

    add_caption(doc, "Inspect the new parent's user row")
    add_code(doc,
        "SELECT u.UserName, u.Email, p.FirstName, p.LastName, p.UserId, u.Id\n"
        "FROM Parents p\n"
        "JOIN AspNetUsers u ON u.Id = p.UserId\n"
        "WHERE p.FirstName = '<first>' AND p.LastName = '<last>';")

    add_caption(doc, "Confirm the Parent role is assigned")
    add_code(doc,
        "SELECT u.UserName, r.Name AS RoleName\n"
        "FROM AspNetUsers u\n"
        "JOIN AspNetUserRoles ur ON ur.UserId = u.Id\n"
        "JOIN AspNetRoles r ON r.Id = ur.RoleId\n"
        "WHERE u.UserName = '<username>';")

    add_caption(doc, "Same shape for the Student side")
    add_code(doc,
        "SELECT u.UserName, u.Email, s.AdmissionNumber, s.FirstName, s.LastName, s.UserId\n"
        "FROM Students s\n"
        "JOIN AspNetUsers u ON u.Id = s.UserId\n"
        "WHERE s.AdmissionNumber = '<admission>';")

    add_page_break(doc)


# ---------- Chapter 7 -------------------------------------------------------

def chapter7_followups(doc):
    add_heading(doc, "7. Follow-ups & known limitations", 1)
    add_bullets(doc, [
        "Existing Parent / Student rows that pre-date sprint 9 still "
        "have UserId = NULL. The school office can either link them by "
        "hand (Users → create user, edit Parent/Student to set "
        "UserId), or soft-delete them and re-create through the new "
        "flow.",
        "There is no UI to update the parent's or student's username "
        "or password from inside the Parent or Student edit page. "
        "Continue to use the SuperAdmin Users area for that.",
        "The 'sign-in for a young pupil' question is left to the school "
        "to decide policy on — junior pupils may never use the portal "
        "directly, but the account still exists for parental-association "
        "purposes and so that older pupils inherit the account when they "
        "are ready.",
        "An explicit IDbContextTransaction around the user-then-row "
        "save would tighten the atomicity guarantee and is a sensible "
        "future hardening.",
        "Email is now required on Parent. Pre-sprint-9 parent rows "
        "without an email are unaffected — only the create form "
        "enforces it.",
    ])

    add_para(doc,
        "With sprint 9 in place, every load-bearing seeded role in the "
        "system now has a workspace that lights up the moment a real "
        "human is given credentials. The next sprint can move on to the "
        "items on the roadmap (notifications, two-way messaging, "
        "online fee payment, the audit-log viewer) knowing that the "
        "portals are no longer a paper-only feature.")


# ---------- Build ------------------------------------------------------------

def build():
    doc = Document()
    configure_document(doc)
    title_page(doc)
    toc_page(doc)
    chapter1_overview(doc)
    chapter2_pre_sprint(doc)
    chapter3_dtos(doc)
    chapter4_services(doc)
    chapter5_ui(doc)
    chapter6_test(doc)
    chapter7_followups(doc)
    out_path = ROOT / OUTPUT
    doc.save(str(out_path))
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    build()
